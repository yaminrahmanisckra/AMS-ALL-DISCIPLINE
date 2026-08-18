"""Curriculator routes: syllabus documents, parts, course entries, assignments, export."""
import os
import tempfile
from flask import render_template, request, redirect, url_for, flash, jsonify, send_file, current_app
from flask_login import login_required, current_user
from io import BytesIO
from datetime import datetime

from extensions import db
from utils.tenant import current_tenant, load_curriculator_pack, infer_year_term_from_code
from . import curriculator_bp
from .models import (
    SyllabusDocument,
    SyllabusPart,
    SyllabusCourseEntry,
    SyllabusAuthorAssignment,
    CurriculatorEditor,
    SyllabusPartASection,
    SyllabusPartBConfig,
    SyllabusPartDSection,
    SyllabusSectionAssignment,
)

# Lazy imports to avoid circular deps
def _teacher():
    from blueprints.class_management.models import Teacher
    return Teacher

def _course_model():
    from blueprints.course_management.models import Course
    return Course


def _effective_roles():
    from role_utils import get_effective_roles
    return set(get_effective_roles(current_user))


def _is_head():
    """Only Head can manage who has add/remove permission."""
    if not current_user.is_authenticated:
        return False
    return 'head' in _effective_roles()


def _can_add_remove_syllabus():
    """Head, or users Head has granted permission via Curriculator editors."""
    if not current_user.is_authenticated:
        return False
    if 'head' in _effective_roles():
        return True
    exists = CurriculatorEditor.query.filter_by(user_id=current_user.id).first()
    return exists is not None


def _can_edit_syllabus():
    """Head, dean, teaching_assistant, or CurriculatorEditor (Head-granted full edit) can edit all content."""
    if not current_user.is_authenticated:
        return False
    roles = _effective_roles()
    if 'head' in roles or 'dean' in roles or 'teaching_assistant' in roles:
        return True
    return CurriculatorEditor.query.filter_by(user_id=current_user.id).first() is not None


def _can_assign_authors():
    return _can_edit_syllabus()


def _can_edit_part_a_section(part_id, section_key):
    """True if Head/TA or user is assigned as section owner."""
    if not current_user.is_authenticated:
        return False
    if _can_edit_syllabus():
        return True
    a = SyllabusSectionAssignment.query.filter_by(part_id=part_id, section_key=section_key).first()
    return a and a.user_id == current_user.id


# Part A section keys and types (key_value | text | serial_description | peos | plos | mapping | mapping_course_plo)
# mapping_course_plo section_key suffix -> (year, term) for filtering Part C entries
def _curriculator_pack():
    return load_curriculator_pack() or {}


def _syllabus_parts_spec():
    parts = _curriculator_pack().get('parts') or [
        {'key': 'A', 'title': 'Part A'},
        {'key': 'B', 'title': 'Part B'},
        {'key': 'C', 'title': 'Part C'},
        {'key': 'D', 'title': 'Part D'},
    ]
    return parts


def _year_term_grid():
    grid = _curriculator_pack().get('year_term_grid') or []
    return [(str(y), str(t)) for y, t in grid]


def _part_a_sections():
    sections = _curriculator_pack().get('part_a_sections')
    if sections:
        return sections
    return PART_A_SECTIONS_FALLBACK


def _mapping_course_plo_year_term():
    raw = _curriculator_pack().get('mapping_course_plo_year_term') or {}
    result = {}
    for key, val in raw.items():
        if isinstance(val, (list, tuple)) and len(val) >= 2:
            result[str(key)] = (str(val[0]), str(val[1]))
    return result or MAPPING_COURSE_PLO_YEAR_TERM_FALLBACK


MAPPING_COURSE_PLO_YEAR_TERM_FALLBACK = {
    '1_1': ('First', 'First'),
    '1_2': ('First', 'Second'),
    '2_1': ('Second', 'First'),
    '2_2': ('Second', 'Second'),
    '3_1': ('Third', 'First'),
    '4_1': ('Fourth', 'First'),
    '4_2': ('Fourth', 'Second'),
}

PART_A_SECTIONS_FALLBACK = [
    {'key': 'overview', 'label': 'Program Overview', 'type': 'key_value'},
    {'key': 'vision_university', 'label': 'Vision of the University', 'type': 'text'},
    {'key': 'vision_discipline', 'label': 'Vision of the Discipline', 'type': 'text'},
    {'key': 'mission_university', 'label': 'Mission of the University', 'type': 'serial_description'},
    {'key': 'mission_discipline', 'label': 'Mission of the Discipline', 'type': 'serial_description'},
    {'key': 'objectives', 'label': 'Objectives of the Discipline', 'type': 'serial_description'},
    {'key': 'description', 'label': 'Description of the Program', 'type': 'text'},
    {'key': 'graduate_attributes', 'label': 'Graduate Attributes', 'type': 'graduate_attributes'},
    {'key': 'peos', 'label': 'Program Educational Objectives (PEOs)', 'type': 'peos'},
    {'key': 'plos', 'label': 'Program Learning Outcomes (PLOs)', 'type': 'plos'},
    {'key': 'mapping_mission_peo', 'label': 'Mapping Mission ↔ PEO', 'type': 'mapping'},
    {'key': 'mapping_plo_peo', 'label': 'Mapping PLO ↔ PEO', 'type': 'mapping'},
    {'key': 'mapping_course_plo_1_1', 'label': 'Mapping Course ↔ PLO (First Year First Term)', 'type': 'mapping_course_plo'},
    {'key': 'mapping_course_plo_1_2', 'label': 'Mapping Course ↔ PLO (First Year Second Term)', 'type': 'mapping_course_plo'},
    {'key': 'mapping_course_plo_2_1', 'label': 'Mapping Course ↔ PLO (Second Year First Term)', 'type': 'mapping_course_plo'},
    {'key': 'mapping_course_plo_2_2', 'label': 'Mapping Course ↔ PLO (Second Year Second Term)', 'type': 'mapping_course_plo'},
    {'key': 'mapping_course_plo_3_1', 'label': 'Mapping Course ↔ PLO (Third Year First Term)', 'type': 'mapping_course_plo'},
    {'key': 'mapping_course_plo_4_1', 'label': 'Mapping Course ↔ PLO (Fourth Year First Term)', 'type': 'mapping_course_plo'},
    {'key': 'mapping_course_plo_4_2', 'label': 'Mapping Course ↔ PLO (Fourth Year Second Term)', 'type': 'mapping_course_plo'},
]


@curriculator_bp.route('/')
@login_required
def index():
    """List syllabus documents."""
    docs = SyllabusDocument.query.order_by(SyllabusDocument.updated_at.desc()).all()
    can_edit = _can_edit_syllabus()
    can_add_remove = _can_add_remove_syllabus()
    is_head = _is_head()
    return render_template(
        'curriculator/index.html',
        documents=docs,
        can_edit=can_edit,
        can_add_remove=can_add_remove,
        is_head=is_head,
    )


@curriculator_bp.route('/doc/<int:doc_id>')
@login_required
def document_detail(doc_id):
    """View a syllabus document; tabs for Part A, B, C, D."""
    doc = SyllabusDocument.query.get_or_404(doc_id)
    parts = SyllabusPart.query.filter_by(document_id=doc_id).order_by(SyllabusPart.sort_order).all()
    part_map = {p.part_key: p for p in parts}
    can_edit = _can_edit_syllabus()
    can_assign = _can_assign_authors()

    # Part C: course entries for list view
    part_c = part_map.get('C')
    course_entries = []
    if part_c:
        course_entries = SyllabusCourseEntry.query.filter_by(part_id=part_c.id).order_by(
            SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id
        ).all()

    return render_template(
        'curriculator/document_detail.html',
        document=doc,
        parts=parts,
        part_map=part_map,
        course_entries=course_entries,
        can_edit=can_edit,
        can_assign=can_assign,
    )


@curriculator_bp.route('/doc/<int:doc_id>/part/<part_key>')
@login_required
def part_view(doc_id, part_key):
    """View/edit a single part (A, B, C, or D)."""
    doc = SyllabusDocument.query.get_or_404(doc_id)
    part = SyllabusPart.query.filter_by(document_id=doc_id, part_key=part_key).first()
    if not part:
        flash(f'Part {part_key} not found.', 'warning')
        return redirect(url_for('curriculator.document_detail', doc_id=doc_id))
    can_edit = _can_edit_syllabus()
    can_assign = _can_assign_authors()
    is_head = _is_head()
    course_entries = []
    part_a_section_list = []
    part_a_eligible_users = []

    part_b_config = None
    part_b_structure_table = None
    part_b_area_wise = []
    part_b_category = []
    part_b_course_distribution = []
    part_b_derived_area = []
    part_b_derived_category = []

    part_d_grading_scale = []
    part_d_theory_eval = []
    part_d_sessional_eval = []
    part_d_approval_records = []

    if part_key == 'C':
        course_entries = SyllabusCourseEntry.query.filter_by(part_id=part.id).order_by(
            SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id
        ).all()
    elif part_key == 'B':
        part_c = SyllabusPart.query.filter_by(document_id=doc.id, part_key='C').first()
        part_c_entries = []
        if part_c:
            part_c_entries = SyllabusCourseEntry.query.filter_by(part_id=part_c.id).order_by(
                SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id
            ).all()
        part_b_derived_area, part_b_derived_category = _derive_part_b_from_c(part_c_entries)
        part_b_course_distribution = _course_distribution_by_term(part_c_entries)
        rec = SyllabusPartBConfig.query.filter_by(part_id=part.id).first()
        if rec:
            cfg = rec.get_config()
            part_b_config = dict(cfg)
            part_b_structure_table = _normalize_structure_curriculum_for_template(cfg)
            oa = cfg.get('area_wise_override')
            oc = cfg.get('category_override')
            part_b_area_wise = oa if isinstance(oa, list) else part_b_derived_area
            part_b_category = oc if isinstance(oc, list) else part_b_derived_category
        else:
            part_b_config = {
                'duration_years': 4,
                'terms': 'Terms',
                'term_duration': {'teaching': 14, 'preparatory': 1, 'exam': 2, 'break': 1, 'total': 18},
            }
            part_b_structure_table = _normalize_structure_curriculum_for_template({})
            part_b_area_wise = part_b_derived_area
            part_b_category = part_b_derived_category
    elif part_key == 'A':
        sections_by_key = {s.section_key: s for s in SyllabusPartASection.query.filter_by(part_id=part.id).all()}
        assign_by_key = {}
        for a in SyllabusSectionAssignment.query.filter_by(part_id=part.id).all():
            assign_by_key[a.section_key] = a
        from user_models import User
        for cfg in _part_a_sections():
            sk = cfg['key']
            sec = sections_by_key.get(sk)
            a = assign_by_key.get(sk)
            owner_name = None
            if a:
                u = User.query.get(a.user_id)
                owner_name = u.full_name if u else None
            part_a_section_list.append({
                'key': sk,
                'label': cfg['label'],
                'type': cfg['type'],
                'has_data': sec is not None and sec.data,
                'can_edit': _can_edit_part_a_section(part.id, sk),
                'owner_name': owner_name,
                'assignment': a,
            })
        if can_assign:
            from role_utils import get_effective_roles
            for u in User.query.order_by(User.full_name).all():
                roles = set(get_effective_roles(u))
                if roles & {'head', 'dean', 'teacher', 'teaching_assistant'}:
                    part_a_eligible_users.append({'id': u.id, 'full_name': u.full_name, 'username': u.username})
    elif part_key == 'D':
        for sk in ('grading_scale', 'theory_evaluation', 'sessional_evaluation', 'approval_records'):
            rec = SyllabusPartDSection.query.filter_by(part_id=part.id, section_key=sk).first()
            d = rec.get_data() if rec and rec.get_data() else []
            arr = d if isinstance(d, list) else []
            if sk == 'grading_scale':
                part_d_grading_scale = arr
            elif sk == 'theory_evaluation':
                part_d_theory_eval = arr
            elif sk == 'sessional_evaluation':
                part_d_sessional_eval = arr
            elif sk == 'approval_records':
                part_d_approval_records = arr

    return render_template(
        'curriculator/part_view.html',
        document=doc,
        part=part,
        course_entries=course_entries,
        can_edit=can_edit,
        can_assign=can_assign,
        is_head=is_head,
        part_a_sections=_part_a_sections(),
        part_a_section_list=part_a_section_list,
        part_a_eligible_users=part_a_eligible_users,
        part_b_config=part_b_config,
        part_b_structure_table=part_b_structure_table,
        part_b_area_wise=part_b_area_wise,
        part_b_category=part_b_category,
        part_b_course_distribution=part_b_course_distribution,
        part_b_derived_area=part_b_derived_area,
        part_b_derived_category=part_b_derived_category,
        part_d_grading_scale=part_d_grading_scale,
        part_d_theory_eval=part_d_theory_eval,
        part_d_sessional_eval=part_d_sessional_eval,
        part_d_approval_records=part_d_approval_records,
    )


def _parse_docx_for_import(file_storage):
    """Parse uploaded DOCX; return dict with name, part_a, part_b, part_d, part_c_entries."""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from io import BytesIO

    buf = BytesIO(file_storage.read())
    doc = Document(buf)
    current_part = None
    out = {'name': (file_storage.filename or 'imported').replace('.docx', '').replace('.DOCX', '')[:80],
           'part_a': {}, 'part_b': {}, 'part_d': {}, 'part_c_entries': []}

    def cell_text(c):
        return (c.text or '').strip()

    def table_rows(tbl):
        return [[cell_text(c) for c in row.cells] for row in tbl.rows]

    for block in doc.element.body:
        tag = block.tag.split('}')[-1] if '}' in block.tag else block.tag
        if tag == 'p':
            p = Paragraph(block, doc)
            t = (p.text or '').strip().upper()
            for label in ('PART A', 'PART B', 'PART C', 'PART D'):
                if label in t:
                    current_part = label.split()[-1]
                    break
        elif tag == 'tbl':
            tbl = Table(block, doc)
            rows = table_rows(tbl)
            if not rows:
                continue
            nr, nc = len(rows), max(len(r) for r in rows) if rows else 0
            if current_part == 'A' and nc >= 2:
                if nr == 1 and nc >= 2:
                    out['part_a'].setdefault('overview', []).append({'key': rows[0][0], 'value': rows[0][1]})
                elif nr >= 2 and nc >= 2:
                    arr = [{'key': (r[0] if len(r) > 0 else ''), 'value': (r[1] if len(r) > 1 else '')} for r in rows[1:] if (r[0] if len(r) > 0 else '') or (r[1] if len(r) > 1 else '')]
                    if arr and not out['part_a'].get('overview'):
                        out['part_a']['overview'] = arr
                elif nr >= 2 and nc >= 3:
                    k0 = (rows[0][0] or '').upper()
                    if 'SERIAL' in k0 and 'DESCRIPTION' in k0:
                        out['part_a']['peos'] = [{'serial': r[0] if len(r) > 0 else '', 'description': r[1] if len(r) > 1 else '', 'domains': r[2] if len(r) > 2 else ''} for r in rows[1:] if (r[0] if len(r) > 0 else '') or (r[1] if len(r) > 1 else '')]
                    elif 'ID' in k0 or 'PLO' in k0:
                        out['part_a']['plos'] = [{'id': r[0] if len(r) > 0 else '', 'text': r[1] if len(r) > 1 else ''} for r in rows[1:] if (r[0] if len(r) > 0 else '') or (r[1] if len(r) > 1 else '')]
            elif current_part == 'B' and nr >= 2:
                def safe_b(r, i):
                    return (r[i] if len(r) > i else '').strip()
                r0, r1 = (rows[0][0] or '').strip().upper(), (rows[1][0] or '').strip().upper() if nr > 1 else ''
                # Structure of the Curriculum: 8x3 table (not Term duration table)
                if nc >= 3 and ('STRUCTURE' in r0 and 'CURRICULUM' in r0 or 'DURATION' in r0 and 'PROGRAM' in r0 or 'ADMISSION' in r0 or ('AVAILABLE' in r0 and 'CREDIT' in r0) or
                               'DURATION' in r1 and 'PROGRAM' in r1 or 'ADMISSION' in r1 or ('AVAILABLE' in r1 and 'CREDIT' in r1)) and \
                               'TEACHING' not in r0 and 'PREPARATORY' not in (rows[0][1] or '').upper():
                    structure_rows = []
                    start = 1 if ('STRUCTURE' in r0 and 'CURRICULUM' in r0 and not safe_b(rows[0], 1)) else 0
                    for r in rows[start:]:
                        lbl, v1, v2 = safe_b(r, 0), safe_b(r, 1), safe_b(r, 2)
                        if not lbl and not v1 and not v2:
                            continue
                        if 'STRUCTURE' in lbl.upper() and 'CURRICULUM' in lbl.upper() and not v1 and not v2:
                            continue
                        structure_rows.append({'label': lbl, 'value_1': v1, 'value_2': v2})
                        if 'DURATION' in lbl.upper() and 'PROGRAM' in lbl.upper() and v1:
                            import re
                            m = re.match(r'^(\d+)\s*years?', v1, re.I)
                            if m:
                                out['part_b']['duration_years'] = int(m.group(1))
                            if v2:
                                out['part_b']['terms'] = v2
                    if structure_rows:
                        out['part_b']['structure_curriculum'] = structure_rows
                # Term duration: 3x5, row2 = "Teaching and Learning | Preparatory Leave | ...", row3 = "14 Weeks | 2 Weeks | ..."
                if nc >= 5 and ('TERM DURATION' in r0 or 'TEACHING' in r0 or 'TEACHING' in (rows[0][1] or '').upper()):
                    def parse_weeks(s):
                        import re
                        m = re.search(r'(\d+)', (s or ''))
                        return int(m.group(1)) if m else 0
                    if nr >= 3:
                        r2 = rows[2] if len(rows) > 2 else []
                        out['part_b']['term_duration'] = {
                            'teaching': parse_weeks(safe_b(r2, 0)),
                            'preparatory': parse_weeks(safe_b(r2, 1)),
                            'exam': parse_weeks(safe_b(r2, 2)),
                            'break': parse_weeks(safe_b(r2, 3)),
                            'total': parse_weeks(safe_b(r2, 4)),
                        }
                    if not out['part_b'].get('term_duration') or not out['part_b']['term_duration'].get('total'):
                        out['part_b']['term_duration'] = {'teaching': 14, 'preparatory': 1, 'exam': 2, 'break': 1, 'total': 18}
                # Area-wise Credit Distribution: header "Area | Type | Number of Courses | Credits | Total Credit"
                if nc >= 5 and ('AREA' in (rows[0][0] or '').upper() and 'TYPE' in (rows[0][1] or '').upper() and 'NUMBER' in (rows[0][2] or '').upper()):
                    area_wise = []
                    for r in rows[1:]:
                        area, typ, num, cred, total = safe_b(r, 0), safe_b(r, 1), safe_b(r, 2), safe_b(r, 3), safe_b(r, 4)
                        if not area and not typ:
                            continue
                        try:
                            area_wise.append({
                                'area': area, 'type': typ,
                                'num_courses': int(num) if num else 0,
                                'credits': float(cred) if cred else 0,
                                'total_credit': float(total) if total else 0,
                            })
                        except (ValueError, TypeError):
                            area_wise.append({'area': area, 'type': typ, 'num_courses': 0, 'credits': 0, 'total_credit': 0})
                    if area_wise:
                        out['part_b']['area_wise_override'] = area_wise
                # Category of Courses: "Area | Course Type | Course Title | Credit"
                if nc >= 4 and ('AREA' in (rows[0][0] or '').upper() and 'COURSE TYPE' in (rows[0][1] or '').upper() and 'COURSE TITLE' in (rows[0][2] or '').upper()):
                    category = []
                    for r in rows[1:]:
                        area, ctype, title, cred = safe_b(r, 0), safe_b(r, 1), safe_b(r, 2), safe_b(r, 3)
                        if not area and not ctype and not title:
                            continue
                        try:
                            category.append({
                                'area': area or '', 'course_type': ctype or '',
                                'course_title': title or '', 'credit': float(cred) if cred else 0,
                            })
                        except (ValueError, TypeError):
                            category.append({'area': area or '', 'course_type': ctype or '', 'course_title': title or '', 'credit': 0})
                    if category:
                        out['part_b']['category_override'] = category
            elif current_part == 'D' and nr >= 2 and nc >= 3:
                r0 = (rows[0][0] or '').upper()
                def safe(r, i):
                    return (r[i] if len(r) > i else '')
                if 'NUMERICAL' in r0 or 'LETTER' in r0 or 'GRADE' in r0:
                    out['part_d']['grading_scale'] = [{'numerical': safe(r, 0), 'letter': safe(r, 1), 'grade_point': safe(r, 2)} for r in rows[1:] if safe(r, 0) or safe(r, 1)]
                elif 'SL' in r0 or 'ITEMS' in r0:
                    row_list = [{'sl_no': safe(r, 0), 'items': safe(r, 1), 'marks': safe(r, 2)} for r in rows[1:] if safe(r, 0) or safe(r, 1)]
                    if 'theory_evaluation' not in out['part_d']:
                        out['part_d']['theory_evaluation'] = row_list
                    else:
                        out['part_d']['sessional_evaluation'] = row_list

    return out


@curriculator_bp.route('/doc/import', methods=['GET', 'POST'])
@login_required
def document_import():
    """Import syllabus from DOCX. Creates new document and pre-fills Part A/B/D (and Part C where detected)."""
    if not _can_add_remove_syllabus():
        flash('You do not have permission to import. Only Head or assigned editors can add syllabus.', 'error')
        return redirect(url_for('curriculator.index'))
    if request.method == 'GET':
        return render_template('curriculator/document_import.html')
    f = request.files.get('file')
    if not f or not f.filename or not f.filename.lower().endswith('.docx'):
        flash('Upload a .docx file.', 'error')
        return redirect(url_for('curriculator.document_import'))
    try:
        parsed = _parse_docx_for_import(f)
    except Exception as e:
        current_app.logger.exception('DOCX import parse failed')
        flash('Could not parse DOCX: %s' % str(e), 'error')
        return redirect(url_for('curriculator.document_import'))
    name = (request.form.get('name') or '').strip() or parsed['name']
    batches = (request.form.get('applicable_batches') or '').strip() or None
    doc = SyllabusDocument(name=name, applicable_batches=batches)
    db.session.add(doc)
    db.session.flush()
    _seed_syllabus_parts(doc.id)
    db.session.flush()
    part_map = {p.part_key: p for p in SyllabusPart.query.filter_by(document_id=doc.id).all()}
    part_a, part_b, part_d = part_map.get('A'), part_map.get('B'), part_map.get('D')
    if part_a and parsed.get('part_a'):
        for sk, data in parsed['part_a'].items():
            if not data:
                continue
            rec = SyllabusPartASection(part_id=part_a.id, section_key=sk)
            rec.set_data(data)
            db.session.add(rec)
    if part_b and parsed.get('part_b'):
        cfg = parsed['part_b']
        rec = SyllabusPartBConfig(part_id=part_b.id)
        rec.set_config(cfg)
        db.session.add(rec)
    if part_d and parsed.get('part_d'):
        for sk, data in parsed['part_d'].items():
            if not data:
                continue
            rec = SyllabusPartDSection(part_id=part_d.id, section_key=sk)
            rec.set_data(data)
            db.session.add(rec)
    db.session.commit()
    flash('Syllabus imported. Review and edit Part A/B/C/D as needed.', 'success')
    return redirect(url_for('curriculator.document_detail', doc_id=doc.id))


@curriculator_bp.route('/doc/create', methods=['GET', 'POST'])
@login_required
def document_create():
    """Create a new syllabus document."""
    if not _can_add_remove_syllabus():
        flash('You do not have permission to create syllabus documents. Only Head or assigned editors can add.', 'error')
        return redirect(url_for('curriculator.index'))
    if request.method == 'POST':
        name = (request.form.get('name') or '').strip()
        batches = (request.form.get('applicable_batches') or '').strip() or None
        if not name:
            flash('Name is required.', 'error')
            return redirect(url_for('curriculator.document_create'))
        doc = SyllabusDocument(name=name, applicable_batches=batches)
        db.session.add(doc)
        db.session.flush()
        _seed_syllabus_parts(doc.id)
        db.session.commit()
        flash(f'Syllabus "{name}" created. Add content to each part.', 'success')
        return redirect(url_for('curriculator.document_detail', doc_id=doc.id))
    return render_template('curriculator/document_create.html')


@curriculator_bp.route('/doc/<int:doc_id>/delete', methods=['POST'])
@login_required
def document_delete(doc_id):
    """Delete a syllabus document (and its parts, entries, assignments)."""
    if not _can_add_remove_syllabus():
        flash('You do not have permission to delete syllabus documents. Only Head or assigned editors can remove.', 'error')
        return redirect(url_for('curriculator.index'))
    doc = SyllabusDocument.query.get_or_404(doc_id)
    name = doc.name
    db.session.delete(doc)
    db.session.commit()
    flash(f'Syllabus "{name}" removed.', 'success')
    return redirect(url_for('curriculator.index'))


@curriculator_bp.route('/permissions')
@login_required
def permissions():
    """Head only: manage who has full edit permission (Curriculator editors)."""
    if not _is_head():
        flash('Only Head can manage full edit permissions.', 'error')
        return redirect(url_for('curriculator.index'))
    from user_models import User
    from role_utils import get_effective_roles
    editors = CurriculatorEditor.query.order_by(CurriculatorEditor.created_at.desc()).all()
    editor_user_ids = {e.user_id for e in editors}
    head_role_ids = set()
    for u in User.query.all():
        roles = set(get_effective_roles(u))
        if 'head' in roles:
            head_role_ids.add(u.id)
    # Eligible: teachers only (exclude teaching_assistant), not Head, not already editor
    eligible = []
    for u in User.query.order_by(User.full_name).all():
        if u.id in head_role_ids or u.id in editor_user_ids:
            continue
        roles = set(get_effective_roles(u))
        if 'teacher' in roles and 'teaching_assistant' not in roles:
            eligible.append(u)
    return render_template(
        'curriculator/permissions.html',
        editors=editors,
        eligible_users=eligible,
    )


@curriculator_bp.route('/permissions/add', methods=['POST'])
@login_required
def permissions_add():
    """Head only: grant full edit permission to a user."""
    if not _is_head():
        flash('Only Head can grant full edit permission.', 'error')
        return redirect(url_for('curriculator.index'))
    from user_models import User
    user_id = request.form.get('user_id')
    if not user_id:
        flash('Select a user.', 'error')
        return redirect(url_for('curriculator.permissions'))
    try:
        uid = int(user_id)
    except (TypeError, ValueError):
        flash('Invalid user.', 'error')
        return redirect(url_for('curriculator.permissions'))
    user = User.query.get(uid)
    if not user:
        flash('User not found.', 'error')
        return redirect(url_for('curriculator.permissions'))
    from role_utils import get_effective_roles
    roles = set(get_effective_roles(user))
    if 'head' in roles:
        flash('Head already has permission; no need to add.', 'info')
        return redirect(url_for('curriculator.permissions'))
    if 'teacher' not in roles or 'teaching_assistant' in roles:
        flash('Only teachers can be granted full edit permission.', 'error')
        return redirect(url_for('curriculator.permissions'))
    existing = CurriculatorEditor.query.filter_by(user_id=uid).first()
    if existing:
        flash(f'{user.full_name} already has full edit permission.', 'info')
        return redirect(url_for('curriculator.permissions'))
    e = CurriculatorEditor(user_id=uid)
    db.session.add(e)
    db.session.commit()
    flash(f'Full edit permission granted to {user.full_name}.', 'success')
    return redirect(url_for('curriculator.permissions'))


@curriculator_bp.route('/permissions/remove', methods=['POST'])
@login_required
def permissions_remove():
    """Head only: revoke full edit permission from a user."""
    if not _is_head():
        flash('Only Head can revoke full edit permission.', 'error')
        return redirect(url_for('curriculator.index'))
    user_id = request.form.get('user_id')
    if not user_id:
        flash('Select a user.', 'error')
        return redirect(url_for('curriculator.permissions'))
    try:
        uid = int(user_id)
    except (TypeError, ValueError):
        flash('Invalid user.', 'error')
        return redirect(url_for('curriculator.permissions'))
    rec = CurriculatorEditor.query.filter_by(user_id=uid).first()
    if not rec:
        flash('User is not in the editors list.', 'error')
        return redirect(url_for('curriculator.permissions'))
    from user_models import User
    u = User.query.get(uid)
    name = u.full_name if u else 'User'
    db.session.delete(rec)
    db.session.commit()
    flash(f'Full edit permission removed from {name}.', 'success')
    return redirect(url_for('curriculator.permissions'))


@curriculator_bp.route('/api/part/<int:part_id>/save', methods=['POST'])
@login_required
def part_save(part_id):
    """Save part content (A/B/D rich text)."""
    part = SyllabusPart.query.get_or_404(part_id)
    if not _can_edit_syllabus():
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json() or {}
    content = data.get('content')
    if content is not None:
        part.content = content
        part.updated_at = datetime.utcnow()
        db.session.commit()
    return jsonify({'success': True})


@curriculator_bp.route('/doc/<int:doc_id>/part/A/section/<section_key>')
@login_required
def part_a_section_edit(doc_id, section_key):
    """View/edit a single Part A section."""
    doc = SyllabusDocument.query.get_or_404(doc_id)
    part = SyllabusPart.query.filter_by(document_id=doc_id, part_key='A').first()
    if not part:
        flash('Part A not found.', 'warning')
        return redirect(url_for('curriculator.document_detail', doc_id=doc_id))
    cfg = next((c for c in _part_a_sections() if c['key'] == section_key), None)
    if not cfg:
        flash(f'Unknown section: {section_key}.', 'warning')
        return redirect(url_for('curriculator.part_view', doc_id=doc_id, part_key='A'))
    can_edit = _can_edit_part_a_section(part.id, section_key)
    rec = SyllabusPartASection.query.filter_by(part_id=part.id, section_key=section_key).first()
    data = rec.get_data() if rec else None
    part_c_entries = []
    peos_data = []
    plos_data = []
    mission_row_labels = []
    mission_um_labels = []

    plo_domain_groups = []  # [(domain_label, [plo_idx, ...]), ...] for mapping_course_plo header
    if cfg['type'] == 'mapping_course_plo':
        part_c = SyllabusPart.query.filter_by(document_id=doc_id, part_key='C').first()
        if part_c:
            all_entries = SyllabusCourseEntry.query.filter_by(part_id=part_c.id).order_by(
                SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id
            ).all()
            suffix = section_key.replace('mapping_course_plo_', '') if section_key.startswith('mapping_course_plo_') else ''
            yt = _mapping_course_plo_year_term().get(suffix)
            if yt:
                year, term = yt
                part_c_entries = [e for e in all_entries if (e.year or '').strip() == year and (e.term or '').strip() == term]
            else:
                part_c_entries = all_entries
        plo_rec = SyllabusPartASection.query.filter_by(part_id=part.id, section_key='plos').first()
        if plo_rec and plo_rec.get_data():
            d = plo_rec.get_data()
            plos_data = d if isinstance(d, list) else []
        # Group consecutive PLOs by category (domain) for table header
        if plos_data:
            cur_cat = None
            cur_indices = []
            for i, plo in enumerate(plos_data):
                cat = (plo.get('category') or '').strip()
                if cat != cur_cat:
                    if cur_indices:
                        plo_domain_groups.append((cur_cat or '—', cur_indices))
                    cur_cat = cat
                    cur_indices = [i]
                else:
                    cur_indices.append(i)
            if cur_indices:
                plo_domain_groups.append((cur_cat or '—', cur_indices))

    if cfg['type'] == 'mapping':
        peo_rec = SyllabusPartASection.query.filter_by(part_id=part.id, section_key='peos').first()
        if peo_rec and peo_rec.get_data():
            d = peo_rec.get_data()
            peos_data = d if isinstance(d, list) else []
        plo_rec = SyllabusPartASection.query.filter_by(part_id=part.id, section_key='plos').first()
        if plo_rec and plo_rec.get_data():
            d = plo_rec.get_data()
            plos_data = d if isinstance(d, list) else []
        if section_key == 'mapping_mission_peo':
            for sk in ('mission_university', 'mission_discipline'):
                sr = SyllabusPartASection.query.filter_by(part_id=part.id, section_key=sk).first()
                if sr and sr.get_data():
                    arr = sr.get_data()
                    if not isinstance(arr, list):
                        arr = []
                    for it in arr:
                        sn = (it.get('serial_no') or it.get('serial') or '').strip()
                        desc = (it.get('description') or '').strip()
                        mission_row_labels.append((sn + ' ' + desc).strip() or ('Item ' + str(len(mission_row_labels) + 1)))
            # University Missions only for column headers (UM 1, UM 2, UM 3)
            mu_rec = SyllabusPartASection.query.filter_by(part_id=part.id, section_key='mission_university').first()
            if mu_rec and mu_rec.get_data():
                arr = mu_rec.get_data()
                if isinstance(arr, list):
                    for i, it in enumerate(arr):
                        mission_um_labels.append((it.get('serial_no') or it.get('serial') or ('UM ' + str(i + 1))).strip() or ('UM ' + str(i + 1)))
            if not mission_um_labels:
                mission_um_labels = ['UM 1', 'UM 2', 'UM 3']

    return render_template(
        'curriculator/part_a_section_edit.html',
        document=doc,
        part=part,
        section_key=section_key,
        section_cfg=cfg,
        can_edit=can_edit,
        data=data,
        part_c_entries=part_c_entries,
        peos_data=peos_data,
        plos_data=plos_data,
        plo_domain_groups=plo_domain_groups,
        mission_row_labels=mission_row_labels,
        mission_um_labels=mission_um_labels,
    )


@curriculator_bp.route('/api/part/<int:part_id>/section/save', methods=['POST'])
@login_required
def part_section_save(part_id):
    """Save Part A (or D) section data. Permission: section owner or Head/TA."""
    part = SyllabusPart.query.get_or_404(part_id)
    data = request.get_json() or {}
    section_key = (data.get('section_key') or '').strip()
    if not section_key:
        return jsonify({'success': False, 'message': 'section_key required'}), 400
    if part.part_key == 'A':
        if not _can_edit_part_a_section(part_id, section_key):
            return jsonify({'success': False, 'message': 'Forbidden'}), 403
        cfg = next((c for c in _part_a_sections() if c['key'] == section_key), None)
        if not cfg:
            return jsonify({'success': False, 'message': 'Unknown section'}), 400
        payload = data.get('data')
        rec = SyllabusPartASection.query.filter_by(part_id=part_id, section_key=section_key).first()
        if not rec:
            rec = SyllabusPartASection(part_id=part_id, section_key=section_key)
            db.session.add(rec)
        rec.set_data(payload)
        rec.updated_at = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True})
    if part.part_key == 'D':
        if not _can_edit_syllabus():
            return jsonify({'success': False, 'message': 'Forbidden'}), 403
        payload = data.get('data')
        rec = SyllabusPartDSection.query.filter_by(part_id=part_id, section_key=section_key).first()
        if not rec:
            rec = SyllabusPartDSection(part_id=part_id, section_key=section_key)
            db.session.add(rec)
        rec.set_data(payload)
        rec.updated_at = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True})
    return jsonify({'success': False, 'message': 'Only Part A/D sections supported'}), 400


@curriculator_bp.route('/api/part/<int:part_id>/section/assign', methods=['POST'])
@login_required
def part_section_assign(part_id):
    """Assign section owner (Head, Dean, TA, or CurriculatorEditor)."""
    if not _can_edit_syllabus():
        return jsonify({'success': False, 'message': 'Only Head, Dean, TA, or full-edit editors can assign section owners'}), 403
    part = SyllabusPart.query.get_or_404(part_id)
    data = request.get_json() or {}
    section_key = (data.get('section_key') or '').strip()
    user_id = data.get('user_id')
    if not section_key:
        return jsonify({'success': False, 'message': 'section_key required'}), 400
    if user_id is None or user_id == '':
        # Revoke assignment
        rec = SyllabusSectionAssignment.query.filter_by(part_id=part_id, section_key=section_key).first()
        if rec:
            db.session.delete(rec)
            db.session.commit()
        return jsonify({'success': True})
    try:
        uid = int(user_id)
    except (TypeError, ValueError):
        return jsonify({'success': False, 'message': 'Invalid user_id'}), 400
    from user_models import User
    u = User.query.get(uid)
    if not u:
        return jsonify({'success': False, 'message': 'User not found'}), 400
    rec = SyllabusSectionAssignment.query.filter_by(part_id=part_id, section_key=section_key).first()
    if not rec:
        rec = SyllabusSectionAssignment(part_id=part_id, section_key=section_key, user_id=uid, assigned_by_id=current_user.id)
        db.session.add(rec)
    else:
        rec.user_id = uid
        rec.assigned_by_id = current_user.id
    db.session.commit()
    return jsonify({'success': True})


# Default rows for Part B "Structure of the Curriculum" table (LLB DOCX layout).
DEFAULT_STRUCTURE_CURRICULUM = [
    {'label': 'Duration of the Program', 'value_1': '4 Years', 'value_2': 'Terms'},
    {'label': 'Admission Requirements', 'value_1': '', 'value_2': ''},
    {'label': 'Available Credits', 'value_1': '170', 'value_2': '170'},
    {'label': 'Graduating Credits', 'value_1': '146', 'value_2': '146'},
    {'label': 'Total Class Weeks in a Term*', 'value_1': '14', 'value_2': '14'},
    {'label': 'Minimum CGPA Requirements for Graduation', 'value_1': '2.50', 'value_2': '2.50'},
    {'label': 'Maximum Years of Completion', 'value_1': '7 Years', 'value_2': '7 Years'},
]


def _normalize_structure_curriculum_for_template(cfg):
    """Return {column_headers: [...], rows: [{label, values: [...]}]} for template. Supports legacy list format."""
    sc = cfg.get('structure_curriculum')
    years = cfg.get('duration_years', 4)
    terms = (cfg.get('terms') or 'Terms').strip() or 'Terms'
    if isinstance(sc, dict) and 'column_headers' in sc and 'rows' in sc:
        return {'column_headers': sc['column_headers'], 'rows': sc['rows']}
    if isinstance(sc, list) and sc:
        headers = ['Value 1', 'Value 2']
        rows = []
        for r in sc:
            if isinstance(r, dict):
                vals = [r.get('value_1', ''), r.get('value_2', '')]
                if len(vals) < 2:
                    vals.extend([''] * (2 - len(vals)))
                rows.append({'label': r.get('label', ''), 'values': vals})
        return {'column_headers': headers, 'rows': rows}
    default = [dict(r) for r in DEFAULT_STRUCTURE_CURRICULUM]
    default[0]['value_1'] = f'{years} Years'
    default[0]['value_2'] = terms
    return {
        'column_headers': ['Value 1', 'Value 2'],
        'rows': [{'label': r.get('label', ''), 'values': [r.get('value_1', ''), r.get('value_2', '')]} for r in default]
    }


def _get_part_b_structure_curriculum(cfg):
    """Return structure_curriculum for config (legacy list or new dict). Used when no override."""
    sc = cfg.get('structure_curriculum')
    if isinstance(sc, dict) and 'column_headers' in sc and 'rows' in sc:
        return sc
    if isinstance(sc, list) and sc:
        return sc
    years = cfg.get('duration_years', 4)
    terms = (cfg.get('terms') or 'Terms').strip() or 'Terms'
    default = [dict(r) for r in DEFAULT_STRUCTURE_CURRICULUM]
    default[0]['value_1'] = f'{years} Years'
    default[0]['value_2'] = terms
    return default


def _derive_part_b_from_c(part_c_entries):
    """Derive area_wise and category from Part C entries. Returns (area_wise, category)."""
    from collections import defaultdict
    area_wise_map = defaultdict(lambda: {'count': 0, 'credits': 0})
    category = []
    for e in part_c_entries:
        area = (e.status or 'Core').strip()
        t = (e.entry_type or 'Theory').strip()
        key = (area, t)
        area_wise_map[key]['count'] += 1
        cr = float(e.credit) if e.credit is not None else 0
        area_wise_map[key]['credits'] += cr
        category.append({
            'area': area,
            'course_type': t,
            'course_title': (e.course_name or ''),
            'credit': float(e.credit) if e.credit is not None else 0,
        })
    area_wise = []
    for (area, t), v in sorted(area_wise_map.items()):
        area_wise.append({
            'area': area,
            'type': t,
            'num_courses': v['count'],
            'credits': v['credits'],
            'total_credit': v['credits'],
        })
    return area_wise, category


def _seed_syllabus_parts(doc_id):
    for i, spec in enumerate(_syllabus_parts_spec()):
        key = spec.get('key') or spec.get('part_key')
        if not key:
            continue
        title = spec.get('title') or f'Part {key}'
        db.session.add(SyllabusPart(
            document_id=doc_id, part_key=key, title=title, sort_order=spec.get('sort_order', i)
        ))


def _course_distribution_by_term(part_c_entries):
    """Group Part C entries by year-term. Returns dict (year, term) -> list of entries."""
    from collections import defaultdict
    groups = defaultdict(list)
    for e in part_c_entries:
        y = (e.year or '').strip()
        t = (e.term or '').strip()
        key = (y or '—', t or '—')
        groups[key].append(e)
    order = _year_term_grid() or [
        ('First', 'First'), ('First', 'Second'), ('Second', 'First'), ('Second', 'Second'),
        ('Third', 'First'), ('Fourth', 'First'), ('Fourth', 'Second'), ('LLM', 'First'), ('LLM', 'Second'),
    ]
    result = []
    seen = set()
    for y, t in order:
        k = (y, t)
        if k not in seen and k in groups:
            seen.add(k)
            result.append({'year': y, 'term': t, 'entries': groups[k]})
    for k, v in groups.items():
        if k not in seen:
            result.append({'year': k[0], 'term': k[1], 'entries': v})
    return result


@curriculator_bp.route('/api/part/<int:part_id>/part-b/save', methods=['POST'])
@login_required
def part_b_config_save(part_id):
    """Save Part B config (duration, term_duration, area_wise_override, category_override)."""
    part = SyllabusPart.query.get_or_404(part_id)
    if part.part_key != 'B':
        return jsonify({'success': False, 'message': 'Not Part B'}), 400
    if not _can_edit_syllabus():
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json() or {}
    rec = SyllabusPartBConfig.query.filter_by(part_id=part_id).first()
    if not rec:
        rec = SyllabusPartBConfig(part_id=part_id)
        db.session.add(rec)
    cfg = rec.get_config()
    if 'duration_years' in data:
        cfg['duration_years'] = data['duration_years']
    if 'terms' in data:
        cfg['terms'] = data['terms']
    if 'term_duration' in data:
        cfg['term_duration'] = data['term_duration']
    if 'area_wise_override' in data:
        cfg['area_wise_override'] = data['area_wise_override']
    if 'category_override' in data:
        cfg['category_override'] = data['category_override']
    if 'structure_curriculum' in data:
        sc = data['structure_curriculum']
        import re
        if isinstance(sc, dict) and 'column_headers' in sc and 'rows' in sc:
            cfg['structure_curriculum'] = sc
            for row in sc.get('rows') or []:
                if isinstance(row, dict) and (row.get('label') or '').strip():
                    lbl = (row.get('label') or '').strip().lower()
                    if 'duration' in lbl and 'program' in lbl:
                        vals = row.get('values') or []
                        if vals:
                            m = re.match(r'^(\d+)\s*years?', (vals[0] or '').strip(), re.I)
                            if m:
                                cfg['duration_years'] = int(m.group(1))
                        if len(vals) > 1 and (vals[1] or '').strip():
                            cfg['terms'] = (vals[1] or '').strip()
                        break
        elif isinstance(sc, list) and sc:
            cfg['structure_curriculum'] = sc
            for row in sc:
                if isinstance(row, dict) and (row.get('label') or '').strip():
                    lbl = (row.get('label') or '').strip().lower()
                    if 'duration' in lbl and 'program' in lbl:
                        v1 = (row.get('value_1') or '').strip()
                        m = re.match(r'^(\d+)\s*years?', v1, re.I)
                        if m:
                            cfg['duration_years'] = int(m.group(1))
                        v2 = (row.get('value_2') or '').strip()
                        if v2:
                            cfg['terms'] = v2
                        break
    rec.set_config(cfg)
    rec.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'success': True})


@curriculator_bp.route('/api/part/<int:part_id>/part-b/reset', methods=['POST'])
@login_required
def part_b_reset_from_c(part_id):
    """Reset Part B area_wise and/or category from Part C. Overwrites overrides."""
    part = SyllabusPart.query.get_or_404(part_id)
    if part.part_key != 'B':
        return jsonify({'success': False, 'message': 'Not Part B'}), 400
    if not _can_edit_syllabus():
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json() or {}
    target = (data.get('target') or 'both').strip().lower()
    part_c = SyllabusPart.query.filter_by(document_id=part.document_id, part_key='C').first()
    if not part_c:
        return jsonify({'success': False, 'message': 'Part C not found'}), 400
    entries = SyllabusCourseEntry.query.filter_by(part_id=part_c.id).order_by(
        SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id
    ).all()
    area_wise, category = _derive_part_b_from_c(entries)
    rec = SyllabusPartBConfig.query.filter_by(part_id=part_id).first()
    if not rec:
        rec = SyllabusPartBConfig(part_id=part_id)
        db.session.add(rec)
    cfg = rec.get_config()
    if target in ('area_wise', 'both'):
        cfg['area_wise_override'] = area_wise
    if target in ('category', 'both'):
        cfg['category_override'] = category
    rec.set_config(cfg)
    rec.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'success': True, 'area_wise': area_wise, 'category': category})


@curriculator_bp.route('/api/part/<int:part_id>/course-entry/add', methods=['POST'])
@login_required
def course_entry_add(part_id):
    """Add a Part C course entry."""
    if not _can_edit_syllabus():
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    part = SyllabusPart.query.get_or_404(part_id)
    if part.part_key != 'C':
        return jsonify({'success': False, 'message': 'Only Part C can have course entries'}), 400
    data = request.get_json() or {}
    code = (data.get('course_code') or '').strip() or None
    name = (data.get('course_name') or '').strip() or None
    credit = data.get('credit')
    try:
        credit = float(credit) if credit not in (None, '') else None
    except (TypeError, ValueError):
        credit = None
    year_term = (data.get('year_term') or '').strip() or None
    year = (data.get('year') or '').strip() or None
    term = (data.get('term') or '').strip() or None
    entry_type = (data.get('entry_type') or '').strip() or None
    status = (data.get('status') or '').strip() or None
    prereq_id = data.get('prerequisite_entry_id')
    try:
        prereq_id = int(prereq_id) if prereq_id not in (None, '') else None
    except (TypeError, ValueError):
        prereq_id = None
    course_id = data.get('course_id')
    if course_id is not None:
        try:
            course_id = int(course_id)
        except (TypeError, ValueError):
            course_id = None
    max_order = db.session.query(db.func.max(SyllabusCourseEntry.sort_order)).filter_by(part_id=part_id).scalar() or 0
    entry = SyllabusCourseEntry(
        part_id=part_id,
        course_id=course_id,
        course_code=code,
        course_name=name,
        credit=credit,
        year_term=year_term,
        year=year,
        term=term,
        entry_type=entry_type,
        status=status,
        prerequisite_entry_id=prereq_id,
        sort_order=int(max_order) + 1,
    )
    db.session.add(entry)
    db.session.commit()
    return jsonify({'success': True, 'id': entry.id})


@curriculator_bp.route('/api/course-entry/<int:entry_id>/save', methods=['POST'])
@login_required
def course_entry_save(entry_id):
    """Save Part C course entry. Permission: assigned teacher or head/TA."""
    entry = SyllabusCourseEntry.query.get_or_404(entry_id)
    teacher = _teacher().query.filter_by(name=current_user.full_name).first()
    is_assigned = any(
        a.teacher_id == teacher.id for a in entry.author_assignments
    ) if teacher else False
    if not _can_edit_syllabus() and not is_assigned:
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json() or {}
    if 'course_code' in data:
        entry.course_code = (data.get('course_code') or '').strip() or None
    if 'course_name' in data:
        entry.course_name = (data.get('course_name') or '').strip() or None
    if 'credit' in data:
        try:
            entry.credit = float(data['credit']) if data['credit'] not in (None, '') else None
        except (TypeError, ValueError):
            pass
    if 'year_term' in data:
        entry.year_term = (data.get('year_term') or '').strip() or None
    if 'year' in data:
        entry.year = (data.get('year') or '').strip() or None
    if 'term' in data:
        entry.term = (data.get('term') or '').strip() or None
    if 'entry_type' in data:
        entry.entry_type = (data.get('entry_type') or '').strip() or None
    if 'status' in data:
        entry.status = (data.get('status') or '').strip() or None
    if 'prerequisite_entry_id' in data:
        v = data.get('prerequisite_entry_id')
        try:
            pid = int(v) if v not in (None, '') else None
        except (TypeError, ValueError):
            pid = None
        if pid == entry.id:
            pid = None
        entry.prerequisite_entry_id = pid
    if 'content_json' in data:
        entry.set_content_dict(data['content_json'])
    entry.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'success': True})


def _infer_year_term_from_code(course_code):
    """Infer year and term from last 4 digits of course code (matches course_management)."""
    return infer_year_term_from_code(course_code)


def _map_course_type_to_entry_type(ct):
    """Map Course.course_type to Theory|Sessional|Viva|Capstone."""
    if not ct:
        return 'Theory'
    c = (ct or '').strip().lower()
    if 'sessional' in c:
        return 'Sessional'
    if 'viva' in c:
        return 'Viva'
    if 'thesis' in c or 'dissertation' in c:
        return 'Capstone'
    return 'Theory'


@curriculator_bp.route('/api/course-entry/<int:entry_id>/content')
@login_required
def api_entry_content(entry_id):
    """Return content_json (section_a, section_b, clos) for a Part C entry."""
    entry = SyllabusCourseEntry.query.get_or_404(entry_id)
    return jsonify(entry.get_content_dict() or {})


@curriculator_bp.route('/api/course-code-info')
@login_required
def api_course_code_info():
    """Return type, year, term inferred from course code (Curriculum lookup or digit parsing)."""
    code = (request.args.get('code') or '').strip()
    if not code:
        return jsonify({'success': True, 'year': '', 'term': '', 'type': 'Theory'})
    Course = _course_model()
    normalized = ''.join(ch for ch in code if not ch.isspace())
    course = Course.query.filter(Course.course_code == code).first()
    if not course and normalized != code:
        course = Course.query.filter(Course.course_code == normalized).first()
    if course:
        year = (course.display_year or course.year or '') or ''
        term = (course.display_term or course.term or '') or ''
        entry_type = _map_course_type_to_entry_type(getattr(course, 'course_type', None))
        return jsonify({'success': True, 'year': year, 'term': term, 'type': entry_type})
    y, t = _infer_year_term_from_code(code)
    return jsonify({'success': True, 'year': y, 'term': t, 'type': 'Theory'})


@curriculator_bp.route('/api/teachers')
@login_required
def api_teachers():
    """Return only teachers (excluding Head, TA, Admin) for assign-author dropdown."""
    try:
        from role_utils import get_teachers_excluding_head
        teachers = get_teachers_excluding_head()
    except ImportError:
        Teacher = _teacher()
        teachers = Teacher.query.order_by(Teacher.name).all()
    return jsonify({
        'success': True,
        'teachers': [{'id': t.id, 'name': t.name, 'short_name': t.short_name or ''} for t in teachers],
    })


@curriculator_bp.route('/api/assign-author', methods=['POST'])
@login_required
def assign_author():
    """Assign a teacher to a Part C course entry."""
    if not _can_assign_authors():
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json() or {}
    entry_id = data.get('course_entry_id')
    teacher_id = data.get('teacher_id')
    if not entry_id or not teacher_id:
        return jsonify({'success': False, 'message': 'course_entry_id and teacher_id required'}), 400
    entry = SyllabusCourseEntry.query.get_or_404(int(entry_id))
    existing = SyllabusAuthorAssignment.query.filter_by(
        course_entry_id=entry.id, teacher_id=int(teacher_id)
    ).first()
    if existing:
        return jsonify({'success': True, 'message': 'Already assigned'})
    a = SyllabusAuthorAssignment(course_entry_id=entry.id, teacher_id=int(teacher_id))
    a.assigned_by_id = current_user.id
    db.session.add(a)
    db.session.commit()
    return jsonify({'success': True})


@curriculator_bp.route('/api/assign-author/withdraw', methods=['POST'])
@login_required
def withdraw_author():
    """Withdraw a teacher's Part C course entry author assignment."""
    if not _can_assign_authors():
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json() or {}
    entry_id = data.get('course_entry_id')
    teacher_id = data.get('teacher_id')
    if not entry_id or not teacher_id:
        return jsonify({'success': False, 'message': 'course_entry_id and teacher_id required'}), 400
    entry = SyllabusCourseEntry.query.get_or_404(int(entry_id))
    rec = SyllabusAuthorAssignment.query.filter_by(
        course_entry_id=entry.id, teacher_id=int(teacher_id)
    ).first()
    if rec:
        db.session.delete(rec)
        db.session.commit()
    return jsonify({'success': True})


@curriculator_bp.route('/doc/<int:doc_id>/export', methods=['GET', 'POST'])
@login_required
def export(doc_id):
    """Export syllabus (full or selected parts/courses) as DOCX or PDF."""
    doc = SyllabusDocument.query.get_or_404(doc_id)
    if request.method == 'GET':
        parts = SyllabusPart.query.filter_by(document_id=doc_id).order_by(SyllabusPart.sort_order).all()
        part_c = next((p for p in parts if p.part_key == 'C'), None)
        course_entries = []
        if part_c:
            course_entries = SyllabusCourseEntry.query.filter_by(part_id=part_c.id).order_by(
                SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id
            ).all()
        return render_template(
            'curriculator/export.html',
            document=doc,
            parts=parts,
            course_entries=course_entries,
        )
    # POST: do export
    parts = request.form.getlist('parts') or []
    course_ids = request.form.getlist('course_ids') or []
    fmt = (request.form.get('format') or 'docx').strip().lower()
    if fmt not in ('docx', 'pdf'):
        fmt = 'docx'
    if not parts and not course_ids:
        flash('Select at least one part or specific courses.', 'warning')
        return redirect(url_for('curriculator.export', doc_id=doc_id))
    if not parts and course_ids:
        parts = ['C']  # export only Part C for selected courses

    try:
        if fmt == 'docx':
            buf, filename = _export_docx(doc, parts, course_ids)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            buf, filename = _export_pdf(doc, parts, course_ids, request)
            mimetype = 'application/pdf'
        if buf is None:
            flash('Export failed: no data generated.', 'error')
            return redirect(url_for('curriculator.export', doc_id=doc_id))
        buf.seek(0)
        # Write to temp file so cPanel/Passenger can send reliably (BytesIO streaming often fails there)
        suffix = '.docx' if fmt == 'docx' else '.pdf'
        fd, path = tempfile.mkstemp(suffix=suffix)
        try:
            os.write(fd, buf.getvalue())
            os.close(fd)
            fd = None
            resp = send_file(
                path,
                mimetype=mimetype,
                as_attachment=True,
                download_name=filename,
            )
            resp.call_on_close(lambda: os.unlink(path) if os.path.exists(path) else None)
            return resp
        finally:
            if fd is not None:
                try:
                    os.close(fd)
                except OSError:
                    pass
                if os.path.exists(path):
                    try:
                        os.unlink(path)
                    except OSError:
                        pass
    except Exception as e:
        current_app.logger.exception('Curriculator export failed')
        err_msg = str(e)
        if len(err_msg) > 200:
            err_msg = err_msg[:200] + '…'
        flash(f'Export failed: {err_msg}', 'error')
        return redirect(url_for('curriculator.export', doc_id=doc_id))


def _docx_add_page_number_footer(document):
    """Add footer with 'Page X of Y' to the document's first section using OOXML fields."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_field(run_el, instr_text):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = instr_text
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_el.append(fldChar1)
        run_el.append(instr)
        run_el.append(fldChar2)
        run_el.append(fldChar3)

    p.add_run('Page ')
    run_page = p.add_run()
    add_field(run_page._r, 'PAGE')
    p.add_run(' of ')
    run_num = p.add_run()
    add_field(run_num._r, 'NUMPAGES')


def _export_docx(doc, parts, course_ids):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    d.add_heading(doc.name, 0)
    part_map = {p.part_key: p for p in doc.parts}
    part_a = part_map.get('A')
    part_b = part_map.get('B')
    part_c = part_map.get('C')
    part_d = part_map.get('D')
    all_c = []
    if part_c:
        all_c = SyllabusCourseEntry.query.filter_by(part_id=part_c.id).order_by(
            SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id
        ).all()
    if part_c and course_ids:
        ids_ok = []
        for x in course_ids:
            try:
                ids_ok.append(int(x))
            except (TypeError, ValueError):
                pass
        if ids_ok:
            all_c = SyllabusCourseEntry.query.filter(
                SyllabusCourseEntry.part_id == part_c.id,
                SyllabusCourseEntry.id.in_(ids_ok),
            ).order_by(SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id).all()

    def add_table(doc, rows, cols):
        t = doc.add_table(rows=rows, cols=cols)
        t.style = 'Table Grid'
        return t

    def set_cell(cell, val):
        cell.text = (val if val is not None else '')[:8000]

    def merge_h(table, row_idx, col_start, col_end):
        """Merge cells horizontally in a row (inclusive)."""
        if col_end <= col_start:
            return table.cell(row_idx, col_start)
        a = table.cell(row_idx, col_start)
        for c in range(col_start + 1, col_end + 1):
            a = a.merge(table.cell(row_idx, c))
        return a

    def merge_v(table, col_idx, row_start, row_end):
        """Merge cells vertically in a column (inclusive)."""
        if row_end <= row_start:
            return table.cell(row_start, col_idx)
        a = table.cell(row_start, col_idx)
        for r in range(row_start + 1, row_end + 1):
            a = a.merge(table.cell(r, col_idx))
        return a

    def table_sep():
        """Add paragraphs after a table so the next table is clearly separate (LLB design)."""
        d.add_paragraph()
        d.add_paragraph()

    part_order = ['A', 'B', 'C', 'D']
    for key in part_order:
        if key not in parts and not course_ids:
            continue
        part = part_map.get(key)
        if not part:
            continue
        d.add_heading(f'Part {key}', level=1)

        if key == 'A' and part_a:
            sections_by_key = {s.section_key: s for s in SyllabusPartASection.query.filter_by(part_id=part_a.id).all()}
            plos_data = []
            peos_data = []
            plo_rec = SyllabusPartASection.query.filter_by(part_id=part_a.id, section_key='plos').first()
            if plo_rec and plo_rec.get_data():
                plos_data = plo_rec.get_data() if isinstance(plo_rec.get_data(), list) else []
            peo_rec = SyllabusPartASection.query.filter_by(part_id=part_a.id, section_key='peos').first()
            if peo_rec and peo_rec.get_data():
                peos_data = peo_rec.get_data() if isinstance(peo_rec.get_data(), list) else []
            for cfg in _part_a_sections():
                sk = cfg['key']
                rec = sections_by_key.get(sk)
                data = rec.get_data() if rec else None
                if cfg['type'] == 'key_value' and data and isinstance(data, list):
                    # LLB-style: title row merged across 2 cols, then key/value rows (no header row)
                    t = add_table(d, 1 + len(data), 2)
                    title = merge_h(t, 0, 0, 1)
                    set_cell(title, cfg['label'].upper())
                    for i, row in enumerate(data):
                        set_cell(t.rows[i + 1].cells[0], row.get('key'))
                        set_cell(t.rows[i + 1].cells[1], row.get('value'))
                    table_sep()
                elif cfg['type'] == 'text' and data:
                    txt = data.get('text', data) if isinstance(data, dict) else str(data)
                    if txt:
                        d.add_paragraph(cfg['label'])
                        d.add_paragraph(str(txt)[:5000])
                elif cfg['type'] == 'serial_description' and data and isinstance(data, list):
                    # LLB-style: title row merged, then header row, then data rows
                    t = add_table(d, 2 + len(data), 2)
                    title = merge_h(t, 0, 0, 1)
                    set_cell(title, cfg['label'].upper())
                    t.rows[1].cells[0].text = 'Serial No.'
                    t.rows[1].cells[1].text = 'Description'
                    for i, row in enumerate(data):
                        set_cell(t.rows[i + 2].cells[0], row.get('serial_no') or row.get('serial'))
                        set_cell(t.rows[i + 2].cells[1], row.get('description'))
                    table_sep()
                elif cfg['type'] == 'peos' and data and isinstance(data, list):
                    # LLB-style: title row merged across 3, then header row
                    t = add_table(d, 2 + len(data), 3)
                    title = merge_h(t, 0, 0, 2)
                    set_cell(title, cfg['label'].upper())
                    t.rows[1].cells[0].text = 'Serial No.'
                    t.rows[1].cells[1].text = 'Description'
                    t.rows[1].cells[2].text = 'Domains'
                    for i, row in enumerate(data):
                        set_cell(t.rows[i + 2].cells[0], row.get('serial'))
                        set_cell(t.rows[i + 2].cells[1], row.get('description'))
                        set_cell(t.rows[i + 2].cells[2], row.get('domains'))
                    table_sep()
                elif cfg['type'] == 'plos' and data and isinstance(data, list):
                    # LLB-style: title, desc, then optional category rows + two-column PLO rows (id | text)
                    n_extra = sum(1 for r in data if (r.get('category') or '').strip())
                    n_rows = 2 + len(data) + n_extra
                    t = add_table(d, n_rows, 2)
                    title = merge_h(t, 0, 0, 1)
                    set_cell(title, cfg['label'].upper())
                    desc = merge_h(t, 1, 0, 1)
                    set_cell(desc, 'After successful completion of the degree, the learners will be able to:')
                    row_idx = 2
                    for r in data:
                        cat = (r.get('category') or '').strip()
                        if cat:
                            set_cell(merge_h(t, row_idx, 0, 1), cat)
                            row_idx += 1
                        set_cell(t.rows[row_idx].cells[0], r.get('id'))
                        set_cell(t.rows[row_idx].cells[1], r.get('text'))
                        row_idx += 1
                    table_sep()
                elif cfg['type'] == 'graduate_attributes' and data and isinstance(data, list):
                    # LLB-style: title row, header "Graduate Attributes" | "Domain", then category rows (merged) + GA rows (id: description | domain)
                    n_extra = sum(1 for r in data if (r.get('category') or '').strip())
                    n_rows = 2 + len(data) + n_extra
                    t = add_table(d, n_rows, 2)
                    title = merge_h(t, 0, 0, 1)
                    set_cell(title, cfg['label'].upper())
                    t.rows[1].cells[0].text = 'Graduate Attributes'
                    t.rows[1].cells[1].text = 'Domain'
                    row_idx = 2
                    for r in data:
                        cat = (r.get('category') or '').strip()
                        if cat:
                            set_cell(merge_h(t, row_idx, 0, 1), cat)
                            row_idx += 1
                        ga_id = (r.get('id') or '').strip()
                        desc = (r.get('description') or '').strip()
                        left_text = (ga_id + ': ' + desc) if ga_id else desc
                        set_cell(t.rows[row_idx].cells[0], left_text)
                        set_cell(t.rows[row_idx].cells[1], r.get('domain') or '')
                        row_idx += 1
                    table_sep()
                    d.add_paragraph('GA = Graduate Attributes')
                    table_sep()
                elif sk == 'mapping_plo_peo' and plos_data and peos_data:
                    # LLB-style: Program Learning Outcomes (PLOs) | Program Educational Objectives (PEOs); Domains column; binary • in cells
                    data = data if isinstance(data, dict) else {}
                    cells = data.get('cells') or []
                    row_domains = data.get('row_domains') or []
                    nc = 2 + len(peos_data)
                    nr = 3 + len(plos_data)
                    t = add_table(d, nr, nc)
                    merge_h(t, 0, 0, 1).text = 'Program Learning Outcomes (PLOs)'
                    if nc > 2:
                        merge_h(t, 0, 2, nc - 1).text = 'Program Educational Objectives (PEOs)'
                    t.rows[1].cells[0].text = 'Domains'
                    t.rows[1].cells[1].text = 'PLO'
                    for j, peo in enumerate(peos_data):
                        if j + 2 < nc:
                            set_cell(t.rows[1].cells[j + 2], peo.get('serial') or ('PEO' + str(j + 1)))
                    for i, plo in enumerate(plos_data):
                        r = 2 + i
                        dom = (row_domains[i] if i < len(row_domains) else '') or (plo.get('category') or '')
                        set_cell(t.rows[r].cells[0], dom)
                        set_cell(t.rows[r].cells[1], plo.get('id') or ('PLO' + str(i + 1)))
                        for j in range(len(peos_data)):
                            v = ''
                            if i < len(cells) and j < len(cells[i]) and cells[i][j]:
                                v = '•'
                            if j + 2 < nc:
                                set_cell(t.rows[r].cells[j + 2], v)
                    # Merge vertically consecutive same domain in column 0
                    i = 0
                    while i < len(plos_data):
                        dom = (row_domains[i] if i < len(row_domains) else '') or (plos_data[i].get('category') or '')
                        j = i + 1
                        while j < len(plos_data) and ((row_domains[j] if j < len(row_domains) else '') or (plos_data[j].get('category') or '')) == dom:
                            j += 1
                        if j > i + 1:
                            merge_v(t, 0, 2 + i, 2 + j - 1)
                        i = j
                    table_sep()
                elif cfg['type'] == 'mapping' and sk != 'mapping_plo_peo' and data and isinstance(data, dict):
                    rl = data.get('row_labels') or []
                    cl = data.get('col_labels') or []
                    cells = data.get('cells') or []
                    # LLB-style: title row merged across all columns, then header row
                    nr = 2 + len(rl)
                    nc = 1 + len(cl)
                    if nc < 2:
                        nc = 2
                    t = add_table(d, nr, nc)
                    title = merge_h(t, 0, 0, nc - 1)
                    set_cell(title, cfg['label'].upper())
                    # header row
                    set_cell(t.rows[1].cells[0], 'PEOs' if sk == 'mapping_mission_peo' else 'Row')
                    for j, c in enumerate(cl):
                        if j + 1 < nc:
                            set_cell(t.rows[1].cells[j + 1], c)
                    # data rows
                    for i, r in enumerate(rl):
                        set_cell(t.rows[i + 2].cells[0], r)
                        for j in range(len(cl)):
                            if j + 1 < nc:
                                v = ''
                                if i < len(cells) and j < len(cells[i]):
                                    v = cells[i][j]
                                set_cell(t.rows[i + 2].cells[j + 1], v)
                    table_sep()
                    if sk == 'mapping_mission_peo':
                        d.add_paragraph('Level of association: 3=High, 2=Medium, 1=Low')
                        table_sep()
                elif cfg['type'] == 'mapping_course_plo' and all_c and plos_data:
                    suffix = sk.replace('mapping_course_plo_', '') if sk.startswith('mapping_course_plo_') else ''
                    yt = _mapping_course_plo_year_term().get(suffix)
                    if yt:
                        year, term = yt
                        entries = [e for e in all_c if (e.year or '').strip() == year and (e.term or '').strip() == term]
                    else:
                        entries = all_c
                    stored = (data or {}).get('cells', []) if isinstance(data, dict) else []
                    nc = 1 + len(plos_data)
                    # Domain groups: consecutive PLOs with same category
                    plo_domain_groups = []
                    cur_cat, cur_indices = None, []
                    for i, plo in enumerate(plos_data):
                        cat = (plo.get('category') or '').strip()
                        if cat != cur_cat:
                            if cur_indices:
                                plo_domain_groups.append((cur_cat or '—', cur_indices))
                            cur_cat = cat
                            cur_indices = [i]
                        else:
                            cur_indices.append(i)
                    if cur_indices:
                        plo_domain_groups.append((cur_cat or '—', cur_indices))
                    nr = 4 + len(entries)
                    if nr > 4 and nc > 1:
                        t = add_table(d, nr, nc)
                        merge_h(t, 0, 0, nc - 1).text = cfg['label'].upper()
                        t.rows[1].cells[0].text = 'Course Code and Course Title'
                        merge_h(t, 1, 1, nc - 1).text = 'Program Learning Outcomes (PLOs)'
                        # Domain row
                        col = 1
                        for dom_label, indices in plo_domain_groups:
                            if col + len(indices) <= nc:
                                if len(indices) > 1:
                                    merge_h(t, 2, col, col + len(indices) - 1).text = dom_label
                                else:
                                    set_cell(t.rows[2].cells[col], dom_label)
                                col += len(indices)
                        # PLO row
                        for j, plo in enumerate(plos_data):
                            if j + 1 < nc:
                                set_cell(t.rows[3].cells[j + 1], plo.get('id'))
                        for i, e in enumerate(entries):
                            set_cell(t.rows[4 + i].cells[0], (e.course_code or '') + ' / ' + (e.course_name or ''))
                            for j in range(nc - 1):
                                v = ''
                                if i < len(stored) and j < len(stored[i]) and stored[i][j]:
                                    v = '•'
                                set_cell(t.rows[4 + i].cells[j + 1], v)
                        table_sep()

        elif key == 'B' and part_b:
            cfg_rec = SyllabusPartBConfig.query.filter_by(part_id=part_b.id).first()
            cfg = cfg_rec.get_config() if cfg_rec else {}
            td = cfg.get('term_duration') or {}
            area_wise = cfg.get('area_wise_override')
            category = cfg.get('category_override')
            if area_wise is None or not isinstance(area_wise, list):
                area_wise, _ = _derive_part_b_from_c(all_c)
            if category is None or not isinstance(category, list):
                _, category = _derive_part_b_from_c(all_c)
            # Structure of the Curriculum table (Label + dynamic value columns)
            structure_curriculum = cfg.get('structure_curriculum')
            if isinstance(structure_curriculum, dict) and structure_curriculum.get('column_headers') and structure_curriculum.get('rows'):
                headers = structure_curriculum['column_headers']
                rows_data = structure_curriculum['rows']
                ncols = 1 + len(headers)
                t = add_table(d, 1 + len(rows_data), ncols)
                set_cell(t.rows[0].cells[0], 'Structure of the Curriculum')
                for j, h in enumerate(headers):
                    if j + 1 < ncols:
                        set_cell(t.rows[0].cells[j + 1], h)
                for i, row in enumerate(rows_data):
                    vals = row.get('values') or []
                    set_cell(t.rows[i + 1].cells[0], row.get('label'))
                    for j in range(len(headers)):
                        if j + 1 < ncols:
                            set_cell(t.rows[i + 1].cells[j + 1], vals[j] if j < len(vals) else '')
                table_sep()
            elif isinstance(structure_curriculum, list) and structure_curriculum:
                t = add_table(d, 1 + len(structure_curriculum), 3)
                t.rows[0].cells[0].text = 'Structure of the Curriculum'
                t.rows[0].cells[1].text = 'Structure of the Curriculum'
                t.rows[0].cells[2].text = 'Structure of the Curriculum'
                for i, row in enumerate(structure_curriculum):
                    set_cell(t.rows[i + 1].cells[0], row.get('label'))
                    set_cell(t.rows[i + 1].cells[1], row.get('value_1'))
                    set_cell(t.rows[i + 1].cells[2], row.get('value_2'))
                table_sep()
            else:
                t = add_table(d, 2, 3)
                t.rows[0].cells[0].text = 'Duration of the Program'
                t.rows[0].cells[1].text = str(cfg.get('duration_years', 4)) + ' Years'
                t.rows[0].cells[2].text = str(cfg.get('terms', 'Terms'))
                t.rows[1].cells[0].text = 'Term duration (weeks)'
                t.rows[1].cells[1].text = 'Teaching: %s, Preparatory: %s, Exam: %s, Break: %s' % (
                    td.get('teaching', ''), td.get('preparatory', ''), td.get('exam', ''), td.get('break', ''))
                t.rows[1].cells[2].text = 'Total: %s' % td.get('total', '')
                table_sep()
            # Term duration (weeks) table: header + one row
            t = add_table(d, 3, 5)
            t.rows[0].cells[0].text = '*Term duration (weeks)'
            t.rows[0].cells[1].text = '*Term duration (weeks)'
            t.rows[0].cells[2].text = '*Term duration (weeks)'
            t.rows[0].cells[3].text = '*Term duration (weeks)'
            t.rows[0].cells[4].text = '*Term duration (weeks)'
            t.rows[1].cells[0].text = 'Teaching and Learning'
            t.rows[1].cells[1].text = 'Preparatory Leave'
            t.rows[1].cells[2].text = 'Term Final Examination'
            t.rows[1].cells[3].text = 'Term Break'
            t.rows[1].cells[4].text = 'Total'
            t.rows[2].cells[0].text = '%s Weeks' % td.get('teaching', 14)
            t.rows[2].cells[1].text = '%s Weeks' % td.get('preparatory', 1)
            t.rows[2].cells[2].text = '%s Weeks' % td.get('exam', 2)
            t.rows[2].cells[3].text = '%s Weeks' % td.get('break', 1)
            t.rows[2].cells[4].text = '%s Weeks' % td.get('total', 18)
            table_sep()
            if area_wise:
                t = add_table(d, 1 + len(area_wise), 5)
                t.rows[0].cells[0].text, t.rows[0].cells[1].text = 'Area', 'Type'
                t.rows[0].cells[2].text, t.rows[0].cells[3].text, t.rows[0].cells[4].text = 'Number of Courses', 'Credits', 'Total Credit'
                for i, r in enumerate(area_wise):
                    set_cell(t.rows[i + 1].cells[0], r.get('area'))
                    set_cell(t.rows[i + 1].cells[1], r.get('type'))
                    set_cell(t.rows[i + 1].cells[2], str(r.get('num_courses', '')))
                    set_cell(t.rows[i + 1].cells[3], str(r.get('credits', '')))
                    set_cell(t.rows[i + 1].cells[4], str(r.get('total_credit', '')))
                table_sep()
            if category:
                t = add_table(d, 1 + len(category), 4)
                t.rows[0].cells[0].text, t.rows[0].cells[1].text = 'Area', 'Course Type'
                t.rows[0].cells[2].text, t.rows[0].cells[3].text = 'Course Title', 'Credit'
                for i, r in enumerate(category):
                    set_cell(t.rows[i + 1].cells[0], r.get('area'))
                    set_cell(t.rows[i + 1].cells[1], r.get('course_type'))
                    set_cell(t.rows[i + 1].cells[2], r.get('course_title'))
                    set_cell(t.rows[i + 1].cells[3], str(r.get('credit', '')))
                table_sep()
            # Course distribution by term (LLB-style 8-column tables)
            dist = _course_distribution_by_term(all_c)
            for g in dist:
                entries = g.get('entries') or []
                t = add_table(d, 4 + len(entries), 8)
                merge_h(t, 0, 0, 7).text = 'COURSE DISTRIBUTION BY TERM'
                merge_h(t, 1, 0, 7).text = (g['year'] + ' YEAR ' + g['term'] + ' TERM').upper()
                # Two-row header
                headers1 = ['Course Code', 'Course Title', 'Course Status', 'Notional Hours', 'Contact Hours/Week', 'Contact Hours/Week', 'Credits', 'Prerequisites']
                headers2 = ['Course Code', 'Course Title', 'Course Status', 'Notional Hours', 'Theory', 'Sessional', 'Credits', 'Prerequisites']
                for j, h in enumerate(headers1):
                    set_cell(t.rows[2].cells[j], h)
                for j, h in enumerate(headers2):
                    set_cell(t.rows[3].cells[j], h)

                def notional_hours(entry):
                    try:
                        cr = float(entry.credit) if entry.credit is not None else 0
                    except Exception:
                        cr = 0
                    et = (entry.entry_type or '').strip().lower()
                    if 'sessional' in et or et == 'sessional':
                        return 60
                    return int(round(cr * 40))

                def contact_hours(entry):
                    try:
                        cr = float(entry.credit) if entry.credit is not None else 0
                    except Exception:
                        cr = 0
                    et = (entry.entry_type or 'Theory').strip().lower()
                    if 'sessional' in et:
                        return '', f'{cr * 1.5:.1f}'.rstrip('0').rstrip('.')
                    if 'theory' in et or not et:
                        return f'{cr:.1f}'.rstrip('0').rstrip('.'), '-'
                    return '-', '-'

                for i, e in enumerate(entries):
                    r = 4 + i
                    set_cell(t.rows[r].cells[0], e.course_code or '')
                    set_cell(t.rows[r].cells[1], e.course_name or '')
                    set_cell(t.rows[r].cells[2], e.status or '')
                    set_cell(t.rows[r].cells[3], str(notional_hours(e)))
                    th, se = contact_hours(e)
                    set_cell(t.rows[r].cells[4], th)
                    set_cell(t.rows[r].cells[5], se)
                    set_cell(t.rows[r].cells[6], str(e.credit) if e.credit is not None else '')
                    prereq = 'None'
                    if e.prerequisite_entry_id:
                        pe = SyllabusCourseEntry.query.get(e.prerequisite_entry_id)
                        if pe and pe.course_code:
                            prereq = pe.course_code
                    set_cell(t.rows[r].cells[7], prereq)
                table_sep()

        elif key == 'C' and part_c and all_c:
            dist = _course_distribution_by_term(all_c)
            # LLB-style heading tables
            t_head = add_table(d, 1, 1)
            set_cell(t_head.rows[0].cells[0], 'COURSE DESCRIPTION BY TERM')
            table_sep()
            for g in dist:
                t_term = add_table(d, 1, 1)
                set_cell(t_term.rows[0].cells[0], (g['year'] + ' YEAR ' + g['term'] + ' TERM').upper())
                table_sep()
                for e in g['entries']:
                    d.add_page_break()
                    cnt = e.get_content_dict()
                    prerequisite = (cnt.get('prerequisite') or '').strip()
                    if not prerequisite and e.prerequisite_entry_id:
                        pe = SyllabusCourseEntry.query.get(e.prerequisite_entry_id)
                        prerequisite = (pe.course_code or 'None') if pe else 'None'
                    if not prerequisite:
                        prerequisite = 'None'
                    rationale = cnt.get('rationale') or ''

                    # Course header (LLB-style 4x5 with merges)
                    t = add_table(d, 4, 5)
                    c00 = merge_h(t, 0, 0, 1)
                    set_cell(c00, f"Course No.: {e.course_code or ''}".strip())
                    set_cell(t.rows[0].cells[2], f"Credit: {str(e.credit) if e.credit is not None else ''}".strip())
                    set_cell(t.rows[0].cells[3], f"Year: {e.year or ''}".strip())
                    set_cell(t.rows[0].cells[4], f"Term: {e.term or ''}".strip())
                    c10 = merge_h(t, 1, 0, 2)
                    set_cell(c10, f"Course Title: {e.course_name or ''}".strip())
                    c13 = merge_h(t, 1, 3, 4)
                    set_cell(c13, f"Course Status: {e.status or ''}".strip())
                    set_cell(t.rows[2].cells[0], 'Prerequisite')
                    c21 = merge_h(t, 2, 1, 4)
                    set_cell(c21, prerequisite)
                    set_cell(t.rows[3].cells[0], 'Rationale')
                    c31 = merge_h(t, 3, 1, 4)
                    set_cell(c31, str(rationale)[:8000])
                    table_sep()

                    # COURSE CONTENTS (LLB-style: full-width header, Section A|CLOs + rows, Section B|CLOs + rows, 3 cols)
                    items_a_raw = cnt.get('section_a_items') or []
                    items_b_raw = cnt.get('section_b_items') or []
                    if not items_a_raw and not items_b_raw:
                        raw_a = (cnt.get('section_a') or '').strip()
                        raw_b = (cnt.get('section_b') or '').strip()
                        items_a = [{'content': line.strip(), 'clos': ''} for line in raw_a.split('\n') if line.strip()]
                        items_b = [{'content': line.strip(), 'clos': ''} for line in raw_b.split('\n') if line.strip()]
                        if not items_a:
                            items_a = [{'content': '—', 'clos': ''}]
                        if not items_b:
                            items_b = [{'content': '—', 'clos': ''}]
                    else:
                        items_a = [{'content': (x.get('content') or '').strip(), 'clos': (x.get('clos') or '').strip()} for x in items_a_raw] if items_a_raw else [{'content': '—', 'clos': ''}]
                        items_b = [{'content': (x.get('content') or '').strip(), 'clos': (x.get('clos') or '').strip()} for x in items_b_raw] if items_b_raw else [{'content': '—', 'clos': ''}]
                    n_a, n_b = len(items_a), len(items_b)
                    n_rows = 1 + 1 + n_a + 1 + n_b
                    tcnt = add_table(d, n_rows, 3)
                    merge_h(tcnt, 0, 0, 2).text = 'COURSE CONTENTS'
                    merge_h(tcnt, 1, 0, 1).text = 'Section A'
                    tcnt.rows[1].cells[2].text = 'CLOs'
                    for i, it in enumerate(items_a):
                        set_cell(tcnt.rows[2 + i].cells[0], str(i + 1))
                        set_cell(tcnt.rows[2 + i].cells[1], it.get('content') or '')
                        set_cell(tcnt.rows[2 + i].cells[2], it.get('clos') or '')
                    merge_h(tcnt, 2 + n_a, 0, 1).text = 'Section B'
                    set_cell(tcnt.rows[2 + n_a].cells[2], '')
                    for i, it in enumerate(items_b):
                        set_cell(tcnt.rows[3 + n_a + i].cells[0], str(n_a + i + 1))
                        set_cell(tcnt.rows[3 + n_a + i].cells[1], it.get('content') or '')
                        set_cell(tcnt.rows[3 + n_a + i].cells[2], it.get('clos') or '')
                    table_sep()

                    # CLO table (LLB-style: first column vertically merged; header col1-2 merged)
                    clos = cnt.get('clos') or []
                    if not isinstance(clos, list):
                        clos = []
                    n_clo_rows = max(1, len(clos))
                    tclo = add_table(d, 1 + n_clo_rows, 4)
                    merge_v(tclo, 0, 0, n_clo_rows)
                    set_cell(tclo.rows[0].cells[0], 'Course Learning Outcomes (CLOs)')
                    merge_h(tclo, 0, 1, 2)
                    set_cell(tclo.rows[0].cells[1], 'Upon completion of this course, the student will be able to:')
                    set_cell(tclo.rows[0].cells[3], 'Mapping with PLO')
                    if clos:
                        for i, c in enumerate(clos):
                            set_cell(tclo.rows[i + 1].cells[1], f'CLO {i + 1}')
                            set_cell(tclo.rows[i + 1].cells[2], c.get('text') or c.get('upon_completion') or '')
                            set_cell(tclo.rows[i + 1].cells[3], c.get('plo') or '')
                    else:
                        set_cell(tclo.rows[1].cells[1], 'CLO 1')
                    table_sep()

                    # Mapping CLOs with Teaching-Learning & Assessment (from mapping_clos or placeholder)
                    mapping_clos = cnt.get('mapping_clos') or []
                    if not isinstance(mapping_clos, list):
                        mapping_clos = []
                    n_map = max(1, len(clos))
                    tmap = add_table(d, 2 + n_map, 3)
                    merge_h(tmap, 0, 0, 2).text = 'MAPPING CLOs WITH THE TEACHING-LEARNING STRATEGY AND ASSESSMENT STRATEGY'
                    tmap.rows[1].cells[0].text = 'CLOs'
                    tmap.rows[1].cells[1].text = 'Teaching-Learning Strategy'
                    tmap.rows[1].cells[2].text = 'Assessment Strategy'
                    for i in range(n_map):
                        row = mapping_clos[i] if i < len(mapping_clos) else {}
                        tl = (row.get('teaching_learning') or '').strip()
                        ass = (row.get('assessment') or '').strip()
                        set_cell(tmap.rows[i + 2].cells[0], str(i + 1))
                        set_cell(tmap.rows[i + 2].cells[1], tl)
                        set_cell(tmap.rows[i + 2].cells[2], ass)
                    table_sep()

                    # Indicative learning materials (placeholder)
                    tmat = add_table(d, 3, 2)
                    merge_h(tmat, 0, 0, 1).text = 'INDICATIVE LEARNING MATERIALS'
                    tmat.rows[1].cells[0].text = 'Recommended Readings'
                    tmat.rows[2].cells[0].text = 'Supplementary Readings'
                    set_cell(tmat.rows[1].cells[1], cnt.get('recommended_readings') or '')
                    set_cell(tmat.rows[2].cells[1], cnt.get('supplementary_readings') or '')
                    table_sep()

        elif key == 'D' and part_d:
            d.add_paragraph()
            p = d.add_paragraph()
            p.add_run('20 ').bold = True
            p.add_run('Grading and Evaluation')
            # 20.1 Grading Scale
            d.add_paragraph()
            p = d.add_paragraph()
            p.add_run('20.1 Grading Scale').bold = True
            d.add_paragraph('a) Letter Grades and corresponding Grade Points will be awarded following provisions shown below:')
            rec = SyllabusPartDSection.query.filter_by(part_id=part_d.id, section_key='grading_scale').first()
            data = rec.get_data() if rec else []
            if isinstance(data, list) and data:
                t = add_table(d, 1 + len(data), 3)
                t.rows[0].cells[0].text, t.rows[0].cells[1].text, t.rows[0].cells[2].text = 'Numerical Grade', 'Letter Grade', 'Grade Point'
                for i, r in enumerate(data):
                    set_cell(t.rows[i + 1].cells[0], r.get('numerical'))
                    set_cell(t.rows[i + 1].cells[1], r.get('letter'))
                    set_cell(t.rows[i + 1].cells[2], r.get('grade_point'))
                table_sep()
            # 20.1.1 Evaluation of Theory Courses
            p = d.add_paragraph()
            p.add_run('20.1.1 Evaluation of Theory Courses').bold = True
            d.add_paragraph('a) All theory courses will be evaluated out of 100 marks, the distribution of which is given below:')
            rec = SyllabusPartDSection.query.filter_by(part_id=part_d.id, section_key='theory_evaluation').first()
            data = rec.get_data() if rec else []
            if isinstance(data, list) and data:
                t = add_table(d, 1 + len(data), 3)
                t.rows[0].cells[0].text, t.rows[0].cells[1].text, t.rows[0].cells[2].text = 'Sl. No.', 'Items', 'Marks'
                for i, r in enumerate(data):
                    set_cell(t.rows[i + 1].cells[0], r.get('sl_no'))
                    set_cell(t.rows[i + 1].cells[1], r.get('items'))
                    set_cell(t.rows[i + 1].cells[2], r.get('marks'))
                table_sep()
            # 20.1.2 Evaluation of Sessional Courses
            p = d.add_paragraph()
            p.add_run('20.1.2 Evaluation of Sessional Courses').bold = True
            d.add_paragraph('a) All sessional courses will be evaluated out of 100 marks, the distribution of which is given below:')
            rec = SyllabusPartDSection.query.filter_by(part_id=part_d.id, section_key='sessional_evaluation').first()
            data = rec.get_data() if rec else []
            if isinstance(data, list) and data:
                t = add_table(d, 1 + len(data), 3)
                t.rows[0].cells[0].text, t.rows[0].cells[1].text, t.rows[0].cells[2].text = 'Sl. No.', 'Items', 'Marks'
                for i, r in enumerate(data):
                    set_cell(t.rows[i + 1].cells[0], r.get('sl_no'))
                    set_cell(t.rows[i + 1].cells[1], r.get('items'))
                    set_cell(t.rows[i + 1].cells[2], r.get('marks'))
                table_sep()
            # Approval Records
            rec = SyllabusPartDSection.query.filter_by(part_id=part_d.id, section_key='approval_records').first()
            data = rec.get_data() if rec else []
            if isinstance(data, list) and data:
                d.add_paragraph()
                p = d.add_paragraph('Approval Records')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if hasattr(WD_ALIGN_PARAGRAPH, 'CENTER') else 1
                t = add_table(d, 1 + len(data), 2)
                t.rows[0].cells[0].text, t.rows[0].cells[1].text = 'Approving Authority', 'Date of Approval'
                for i, r in enumerate(data):
                    set_cell(t.rows[i + 1].cells[0], r.get('authority'))
                    set_cell(t.rows[i + 1].cells[1], r.get('date'))
                table_sep()

    _docx_add_page_number_footer(d)
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    name = (doc.name or 'syllabus').replace(' ', '_')[:50]
    fn = f'{name}_export_{datetime.utcnow().strftime("%Y%m%d_%H%M")}.docx'
    return buf, fn


def _export_pdf(doc, parts, course_ids, req):
    from weasyprint import HTML
    from markupsafe import escape
    from utils.pdf_fonts import resolve_formal_pdf_fonts, formal_font_face_css

    formal_fonts = resolve_formal_pdf_fonts()
    if not formal_fonts:
        raise RuntimeError(
            'PDF fonts missing. Upload LiberationSerif-Regular.ttf and LiberationSerif-Bold.ttf to static/fonts/.'
        )

    part_map = {p.part_key: p for p in doc.parts}
    part_a = part_map.get('A')
    part_b = part_map.get('B')
    part_c = part_map.get('C')
    part_d = part_map.get('D')
    all_c = []
    if part_c:
        all_c = SyllabusCourseEntry.query.filter_by(part_id=part_c.id).order_by(
            SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id
        ).all()
    if part_c and course_ids:
        ids_ok = []
        for x in course_ids:
            try:
                ids_ok.append(int(x))
            except (TypeError, ValueError):
                pass
        if ids_ok:
            all_c = SyllabusCourseEntry.query.filter(
                SyllabusCourseEntry.part_id == part_c.id,
                SyllabusCourseEntry.id.in_(ids_ok),
            ).order_by(SyllabusCourseEntry.sort_order, SyllabusCourseEntry.id).all()

    def esc(s):
        return escape(str(s)[:8000] if s is not None else '')

    def pdf_table_sep():
        chunks.append('<p class="table-gap"></p>')

    chunks = []
    chunks.append('<!DOCTYPE html><html><head><meta charset="utf-8"/><style>')
    chunks.append(formal_font_face_css(formal_fonts, 'PDFSerif'))
    chunks.append(
        "body{ font-family: 'PDFSerif', 'Times New Roman', Times, serif; padding: 1in; } "
        "h1{ font-size: 18pt; } h2{ font-size: 14pt; margin-top: 1em; }"
    )
    chunks.append('p{ margin: 0.5em 0; } table{ border-collapse: collapse; margin: 0 0 1.2em 0; display: block; width: 100%; } th,td{ border: 1px solid #333; padding: 6px 8px; text-align: left; } .table-gap{ margin: 0.6em 0; height: 0; overflow: hidden; } .section-caption{ font-weight: bold; margin: 0.8em 0 0.3em 0; } .course-page{ page-break-before: always; }')
    chunks.append("@page{ margin: 1in; @bottom-center{ content: counter(page) \" / \" counter(pages); font-size: 10pt; font-family: 'PDFSerif', serif; } }")
    chunks.append('</style></head><body><h1>%s</h1>' % esc(doc.name))

    part_order = ['A', 'B', 'C', 'D']
    for key in part_order:
        if key not in parts and not course_ids:
            continue
        part = part_map.get(key)
        if not part:
            continue
        chunks.append('<h2>Part %s</h2>' % key)

        if key == 'A' and part_a:
            sections_by_key = {s.section_key: s for s in SyllabusPartASection.query.filter_by(part_id=part_a.id).all()}
            plos_data = []
            peos_data = []
            plo_rec = SyllabusPartASection.query.filter_by(part_id=part_a.id, section_key='plos').first()
            if plo_rec and plo_rec.get_data():
                plos_data = plo_rec.get_data() if isinstance(plo_rec.get_data(), list) else []
            peo_rec = SyllabusPartASection.query.filter_by(part_id=part_a.id, section_key='peos').first()
            if peo_rec and peo_rec.get_data():
                peos_data = peo_rec.get_data() if isinstance(peo_rec.get_data(), list) else []
            for cfg in _part_a_sections():
                sk = cfg['key']
                rec = sections_by_key.get(sk)
                data = rec.get_data() if rec else None
                if cfg['type'] == 'key_value' and data and isinstance(data, list):
                    chunks.append('<table><tr><th colspan=\"2\">%s</th></tr>' % esc(cfg['label']).upper())
                    for r in data:
                        chunks.append('<tr><td>%s</td><td>%s</td></tr>' % (esc(r.get('key')), esc(r.get('value'))))
                    chunks.append('</table>')
                    pdf_table_sep()
                elif cfg['type'] == 'text' and data:
                    txt = data.get('text', data) if isinstance(data, dict) else str(data)
                    if txt:
                        chunks.append('<p><strong>%s</strong></p><p>%s</p>' % (esc(cfg['label']), esc(str(txt)[:5000])))
                elif cfg['type'] == 'serial_description' and data and isinstance(data, list):
                    chunks.append('<table><tr><th colspan=\"2\">%s</th></tr><tr><th>Serial No.</th><th>Description</th></tr>' % esc(cfg['label']).upper())
                    for r in data:
                        chunks.append('<tr><td>%s</td><td>%s</td></tr>' % (esc(r.get('serial_no') or r.get('serial')), esc(r.get('description'))))
                    chunks.append('</table>')
                    pdf_table_sep()
                elif cfg['type'] == 'peos' and data and isinstance(data, list):
                    chunks.append('<table><tr><th colspan=\"3\">%s</th></tr><tr><th>Serial No.</th><th>Description</th><th>Domains</th></tr>' % esc(cfg['label']).upper())
                    for r in data:
                        chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td></tr>' % (esc(r.get('serial')), esc(r.get('description')), esc(r.get('domains'))))
                    chunks.append('</table>')
                    pdf_table_sep()
                elif cfg['type'] == 'plos' and data and isinstance(data, list):
                    chunks.append('<table><tr><th colspan="2">%s</th></tr>' % esc(cfg['label']).upper())
                    chunks.append('<tr><td colspan="2">After successful completion of the degree, the learners will be able to:</td></tr>')
                    for r in data:
                        cat = (r.get('category') or '').strip()
                        if cat:
                            chunks.append('<tr><td colspan="2" class="fw-bold">%s</td></tr>' % esc(cat))
                        chunks.append('<tr><td style="width:6rem;">%s</td><td>%s</td></tr>' % (esc(r.get('id')), esc(r.get('text'))))
                    chunks.append('</table>')
                    pdf_table_sep()
                elif cfg['type'] == 'graduate_attributes' and data and isinstance(data, list):
                    chunks.append('<table><tr><th colspan="2">%s</th></tr><tr><th>Graduate Attributes</th><th style="text-align:right;">Domain</th></tr>' % esc(cfg['label']).upper())
                    for r in data:
                        cat = (r.get('category') or '').strip()
                        if cat:
                            chunks.append('<tr><td colspan="2" class="fw-bold">%s</td></tr>' % esc(cat))
                        ga_id = (r.get('id') or '').strip()
                        desc = (r.get('description') or '').strip()
                        left_text = (ga_id + ': ' + desc) if ga_id else desc
                        chunks.append('<tr><td>%s</td><td style="text-align:right;">%s</td></tr>' % (esc(left_text), esc(r.get('domain') or '')))
                    chunks.append('</table>')
                    chunks.append('<p class="small mb-0 mt-1">GA = Graduate Attributes</p>')
                    pdf_table_sep()
                elif sk == 'mapping_plo_peo' and plos_data and peos_data:
                    data = data if isinstance(data, dict) else {}
                    cells = data.get('cells') or []
                    row_domains = data.get('row_domains') or []
                    npeo = len(peos_data)
                    chunks.append('<table><tr><th colspan="2">Program Learning Outcomes (PLOs)</th><th colspan="%d">Program Educational Objectives (PEOs)</th></tr>' % npeo)
                    chunks.append('<tr><th>Domains</th><th>PLO</th>' + ''.join('<th>%s</th>' % esc(p.get('serial') or ('PEO' + str(j + 1))) for j, p in enumerate(peos_data)) + '</tr>')
                    i = 0
                    while i < len(plos_data):
                        dom = (row_domains[i] if i < len(row_domains) else '') or (plos_data[i].get('category') or '')
                        j = i + 1
                        while j < len(plos_data) and ((row_domains[j] if j < len(row_domains) else '') or (plos_data[j].get('category') or '')) == dom:
                            j += 1
                        rowspan = j - i
                        row_cells = cells[i] if i < len(cells) else []
                        plo_id = plos_data[i].get('id') or ('PLO' + str(i + 1))
                        cells_html = ''.join('<td class="text-center">%s</td>' % ('•' if (k < len(row_cells) and row_cells[k]) else '') for k in range(npeo))
                        if rowspan > 1:
                            chunks.append('<tr><td rowspan="%d">%s</td><td>%s</td>%s</tr>' % (rowspan, esc(dom), esc(plo_id), cells_html))
                            for k in range(i + 1, j):
                                row_cells_k = cells[k] if k < len(cells) else []
                                plo_id_k = plos_data[k].get('id') or ('PLO' + str(k + 1))
                                cells_html_k = ''.join('<td class="text-center">%s</td>' % ('•' if (c < len(row_cells_k) and row_cells_k[c]) else '') for c in range(npeo))
                                chunks.append('<tr><td>%s</td>%s</tr>' % (esc(plo_id_k), cells_html_k))
                        else:
                            chunks.append('<tr><td>%s</td><td>%s</td>%s</tr>' % (esc(dom), esc(plo_id), cells_html))
                        i = j
                    chunks.append('</table>')
                    pdf_table_sep()
                elif cfg['type'] == 'mapping' and sk != 'mapping_plo_peo' and data and isinstance(data, dict):
                    rl, cl, cells = data.get('row_labels') or [], data.get('col_labels') or [], data.get('cells') or []
                    if cl:
                        chunks.append('<table><tr><th colspan=\"%d\">%s</th></tr>' % (1 + len(cl), esc(cfg['label']).upper()))
                        chunks.append('<tr><th>%s</th>' % ('PEOs' if sk == 'mapping_mission_peo' else '') + ''.join('<th>%s</th>' % esc(c) for c in cl) + '</tr>')
                        for i, r in enumerate(rl):
                            row = [esc(r)]
                            for j in range(len(cl)):
                                v = cells[i][j] if i < len(cells) and j < len(cells[i]) else ''
                                row.append(esc(v))
                            chunks.append('<tr>' + ''.join('<td>%s</td>' % c for c in row) + '</tr>')
                        chunks.append('</table>')
                        if sk == 'mapping_mission_peo':
                            chunks.append('<p class="small mb-0 mt-1"><strong>Level of association: 3=High, 2=Medium, 1=Low</strong></p>')
                    pdf_table_sep()
                elif cfg['type'] == 'mapping_course_plo' and all_c and plos_data:
                    suffix = sk.replace('mapping_course_plo_', '') if sk.startswith('mapping_course_plo_') else ''
                    yt = _mapping_course_plo_year_term().get(suffix)
                    entries = [e for e in all_c if (e.year or '').strip() == yt[0] and (e.term or '').strip() == yt[1]] if yt else all_c
                    stored = (data or {}).get('cells', []) if isinstance(data, dict) else []
                    nplo = len(plos_data)
                    plo_domain_groups = []
                    cur_cat, cur_indices = None, []
                    for i, plo in enumerate(plos_data):
                        cat = (plo.get('category') or '').strip()
                        if cat != cur_cat:
                            if cur_indices:
                                plo_domain_groups.append((cur_cat or '—', cur_indices))
                            cur_cat = cat
                            cur_indices = [i]
                        else:
                            cur_indices.append(i)
                    if cur_indices:
                        plo_domain_groups.append((cur_cat or '—', cur_indices))
                    chunks.append('<table><tr><th colspan="%d">%s</th></tr>' % (1 + nplo, esc(cfg['label']).upper()))
                    chunks.append('<tr><th>Course Code and Course Title</th><th colspan="%d" class="text-center">Program Learning Outcomes (PLOs)</th></tr>' % nplo)
                    domain_cells = ''.join('<th colspan="%d" class="text-center">%s</th>' % (len(idx), esc(lbl)) for lbl, idx in plo_domain_groups) if plo_domain_groups else '<th colspan="%d">—</th>' % nplo
                    chunks.append('<tr><th></th>%s</tr>' % domain_cells)
                    chunks.append('<tr><th></th>' + ''.join('<th class="text-center">%s</th>' % esc(p.get('id')) for p in plos_data) + '</tr>')
                    for i, e in enumerate(entries):
                        row = [esc((e.course_code or '') + ' / ' + (e.course_name or ''))]
                        for j in range(nplo):
                            v = '•' if (i < len(stored) and j < len(stored[i]) and stored[i][j]) else ''
                            row.append('<td class="text-center">%s</td>' % esc(v))
                        chunks.append('<tr><td>%s</td>' % row[0] + ''.join(row[1:]) + '</tr>')
                    chunks.append('</table>')
                    pdf_table_sep()

        elif key == 'B' and part_b:
            cfg_rec = SyllabusPartBConfig.query.filter_by(part_id=part_b.id).first()
            cfg = cfg_rec.get_config() if cfg_rec else {}
            td = cfg.get('term_duration') or {}
            area_wise = cfg.get('area_wise_override') if isinstance(cfg.get('area_wise_override'), list) else None
            category = cfg.get('category_override') if isinstance(cfg.get('category_override'), list) else None
            if not area_wise or not category:
                aw, cat = _derive_part_b_from_c(all_c)
                if not area_wise:
                    area_wise = aw
                if not category:
                    category = cat
            structure_curriculum = cfg.get('structure_curriculum')
            if isinstance(structure_curriculum, dict) and structure_curriculum.get('column_headers') and structure_curriculum.get('rows'):
                headers = structure_curriculum['column_headers']
                rows_data = structure_curriculum['rows']
                chunks.append('<table><tr><th>Structure of the Curriculum</th>' + ''.join('<th>%s</th>' % esc(h) for h in headers) + '</tr>')
                for row in rows_data:
                    vals = row.get('values') or []
                    cells = '<td>%s</td>' % esc(row.get('label'))
                    for j, h in enumerate(headers):
                        cells += '<td>%s</td>' % esc(vals[j] if j < len(vals) else '')
                    chunks.append('<tr>%s</tr>' % cells)
                chunks.append('</table>')
                pdf_table_sep()
            elif isinstance(structure_curriculum, list) and structure_curriculum:
                chunks.append('<table><tr><th>Structure of the Curriculum</th><th>Value 1</th><th>Value 2</th></tr>')
                for row in structure_curriculum:
                    chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td></tr>' % (
                        esc(row.get('label')), esc(row.get('value_1')), esc(row.get('value_2'))))
                chunks.append('</table>')
                pdf_table_sep()
            else:
                chunks.append('<table><tr><th>Duration</th><td>%s Years</td><td>%s</td></tr></table>' % (esc(cfg.get('duration_years', 4)), esc(cfg.get('terms', 'Terms'))))
                pdf_table_sep()
            chunks.append('<table><tr><th>Term duration (weeks)</th></tr><tr><td>Teaching %s, Prep %s, Exam %s, Break %s, Total %s</td></tr></table>' % (
                esc(td.get('teaching')), esc(td.get('preparatory')), esc(td.get('exam')), esc(td.get('break')), esc(td.get('total'))))
            pdf_table_sep()
            if area_wise:
                chunks.append('<table><tr><th>Area</th><th>Type</th><th>No. Courses</th><th>Credits</th><th>Total</th></tr>')
                for r in area_wise:
                    chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>' % (
                        esc(r.get('area')), esc(r.get('type')), esc(r.get('num_courses')), esc(r.get('credits')), esc(r.get('total_credit'))))
                chunks.append('</table>')
                pdf_table_sep()
            if category:
                chunks.append('<table><tr><th>Area</th><th>Type</th><th>Title</th><th>Credit</th></tr>')
                for r in category:
                    chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>' % (
                        esc(r.get('area')), esc(r.get('course_type')), esc(r.get('course_title')), esc(r.get('credit'))))
                chunks.append('</table>')
                pdf_table_sep()
            # Course distribution by term (LLB-style 8-column tables)
            for g in _course_distribution_by_term(all_c):
                entries = g.get('entries') or []
                chunks.append('<table>')
                chunks.append('<tr><th colspan="8">COURSE DISTRIBUTION BY TERM</th></tr>')
                chunks.append('<tr><th colspan="8">%s</th></tr>' % esc((g['year'] + ' YEAR ' + g['term'] + ' TERM').upper()))
                chunks.append('<tr><th>Course Code</th><th>Course Title</th><th>Course Status</th><th>Notional Hours</th><th>Contact Hours/Week</th><th>Contact Hours/Week</th><th>Credits</th><th>Prerequisites</th></tr>')
                chunks.append('<tr><th>Course Code</th><th>Course Title</th><th>Course Status</th><th>Notional Hours</th><th>Theory</th><th>Sessional</th><th>Credits</th><th>Prerequisites</th></tr>')
                for e in entries:
                    try:
                        cr = float(e.credit) if e.credit is not None else 0
                    except Exception:
                        cr = 0
                    et = (e.entry_type or 'Theory').strip().lower()
                    notional = '60' if 'sessional' in et else str(int(round(cr * 40)))
                    th = ('%.1f' % cr).rstrip('0').rstrip('.') if ('theory' in et or not et) else '-'
                    se = ('%.1f' % (cr * 1.5)).rstrip('0').rstrip('.') if 'sessional' in et else ('-' if th == '-' else '-')
                    prereq = 'None'
                    if e.prerequisite_entry_id:
                        pe = SyllabusCourseEntry.query.get(e.prerequisite_entry_id)
                        if pe and pe.course_code:
                            prereq = pe.course_code
                    chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>' % (
                        esc(e.course_code or ''), esc(e.course_name or ''), esc(e.status or ''), esc(notional),
                        esc(th), esc(se), esc(e.credit if e.credit is not None else ''), esc(prereq)
                    ))
                chunks.append('</table>')
                pdf_table_sep()

        elif key == 'C' and part_c and all_c:
            chunks.append('<table><tr><th>COURSE DESCRIPTION BY TERM</th></tr></table>')
            pdf_table_sep()
            for g in _course_distribution_by_term(all_c):
                chunks.append('<table><tr><th>%s</th></tr></table>' % esc((g['year'] + ' YEAR ' + g['term'] + ' TERM').upper()))
                pdf_table_sep()
                for e in g['entries']:
                    chunks.append('<div class="course-page">')
                    c = e.get_content_dict() or {}
                    prerequisite = (c.get('prerequisite') or '').strip()
                    if not prerequisite and e.prerequisite_entry_id:
                        pe = SyllabusCourseEntry.query.get(e.prerequisite_entry_id)
                        prerequisite = (pe.course_code or 'None') if pe else 'None'
                    if not prerequisite:
                        prerequisite = 'None'
                    rationale = c.get('rationale') or ''

                    # Course header (LLB-style 4x5 with colspans)
                    chunks.append('<table>')
                    chunks.append('<tr><td colspan="2">%s</td><td>%s</td><td>%s</td><td>%s</td></tr>' % (
                        esc('Course No.: ' + (e.course_code or '')),
                        esc('Credit: ' + (str(e.credit) if e.credit is not None else '')),
                        esc('Year: ' + (e.year or '')),
                        esc('Term: ' + (e.term or '')),
                    ))
                    chunks.append('<tr><td colspan="3">%s</td><td colspan="2">%s</td></tr>' % (
                        esc('Course Title: ' + (e.course_name or '')),
                        esc('Course Status: ' + (e.status or '')),
                    ))
                    chunks.append('<tr><td>Prerequisite</td><td colspan="4">%s</td></tr>' % esc(prerequisite))
                    chunks.append('<tr><td>Rationale</td><td colspan="4">%s</td></tr>' % esc(rationale))
                    chunks.append('</table>')
                    pdf_table_sep()

                    # COURSE CONTENTS (LLB-style: full-width header, Section A|CLOs + rows, Section B|CLOs + rows)
                    items_a_raw = c.get('section_a_items') or []
                    items_b_raw = c.get('section_b_items') or []
                    if not items_a_raw and not items_b_raw:
                        raw_a = (c.get('section_a') or '').strip()
                        raw_b = (c.get('section_b') or '').strip()
                        items_a = [{'content': line.strip(), 'clos': ''} for line in raw_a.split('\n') if line.strip()]
                        items_b = [{'content': line.strip(), 'clos': ''} for line in raw_b.split('\n') if line.strip()]
                        if not items_a:
                            items_a = [{'content': '—', 'clos': ''}]
                        if not items_b:
                            items_b = [{'content': '—', 'clos': ''}]
                    else:
                        items_a = [{'content': (x.get('content') or '').strip(), 'clos': (x.get('clos') or '').strip()} for x in items_a_raw] if items_a_raw else [{'content': '—', 'clos': ''}]
                        items_b = [{'content': (x.get('content') or '').strip(), 'clos': (x.get('clos') or '').strip()} for x in items_b_raw] if items_b_raw else [{'content': '—', 'clos': ''}]
                    n_a, n_b = len(items_a), len(items_b)
                    chunks.append('<table>')
                    chunks.append('<tr><th colspan="3">COURSE CONTENTS</th></tr>')
                    chunks.append('<tr><th colspan="2">Section A</th><th>CLOs</th></tr>')
                    for i, it in enumerate(items_a):
                        chunks.append('<tr><td>%d</td><td>%s</td><td class="text-end">%s</td></tr>' % (i + 1, esc(it.get('content') or ''), esc(it.get('clos') or '')))
                    chunks.append('<tr><th colspan="2">Section B</th><th></th></tr>')
                    for i, it in enumerate(items_b):
                        chunks.append('<tr><td>%d</td><td>%s</td><td class="text-end">%s</td></tr>' % (n_a + i + 1, esc(it.get('content') or ''), esc(it.get('clos') or '')))
                    chunks.append('</table>')
                    pdf_table_sep()

                    # CLO table (LLB-style)
                    clos = c.get('clos') or []
                    if not isinstance(clos, list):
                        clos = []
                    nrows = max(1, len(clos))
                    chunks.append('<table>')
                    chunks.append('<tr><th rowspan="%d">Course Learning Outcomes (CLOs)</th><th colspan="2">Upon completion of this course, the student will be able to:</th><th>Mapping with PLO</th></tr>' % (1 + nrows))
                    if clos:
                        for i, cl in enumerate(clos):
                            chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td></tr>' % (
                                esc('CLO %d' % (i + 1)),
                                esc(cl.get('text') or cl.get('upon_completion') or ''),
                                esc(cl.get('plo') or ''),
                            ))
                    else:
                        chunks.append('<tr><td>CLO 1</td><td></td><td></td></tr>')
                    chunks.append('</table>')
                    pdf_table_sep()

                    # Mapping CLOs table (from mapping_clos or placeholder)
                    mapping_clos = c.get('mapping_clos') or []
                    if not isinstance(mapping_clos, list):
                        mapping_clos = []
                    n_map = max(1, len(clos))
                    chunks.append('<table><tr><th colspan="3">MAPPING CLOs WITH THE TEACHING-LEARNING STRATEGY AND ASSESSMENT STRATEGY</th></tr>'
                                  '<tr><th>CLOs</th><th>Teaching-Learning Strategy</th><th>Assessment Strategy</th></tr>')
                    for i in range(n_map):
                        row = mapping_clos[i] if i < len(mapping_clos) else {}
                        tl = esc((row.get('teaching_learning') or '').strip())
                        ass = esc((row.get('assessment') or '').strip())
                        chunks.append('<tr><td>%d</td><td>%s</td><td>%s</td></tr>' % (i + 1, tl, ass))
                    chunks.append('</table>')
                    pdf_table_sep()

                    # Indicative learning materials (placeholder)
                    chunks.append('<table><tr><th colspan="2">INDICATIVE LEARNING MATERIALS</th></tr>'
                                  '<tr><td>Recommended Readings</td><td>%s</td></tr>'
                                  '<tr><td>Supplementary Readings</td><td>%s</td></tr></table>' % (
                                      esc(c.get('recommended_readings') or ''), esc(c.get('supplementary_readings') or '')
                                  ))
                    pdf_table_sep()
                    chunks.append('</div>')

        elif key == 'D' and part_d:
            chunks.append('<p><strong>20 Grading and Evaluation</strong></p>')
            chunks.append('<p><strong>20.1 Grading Scale</strong></p>')
            chunks.append('<p>a) Letter Grades and corresponding Grade Points will be awarded following provisions shown below:</p>')
            rec = SyllabusPartDSection.query.filter_by(part_id=part_d.id, section_key='grading_scale').first()
            data = rec.get_data() if rec else []
            if isinstance(data, list) and data:
                chunks.append('<table><tr><th>Numerical Grade</th><th>Letter Grade</th><th>Grade Point</th></tr>')
                for r in data:
                    chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td></tr>' % (esc(r.get('numerical')), esc(r.get('letter')), esc(r.get('grade_point'))))
                chunks.append('</table>')
                pdf_table_sep()
            chunks.append('<p><strong>20.1.1 Evaluation of Theory Courses</strong></p>')
            chunks.append('<p>a) All theory courses will be evaluated out of 100 marks, the distribution of which is given below:</p>')
            rec = SyllabusPartDSection.query.filter_by(part_id=part_d.id, section_key='theory_evaluation').first()
            data = rec.get_data() if rec else []
            if isinstance(data, list) and data:
                chunks.append('<table><tr><th>Sl. No.</th><th>Items</th><th>Marks</th></tr>')
                for r in data:
                    chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td></tr>' % (esc(r.get('sl_no')), esc(r.get('items')), esc(r.get('marks'))))
                chunks.append('</table>')
                pdf_table_sep()
            chunks.append('<p><strong>20.1.2 Evaluation of Sessional Courses</strong></p>')
            chunks.append('<p>a) All sessional courses will be evaluated out of 100 marks, the distribution of which is given below:</p>')
            rec = SyllabusPartDSection.query.filter_by(part_id=part_d.id, section_key='sessional_evaluation').first()
            data = rec.get_data() if rec else []
            if isinstance(data, list) and data:
                chunks.append('<table><tr><th>Sl. No.</th><th>Items</th><th>Marks</th></tr>')
                for r in data:
                    chunks.append('<tr><td>%s</td><td>%s</td><td>%s</td></tr>' % (esc(r.get('sl_no')), esc(r.get('items')), esc(r.get('marks'))))
                chunks.append('</table>')
                pdf_table_sep()
            rec = SyllabusPartDSection.query.filter_by(part_id=part_d.id, section_key='approval_records').first()
            data = rec.get_data() if rec else []
            if isinstance(data, list) and data:
                chunks.append('<p class="text-center"><strong>Approval Records</strong></p>')
                chunks.append('<table><tr><th>Approving Authority</th><th>Date of Approval</th></tr>')
                for r in data:
                    chunks.append('<tr><td>%s</td><td>%s</td></tr>' % (esc(r.get('authority')), esc(r.get('date'))))
                chunks.append('</table>')
                pdf_table_sep()

    chunks.append('</body></html>')
    html = ''.join(chunks)
    try:
        h = HTML(string=html, base_url=formal_fonts['fonts_dir'].as_uri() + '/')
        buf = BytesIO()
        h.write_pdf(buf, presentational_hints=True)
        buf.seek(0)
    except Exception as e:
        current_app.logger.exception('Curriculator PDF generation failed')
        raise RuntimeError('PDF generation failed: %s' % str(e)) from e
    name = (doc.name or 'syllabus').replace(' ', '_')[:50]
    fn = f'{name}_export_{datetime.utcnow().strftime("%Y%m%d_%H%M")}.pdf'
    return buf, fn
