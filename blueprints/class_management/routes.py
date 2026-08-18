from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, current_app, Response, jsonify
from flask_login import login_required, current_user
from utils.timezone import format_bd
from utils.tenant import current_tenant
from utils.academic_rules import assessment_cfg, take_best_marks, scale_pg_total, result_split
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from sqlalchemy import or_, text, func
from sqlalchemy.orm import aliased
from .models import (
    db,
    Teacher,
    Session,
    ClassStudent,
    ClassAttendance,
    ClassSplitInvite,
    CourseReview,
    EvaluationInvite,
    EvaluationSubmission,
    ExamScrutinizerInvite,
    ExamPaperEvaluation,
    StudentFeedbackLink,
    StudentFeedbackResponse,
    CourseOutline,
    CourseQuestionThread,
    CourseQuestionMessage,
    CourseQuestionAttachment,
    StudentNotification,
    QuestionBankFile,
    QuestionBankFolder,
)
from user_models import User  # Same User as login (users table); required for StudentNotification.user_id
try:
    from blueprints.student_management.models import Student
except ImportError:
    Student = None

try:
    from blueprints.course_management.models import Curriculum, Course, CurriculumYearTerm, CourseSessionAssignment, StudentCourseRegistration
except ImportError:
    Curriculum = None
    CourseSessionAssignment = None
    Course = None
    CurriculumYearTerm = None
    StudentCourseRegistration = None
import pandas as pd
import os
from datetime import datetime, date
from utils.timezone import format_bd, bd_now
from decimal import Decimal, ROUND_HALF_UP
import secrets
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from collections import Counter, defaultdict
from reportlab.lib.units import inch
import io
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import json
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
# docx imports moved to lazy imports (only when needed) to prevent startup hang
# from docx import Document
# from docx.shared import Pt, Inches
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from uuid import uuid4
from role_utils import has_teacher_privileges, is_admin, parse_roles
from utils.ownership import user_owns_class_session
try:
    from utils.semester_utils import filter_by_active_semester
except ImportError:
    filter_by_active_semester = None
try:
    from utils.window_utils import filter_by_active_window, get_effective_window_id, query_for_window, get_for_window, get_or_404_for_window, filter_offered_courses, stamp_window_id
except ImportError:
    filter_by_active_window = None
    get_effective_window_id = None
    query_for_window = None
    filter_offered_courses = None
    stamp_window_id = None


def _class_students_for_session(session_id, order=True):
    """ClassStudent list for a session, excluding students deleted from Student Management."""
    q = ClassStudent.query.filter(ClassStudent.session_id == session_id)
    session = Session.query.get(session_id)
    if Student and not getattr(session, 'is_external_course', False):
        q = q.filter(ClassStudent.student_id.in_(db.session.query(Student.student_id)))
    if order:
        q = q.order_by(ClassStudent.student_id)
    return q.all()


def _class_students_for_sessions(session_ids, order=True):
    """ClassStudent list for given sessions, excluding deleted students."""
    if not session_ids:
        return []
    q = ClassStudent.query.filter(ClassStudent.session_id.in_(session_ids))
    if Student:
        has_external = Session.query.filter(
            Session.id.in_(session_ids),
            Session.is_external_course.is_(True),
        ).first() is not None
        if not has_external:
            q = q.filter(ClassStudent.student_id.in_(db.session.query(Student.student_id)))
    if order:
        q = q.order_by(ClassStudent.student_id)
    return q.all()


def _notify_students_in_session(session_id, notif_type, title, link_url):
    """Create a StudentNotification for each student in the given session (by student_id -> User.username)."""
    try:
        class_students = _class_students_for_session(session_id, order=False)
        seen_user_ids = set()
        for cs in class_students:
            user = User.query.filter_by(username=cs.student_id).first()
            if user and user.id not in seen_user_ids:
                seen_user_ids.add(user.id)
                n = StudentNotification(user_id=user.id, type=notif_type, title=title, link_url=link_url)
                db.session.add(n)
        db.session.commit()
    except Exception as e:
        current_app.logger.warning(f"Could not create student notifications: {e}")
        db.session.rollback()


def _notify_student_by_username(student_id_username, notif_type, title, link_url):
    """Create a single StudentNotification for the user with username=student_id_username."""
    try:
        user = User.query.filter_by(username=student_id_username).first()
        if user:
            n = StudentNotification(user_id=user.id, type=notif_type, title=title, link_url=link_url)
            db.session.add(n)
            db.session.commit()
    except Exception as e:
        current_app.logger.warning(f"Could not create student notification: {e}")
        db.session.rollback()


def _absolute_url_student_my_scores():
    """Absolute URL for My Scores (reveal emails). Prefer PUBLIC_APP_URL behind reverse proxies."""
    public = (current_app.config.get('PUBLIC_APP_URL') or os.environ.get('PUBLIC_APP_URL') or '').strip().rstrip('/')
    if public:
        return f"{public}/class-management/student/view-scores"
    try:
        return url_for('class_management.student_view_scores', _external=True)
    except Exception:
        pass
    try:
        rel = url_for('class_management.student_view_scores')
        if request and getattr(request, 'host_url', None):
            return request.host_url.rstrip('/') + rel
    except Exception:
        pass
    return url_for('class_management.student_view_scores')


def _send_marks_revealed_email_to_session_students(session_id, course_label, assessment_type=None):
    """Send marks-revealed email via NOTIFICATION_MAIL_* (noreply@) only."""
    from utils.notification_email import _notification_smtp_configured, send_notification_batch

    if not _notification_smtp_configured():
        current_app.logger.error(
            'marks_revealed email skipped: set NOTIFICATION_MAIL_USERNAME/PASSWORD/SENDER'
        )
        return

    try:
        class_students = _class_students_for_session(session_id, order=False)
        seen_emails = set()
        view_scores_url = _absolute_url_student_my_scores()
        assessment_label = (assessment_type or 'assessment').replace('_', ' ').title()
        skipped_no_user = 0
        skipped_no_email = 0

        entries = []
        for cs in class_students:
            user = User.query.filter_by(username=cs.student_id).first()
            if not user:
                skipped_no_user += 1
                continue
            if not (user.email and str(user.email).strip()):
                skipped_no_email += 1
                continue
            email_key = user.email.strip().lower()
            if email_key in seen_emails:
                continue
            seen_emails.add(email_key)

            student_name = (user.full_name or cs.student_id or 'Student').strip()
            student_id = (cs.student_id or user.username or '').strip()
            student_email = user.email.strip()
            # Hosting Bangladesh: personalized per student; specific subject;
            # clear link purpose; HTML + plain-text; no identical mass template.
            subject = (
                f"Your {assessment_label} result is now visible in AMS "
                f"for {course_label} — {student_name} ({student_id})"
            )
            t = current_tenant()
            text_body = (
                f"Dear {student_name},\n\n"
                "This academic notification is from the Academic Management System (AMS) "
                f"of the {t.display_with_university}.\n\n"
                "Your course teacher has published / revealed marks for an assessment in "
                "your enrolled course session. The details for your account are:\n\n"
                f"- Student name: {student_name}\n"
                f"- Student ID: {student_id}\n"
                f"- Student email: {student_email}\n"
                f"- Course / class session: {course_label}\n"
                f"- Assessment name: {assessment_label}\n\n"
                "Purpose of the link below:\n"
                "The link opens your personal My Scores page in AMS so you can view the "
                f"{assessment_label} marks for {course_label}. It is not a generic website "
                "link; it is the scores page for students of this system. You must sign in "
                "with your AMS student username to see your own marks.\n\n"
                f"My Scores page address:\n{view_scores_url}\n\n"
                "How to view your result:\n"
                "1. Open the My Scores address above in your browser.\n"
                "2. Sign in with your AMS student account if asked.\n"
                f"3. Check the marks shown for {assessment_label} under {course_label}.\n\n"
                "If the address does not open from a click, copy the full line and paste it "
                "into your browser address bar.\n\n"
                "If you are not enrolled in this course or believe this message was sent "
                f"by mistake, contact your course teacher or the {t.office_label}. "
                "No action is required if you already reviewed your marks in AMS.\n\n"
                "Regards,\n"
                "Academic Management System\n"
                f"{t.display_with_university}\n"
                f"Sender: {(current_app.config.get('NOTIFICATION_MAIL_SENDER') or current_app.config.get('NOTIFICATION_MAIL_USERNAME') or '').strip()}\n"
            )
            html_body = (
                f"<p>Dear {student_name},</p>"
                "<p>This academic notification is from the Academic Management System (AMS) "
                f"of the {t.display_with_university}.</p>"
                "<p>Your course teacher has published / revealed marks for an assessment in "
                "your enrolled course session. The details for your account are:</p>"
                "<ul>"
                f"<li>Student name: {student_name}</li>"
                f"<li>Student ID: {student_id}</li>"
                f"<li>Student email: {student_email}</li>"
                f"<li>Course / class session: {course_label}</li>"
                f"<li>Assessment name: {assessment_label}</li>"
                "</ul>"
                "<p><strong>Purpose of the link below:</strong><br>"
                "The link opens your personal My Scores page in AMS so you can view the "
                f"{assessment_label} marks for {course_label}. You must sign in with your "
                "AMS student username to see your own marks.</p>"
                f"<p><a href=\"{view_scores_url}\">"
                f"Open My Scores to view {assessment_label} marks for {student_name} ({student_id})"
                "</a></p>"
                f"<p>Full address (copy and paste if needed):<br>"
                f"<span style=\"word-break:break-all;\">{view_scores_url}</span></p>"
                "<p><strong>How to view your result:</strong></p>"
                "<ol>"
                "<li>Open the My Scores address above in your browser.</li>"
                "<li>Sign in with your AMS student account if asked.</li>"
                f"<li>Check the marks shown for {assessment_label} under {course_label}.</li>"
                "</ol>"
                "<p>If you are not enrolled in this course or believe this message was sent "
                f"by mistake, contact your course teacher or the {t.office_label}.</p>"
                "<p>Regards,<br>"
                "Academic Management System<br>"
                f"{t.display_with_university}<br>"
                f"Sender: {(current_app.config.get('NOTIFICATION_MAIL_SENDER') or current_app.config.get('NOTIFICATION_MAIL_USERNAME') or '').strip()}</p>"
            )
            entries.append({
                'recipient': student_email,
                'subject': subject,
                'text_body': text_body,
                'html_body': html_body,
            })

        if not entries:
            current_app.logger.warning(
                f"marks_revealed email: session {session_id} has no recipients "
                f"(class_students={len(class_students)}, no_user={skipped_no_user}, no_email={skipped_no_email})"
            )
            return

        sent_count = send_notification_batch(None, entries)
        current_app.logger.info(
            f"Marks-revealed email for session {session_id}: {sent_count} message(s) sent "
            f"(unique recipients={len(entries)})"
        )
    except Exception as e:
        current_app.logger.error(f"marks_revealed email unexpected error: {e}", exc_info=True)


# WeasyPrint lazy import - only import when needed to prevent startup hang
# Module-level import removed because it causes startup hang on macOS
# This prevents the app from hanging during startup
_WEASYPRINT_HTML = None
_WEASYPRINT_AVAILABLE = None

def _resolve_formal_pdf_fonts():
    """Compatibility wrapper — prefer utils.pdf_fonts.resolve_formal_pdf_fonts."""
    from utils.pdf_fonts import resolve_formal_pdf_fonts
    return resolve_formal_pdf_fonts()


def _resolve_dejavu_pdf_fonts():
    """Compatibility wrapper — prefer utils.pdf_fonts.resolve_dejavu_pdf_fonts."""
    from utils.pdf_fonts import resolve_dejavu_pdf_fonts
    return resolve_dejavu_pdf_fonts()


def _get_weasyprint_html():
    """Lazy import WeasyPrint HTML - only import when actually needed"""
    global _WEASYPRINT_HTML, _WEASYPRINT_AVAILABLE
    
    if _WEASYPRINT_AVAILABLE is None:
        # First time - try to import
        import logging
        import os
        import platform
        import ctypes
        from ctypes import util as ctypes_util
        
        logger = logging.getLogger(__name__)
        
        # Setup library paths for macOS BEFORE importing WeasyPrint
        if platform.system() == 'Darwin':
            homebrew_lib_path = '/opt/homebrew/lib'
            if os.path.exists(homebrew_lib_path):
                # Set environment variables
                os.environ['DYLD_FALLBACK_LIBRARY_PATH'] = f"{homebrew_lib_path}:{os.environ.get('DYLD_FALLBACK_LIBRARY_PATH', '')}"
                os.environ['PKG_CONFIG_PATH'] = f"/opt/homebrew/lib/pkgconfig:{os.environ.get('PKG_CONFIG_PATH', '')}"
                
                # Monkey-patch ctypes.util.find_library
                original_find_library = ctypes_util.find_library
                def patched_find_library(name):
                    lib_mappings = {
                        'gobject-2.0-0': 'libgobject-2.0.0.dylib',
                        'gobject-2.0': 'libgobject-2.0.dylib',
                    }
                    if name in lib_mappings:
                        lib_path = os.path.join(homebrew_lib_path, lib_mappings[name])
                        if os.path.exists(lib_path):
                            return lib_path
                    for pattern in [f'lib{name}.dylib', f'lib{name}.0.dylib']:
                        lib_path = os.path.join(homebrew_lib_path, pattern)
                        if os.path.exists(lib_path):
                            return lib_path
                    result = original_find_library(name)
                    return result if result else None
                
                ctypes_util.find_library = patched_find_library
                
                # Pre-load libraries
                try:
                    lib_path = os.path.join(homebrew_lib_path, 'libgobject-2.0.0.dylib')
                    if os.path.exists(lib_path):
                        ctypes.CDLL(lib_path, mode=ctypes.RTLD_GLOBAL)
                except:
                    pass
        
        try:
            logger.info("Attempting to import WeasyPrint (lazy import)...")
            from weasyprint import HTML
            _WEASYPRINT_HTML = HTML
            _WEASYPRINT_AVAILABLE = True
            logger.info("✓ WeasyPrint imported successfully (lazy import)")
        except ImportError as e:
            _WEASYPRINT_AVAILABLE = False
            _WEASYPRINT_HTML = None
            logger.error(f"✗ WeasyPrint ImportError: {e}")
            logger.warning("PDF generation features will be disabled.")
        except Exception as e:
            _WEASYPRINT_AVAILABLE = False
            _WEASYPRINT_HTML = None
            logger.error(f"✗ WeasyPrint import error: {e}", exc_info=True)
            logger.warning("PDF generation features will be disabled.")
    
    if _WEASYPRINT_AVAILABLE and _WEASYPRINT_HTML is None:
        # Retry if somehow HTML is None but available is True
        try:
            from weasyprint import HTML
            _WEASYPRINT_HTML = HTML
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Failed to re-import WeasyPrint: {e}")
            _WEASYPRINT_AVAILABLE = False
    
    return _WEASYPRINT_HTML if _WEASYPRINT_AVAILABLE else None

def _is_weasyprint_available():
    """Check if WeasyPrint is available (lazy check)"""
    if _WEASYPRINT_AVAILABLE is None:
        _get_weasyprint_html()  # Trigger lazy import
    return _WEASYPRINT_AVAILABLE is True


COURSE_REVIEW_GRADE_ROWS = [
    {'key': 'grade_a_plus', 'scale': '80% and above', 'letter': 'A+'},
    {'key': 'grade_a', 'scale': '75% to less than 80%', 'letter': 'A'},
    {'key': 'grade_a_minus', 'scale': '70% to less than 75%', 'letter': 'A-'},
    {'key': 'grade_b_plus', 'scale': '65% to less than 70%', 'letter': 'B+'},
    {'key': 'grade_b', 'scale': '60% to less than 65%', 'letter': 'B'},
    {'key': 'grade_b_minus', 'scale': '55% to less than 60%', 'letter': 'B-'},
    {'key': 'grade_c_plus', 'scale': '50% to less than 55%', 'letter': 'C+'},
    {'key': 'grade_c', 'scale': '45% to less than 50%', 'letter': 'C'},
    {'key': 'grade_d', 'scale': '40% to less than 45%', 'letter': 'D'},
    {'key': 'grade_f', 'scale': 'Less than 40%', 'letter': 'F'},
    {'key': 'grade_i', 'scale': 'Incomplete', 'letter': 'I'},
    {'key': 'grade_w', 'scale': 'Withdrawal', 'letter': 'W'},
]

COURSE_REVIEW_COMMENT_FIELDS = [
    {'key': 'comment_student_questionnaires', 'label': '1) Student (Course Evaluation) Questionnaires:'},
    {'key': 'comment_external_examiners', 'label': '2) External Examiners or Moderators (if any):'},
    {'key': 'comment_curriculum', 'label': '3) Curriculum: Comment on the continuing appropriateness of the Course curriculum in relation to the intended learning outcomes and its compliance with the National Qualification Framework'},
    {'key': 'comment_assessment', 'label': '4) Assessment: Comment on the continuing effectiveness of method(s) of assessment in relation to the intended learning outcomes (Course objectives)'},
    {'key': 'comment_enhancement', 'label': '5) Enhancement: Comment on the implementation of changes proposed in earlier Faculty Course Review Reports'},
    {'key': 'comment_future_changes', 'label': "6) Outline any changes in the future delivery or structure of the Course that this semester/term's experience may prompt"},
]

SCOPE_FULL = 'full'
SCOPE_PART_A = 'part_a'
SCOPE_PART_B = 'part_b'
SPLIT_PARTS = {SCOPE_PART_A, SCOPE_PART_B}
COURSE_SCOPE_LABELS = {
    SCOPE_FULL: 'Full Course',
    SCOPE_PART_A: 'Part A',
    SCOPE_PART_B: 'Part B',
}


def _normalize_offering_text(value):
    return (value or '').strip().lower()


def _session_offering_key(session):
    """
    Identity key for one teacher offering on Active Courses.
    Part A/B stay distinct via course_scope.
    window_id / academic_session are intentionally omitted: reassignment often
    creates a second row with a newer window or filled session string that still
    represents the same offering in the current dashboard list.
    """
    return (
        session.teacher_id,
        _normalize_offering_text(session.course_code),
        _normalize_offering_text(session.year),
        _normalize_offering_text(session.term),
        _normalize_offering_text(getattr(session, 'course_scope', None) or SCOPE_FULL),
    )


def _pick_canonical_session(candidates):
    """Prefer CSA-linked session, then most students, then oldest id."""
    def score(session):
        has_csa = 0
        if CourseSessionAssignment:
            has_csa = 1 if CourseSessionAssignment.query.filter_by(session_id=session.id).first() else 0
        student_count = ClassStudent.query.filter_by(session_id=session.id).count()
        # Prefer older session when other scores tie (keeps history).
        return (has_csa, student_count, -(session.id or 0))

    return max(candidates, key=score)


def _dedupe_active_sessions(sessions):
    """
    Collapse duplicate active offerings for the same teacher/course/scope.
    Archives losers after moving students + CSA links onto the winner.
    Always returns a deduped list for the UI even if archive commit fails.
    """
    if not sessions:
        return sessions

    groups = defaultdict(list)
    for session in sessions:
        groups[_session_offering_key(session)].append(session)

    kept = []
    changed = False
    for key, group in groups.items():
        if len(group) == 1:
            kept.append(group[0])
            continue

        winner = _pick_canonical_session(group)
        kept.append(winner)
        current_app.logger.warning(
            'Duplicate Active Courses detected for key=%s count=%s winner=%s losers=%s',
            key,
            len(group),
            winner.id,
            [s.id for s in group if s.id != winner.id],
        )
        for loser in group:
            if loser.id == winner.id:
                continue
            # Heal winner fields from loser when winner is missing them
            if not (winner.academic_session or '').strip() and (loser.academic_session or '').strip():
                winner.academic_session = loser.academic_session
            if winner.window_id is None and loser.window_id is not None:
                winner.window_id = loser.window_id

            # Move unique students onto the canonical session
            for class_student in ClassStudent.query.filter_by(session_id=loser.id).all():
                already = ClassStudent.query.filter_by(
                    session_id=winner.id,
                    student_id=class_student.student_id,
                ).first()
                if already:
                    continue
                class_student.session_id = winner.id
                class_student.teacher_id = winner.teacher_id

            if CourseSessionAssignment:
                for assignment in CourseSessionAssignment.query.filter_by(session_id=loser.id).all():
                    assignment.session_id = winner.id
                    assignment.session_created = True

            loser.archived = True
            changed = True
            current_app.logger.warning(
                'Archived duplicate class session %s in favor of %s (%s / %s / %s / %s)',
                loser.id,
                winner.id,
                loser.course_code,
                loser.year,
                loser.term,
                getattr(loser, 'course_scope', SCOPE_FULL),
            )

    if changed:
        try:
            db.session.commit()
        except Exception as dedupe_error:
            db.session.rollback()
            current_app.logger.error(
                f'Error archiving duplicate sessions (UI will still hide duplicates): {dedupe_error}',
                exc_info=True,
            )

    # Preserve newest-first ordering used by the dashboard
    kept.sort(key=lambda s: s.created_at or datetime(1970, 1, 1), reverse=True)
    return kept


def _assessment_column_labels():
    cfg = assessment_cfg()
    return {
        'best_three': f'Total of Best {cfg["take_best"]} ({cfg["ug_out_of"]})',
        'part_a_b_15': (
            f'Part A ({cfg["part_a_b_each"]}) + Part B ({cfg["part_a_b_each"]}) '
            f'({cfg["ug_out_of"]})'
        ),
        'best_three_40': f'Total ({cfg["pg_out_of"]})',
    }


def _external_assessment_modes():
    cfg = assessment_cfg()
    return {
        'best_three': f'Best {cfg["take_best"]}',
        'part_a_b_15': (
            f'Best of Part A ({cfg["part_a_b_each"]}) + Best of Part B ({cfg["part_a_b_each"]})'
        ),
        'best_three_40': f'Best {cfg["take_best"]} converted into {cfg["pg_out_of"]}',
    }


EXTERNAL_ASSESSMENT_MODES = {
    'best_three': 'Best three',
    'part_a_b_15': 'Best of Part A (15) + Best of Part B (15)',
    'best_three_40': 'Best three converted into 40',
}
DEFAULT_EXTERNAL_ASSESSMENT_MODE = 'best_three'
EXTERNAL_ASSESSMENT_COLUMN_LABELS = {
    'best_three': 'Total of Best 3 (30)',
    'part_a_b_15': 'Part A (15) + Part B (15) (30)',
    'best_three_40': 'Total (40)',
}


def _normalize_external_assessment_mode(mode):
    normalized = (mode or DEFAULT_EXTERNAL_ASSESSMENT_MODE).strip().lower()
    if normalized not in EXTERNAL_ASSESSMENT_MODES:
        return DEFAULT_EXTERNAL_ASSESSMENT_MODE
    return normalized


def _external_assessment_mark_max(mode, assess_idx=None):
    cfg = assessment_cfg()
    if _normalize_external_assessment_mode(mode) == 'part_a_b_15':
        return cfg['part_a_b_each']
    return cfg['slot_max']


def _external_assessment_column_header(assess_idx, mode):
    return f'Assessment {assess_idx} ({_external_assessment_mark_max(mode)})'


def _parse_external_assessment_value(raw_value, session):
    """Parse and clamp assessment input for external theory sessions."""
    if raw_value in (None, ''):
        return None
    try:
        value = float(raw_value)
    except (TypeError, ValueError):
        return None
    if _is_external_theory_session(session):
        mark_max = _external_assessment_mark_max(getattr(session, 'external_assessment_mode', None))
        if value < 0:
            return None
        if value > mark_max:
            value = float(mark_max)
    return value


def _compute_external_assessment_total(combined, mode=None):
    """Compute display/persisted totals for external course assessment modes."""
    mode = _normalize_external_assessment_mode(mode)
    labels = _assessment_column_labels()
    column_label = labels.get(mode, labels['best_three'])
    empty = {
        'display_total': None,
        'assessment_total': None,
        'assessment_total_40': None,
        'assessment_avg': None,
        'column_label': column_label,
    }

    if mode == 'part_a_b_15':
        part_a = [combined.get(i) for i in (1, 2) if combined.get(i) is not None]
        part_b = [combined.get(i) for i in (3, 4) if combined.get(i) is not None]
        if not part_a and not part_b:
            return empty
        part_a_score = max(part_a) if part_a else 0
        part_b_score = max(part_b) if part_b else 0
        display_total = part_a_score + part_b_score
        return {
            'display_total': display_total,
            'assessment_total': display_total,
            'assessment_total_40': None,
            'assessment_avg': None,
            'column_label': column_label,
        }

    valid_marks = [v for v in combined.values() if v is not None]
    if not valid_marks:
        return empty

    valid_marks.sort(reverse=True)
    best = take_best_marks(valid_marks)
    best_sum = sum(best)

    if mode == 'best_three_40':
        total_40 = _round_half_up_int(scale_pg_total(best_sum))
        return {
            'display_total': total_40,
            'assessment_total': None,
            'assessment_total_40': total_40,
            'assessment_avg': round(sum(best) / len(best), 2),
            'column_label': column_label,
        }

    return {
        'display_total': best_sum,
        'assessment_total': best_sum,
        'assessment_total_40': None,
        'assessment_avg': None,
        'column_label': column_label,
    }


def _combined_dict_from_entries(entries):
    combined = {1: None, 2: None, 3: None, 4: None}
    for entry in entries:
        for idx in range(1, 5):
            val = getattr(entry, f'assessment{idx}', None)
            if val is not None and combined[idx] is None:
                try:
                    combined[idx] = float(val)
                except (ValueError, TypeError):
                    pass
    return combined


def _is_external_theory_session(session):
    return bool(session and getattr(session, 'is_external_course', False) and session.course_type == 'theory')


def _get_external_assessment_display_total(session, student_id, combined_values, combined_best3=None, combined_pg_total=None):
    """Resolved display total for external theory exports (PDF/Excel/combined report)."""
    mode = _normalize_external_assessment_mode(getattr(session, 'external_assessment_mode', None))
    combined = combined_values.get(student_id, {})
    display_total = _compute_external_assessment_total(combined, mode)['display_total']
    if display_total is None:
        if mode == 'best_three_40' and combined_pg_total is not None:
            display_total = combined_pg_total.get(student_id)
        elif combined_best3 is not None:
            display_total = combined_best3.get(student_id)
    if mode != 'best_three_40':
        display_total = _maybe_round_assessment_total(session, display_total)
    return display_total, mode


def _combined_pdf_assessment_header(session):
    """Assessment column header for combined attendance + assessment PDF."""
    if _is_external_theory_session(session):
        mode = _normalize_external_assessment_mode(getattr(session, 'external_assessment_mode', None))
        if mode == 'best_three_40':
            return f'Continuous Assessment ({assessment_cfg()["pg_out_of"]})'
        return f'Continuous Assessment ({assessment_cfg()["ug_out_of"]})'
    if session.course_type == 'theory' and session.category == 'pg':
        return f'Continuous Assessment ({assessment_cfg()["pg_out_of"]})'
    return f'Continuous Assessment ({assessment_cfg()["ug_out_of"]})'


def _combined_pdf_title(session):
    """Main title for combined attendance + assessment PDF."""
    if session.course_type == 'sessional':
        return 'Sessional Assessment and Attendance Marks'
    return 'Continuous Assessment and Attendance Marks'


def _normalize_session_course_type(raw_course_type):
    """Map course types to compact values fitting class_session.course_type."""
    value = (raw_course_type or '').strip().lower()
    if not value:
        return 'theory'
    if 'dissertation' in value:
        return 'dissertation'
    if 'thesis' in value:
        return 'thesis'
    if 'sessional' in value:
        return 'sessional'
    if 'viva' in value:
        return 'viva'
    if 'theory' in value:
        return 'theory'
    return value[:20]


def _round_half_up_int(value):
    """Round to nearest integer; fractional part exactly .5 rounds up (grades convention)."""
    if value is None:
        return None
    try:
        return int(Decimal(str(float(value))).quantize(Decimal('1'), rounding=ROUND_HALF_UP))
    except (TypeError, ValueError, ArithmeticError):
        return None


def _session_rounds_assessment_total(session):
    return bool(session and getattr(session, 'round_assessment_total', False))


def _maybe_round_assessment_total(session, value):
    """Optionally round an assessment total when the session toggle is enabled."""
    if value is None:
        return None
    if _session_rounds_assessment_total(session):
        return _round_half_up_int(value)
    return value


def _sync_round_assessment_total(session, enabled):
    """Persist round-total flag on this session and split partners."""
    enabled = bool(enabled)
    related = _resolve_attendance_related_sessions(session, include_archived=False) or [session]
    for related_session in related:
        if related_session is None:
            continue
        related_session.round_assessment_total = enabled
    return enabled


def _generate_feedback_code():
    """Generate a short, URL-friendly code for feedback access."""
    while True:
        raw = secrets.token_urlsafe(8)
        code = ''.join(ch for ch in raw if ch.isalnum()).upper()
        code = code[:10]
        if len(code) < 6:
            continue
        if not StudentFeedbackLink.query.filter_by(access_code=code).first():
            return code


def _external_course_conflict(course_code, teacher_id, exclude_session_id=None):
    """Return an error message if the teacher already has this external course code."""
    if not course_code:
        return None
    conflict_query = query_for_window(Session).filter(
        Session.course_code == course_code,
        Session.teacher_id == teacher_id,
        Session.archived.is_(False),
        Session.is_external_course.is_(True),
    )
    if exclude_session_id:
        conflict_query = conflict_query.filter(Session.id != exclude_session_id)
    if conflict_query.first():
        return f'An external course with code "{course_code}" already exists.'
    return None


def _normalize_excel_student_id(value):
    """Normalize Excel cell values into a clean student ID string."""
    if value is None or (hasattr(pd, 'isna') and pd.isna(value)):
        return ''
    if isinstance(value, float) and value.is_integer():
        return str(int(value)).strip()
    text = str(value).strip()
    if text.lower() == 'nan':
        return ''
    if text.endswith('.0') and text[:-2].replace('.', '', 1).isdigit():
        return text[:-2]
    return text


def _resolve_excel_student_columns(df):
    """Map common Excel header variants to student_id and name columns."""
    column_mapping = {
        'student_id': ['student_id', 'id', 'studentid', 'roll', 'roll_no', 'registration_no', 'reg_no'],
        'name': ['name', 'student_name', 'full_name', 'fullname'],
    }
    actual_columns = {}
    for key, possible_names in column_mapping.items():
        for col in df.columns:
            if col in possible_names:
                actual_columns[key] = col
                break
    return actual_columns


def _import_students_from_excel(session, file, teacher):
    """Import students from an Excel file into a class session.

    Returns (added_count, skipped_count, error_message).
    """
    if not file or not getattr(file, 'filename', ''):
        return 0, 0, None

    filename = file.filename.lower()
    if not (filename.endswith('.xlsx') or filename.endswith('.xls')):
        return 0, 0, 'Please upload an Excel file (.xls or .xlsx)!'

    try:
        df = pd.read_excel(file)
        df.columns = [str(col).strip().lower().replace(' ', '_') for col in df.columns]
        actual_columns = _resolve_excel_student_columns(df)
        if 'student_id' not in actual_columns or 'name' not in actual_columns:
            found = ', '.join(df.columns) if len(df.columns) else '(none)'
            return 0, 0, f'Excel file must have columns: Student ID, Name. Found: {found}'

        existing_student_ids = {
            s.student_id for s in ClassStudent.query.filter_by(session_id=session.id).all()
        }
        added_count = 0
        skipped_count = 0
        empty_rows = 0

        for _, row in df.iterrows():
            student_id = _normalize_excel_student_id(row[actual_columns['student_id']])
            name_raw = row[actual_columns['name']]
            name = '' if (name_raw is None or (hasattr(pd, 'isna') and pd.isna(name_raw))) else str(name_raw).strip()
            if not student_id or not name or name.lower() == 'nan':
                empty_rows += 1
                continue
            if student_id in existing_student_ids:
                skipped_count += 1
                continue

            class_student = ClassStudent(
                student_id=student_id,
                name=name,
                session_id=session.id,
                teacher_id=session.teacher_id or teacher.id,
            )
            db.session.add(class_student)
            _replicate_student_to_peers(session, class_student)
            existing_student_ids.add(student_id)
            added_count += 1

        if added_count == 0 and skipped_count == 0 and empty_rows == 0 and len(df.index) > 0:
            return 0, 0, 'No valid student rows found in the Excel file.'
        return added_count, skipped_count, None
    except Exception as exc:
        return 0, 0, str(exc)


def _student_upload_redirect(session_id, return_to=None):
    if return_to == 'index':
        return redirect(url_for('class_management.index'))
    return redirect(url_for('class_management.students_list', session_id=session_id))


def find_course_from_curriculum(session_course_code, session_course_name=None, session=None):
    """
    Find a Course from the curriculum that matches the session's course code or name.
    Handles various formats like "0421 28 Law 4103" -> "Law 4103" or "Law4103"
    
    Returns: Course object if found, None otherwise
    """
    if session is not None and getattr(session, 'is_external_course', False):
        return None
    import re
    import logging
    logger = logging.getLogger(__name__)
    
    if not Course:
        return None
    
    course_data = None
    
    def extract_core_code(code_str):
        """Extract the core code pattern (e.g., 'Law4103' without space) from various formats"""
        if not code_str:
            return None
        # Try to find pattern like "Law 4103", "Law4103", "CSE 1101", etc.
        match = re.search(r'([A-Za-z]+)\s*(\d{4})', code_str)
        if match:
            return f"{match.group(1)}{match.group(2)}"  # No space: "Law4103"
        return None
    
    # Extract core code pattern from session course code
    session_core_code = extract_core_code(session_course_code)
    logger.debug(f"find_course_from_curriculum: session_code='{session_course_code}', session_core='{session_core_code}', session_name='{session_course_name}'")
    
    # Try exact match by course code
    if session_course_code:
        course_data = Course.query.filter_by(course_code=session_course_code).first()
        if course_data:
            logger.debug(f"Found by exact course_code match: {course_data.course_code}")
            return course_data
    
    # Try case-insensitive match by course code
    if session_course_code:
        course_data = Course.query.filter(func.lower(Course.course_code) == func.lower(session_course_code)).first()
        if course_data:
            logger.debug(f"Found by case-insensitive course_code match: {course_data.course_code}")
            return course_data
    
    # Try whitespace-normalized match (handles extra spaces, tabs, etc.)
    if session_course_code:
        session_code_normalized = ' '.join(session_course_code.strip().split())  # Normalize whitespace
        all_courses = Course.query.all()
        for course in all_courses:
            if course.course_code:
                curriculum_code_normalized = ' '.join(course.course_code.strip().split())
                if session_code_normalized.lower() == curriculum_code_normalized.lower():
                    logger.debug(f"Found by whitespace-normalized match: {course.course_code}")
                    return course
    
    # Try with extracted course code pattern (with space)
    if session_core_code:
        # Try "Law 4103" format (with space)
        extracted_with_space = re.sub(r'([A-Za-z]+)(\d{4})', r'\1 \2', session_core_code)
        course_data = Course.query.filter(func.lower(Course.course_code) == func.lower(extracted_with_space)).first()
        if course_data:
            logger.debug(f"Found by extracted code with space: {course_data.course_code}")
            return course_data
        
        # Try "Law4103" format (without space)
        course_data = Course.query.filter(func.lower(Course.course_code) == func.lower(session_core_code)).first()
        if course_data:
            logger.debug(f"Found by extracted code without space: {course_data.course_code}")
            return course_data
    
    # Try normalized matching - compare core codes from both session and curriculum
    if session_core_code:
        all_courses = Course.query.all()
        session_core_lower = session_core_code.lower()
        for course in all_courses:
            if course.course_code:
                # Extract core code from curriculum course code
                curriculum_core = extract_core_code(course.course_code)
                if curriculum_core and curriculum_core.lower() == session_core_lower:
                    logger.debug(f"Found by normalized core code match: {course.course_code} (core: {curriculum_core})")
                    return course
    
    # Try partial match - check if curriculum course code is contained in session code
    if session_course_code:
        all_courses = Course.query.all()
        session_code_lower = session_course_code.lower()
        # Also try without spaces for partial matching
        session_code_no_space = session_code_lower.replace(' ', '')
        for course in all_courses:
            if course.course_code:
                course_code_lower = course.course_code.lower()
                course_code_no_space = course_code_lower.replace(' ', '')
                # Check with and without spaces
                if course_code_lower in session_code_lower or course_code_no_space in session_code_no_space:
                    logger.debug(f"Found by partial code match: {course.course_code}")
                    return course
    
    # Try exact match by course name
    if session_course_name:
        course_data = Course.query.filter_by(course_name=session_course_name).first()
        if course_data:
            logger.debug(f"Found by exact course_name match: {course_data.course_name}")
            return course_data
    
    # Try case-insensitive partial match by course name
    if session_course_name:
        course_data = Course.query.filter(func.lower(Course.course_name).like(f'%{session_course_name.lower()}%')).first()
        if course_data:
            logger.debug(f"Found by partial course_name match: {course_data.course_name}")
            return course_data
    
    # Try reverse match (session name contains course name or vice versa)
    if session_course_name:
        all_courses = Course.query.all()
        session_name_lower = session_course_name.lower()
        for course in all_courses:
            if course.course_name:
                course_name_lower = course.course_name.lower()
                if session_name_lower in course_name_lower or course_name_lower in session_name_lower:
                    logger.debug(f"Found by reverse course_name match: {course.course_name}")
                    return course
    
    logger.debug(f"No course found for session_code='{session_course_code}', session_name='{session_course_name}'")
    return None


class_management_bp = Blueprint(
    'class_management', __name__,
    template_folder='templates',
    static_folder='static'
)


@class_management_bp.before_request
def restrict_to_teaching_roles():
    if not current_user.is_authenticated:
        return
    
    # Allow student routes for all authenticated users
    student_routes = [
        'class_management.student_view_scores',
        'class_management.student_course_files',
        'class_management.student_download_course_outline_pdf',
        'class_management.student_download_uploaded_file',
        'class_management.student_create_course_question',
        'class_management.student_reply_course_question',
        'class_management.download_course_question_attachment',
        'class_management.delete_course_question_thread',
        'class_management.delete_course_question_message',
        'class_management.question_bank',
        'class_management.upload_question_bank_file',
        'class_management.download_question_bank_file',
        'class_management.download_question_bank_folder',
        'class_management.download_question_bank_zip',
        'class_management.student_notification_read',
        'class_management.student_notifications_page',
        'class_management.student_notifications_mark_all_read',
    ]
    if request.endpoint in student_routes:
        return
    
    if has_teacher_privileges(current_user):
        return
    # Students: redirect to their dashboard without showing the "teaching staff" message
    if parse_roles(current_user.role) and 'student' in parse_roles(current_user.role):
        return redirect(url_for('student_dashboard'))
    flash('Class Management is available only to teaching staff.', 'danger')
    return redirect(url_for('index'))


def _ensure_current_teacher():
    """Return a Teacher instance for the logged-in user, creating one if needed."""
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if teacher:
        return teacher

    base = (current_user.full_name or 'teacher').split(' ')[0].lower()
    base = ''.join(ch for ch in base if ch.isalnum()) or 'teacher'
    base = base[:10]
    candidate = base
    counter = 1
    while Teacher.query.filter_by(short_name=candidate).first():
        suffix = str(counter)
        candidate = f"{base[:10-len(suffix)]}{suffix}"
        counter += 1

    teacher = Teacher(name=current_user.full_name, short_name=candidate, institute=current_tenant().institute_label)
    db.session.add(teacher)
    db.session.commit()
    return teacher


def _get_related_sessions(session, include_archived=False):
    """Return all sessions that belong to the same split group."""
    if not session or not session.split_group_id:
        return [session] if session else []

    query = query_for_window(Session).filter_by(split_group_id=session.split_group_id)
    if not include_archived:
        query = query.filter_by(archived=False)
    related = query.order_by(Session.id.asc()).all()
    return related or [session]


def _resolve_attendance_related_sessions(session, include_archived=False):
    """Resolve related sessions for attendance calculations and reports."""
    related_sessions = _get_related_sessions(session, include_archived=include_archived)
    if not session or session.course_scope not in SPLIT_PARTS:
        return related_sessions

    # Fallback for reassign/unassign history where split_group linkage may drift.
    query = query_for_window(Session).filter(
        Session.course_code == session.course_code,
        Session.year == session.year,
        Session.term == session.term,
        Session.course_scope.in_(list(SPLIT_PARTS))
    )
    if session.academic_session:
        query = query.filter(Session.academic_session == session.academic_session)
    else:
        query = query.filter(or_(Session.academic_session.is_(None), Session.academic_session == ''))
    if not include_archived:
        query = query.filter_by(archived=False)

    merged = {s.id: s for s in related_sessions if s}
    for peer in query.order_by(Session.id.asc()).all():
        merged[peer.id] = peer
    return sorted(merged.values(), key=lambda s: s.id)


def _carry_on_assessment_marks(class_student, session):
    """Carry on previous assessment marks for retake students if carry_on is enabled in registration"""
    if getattr(session, 'is_external_course', False):
        return
    if not StudentCourseRegistration or not Student:
        return
    
    try:
        # Get student record from Students Management
        student_record = Student.query.filter_by(student_id=class_student.student_id).first()
        if not student_record:
            return
        
        # Find registration for this course and session
        registration = StudentCourseRegistration.query.filter_by(
            student_id=student_record.id,
            course_code=session.course_code,
            academic_session=session.academic_session,
            year=session.year,
            term=session.term
        ).first()
        
        if not registration or not registration.carry_on:
            return
        
        # Only carry on for retake/re-retake students
        if registration.remark not in ['Retake', 'Re-retake']:
            return
        
        # Find previous session with same course_code and student_id
        # Look for sessions with different academic_session/year/term
        previous_sessions = query_for_window(Session).filter(
            Session.course_code == session.course_code,
            Session.id != session.id,
            Session.archived == False
        ).order_by(Session.academic_session.desc(), Session.created_at.desc()).all()
        
        for prev_session in previous_sessions:
            # Find student in previous session
            prev_student = ClassStudent.query.filter_by(
                session_id=prev_session.id,
                student_id=class_student.student_id
            ).first()
            
            if prev_student:
                # Copy assessment marks
                if prev_student.assessment1 is not None:
                    class_student.assessment1 = prev_student.assessment1
                if prev_student.assessment2 is not None:
                    class_student.assessment2 = prev_student.assessment2
                if prev_student.assessment3 is not None:
                    class_student.assessment3 = prev_student.assessment3
                if prev_student.assessment4 is not None:
                    class_student.assessment4 = prev_student.assessment4
                
                current_app.logger.info(
                    f'Carried on assessment marks for student {class_student.student_id} '
                    f'from session {prev_session.id} to session {session.id}'
                )
                break  # Only carry from the most recent previous session
    except Exception as e:
        current_app.logger.error(f'Error carrying on assessment marks: {str(e)}', exc_info=True)


def _replicate_student_to_peers(session, source_student, *, old_identifier=None):
    """Create or update a student record across split peer sessions."""
    if not session or not session.split_group_id:
        return

    peer_sessions = [s for s in _get_related_sessions(session) if s.id != session.id]
    for peer in peer_sessions:
        identifier = old_identifier or source_student.student_id
        peer_student = ClassStudent.query.filter_by(session_id=peer.id, student_id=identifier).first()
        if not peer_student:
            peer_student = ClassStudent(
                student_id=source_student.student_id,
                name=source_student.name,
                session_id=peer.id,
                teacher_id=peer.teacher_id
            )
            db.session.add(peer_student)
        else:
            peer_student.student_id = source_student.student_id
            peer_student.name = source_student.name


def _delete_student_from_peers(session, student_identifier):
    """Remove a student from all peer sessions within the split group."""
    if not session or not session.split_group_id:
        return

    for peer in _get_related_sessions(session):
        if peer.id == session.id:
            continue
        peer_student = ClassStudent.query.filter_by(session_id=peer.id, student_id=student_identifier).first()
        if peer_student:
            db.session.delete(peer_student)


def _gather_split_student_map(session):
    """Return related sessions and a map of student_id -> [ClassStudent,...] (excluding deleted students)."""
    related_sessions = _get_related_sessions(session)
    session_ids = [s.id for s in related_sessions if s]
    if not session_ids:
        return related_sessions, {}

    students = _class_students_for_sessions(session_ids)
    student_map = defaultdict(list)
    for stu in students:
        student_map[stu.student_id].append(stu)
    return related_sessions, student_map


def _recalculate_assessment_totals(session):
    """Recompute assessment aggregates across split sessions."""
    if not session or session.course_type != 'theory':
        return

    _, student_map = _gather_split_student_map(session)
    if _is_external_theory_session(session):
        mode = _normalize_external_assessment_mode(getattr(session, 'external_assessment_mode', None))
        for entries in student_map.values():
            combined = _combined_dict_from_entries(entries)
            result = _compute_external_assessment_total(combined, mode)
            display_total = _maybe_round_assessment_total(session, result['display_total'])
            assessment_total = result['assessment_total']
            assessment_total_40 = result['assessment_total_40']
            if mode != 'best_three_40':
                assessment_total = _maybe_round_assessment_total(session, assessment_total)
            else:
                assessment_total_40 = display_total
            for entry in entries:
                entry.assessment_total = assessment_total
                entry.assessment_total_40 = assessment_total_40
                entry.assessment_avg = result['assessment_avg']
        return

    for entries in student_map.values():
        marks = []
        for entry in entries:
            for idx in range(1, 5):
                value = getattr(entry, f'assessment{idx}')
                if value is not None:
                    marks.append(value)
        marks.sort(reverse=True)

        if session.category == 'pg':
            if marks:
                best = take_best_marks(marks)
                avg = sum(best) / len(best)
                best_sum = sum(best)
                total_40 = int(round(scale_pg_total(best_sum)))
                avg_value = round(avg, 2)
            else:
                avg_value = None
                total_40 = None
            for entry in entries:
                entry.assessment_avg = avg_value
                entry.assessment_total_40 = total_40
                entry.assessment_total = None
        else:
            if marks:
                best = take_best_marks(marks)
                total = _maybe_round_assessment_total(session, sum(best))
            else:
                total = None
            for entry in entries:
                entry.assessment_total = total
                entry.assessment_avg = None
                entry.assessment_total_40 = None


def _collect_combined_assessment_marks(session):
    """Return a map of student_id -> list of assessment marks across split sessions."""
    _, student_map = _gather_split_student_map(session)
    marks_map = {}
    for student_id, entries in student_map.items():
        values = []
        for entry in entries:
            for idx in range(1, 5):
                val = getattr(entry, f'assessment{idx}')
                if val is not None:
                    values.append(val)
        marks_map[student_id] = values
    return marks_map


def _build_combined_assessment_values(session):
    """
    Combine assessment values from all related sessions in a split course.
    
    Returns:
        value_map: {student_id: {1: val1, 2: val2, 3: val3, 4: val4}}
        ug_best3: {student_id: best3_total}
        pg_avg_map: {student_id: average}
        pg_total_map: {student_id: total_40_scale}
    """
    # Step 1: Get all related sessions (split course parts).
    # Must match _resolve_attendance_related_sessions so marks stay visible when split_group_id
    # is missing or out of sync after reassign/unassign (same fallback as attendance).
    related_sessions = _resolve_attendance_related_sessions(session, include_archived=False)
    if not related_sessions:
        return {}, {}, {}, {}
    
    # Step 2: Get ALL ClassStudent records from ALL related sessions
    session_ids = [s.id for s in related_sessions if s]
    if not session_ids:
        return {}, {}, {}, {}
    all_class_students = ClassStudent.query.filter(
        ClassStudent.session_id.in_(session_ids)
    ).all()
    
    # Step 3: Group by student_id (STRING) - same student from different sessions
    # Structure: {student_id: [ClassStudent_from_session_A, ClassStudent_from_session_B, ...]}
    student_groups = defaultdict(list)
    for cs in all_class_students:
        student_groups[cs.student_id].append(cs)
    
    # Step 4: Build combined values for each student
    value_map = {}
    ug_best3 = {}
    pg_avg_map = {}
    pg_total_map = {}
    
    for student_id, student_records in student_groups.items():
        # Initialize: {1: None, 2: None, 3: None, 4: None}
        combined = {1: None, 2: None, 3: None, 4: None}
        
        # Combine values from ALL records (sessions)
        # Example: Part A session has assessment1=10, assessment2=5
        #          Part B session has assessment3=7, assessment4=8
        # Result: {1: 10, 2: 5, 3: 7, 4: 8}
        for record in student_records:
            for idx in [1, 2, 3, 4]:
                val = getattr(record, f'assessment{idx}', None)
                # Set value if it exists and we don't have one yet
                if val is not None and combined[idx] is None:
                    try:
                        combined[idx] = float(val)
                    except (ValueError, TypeError):
                        pass
        
        value_map[student_id] = combined

        if _is_external_theory_session(session):
            mode = _normalize_external_assessment_mode(getattr(session, 'external_assessment_mode', None))
            result = _compute_external_assessment_total(combined, mode)
            display_total = result['display_total']
            if mode != 'best_three_40':
                display_total = _maybe_round_assessment_total(session, display_total)
            if mode == 'best_three_40':
                pg_total_map[student_id] = display_total
                pg_avg_map[student_id] = result['assessment_avg']
            else:
                ug_best3[student_id] = display_total
            continue

        # Step 5: Calculate Best 3 / PG Total
        valid_marks = [v for v in combined.values() if v is not None]
        valid_marks.sort(reverse=True)  # Descending
        
        if session.category == 'pg':
            if valid_marks:
                best = take_best_marks(valid_marks)
                avg = sum(best) / len(best)
                pg_avg_map[student_id] = round(avg, 2)
                best_sum = sum(best)
                pg_total_map[student_id] = int(round(scale_pg_total(best_sum)))
            else:
                pg_avg_map[student_id] = None
                pg_total_map[student_id] = None
        else:
            # UG: Best 3 total (optionally rounded via session toggle)
            if valid_marks:
                best = take_best_marks(valid_marks)
                ug_best3[student_id] = _maybe_round_assessment_total(session, sum(best))
            else:
                ug_best3[student_id] = None
    
    return value_map, ug_best3, pg_avg_map, pg_total_map


def _get_editable_assessment_indices(session):
    """Return which assessment inputs current teacher can edit."""
    if session.course_scope == SCOPE_PART_A:
        return {1, 2}
    if session.course_scope == SCOPE_PART_B:
        return {3, 4}
    return {1, 2, 3, 4}


def _get_editable_sessional_fields(session):
    """Return which sessional fields current teacher can edit.

    Sessional split courses should remain collaboratively editable
    across partner teachers.
    """
    return {'sessional_report', 'sessional_viva'}


def _build_combined_sessional_values(session):
    """
    Combine sessional report/viva values across related split sessions.

    Returns:
        value_map: {student_id: {'sessional_report': val, 'sessional_viva': val}}
        absent_map: {student_id: {'sessional_report': bool, 'sessional_viva': bool}}
    """
    related_sessions = _resolve_attendance_related_sessions(session, include_archived=False)
    if not related_sessions:
        return {}, {}

    session_ids = [s.id for s in related_sessions if s]
    if not session_ids:
        return {}, {}

    all_class_students = ClassStudent.query.filter(
        ClassStudent.session_id.in_(session_ids)
    ).all()

    student_groups = defaultdict(list)
    for cs in all_class_students:
        student_groups[cs.student_id].append(cs)

    value_map = {}
    absent_map = {}

    for student_id, student_records in student_groups.items():
        # Prioritize current session first, then partner sessions.
        ordered_records = sorted(
            student_records,
            key=lambda rec: (0 if rec.session_id == session.id else 1, rec.session_id)
        )

        report_value = None
        viva_value = None
        combined_absent = {'sessional_report': False, 'sessional_viva': False}

        for record in ordered_records:
            if report_value is None and record.sessional_report is not None:
                report_value = float(record.sessional_report)
            if viva_value is None and record.sessional_viva is not None:
                viva_value = float(record.sessional_viva)

            if record.assessment_absent:
                try:
                    record_absent = json.loads(record.assessment_absent)
                    combined_absent['sessional_report'] = combined_absent['sessional_report'] or bool(
                        record_absent.get('sessional_report', False)
                    )
                    combined_absent['sessional_viva'] = combined_absent['sessional_viva'] or bool(
                        record_absent.get('sessional_viva', False)
                    )
                except Exception:
                    pass

        value_map[student_id] = {
            'sessional_report': report_value,
            'sessional_viva': viva_value
        }
        absent_map[student_id] = combined_absent

    return value_map, absent_map


ATTENDANCE_LABEL_TO_STATUS = {
    'P': ClassAttendance.STATUS_PRESENT,
    'A': ClassAttendance.STATUS_ABSENT,
    'S': ClassAttendance.STATUS_SKIP,
    '-': ClassAttendance.STATUS_NONE,
}
ATTENDANCE_STATUS_TO_LABEL = {value: key for key, value in ATTENDANCE_LABEL_TO_STATUS.items()}
MAX_CLASSES_PER_DAY = 2
ATTENDANCE_CYCLE = {
    'P': 'A',
    'A': 'S',
    'S': '-',
    '-': 'P',
}


def _normalize_attendance_label(label, allow_blank=False):
    normalized = (label or '').strip().upper()
    if normalized in ATTENDANCE_LABEL_TO_STATUS:
        return normalized
    return None


def _attendance_status_from_record(record):
    status = (getattr(record, 'status', None) or '').strip().lower()
    if status in ATTENDANCE_STATUS_TO_LABEL:
        return status
    return ClassAttendance.STATUS_PRESENT if bool(record.is_present) else ClassAttendance.STATUS_ABSENT


def _attendance_label_from_record(record):
    return ATTENDANCE_STATUS_TO_LABEL.get(_attendance_status_from_record(record), 'A')


def _set_attendance_record_status(record, label):
    normalized_label = _normalize_attendance_label(label, allow_blank=True)
    if not normalized_label:
        raise ValueError('Invalid attendance status label.')
    status_value = ATTENDANCE_LABEL_TO_STATUS[normalized_label]
    if hasattr(record, 'set_status'):
        record.set_status(status_value)
    else:
        record.status = status_value
        record.is_present = status_value == ClassAttendance.STATUS_PRESENT


def _status_payload_for_response(student_stats, status_label, record_id=None, student_db_id=None):
    return {
        'success': True,
        'record_id': record_id,
        'status': status_label,
        'present_count': student_stats.get('present', 0),
        'percentage': f"{student_stats.get('percentage', 0):.2f}%",
        'marks': student_stats.get('marks', 0),
        'marks_manual': student_stats.get('marks_manual', False),
        'total_classes': student_stats.get('effective_total_classes', student_stats.get('base_total_classes', 0)),
        'student_db_id': student_db_id
    }


def _attendance_rotate_up_header_png(text, font_path, font_size=18, bg_rgb=(216, 228, 188), scale=4):
    """Bake Excel-style Rotate Text Up into a high-DPI PNG for sharp WeasyPrint output.

    Draws text LTR at ``scale``× resolution, rotates 90° CCW (bottom → top), and
    returns a data URI sized for CSS display so the PDF keeps crisp vector-like edges.
    Returns (data_uri, display_width_px, display_height_px).
    """
    import base64
    from PIL import Image, ImageDraw, ImageFont

    render_size = max(12, int(round(font_size * scale)))
    try:
        font = ImageFont.truetype(str(font_path), render_size)
    except Exception:
        font = ImageFont.load_default()

    dummy = Image.new('RGBA', (1, 1), (0, 0, 0, 0))
    draw = ImageDraw.Draw(dummy)
    bbox = draw.textbbox((0, 0), text, font=font)
    text_w = max(1, bbox[2] - bbox[0])
    text_h = max(1, bbox[3] - bbox[1])
    pad = max(4, render_size // 5)

    # Transparent canvas — table cell supplies the green background
    img = Image.new('RGBA', (text_w + pad * 2, text_h + pad * 2), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    draw.text((pad - bbox[0], pad - bbox[1]), text, font=font, fill=(17, 17, 17, 255))

    # PIL rotate(90) = counter-clockwise = Rotate Text Up
    rotated = img.rotate(90, expand=True, fillcolor=(0, 0, 0, 0), resample=Image.Resampling.BICUBIC)

    # Keep full pixel density; CSS width/height will be display size (÷ scale)
    display_w = max(1, int(round(rotated.width / scale)))
    display_h = max(1, int(round(rotated.height / scale)))

    buf = io.BytesIO()
    rotated.save(buf, format='PNG', compress_level=1)
    data_uri = 'data:image/png;base64,' + base64.b64encode(buf.getvalue()).decode('ascii')
    return data_uri, display_w, display_h


def _cap_classes_per_day(count_value):
    if not count_value:
        return 0
    return min(int(count_value), MAX_CLASSES_PER_DAY)


def _normalize_slot_number(slot_value):
    try:
        slot_int = int(slot_value)
    except (TypeError, ValueError):
        return None
    if 1 <= slot_int <= MAX_CLASSES_PER_DAY:
        return slot_int
    return None


def _records_by_slot(records_for_date):
    """Map attendance records to slot numbers with legacy fallback."""
    if not records_for_date:
        return {}

    slot_map = {}
    next_slot = 1
    for record in sorted(records_for_date, key=lambda r: ((getattr(r, 'slot_number', None) or 99), r.id)):
        explicit_slot = _normalize_slot_number(getattr(record, 'slot_number', None))
        if explicit_slot and explicit_slot not in slot_map:
            slot_map[explicit_slot] = record
            continue

        while next_slot in slot_map and next_slot <= MAX_CLASSES_PER_DAY:
            next_slot += 1
        if next_slot <= MAX_CLASSES_PER_DAY:
            slot_map[next_slot] = record
            next_slot += 1

    return slot_map


def _calculate_attendance_mark_from_percentage(percentage):
    """Convert attendance percentage to marks."""
    if percentage >= 90:
        return 10
    if percentage >= 85:
        return 9
    if percentage >= 80:
        return 8
    if percentage >= 75:
        return 7
    if percentage >= 70:
        return 6
    if percentage >= 65:
        return 5
    if percentage >= 60:
        return 4
    return 0


def _build_attendance_summary(session, include_archived=False):
    """Aggregate attendance across split sessions."""
    related_sessions = _resolve_attendance_related_sessions(session, include_archived=include_archived)
    session_ids = [s.id for s in related_sessions if s]
    if not session_ids:
        return {
            'total_classes': 0,
            'per_student': {},
            'per_session_totals': defaultdict(int),
            'per_student_per_session': {},
            'related_sessions': related_sessions,
        }

    attendance_records = ClassAttendance.query.filter(
        ClassAttendance.session_id.in_(session_ids)
    ).order_by(ClassAttendance.date.asc(), ClassAttendance.id.asc()).all()

    students = _class_students_for_sessions(session_ids)
    student_lookup = {stu.id: stu for stu in students}

    per_student_counts = defaultdict(lambda: {'present': 0, 'skip': 0})
    per_student_session_counts = defaultdict(lambda: defaultdict(lambda: {'present': 0, 'skip': 0}))
    per_session_date_counts = defaultdict(lambda: defaultdict(int))
    per_student_date_slot_seen = defaultdict(set)
    per_student_date_fallback_counts = defaultdict(int)

    for record in attendance_records:
        slot_key = (record.session_id, record.student_id, record.date)
        slot_number = _normalize_slot_number(getattr(record, 'slot_number', None))
        if slot_number is not None:
            if slot_number in per_student_date_slot_seen[slot_key]:
                continue
            per_student_date_slot_seen[slot_key].add(slot_number)
        else:
            per_student_date_fallback_counts[slot_key] += 1
            fallback_slot = per_student_date_fallback_counts[slot_key]
            if fallback_slot > MAX_CLASSES_PER_DAY:
                continue
            per_student_date_slot_seen[slot_key].add(fallback_slot)

        per_session_date_counts[(record.session_id, record.date)][record.student_id] += 1
        student = student_lookup.get(record.student_id)
        if not student:
            continue
        record_status = _attendance_status_from_record(record)
        if record_status == ClassAttendance.STATUS_PRESENT:
            per_student_counts[student.student_id]['present'] += 1
            per_student_session_counts[student.student_id][record.session_id]['present'] += 1
        elif record_status == ClassAttendance.STATUS_SKIP:
            per_student_counts[student.student_id]['skip'] += 1
            per_student_session_counts[student.student_id][record.session_id]['skip'] += 1
        per_student_counts[student.student_id]['records'] = per_student_counts[student.student_id].get('records', 0) + 1

    per_session_totals = defaultdict(int)
    total_classes = 0
    for (session_id, _), counts in per_session_date_counts.items():
        class_count = _cap_classes_per_day(max(counts.values()) if counts else 0)
        total_classes += class_count
        per_session_totals[session_id] += class_count

    per_student_result = {}
    # Group students by student_id to handle split courses (multiple sessions, same student_id)
    students_by_id = {}
    for student in students:
        if student.student_id not in students_by_id:
            students_by_id[student.student_id] = []
        students_by_id[student.student_id].append(student)
    
    for student_id, student_list in students_by_id.items():
        stats = per_student_counts.get(student_id, {'present': 0, 'skip': 0, 'records': 0})
        effective_total_classes = max(total_classes - stats.get('skip', 0), 0)
        percentage = (stats['present'] / effective_total_classes * 100) if effective_total_classes else 0
        
        # Check all student records for this student_id to find manual marks
        # (for split courses, manual marks might be on any of the related student records)
        manual_marks = None
        for student in student_list:
            if student.attendance_marks_manual is not None:
                manual_marks = student.attendance_marks_manual
                break
        
        if manual_marks is not None:
            marks = manual_marks
        else:
            marks = _calculate_attendance_mark_from_percentage(percentage)
        per_student_result[student_id] = {
            'present': stats['present'],
            'skip_count': stats.get('skip', 0),
            'base_total_classes': total_classes,
            'effective_total_classes': effective_total_classes,
            'percentage': percentage,
            'marks': marks,
            'marks_manual': manual_marks is not None
        }

    return {
        'total_classes': total_classes,
        'per_student': per_student_result,
        'per_session_totals': per_session_totals,
        'per_student_per_session': {
            public_id: {session_id: dict(counts) for session_id, counts in sessions.items()}
            for public_id, sessions in per_student_session_counts.items()
        },
        'related_sessions': related_sessions
    }


def _split_attendance_part_breakdown(attendance_summary, student_id):
    """Part A / Part B class and present totals for one student on a split course."""
    related_sessions = attendance_summary.get('related_sessions') or []
    per_session_totals = attendance_summary.get('per_session_totals') or {}
    per_student_session = (attendance_summary.get('per_student_per_session') or {}).get(student_id) or {}

    part_a_classes = 0
    part_b_classes = 0
    part_a_present = 0
    part_b_present = 0
    has_split_part = False

    for related in related_sessions:
        if not related or related.course_scope not in SPLIT_PARTS:
            continue
        has_split_part = True
        classes = per_session_totals.get(related.id, 0) or 0
        present = (per_student_session.get(related.id) or {}).get('present', 0) or 0
        if related.course_scope == SCOPE_PART_A:
            part_a_classes += classes
            part_a_present += present
        elif related.course_scope == SCOPE_PART_B:
            part_b_classes += classes
            part_b_present += present

    if not has_split_part:
        return None

    return {
        'part_a_classes': part_a_classes,
        'part_b_classes': part_b_classes,
        'part_a_present': part_a_present,
        'part_b_present': part_b_present,
    }


def _build_split_context(session, attendance_summary=None):
    """Prepare metadata for templates about split courses."""
    if not session or not session.split_group_id or session.course_scope == SCOPE_FULL:
        return None

    peer_info = []
    for peer in _get_related_sessions(session, include_archived=True):
        if not peer or peer.id == session.id:
            continue
        peer_info.append({
            'id': peer.id,
            'teacher_name': peer.teacher.name if peer.teacher else '—',
            'teacher_short': peer.teacher.short_name if peer.teacher else '',
            'course_scope': COURSE_SCOPE_LABELS.get(peer.course_scope, 'Part')
        })

    context = {
        'scope_label': COURSE_SCOPE_LABELS.get(session.course_scope, 'Part'),
        'peers': peer_info
    }

    if attendance_summary:
        totals = attendance_summary.get('per_session_totals', {})
        context['class_totals'] = []
        for related in attendance_summary.get('related_sessions', []):
            if not related:
                continue
            context['class_totals'].append({
                'session_id': related.id,
                'teacher_name': related.teacher.name if related.teacher else '—',
                'teacher_short': related.teacher.short_name if related.teacher else '',
                'classes': totals.get(related.id, 0),
                'is_current': related.id == session.id
            })
        context['total_classes'] = attendance_summary.get('total_classes', 0)

    return context


# Create uploads folder if it doesn't exist
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Q&A uploads folder (separate)
QA_UPLOAD_FOLDER = os.path.join(UPLOAD_FOLDER, 'qa_questions')
if not os.path.exists(QA_UPLOAD_FOLDER):
    os.makedirs(QA_UPLOAD_FOLDER)

QUESTION_BANK_UPLOAD_FOLDER = os.path.join(UPLOAD_FOLDER, 'question_bank')
if not os.path.exists(QUESTION_BANK_UPLOAD_FOLDER):
    os.makedirs(QUESTION_BANK_UPLOAD_FOLDER)


def _get_qa_upload_dir(thread_id):
    """Return (and create) per-thread upload folder."""
    thread_dir = os.path.join(QA_UPLOAD_FOLDER, str(thread_id))
    os.makedirs(thread_dir, exist_ok=True)
    return thread_dir


def _save_qa_attachments(files, thread_id):
    """Save attachments and return metadata list."""
    saved = []
    if not files:
        return saved
    upload_dir = _get_qa_upload_dir(thread_id)
    for file in files:
        if not file or not file.filename:
            continue
        safe_name = secure_filename(file.filename)
        if not safe_name:
            continue
        unique_name = f"{uuid4().hex}_{safe_name}"
        file_path = os.path.join(upload_dir, unique_name)
        file.save(file_path)
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else None
        file_type = file.mimetype or os.path.splitext(safe_name)[1].lstrip('.')
        saved.append({
            'file_name': safe_name,
            'file_path': file_path,
            'file_size': file_size,
            'file_type': file_type
        })
    return saved


def _delete_qa_attachments(attachments):
    """Delete attachment files from disk."""
    for attachment in attachments:
        try:
            if attachment.file_path and os.path.exists(attachment.file_path):
                os.remove(attachment.file_path)
        except Exception:
            pass

@class_management_bp.route('/')
@login_required
def index():
    """Main dashboard for class management"""
    teacher = _ensure_current_teacher()
    
    # Start with base query
    query = query_for_window(Session).filter_by(
        teacher_id=teacher.id,
        archived=False,
        is_external_course=False,
    )
    
    # IMPORTANT: Update sessions with academic_session BEFORE filtering
    # This ensures all sessions have academic_session set for proper filtering
    sessions_before_update = query.all()
    
    # Update sessions with academic_session and batch from CourseSessionAssignment if available
    # Also sync all assignments with curriculum year-term config if missing
    # IMPORTANT: Do this BEFORE auto-creating sessions so assignments have correct academic_session
    if CourseSessionAssignment and Curriculum and CurriculumYearTerm:
        try:
            # First, update all assignments that are missing batch/academic_session from curriculum year-term config
            all_assignments = CourseSessionAssignment.query.all()
            assignment_update_count = 0
            for assignment in all_assignments:
                if assignment.curriculum_id and (not assignment.batch or not assignment.academic_session):
                    try:
                        curriculum = Curriculum.query.get(assignment.curriculum_id)
                        if curriculum:
                            year_term_config = curriculum.get_year_term_config(assignment.year, assignment.term)
                            if year_term_config:
                                updated = False
                                if not assignment.batch and year_term_config.batch and year_term_config.batch != 'None':
                                    assignment.batch = year_term_config.batch
                                    updated = True
                                    current_app.logger.info(f'Updated assignment {assignment.id} batch from year-term config: {year_term_config.batch}')
                                if not assignment.academic_session and year_term_config.academic_session:
                                    assignment.academic_session = year_term_config.academic_session
                                    updated = True
                                    current_app.logger.info(f'Updated assignment {assignment.id} academic_session from year-term config: {year_term_config.academic_session}')
                                
                                if updated:
                                    assignment_update_count += 1
                    except Exception as assign_error:
                        current_app.logger.warning(f'Error updating assignment {assignment.id}: {assign_error}')
            
            if assignment_update_count > 0:
                db.session.commit()
                current_app.logger.info(f'Updated {assignment_update_count} assignments with batch/academic_session from curriculum year-term config')
            
            # Now update sessions with academic_session / window_id from assignments
            # Only FILL missing values — never overwrite a different academic_session.
            # Overwriting (e.g. CSA had a stale/wrong session string) was hiding courses
            # that previously matched Window active-semester filters.
            updated_count = 0
            for session in sessions_before_update:
                if getattr(session, 'is_external_course', False):
                    continue
                # Find CourseSessionAssignment for this session
                assignment = CourseSessionAssignment.query.filter_by(session_id=session.id).first()
                if assignment:
                    if assignment.academic_session and not (session.academic_session or '').strip():
                        session.academic_session = assignment.academic_session
                        updated_count += 1
                        current_app.logger.info(
                            f'Filled session {session.id} ({session.course_name}) academic_session: '
                            f'{assignment.academic_session}'
                        )
                    elif (
                        assignment.academic_session
                        and session.academic_session
                        and session.academic_session.strip() != assignment.academic_session.strip()
                    ):
                        current_app.logger.warning(
                            f'Session {session.id} ({session.course_name}) academic_session '
                            f'"{session.academic_session}" differs from CSA "{assignment.academic_session}" '
                            f'— keeping session value (no overwrite)'
                        )
                    # Stamp window from assignment when session window is missing
                    if assignment.window_id is not None and session.window_id is None:
                        session.window_id = assignment.window_id
                        updated_count += 1
                        current_app.logger.info(
                            f'Updated session {session.id} ({session.course_name}) window_id from assignment: {assignment.window_id}'
                        )
            
            if updated_count > 0:
                db.session.commit()
                current_app.logger.info(f'Updated {updated_count} sessions with academic_session from CourseSessionAssignment')
        except Exception as e:
            current_app.logger.error(f'Error updating sessions from CourseSessionAssignment: {str(e)}', exc_info=True)
            db.session.rollback()
    
    # Auto-create missing Sessions from CourseSessionAssignment
    # This ensures all assigned courses appear in Class Management
    # IMPORTANT: Do this AFTER updating assignments with academic_session from curriculum config
    if CourseSessionAssignment and Course:
        try:
            # Find all assignments for this teacher that don't have sessions yet
            missing_assignments_query = CourseSessionAssignment.query.filter_by(
                teacher_id=teacher.id
            ).filter(
                or_(
                    CourseSessionAssignment.session_created == False,
                    CourseSessionAssignment.session_id.is_(None)
                )
            )
            if get_effective_window_id and not is_admin(current_user):
                eff_window = get_effective_window_id(admin_override=False)
                if eff_window is not None:
                    missing_assignments_query = missing_assignments_query.filter(
                        or_(
                            CourseSessionAssignment.window_id == eff_window,
                            CourseSessionAssignment.window_id.is_(None),
                        )
                    )
            missing_assignments = missing_assignments_query.all()
            
            current_app.logger.info(f'[DEBUG] Teacher {teacher.id} ({teacher.name}): Found {len(missing_assignments)} assignments without sessions')
            
            # Log assignment details for debugging
            for idx, assignment in enumerate(missing_assignments, 1):
                course = Course.query.get(assignment.course_id) if assignment.course_id else None
                course_code = course.course_code if course else f'course_id={assignment.course_id}'
                current_app.logger.info(f'[DEBUG] Missing assignment #{idx}: ID={assignment.id}, Course={course_code}, Year={assignment.year}, Term={assignment.term}, AcademicSession={assignment.academic_session}, Batch={assignment.batch}, Section={assignment.section}')
            
            created_count = 0
            for assignment in missing_assignments:
                try:
                    # Get course details
                    course = Course.query.get(assignment.course_id)
                    if not course:
                        current_app.logger.warning(f'Course {assignment.course_id} not found for assignment {assignment.id}')
                        continue
                    
                    # Determine course_scope based on section
                    if assignment.section == 'A':
                        course_scope = SCOPE_PART_A
                    elif assignment.section == 'B':
                        course_scope = SCOPE_PART_B
                    else:
                        course_scope = SCOPE_FULL
                    
                    # Prefer an existing session with same scope (+ academic_session when set)
                    existing_candidates = query_for_window(Session).filter_by(
                        course_code=course.course_code,
                        teacher_id=teacher.id,
                        year=assignment.year,
                        term=assignment.term,
                        course_scope=course_scope,
                        archived=False,
                    ).all()
                    existing_session = None
                    if existing_candidates:
                        assign_session_norm = _normalize_offering_text(assignment.academic_session)
                        exact = [
                            s for s in existing_candidates
                            if _normalize_offering_text(s.academic_session) == assign_session_norm
                        ]
                        existing_session = exact[0] if exact else existing_candidates[0]
                        if assignment.academic_session and not existing_session.academic_session:
                            existing_session.academic_session = assignment.academic_session
                        if assignment.window_id is not None and existing_session.window_id is None:
                            existing_session.window_id = assignment.window_id
                    
                    if existing_session:
                        # Link the assignment to the existing session
                        assignment.session_id = existing_session.id
                        assignment.session_created = True
                        current_app.logger.info(f'Linked assignment {assignment.id} to existing session {existing_session.id}')
                    else:
                        # Create new Session with academic_session from assignment
                        session_obj = Session(
                            year=assignment.year,
                            term=assignment.term,
                            academic_session=assignment.academic_session,
                            course_code=course.course_code,
                            course_name=course.course_name,
                            teacher_id=teacher.id,
                            course_type=_normalize_session_course_type(course.course_type),
                            category=course.category if course.category else 'ug',
                            course_scope=course_scope,
                            window_id=assignment.window_id,
                        )
                        db.session.add(session_obj)
                        db.session.flush()  # Get session ID
                        
                        # Link assignment to session
                        assignment.session_id = session_obj.id
                        assignment.session_created = True
                        
                        created_count += 1
                        current_app.logger.info(f'Auto-created session {session_obj.id} from assignment {assignment.id} for course {course.course_code} (academic_session: {assignment.academic_session})')
                
                except Exception as create_error:
                    current_app.logger.error(f'Error auto-creating session for assignment {assignment.id}: {create_error}', exc_info=True)
                    continue
            
            if created_count > 0 or missing_assignments:
                db.session.commit()
                if created_count > 0:
                    current_app.logger.info(f'[DEBUG] Auto-created {created_count} sessions from CourseSessionAssignment for teacher {teacher.id}')
                else:
                    current_app.logger.info(f'[DEBUG] No new sessions created for teacher {teacher.id} (all assignments already have sessions or no valid assignments)')
        except Exception as e:
            current_app.logger.error(f'[DEBUG] Error auto-creating sessions from CourseSessionAssignment: {str(e)}', exc_info=True)
            db.session.rollback()
    
    # Now apply active semester filtering AFTER updating academic_session and creating sessions
    # Re-query to get updated sessions including newly created ones
    query = query_for_window(Session).filter_by(
        teacher_id=teacher.id,
        archived=False,
        is_external_course=False,
    )
    
    # Log sessions before filtering
    sessions_before_filter = query.all()
    current_app.logger.info(f'[DEBUG] Teacher {teacher.id}: Found {len(sessions_before_filter)} sessions BEFORE active semester filtering')
    for s in sessions_before_filter:
        current_app.logger.info(f'[DEBUG] Session before filter: ID={s.id}, Course={s.course_code} ({s.course_name}), Year={s.year}, Term={s.term}, AcademicSession={s.academic_session}')
    
    # Apply active semester filtering (if not admin and filter function available)
    if filter_by_active_semester and not is_admin(current_user):
        try:
            # Don't filter by batch - get ALL active semesters
            # This ensures courses from all active semesters are shown, not just from a specific batch
            batch = None
            
            # Log active semester configuration
            try:
                from utils.semester_utils import get_active_semesters_for_user
                active_semesters = get_active_semesters_for_user(admin_override=False)
                active_sem_info = [
                    f"win{s.window_id}-{s.academic_session}-{s.year}-{s.term}-{s.batch or 'ALL'}"
                    for s in active_semesters
                ]
                current_app.logger.info(f'[DEBUG] Active semesters for filtering (window-scoped): {active_sem_info}')
            except Exception as sem_error:
                current_app.logger.warning(f'[DEBUG] Error getting active semesters: {sem_error}')
            
            # Apply active semester filter (batch=None to get all active semesters)
            query = filter_by_active_semester(query, Session, batch=None, admin_override=False)
            current_app.logger.info(f'[DEBUG] Applied active semester filtering for teacher {teacher.id} (using all active semesters)')
        except Exception as filter_error:
            current_app.logger.error(f'[DEBUG] Error applying active semester filter: {filter_error}', exc_info=True)
            # Don't fail the request, but log the error

    if filter_by_active_window and not is_admin(current_user):
        try:
            query = filter_by_active_window(query, Session, admin_override=False)
            current_app.logger.info(f'[DEBUG] Applied active window filtering for teacher {teacher.id}')
        except Exception as window_filter_error:
            current_app.logger.error(f'[DEBUG] Error applying active window filter: {window_filter_error}', exc_info=True)
    
    sessions = query.order_by(Session.created_at.desc()).all()
    sessions = _dedupe_active_sessions(sessions)

    current_app.logger.info(f'[DEBUG] Teacher {teacher.id} ({teacher.name}): Found {len(sessions)} sessions AFTER filtering (was {len(sessions_before_filter)} before filtering)')
    for s in sessions:
        current_app.logger.info(f'[DEBUG] Session after filter: ID={s.id}, Course={s.course_code} ({s.course_name}), Year={s.year}, Term={s.term}, AcademicSession={s.academic_session}')
    for s in sessions:
        current_app.logger.debug(f'Session: ID={s.id}, Name={s.course_name}, Session={s.academic_session}, Year={s.year}, Term={s.term}, Archived={s.archived}, Teacher={s.teacher_id}')

    split_context_map = {session.id: _build_split_context(session) for session in sessions if session.split_group_id}
    # Get teachers excluding Head of the Discipline
    from role_utils import get_teachers_excluding_head
    teachers = get_teachers_excluding_head()
    pending_split_invites = ClassSplitInvite.query.filter_by(invited_teacher_id=teacher.id, status='pending').order_by(ClassSplitInvite.created_at.desc()).all()
    
    # Get all batches from Students Management for the dropdown
    batches = []
    if Student:
        try:
            all_batches = db.session.query(Student.batch).distinct().filter(Student.batch.isnot(None)).order_by(Student.batch.desc()).all()
            batches = [batch[0] for batch in all_batches]
        except Exception:
            batches = []
    
    # Build assignment map for template to access batch and academic_session from CourseSessionAssignment
    assignment_map = {}
    if CourseSessionAssignment and Course:
        try:
            for session in sessions:
                # Try to find assignment by session_id first
                assignment = CourseSessionAssignment.query.filter_by(session_id=session.id).first()
                
                # If not found by session_id, try to find by course_code, teacher_id, year, term
                if not assignment and session.course_code and session.teacher_id and session.year and session.term:
                    try:
                        # Try to match by course_code, teacher_id, year, term
                        # First try exact match
                        assignment = CourseSessionAssignment.query.filter_by(
                            teacher_id=session.teacher_id,
                            year=session.year,
                            term=session.term
                        ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                            Course.course_code == session.course_code
                        ).first()
                        
                        # If not found, try without section matching (for full course sessions)
                        if not assignment:
                            assignment = CourseSessionAssignment.query.filter_by(
                                teacher_id=session.teacher_id,
                                year=session.year,
                                term=session.term
                            ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                                Course.course_code == session.course_code
                            ).filter(
                                or_(
                                    CourseSessionAssignment.section.is_(None),
                                    CourseSessionAssignment.section == ''
                                )
                            ).first()
                        
                        # If found, update the session_id to link them
                        if assignment and not assignment.session_id:
                            assignment.session_id = session.id
                            assignment.session_created = True
                            try:
                                db.session.commit()
                                current_app.logger.info(f'Linked assignment {assignment.id} to session {session.id} for course {session.course_code}')
                            except Exception as commit_error:
                                db.session.rollback()
                                current_app.logger.warning(f'Could not link assignment {assignment.id} to session {session.id}: {commit_error}')
                    except Exception as query_error:
                        current_app.logger.warning(f'Error querying assignment for session {session.id}: {query_error}')
                
                if assignment:
                    # If assignment doesn't have batch/academic_session, try to get from curriculum year-term config
                    batch = assignment.batch
                    academic_session = assignment.academic_session
                    
                    if Curriculum and CurriculumYearTerm and (not batch or not academic_session):
                        try:
                            if assignment.curriculum_id:
                                curriculum = Curriculum.query.get(assignment.curriculum_id)
                                if curriculum:
                                    year_term_config = curriculum.get_year_term_config(assignment.year, assignment.term)
                                    if year_term_config:
                                        if not batch and year_term_config.batch and year_term_config.batch != 'None':
                                            batch = year_term_config.batch
                                            assignment.batch = batch
                                            current_app.logger.info(f'Updated assignment {assignment.id} batch from year-term config: {batch}')
                                        if not academic_session and year_term_config.academic_session:
                                            academic_session = year_term_config.academic_session
                                            assignment.academic_session = academic_session
                                            current_app.logger.info(f'Updated assignment {assignment.id} academic_session from year-term config: {academic_session}')
                                        
                                        if (not assignment.batch and batch) or (not assignment.academic_session and academic_session):
                                            try:
                                                db.session.commit()
                                            except Exception as commit_error:
                                                db.session.rollback()
                                                current_app.logger.warning(f'Could not update assignment {assignment.id}: {commit_error}')
                        except Exception as config_error:
                            current_app.logger.warning(f'Error getting year-term config for assignment {assignment.id}: {config_error}')
                    
                    assignment_map[session.id] = {
                        'batch': batch or '',
                        'academic_session': academic_session or ''
                    }
                    # Also update session's academic_session if assignment has it and session doesn't
                    if academic_session and not session.academic_session:
                        try:
                            session.academic_session = academic_session
                            db.session.commit()
                            current_app.logger.info(f'Updated session {session.id} academic_session from assignment: {academic_session}')
                        except Exception as update_error:
                            db.session.rollback()
                            current_app.logger.warning(f'Could not update session {session.id} academic_session: {update_error}')
                    
                    # Auto-add students from batch if session has no students but batch is available
                    if getattr(session, 'is_external_course', False):
                        assignment_map[session.id] = {
                            'batch': batch or '',
                            'academic_session': academic_session or ''
                        }
                        continue
                    if batch and batch.strip() and batch != 'None' and Student:
                        try:
                            existing_students_count = ClassStudent.query.filter_by(session_id=session.id).count()
                            if existing_students_count == 0:
                                current_app.logger.info(f'Session {session.id} has no students but batch {batch} is available. Attempting to add students...')
                                students_from_batch = Student.query.filter_by(batch=batch).all()
                                if students_from_batch:
                                    added_count = 0
                                    for student in students_from_batch:
                                        # Check if student is registered for this course (finalized registration only)
                                        if not getattr(session, 'is_external_course', False) and StudentCourseRegistration and session.course_code and session.academic_session and session.year and session.term:
                                            registration = StudentCourseRegistration.query.filter_by(
                                                student_id=student.id,
                                                course_code=session.course_code,
                                                academic_session=session.academic_session,
                                                year=session.year,
                                                term=session.term,
                                                status='finalized'
                                            ).first()
                                            
                                            if not registration:
                                                current_app.logger.info(f'Student {student.student_id} ({student.name}) not registered for course {session.course_code}, skipping...')
                                                continue
                                        
                                        class_student = ClassStudent(
                                            student_id=student.student_id,
                                            name=student.name,
                                            session_id=session.id,
                                            teacher_id=session.teacher_id
                                        )
                                        db.session.add(class_student)
                                        db.session.flush()  # Flush to get class_student.id before carry on
                                        
                                        # Carry on assessment marks if enabled in registration
                                        _carry_on_assessment_marks(class_student, session)
                                        
                                        # Replicate to peer sessions for split courses
                                        _replicate_student_to_peers(session, class_student)
                                        
                                        added_count += 1
                                    
                                    if added_count > 0:
                                        db.session.commit()
                                        current_app.logger.info(f'Successfully added {added_count} students from batch {batch} to session {session.id}')
                                else:
                                    current_app.logger.warning(f'No students found in batch {batch} for session {session.id}')
                        except Exception as auto_add_error:
                            db.session.rollback()
                            current_app.logger.error(f'Error auto-adding students to session {session.id} from batch {batch}: {auto_add_error}', exc_info=True)
        except Exception as e:
            current_app.logger.error(f'Error building assignment map: {str(e)}', exc_info=True)
            db.session.rollback()

    # Q&A notifications: count threads with latest message from student
    qa_notification_map = {}
    try:
        from sqlalchemy.orm import selectinload
        session_ids = [s.id for s in sessions]
        if session_ids:
            threads = CourseQuestionThread.query.options(
                selectinload(CourseQuestionThread.messages)
            ).filter(
                CourseQuestionThread.session_id.in_(session_ids),
                CourseQuestionThread.teacher_id == teacher.id
            ).all()
            for thread in threads:
                last_message = None
                if thread.messages:
                    last_message = max(
                        thread.messages,
                        key=lambda m: (m.created_at or datetime.min, m.id or 0)
                    )
                if (
                    last_message
                    and last_message.sender_role == 'student'
                    and thread.teacher_read_at is None
                ):
                    qa_notification_map[thread.session_id] = qa_notification_map.get(thread.session_id, 0) + 1
    except Exception as e:
        current_app.logger.warning(f'Error loading Q&A notifications: {e}')

    external_sessions = query_for_window(Session).filter_by(
        teacher_id=teacher.id,
        archived=False,
        is_external_course=True,
    ).order_by(Session.created_at.desc()).all()
    external_archived_sessions = query_for_window(Session).filter_by(
        teacher_id=teacher.id,
        archived=True,
        is_external_course=True,
    ).order_by(Session.created_at.desc()).all()

    active_semester_json = '{}'
    active_semesters_json = '[]'
    try:
        from utils.semester_utils import get_active_semesters_for_user
        active_semesters = get_active_semesters_for_user(admin_override=is_admin(current_user))
        if active_semesters:
            semester_rows = [{
                'academic_session': sem.academic_session or '',
                'year': sem.year or '',
                'term': sem.term or '',
            } for sem in active_semesters]
            active_semesters_json = json.dumps(semester_rows)
            active_semester_json = json.dumps(semester_rows[0])
    except Exception:
        pass

    return render_template(
        'class_management/index.html',
        sessions=sessions,
        external_sessions=external_sessions,
        external_archived_sessions=external_archived_sessions,
        active_semester_json=active_semester_json,
        active_semesters_json=active_semesters_json,
        teacher=teacher,
        teacher_display_name=(getattr(current_user, 'full_name', '') or teacher.name or current_user.username),
        teachers=teachers,
        course_scope_labels=COURSE_SCOPE_LABELS,
        split_context_map=split_context_map,
        pending_split_invites=pending_split_invites,
        batches=batches,
        assignment_map=assignment_map,
        qa_notification_map=qa_notification_map
    )

@class_management_bp.route('/create_session', methods=['POST'])
@login_required
def create_session():
    """Create a new session"""
    try:
        teacher = _ensure_current_teacher()
        batch = request.form.get('batch', '').strip()
        curriculum_id = request.form.get('curriculum_id', type=int)
        year = request.form.get('year', '').strip()
        term = request.form.get('term', '').strip()
        academic_session = request.form.get('academic_session', '').strip()
        course_id = request.form.get('course_id', type=int)
        course_code = request.form.get('course_code', '').strip()
        course_name = request.form.get('course_name', '').strip()
        course_type = request.form.get('course_type', 'theory')
        category = request.form.get('category', 'ug')
        course_scope = request.form.get('course_scope', SCOPE_FULL)
        partner_teacher_id = request.form.get('partner_teacher_id')
        
        current_app.logger.info(f'Creating session - batch: {batch}, curriculum_id: {curriculum_id}, year: {year}, term: {term}, course_name: {course_name}, course_code: {course_code}')
        
        # If course_id is provided, fetch course details from Course model
        if course_id and Course:
            course = Course.query.get(course_id)
            if course:
                course_code = course.course_code
                course_name = course.course_name
                course_type = course.course_type.lower()
                category = course.category
                current_app.logger.info(f'Fetched course details from Course model: {course_code} - {course_name}')
        
        if not year or not term:
            flash('Year and term are required!', 'error')
            current_app.logger.warning(f'Missing year or term - year: {year}, term: {term}')
            return redirect(url_for('class_management.index'))
        
        if not course_code or not course_name:
            flash('Course code and course name are required!', 'error')
            current_app.logger.warning(f'Missing course_code or course_name - code: {course_code}, name: {course_name}')
            return redirect(url_for('class_management.index'))

        if course_scope not in COURSE_SCOPE_LABELS:
            flash('Invalid course scope selection.', 'error')
            return redirect(url_for('class_management.index'))

        # Prevent more than two teachers (Part A & Part B) from taking the same course simultaneously
        window_id = get_effective_window_id(admin_override=is_admin(current_user)) if get_effective_window_id else None
        active_sessions_query = query_for_window(Session).filter(
            Session.course_code == course_code,
            Session.archived.is_(False),
            Session.is_external_course.is_(False),
        )
        if window_id is not None:
            active_sessions_query = active_sessions_query.filter(
                or_(Session.window_id == window_id, Session.window_id.is_(None))
            )
        active_sessions = active_sessions_query.all()
        full_exists = any(s.course_scope == SCOPE_FULL for s in active_sessions)
        part_a_exists = any(s.course_scope == SCOPE_PART_A for s in active_sessions)
        part_b_exists = any(s.course_scope == SCOPE_PART_B for s in active_sessions)

        if full_exists:
            flash('This course already has a full-course session. Delete the existing session before assigning another teacher.', 'error')
            return redirect(url_for('class_management.index'))

        if part_a_exists and part_b_exists:
            flash('Both Part A and Part B are already assigned to teachers. Delete an existing section before adding another teacher.', 'error')
            return redirect(url_for('class_management.index'))

        if course_scope == SCOPE_FULL and (part_a_exists or part_b_exists):
            flash('This course is already split between teachers. Delete the split sections before creating a full-course session.', 'error')
            return redirect(url_for('class_management.index'))

        if course_scope == SCOPE_PART_A and part_a_exists:
            flash('Part A is already assigned to another teacher. Delete the existing Part A session before reassigning.', 'error')
            return redirect(url_for('class_management.index'))

        if course_scope == SCOPE_PART_B and part_b_exists:
            flash('Part B is already assigned to another teacher. Delete the existing Part B session before reassigning.', 'error')
            return redirect(url_for('class_management.index'))

        if course_scope == SCOPE_FULL:
            session_obj = Session(
                year=year,
                term=term,
                academic_session=academic_session,
                course_code=course_code,
                course_name=course_name,
                teacher_id=teacher.id,
                course_type=course_type,
                category=category,
                course_scope=SCOPE_FULL,
                window_id=window_id,
            )
            db.session.add(session_obj)
            db.session.flush()  # Get session ID before commit
            
            # Automatically add students from Students Management based on batch
            if batch and Student:
                try:
                    students_from_batch = Student.query.filter_by(batch=batch).all()
                    added_count = 0
                    skipped_count = 0
                    
                    # Get existing student IDs for this session to avoid duplicates
                    existing_student_ids = set()
                    
                    for student in students_from_batch:
                        # Check if already exists
                        existing = ClassStudent.query.filter_by(
                            session_id=session_obj.id,
                            student_id=student.student_id
                        ).first()
                        
                        if existing:
                            skipped_count += 1
                            continue
                        
                        class_student = ClassStudent(
                            student_id=student.student_id,
                            name=student.name,
                            session_id=session_obj.id,
                            teacher_id=teacher.id
                        )
                        db.session.add(class_student)
                        db.session.flush()  # Flush to get class_student.id before carry on
                        
                        # Carry on assessment marks if enabled in registration
                        _carry_on_assessment_marks(class_student, session_obj)
                        
                        _replicate_student_to_peers(session_obj, class_student)
                        added_count += 1
                    
                    db.session.commit()
                    
                    # Emit WebSocket event for live update
                    try:
                        from utils.websocket_events import emit_session_created
                        emit_session_created({
                            'session_id': session_obj.id,
                            'course_code': session_obj.course_code,
                            'course_name': session_obj.course_name,
                            'teacher_id': session_obj.teacher_id
                        })
                    except Exception as e:
                        current_app.logger.warning(f'Failed to emit session created event: {e}')
                    
                    if added_count > 0:
                        flash(f'Session created successfully! Automatically added {added_count} students from batch {batch}.', 'success')
                    else:
                        flash('Session created successfully!', 'success')
                        if skipped_count > 0:
                            flash(f'Note: {skipped_count} students were already in the session.', 'info')
                except Exception as e:
                    db.session.rollback()
                    current_app.logger.error(f'Error auto-adding students: {str(e)}', exc_info=True)
                    db.session.commit()  # Commit session even if student addition fails
                    flash('Session created successfully, but there was an error adding students automatically.', 'warning')
            else:
                db.session.commit()
                # Emit WebSocket event for live update
                try:
                    from utils.websocket_events import emit_session_created
                    emit_session_created({
                        'session_id': session_obj.id,
                        'course_code': session_obj.course_code,
                        'course_name': session_obj.course_name,
                        'teacher_id': session_obj.teacher_id
                    })
                except Exception as e:
                    current_app.logger.warning(f'Failed to emit session created event: {e}')
                flash('Session created successfully!', 'success')
            
            current_app.logger.info(f'Session created successfully - ID: {session_obj.id}, Name: {course_name}')
            return redirect(url_for('class_management.index'))

        # Handle split course (Part A/B)
        counterpart_scope = SCOPE_PART_B if course_scope == SCOPE_PART_A else SCOPE_PART_A

        try:
            partner_teacher_id_int = int(partner_teacher_id) if partner_teacher_id else None
        except (TypeError, ValueError):
            partner_teacher_id_int = None

        partner_teacher = None
        if partner_teacher_id_int:
            partner_teacher = Teacher.query.get(partner_teacher_id_int)

        if not partner_teacher:
            flash('Please select the teacher who will take the other part.', 'error')
            return redirect(url_for('class_management.index'))

        if partner_teacher.id == teacher.id:
            flash('Please assign a different teacher for the other part.', 'error')
            return redirect(url_for('class_management.index'))

        split_group_id = str(uuid4())

        current_session = Session(
            year=year,
            term=term,
            academic_session=academic_session,
            course_code=course_code,
            course_name=course_name,
            teacher_id=teacher.id,
            course_type=course_type,
            category=category,
            course_scope=course_scope,
            split_group_id=split_group_id,
            window_id=window_id,
        )
        db.session.add(current_session)
        db.session.flush()

        invite = ClassSplitInvite(
            split_group_id=split_group_id,
            inviter_session_id=current_session.id,
            inviter_teacher_id=teacher.id,
            invited_teacher_id=partner_teacher.id,
            invited_scope=counterpart_scope,
            status='pending'
        )
        db.session.add(invite)
        
        # Automatically add students from Students Management based on batch
        if batch and Student:
            try:
                students_from_batch = Student.query.filter_by(batch=batch).all()
                added_count = 0
                skipped_count = 0
                not_registered_count = 0
                
                for student in students_from_batch:
                    # Check if already exists
                    existing = ClassStudent.query.filter_by(
                        session_id=current_session.id,
                        student_id=student.student_id
                    ).first()
                    
                    if existing:
                        skipped_count += 1
                        continue
                    
                    # Check if student is registered for this course (finalized registration only)
                    if StudentCourseRegistration and current_session.course_code and current_session.academic_session and current_session.year and current_session.term:
                        registration = StudentCourseRegistration.query.filter_by(
                            student_id=student.id,
                            course_code=current_session.course_code,
                            academic_session=current_session.academic_session,
                            year=current_session.year,
                            term=current_session.term,
                            status='finalized'
                        ).first()
                        
                        if not registration:
                            not_registered_count += 1
                            current_app.logger.info(f'Student {student.student_id} ({student.name}) not registered for course {current_session.course_code}, skipping...')
                            continue
                    
                    class_student = ClassStudent(
                        student_id=student.student_id,
                        name=student.name,
                        session_id=current_session.id,
                        teacher_id=teacher.id
                    )
                    db.session.add(class_student)
                    db.session.flush()  # Flush to get class_student.id before carry on
                    
                    # Carry on assessment marks if enabled in registration
                    _carry_on_assessment_marks(class_student, current_session)
                    
                    _replicate_student_to_peers(current_session, class_student)
                    added_count += 1
                
                # Commit session, invite, and students together
                db.session.commit()
                
                # Emit WebSocket event for live update
                try:
                    from utils.websocket_events import emit_session_created
                    emit_session_created({
                        'session_id': current_session.id,
                        'course_code': current_session.course_code,
                        'course_name': current_session.course_name,
                        'teacher_id': current_session.teacher_id
                    })
                except Exception as e:
                    current_app.logger.warning(f'Failed to emit session created event: {e}')
                
                if added_count > 0:
                    message = f'Split course created. Invitation sent. Automatically added {added_count} students from batch {batch}.'
                    if not_registered_count > 0:
                        message += f' Skipped {not_registered_count} student(s) not registered for this course.'
                    flash(message, 'success')
                else:
                    flash('Split course created. Invitation sent to the selected teacher.', 'success')
                    if skipped_count > 0:
                        flash(f'Note: {skipped_count} students were already in the session.', 'info')
            except Exception as e:
                db.session.rollback()
                current_app.logger.error(f'Error auto-adding students to split course: {str(e)}', exc_info=True)
                # Still commit the session and invite even if student addition fails
                db.session.add(current_session)
                db.session.add(invite)
                db.session.commit()
                flash('Split course created. Invitation sent, but there was an error adding students automatically.', 'warning')
        else:
            # Commit session and invite
            db.session.commit()
            flash('Split course created. Invitation sent to the selected teacher.', 'success')
        
        current_app.logger.info(f'Split session created successfully - ID: {current_session.id}, Name: {course_name}')
        return redirect(url_for('class_management.index'))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error creating session: {str(e)}', exc_info=True)
        flash(f'Error creating session: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))


@class_management_bp.route('/create_external_session', methods=['POST'])
@login_required
def create_external_session():
    """Create a manually entered external course session (isolated from curriculum)."""
    try:
        teacher = _ensure_current_teacher()
        course_name = (request.form.get('course_name') or '').strip()
        course_code = (request.form.get('course_code') or '').strip()
        academic_session = (request.form.get('academic_session') or '').strip()
        year = (request.form.get('year') or '').strip()
        term = (request.form.get('term') or '').strip()
        course_type = (request.form.get('course_type') or 'theory').strip().lower()
        category = (request.form.get('category') or 'ug').strip().lower()

        if not course_name or not course_code:
            flash('Course name and course code are required!', 'error')
            return redirect(url_for('class_management.index'))
        if not academic_session or not year or not term:
            flash('Academic Session, Year, and Term are required.', 'error')
            return redirect(url_for('class_management.index'))
        if course_type not in {'theory', 'sessional'}:
            course_type = 'theory'
        if category not in {'ug', 'pg'}:
            category = 'ug'

        conflict_message = _external_course_conflict(course_code, teacher.id)
        if conflict_message:
            flash(conflict_message, 'error')
            return redirect(url_for('class_management.index'))

        session_obj = Session(
            year=year,
            term=term,
            academic_session=academic_session,
            course_code=course_code,
            course_name=course_name,
            teacher_id=teacher.id,
            course_type=course_type,
            category=category,
            course_scope=SCOPE_FULL,
            is_external_course=True,
        )
        stamp_window_id(session_obj)
        db.session.add(session_obj)
        db.session.commit()

        students_file = request.files.get('students_file')
        if students_file and students_file.filename:
            added_count, skipped_count, upload_error = _import_students_from_excel(
                session_obj, students_file, teacher
            )
            if upload_error:
                flash(f'External course created, but student upload failed: {upload_error}', 'warning')
            else:
                db.session.commit()
                message = 'External course created successfully!'
                if added_count:
                    message += f' Added {added_count} student(s).'
                if skipped_count:
                    message += f' Skipped {skipped_count} duplicate student(s).'
                flash(message, 'success')
            return redirect(url_for('class_management.index'))

        flash('External course created successfully!', 'success')
        return redirect(url_for('class_management.index'))
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error creating external session: {e}', exc_info=True)
        flash(f'Error creating external course: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))


@class_management_bp.route('/split_invites/<int:invite_id>/accept', methods=['POST'])
@login_required
def accept_split_invite(invite_id):
    teacher = _ensure_current_teacher()
    invite = ClassSplitInvite.query.get_or_404(invite_id)
    if invite.invited_teacher_id != teacher.id:
        flash('You are not authorized to respond to this invitation.', 'error')
        return redirect(url_for('class_management.index'))
    if invite.status != 'pending':
        flash('This invitation has already been processed.', 'info')
        return redirect(url_for('class_management.index'))

    inviter_session = get_for_window(Session, invite.inviter_session_id)
    if not inviter_session:
        invite.status = 'declined'
        invite.responded_at = datetime.utcnow()
        db.session.commit()
        flash('The original course is no longer available.', 'error')
        return redirect(url_for('class_management.index'))

    new_session = Session(
        year=inviter_session.year,
        term=inviter_session.term,
        academic_session=inviter_session.academic_session,
        course_code=inviter_session.course_code,
        course_name=inviter_session.course_name,
        teacher_id=teacher.id,
        course_type=inviter_session.course_type,
        category=inviter_session.category,
        course_scope=invite.invited_scope,
        split_group_id=invite.split_group_id
    )
    db.session.add(new_session)
    db.session.flush()

    inviter_students = ClassStudent.query.filter_by(session_id=inviter_session.id).all()
    for stu in inviter_students:
        clone = ClassStudent(
            student_id=stu.student_id,
            name=stu.name,
            session_id=new_session.id,
            teacher_id=teacher.id
        )
        db.session.add(clone)

    invite.status = 'accepted'
    invite.responded_at = datetime.utcnow()
    db.session.commit()
    flash('Invitation accepted. The course has been added to your dashboard.', 'success')
    return redirect(url_for('class_management.index'))


@class_management_bp.route('/split_invites/<int:invite_id>/decline', methods=['POST'])
@login_required
def decline_split_invite(invite_id):
    teacher = _ensure_current_teacher()
    invite = ClassSplitInvite.query.get_or_404(invite_id)
    if invite.invited_teacher_id != teacher.id:
        flash('You are not authorized to respond to this invitation.', 'error')
        return redirect(url_for('class_management.index'))
    if invite.status != 'pending':
        flash('This invitation has already been processed.', 'info')
        return redirect(url_for('class_management.index'))

    invite.status = 'declined'
    invite.responded_at = datetime.utcnow()
    db.session.commit()
    flash('Invitation declined.', 'info')
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/upload_students/<int:session_id>', methods=['POST'])
@login_required
def upload_students(session_id):
    """Upload students from Excel file"""
    teacher = _ensure_current_teacher()
    return_to = (request.form.get('return_to') or '').strip()
    session = get_or_404_for_window(Session, session_id)

    if 'file' not in request.files:
        flash('No file uploaded!', 'error')
        return _student_upload_redirect(session_id, return_to=return_to)

    file = request.files['file']
    if file.filename == '':
        flash('No file selected!', 'error')
        return _student_upload_redirect(session_id, return_to=return_to)

    try:
        added_count, skipped_count, upload_error = _import_students_from_excel(session, file, teacher)
        if upload_error:
            flash(upload_error, 'error')
            return _student_upload_redirect(session_id, return_to=return_to)

        db.session.commit()
        if added_count == 0 and skipped_count == 0:
            flash('No students were added. Check that the Excel file has Student ID and Name columns with valid data.', 'warning')
        else:
            message = f'Successfully uploaded {added_count} student(s).'
            if skipped_count:
                message += f' Skipped {skipped_count} duplicate student(s).'
            flash(message, 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error uploading students: {str(e)}', 'error')

    return _student_upload_redirect(session_id, return_to=return_to)

@class_management_bp.route('/take_attendance/<int:session_id>', methods=['GET', 'POST'])
@login_required
def take_attendance(session_id):
    """Take or update attendance for a session."""
    session = get_or_404_for_window(Session, session_id)
    students = _class_students_for_session(session_id)

    if not user_owns_class_session(current_user, session):
        flash('You are not authorized to manage attendance for this session.', 'danger')
        return redirect(url_for('class_management.index'))

    if request.method == 'POST':
        try:
            date_val = datetime.strptime(request.form.get('date'), '%Y-%m-%d').date()
            double_class = request.form.get('double_class') == '1'
            
            # Overwrite logic: Delete existing records for this date first
            ClassAttendance.query.filter_by(session_id=session_id, date=date_val).delete()
            
            # Add new records
            for student in students:
                is_present = request.form.get(f'student_{student.id}') == 'present'
                num_classes = 2 if double_class else 1
                for slot_idx in range(1, num_classes + 1):
                    attendance = ClassAttendance(
                        date=date_val,
                        is_present=is_present,
                        status=ClassAttendance.STATUS_PRESENT if is_present else ClassAttendance.STATUS_ABSENT,
                        slot_number=slot_idx,
                        student_id=student.id,
                        session_id=session_id,
                        teacher_id=session.teacher_id
                    )
                    db.session.add(attendance)
            
            db.session.commit()
            # Emit WebSocket event for live update
            try:
                from utils.websocket_events import emit_attendance_update
                emit_attendance_update(session_id, {
                    'date': date_val.isoformat(),
                    'double_class': double_class
                })
            except Exception as e:
                current_app.logger.warning(f'Failed to emit attendance update event: {e}')
            flash('Attendance saved successfully!', 'success')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error saving attendance for session {session_id}: {e}")
            flash(f'Error saving attendance: {str(e)}', 'error')
            return redirect(url_for('class_management.take_attendance', session_id=session_id, date=request.form.get('date')))

    # GET request logic
    try:
        selected_date_str = request.args.get('date', date.today().strftime('%Y-%m-%d'))
        selected_date = datetime.strptime(selected_date_str, '%Y-%m-%d').date()

        # Fetch existing records for the selected date
        existing_records = ClassAttendance.query.filter_by(
            session_id=session_id,
            date=selected_date
        ).all()
        
        # Prepare data for the template
        attendance_status = {}
        is_double_class = False
        if existing_records:
            # Check if it was a double class
            student_counts = defaultdict(int)
            for record in existing_records:
                student_counts[record.student_id] += 1
            if student_counts and max(student_counts.values()) > 1:
                is_double_class = True
            
            # Get the attendance status for each student
            for student in students:
                # A student is marked 'present' if they have at least one present record on that day
                is_present = any(r.is_present for r in existing_records if r.student_id == student.id)
                attendance_status[student.id] = is_present

        return render_template('class_management/take_attendance.html', 
                                session=session, 
                                students=students, 
                                today=selected_date_str,
                                attendance_status=attendance_status,
                                is_double_class=is_double_class,
                                split_meta=_build_split_context(session))
    except Exception as e:
        current_app.logger.error(f"Error loading attendance page for session {session_id}: {e}")
        flash(f'Error loading attendance page: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/view_attendance/<int:session_id>')
@login_required
def view_attendance(session_id):
    """View attendance for a session and display a detailed report."""
    session = get_or_404_for_window(Session, session_id)
    
    # Check user permissions: admin/head can view any session, regular teachers only their own
    user_roles = set(parse_roles(getattr(current_user, 'role', '')))
    if getattr(current_user, 'active_role', None):
        user_roles = set(parse_roles(current_user.active_role))
    can_view_all = is_admin(current_user) or 'head' in user_roles or 'dean' in user_roles
    
    # If not admin/head, check if this session belongs to the current teacher
    if not can_view_all:
        current_teacher = _ensure_current_teacher()
        if not current_teacher or session.teacher_id != current_teacher.id:
            flash('You do not have permission to view attendance for this session.', 'danger')
            return redirect(url_for('class_management.index'))
    
    attendance_summary = _build_attendance_summary(session, include_archived=True)
    part_total_classes = None
    if session.course_scope in SPLIT_PARTS:
        part_total_classes = attendance_summary.get('per_session_totals', {}).get(session.id, 0)
    
    # Check if this is a split course
    is_split_course = session.split_group_id and session.course_scope in SPLIT_PARTS
    related_sessions = []
    if is_split_course:
        # Include archived peers as well; they may contain historical attendance from
        # earlier assignment cycles. We merge same teacher/scope sessions below.
        related_sessions = _resolve_attendance_related_sessions(session, include_archived=True)
    
    # For split courses, prepare separate attendance data for each teacher
    teacher_attendance_data = []
    if is_split_course and related_sessions:
        # Check user permissions: admin/head can see all parts, regular teachers only their own
        user_roles = set(parse_roles(getattr(current_user, 'role', '')))
        if getattr(current_user, 'active_role', None):
            user_roles = set(parse_roles(current_user.active_role))
        can_view_all = is_admin(current_user) or 'head' in user_roles or 'dean' in user_roles
        
        # Get current teacher if not admin/head
        current_teacher = None
        if not can_view_all:
            current_teacher = _ensure_current_teacher()
        
        # Get all students from all related sessions (for marks calculation; exclude deleted students)
        all_session_ids = [s.id for s in related_sessions]
        all_students = _class_students_for_sessions(all_session_ids)
        students_by_session = defaultdict(list)
        for stu in all_students:
            students_by_session[stu.session_id].append(stu)
        
        # Get combined attendance summary for marks (this stays the same)
        agg_student_map = attendance_summary.get('per_student', {})
        agg_total_classes = attendance_summary.get('total_classes', 0)
        
        # Group sessions by teacher + scope so old/new session splits are rendered as one report.
        grouped_sessions = {}
        for related_session in related_sessions:
            # Skip if user doesn't have permission to view this session
            if not can_view_all:
                if not current_teacher or related_session.teacher_id != current_teacher.id:
                    continue
            key = (related_session.teacher_id, related_session.course_scope)
            grouped_sessions.setdefault(key, []).append(related_session)

        # Prepare merged attendance data for each teacher/scope group.
        for _, grouped in grouped_sessions.items():
            grouped = sorted(grouped, key=lambda s: s.id)
            grouped_session_ids = [s.id for s in grouped]
            current_scope_session = next((s for s in grouped if s.id == session.id), None)
            primary_session = current_scope_session or grouped[-1]

            # Build a canonical student list by public student_id, preferring the current session row.
            group_students = []
            for sid in grouped_session_ids:
                group_students.extend(students_by_session.get(sid, []))
            if not group_students:
                continue

            students_by_public_id = {}
            for stu in sorted(group_students, key=lambda x: (x.student_id, x.id)):
                selected = students_by_public_id.get(stu.student_id)
                if selected is None:
                    students_by_public_id[stu.student_id] = stu
                    continue
                if stu.session_id == primary_session.id and selected.session_id != primary_session.id:
                    students_by_public_id[stu.student_id] = stu

            class_student_to_public = {stu.id: stu.student_id for stu in group_students}
            group_attendance_records = ClassAttendance.query.filter(
                ClassAttendance.session_id.in_(grouped_session_ids)
            ).order_by(ClassAttendance.date, ClassAttendance.id).all()
            if not group_attendance_records:
                continue

            attendance_by_date = defaultdict(list)
            student_attendance_by_public = defaultdict(lambda: defaultdict(list))
            for record in group_attendance_records:
                public_id = class_student_to_public.get(record.student_id)
                if not public_id:
                    continue
                attendance_by_date[record.date].append((public_id, record))
                student_attendance_by_public[public_id][record.date].append(record)
            
            daily_class_counts = {}
            for date, records in attendance_by_date.items():
                student_records_on_date = defaultdict(list)
                for public_id, record in records:
                    student_records_on_date[public_id].append(record)
                student_counts_on_date = {
                    public_id: len(_records_by_slot(student_records))
                    for public_id, student_records in student_records_on_date.items()
                }
                daily_class_counts[date] = _cap_classes_per_day(
                    max(student_counts_on_date.values()) if student_counts_on_date else 0
                )
            
            headers_with_meta = []
            sorted_dates = sorted(daily_class_counts.keys())
            for date in sorted_dates:
                count = daily_class_counts.get(date, 0)
                if count == 1:
                    headers_with_meta.append({'label': date.strftime('%b %d, %Y'), 'date': date.strftime('%Y-%m-%d'), 'slot': 1})
                else:
                    for i in range(1, count + 1):
                        headers_with_meta.append({'label': f"{date.strftime('%b %d')} ({i})", 'date': date.strftime('%Y-%m-%d'), 'slot': i})
            
            student_report_data = []
            for student in sorted(students_by_public_id.values(), key=lambda s: s.student_id):
                student_public_id = student.student_id
                student_records_by_date = student_attendance_by_public.get(student_public_id, {})
                # Use combined stats for marks calculation
                agg_stats = agg_student_map.get(student_public_id, {'present': 0, 'percentage': 0, 'marks': 0})
                
                attendance_row = []
                for date in sorted_dates:
                    records_for_date = student_records_by_date.get(date, [])
                    slot_map = _records_by_slot(records_for_date)
                    num_classes_on_day = daily_class_counts.get(date, 0)
                    for i in range(num_classes_on_day):
                        cell = {
                            'status': '-',
                            'date': date.strftime('%Y-%m-%d'),
                            'slot': i + 1
                        }
                        record = slot_map.get(i + 1)
                        if record:
                            cell['status'] = _attendance_label_from_record(record)
                            cell['record_id'] = record.id
                        attendance_row.append(cell)
                
                student_data = {
                    'info': student,
                    'attendance_row': attendance_row,
                    'total_classes': agg_stats.get('effective_total_classes', agg_total_classes),
                    'present_count': agg_stats['present'],  # Combined present count
                    'percentage': f"{agg_stats['percentage']:.2f}%",  # Combined percentage
                    'marks': agg_stats['marks'],  # Combined marks
                    'marks_manual': agg_stats.get('marks_manual', False)  # Whether marks are manually set
                }
                student_report_data.append(student_data)
            
            teacher_attendance_data.append({
                'session': primary_session,
                'teacher_name': primary_session.teacher.name if primary_session.teacher else 'Unknown',
                'teacher_short': primary_session.teacher.short_name if primary_session.teacher else '',
                'scope_label': COURSE_SCOPE_LABELS.get(primary_session.course_scope, 'Part'),
                'headers_with_meta': headers_with_meta,
                'student_report_data': student_report_data,
                'part_total_classes': sum(
                    attendance_summary.get('per_session_totals', {}).get(sid, 0)
                    for sid in grouped_session_ids
                ),
                'unique_dates': sorted(attendance_by_date.keys(), reverse=True)  # For delete modal
            })
        
        # Get unique dates from all sessions for delete modal
        all_attendance_records = ClassAttendance.query.filter(ClassAttendance.session_id.in_(all_session_ids)).all()
        attendance_by_date_all = defaultdict(list)
        for record in all_attendance_records:
            attendance_by_date_all[record.date].append(record)
        unique_dates_for_modal = sorted(attendance_by_date_all.keys(), reverse=True)
        
        return render_template(
            'class_management/view_attendance.html',
            session=session,
            headers=[],
            headers_with_meta=[],
            student_report_data=[],
            unique_dates=unique_dates_for_modal,
            split_meta=_build_split_context(session, attendance_summary),
            attendance_summary=attendance_summary,
            part_total_classes=part_total_classes,
            is_split_course=True,
            teacher_attendance_data=teacher_attendance_data
        )
    
    # Non-split course: original logic (exclude students deleted from Student Management)
    students = _class_students_for_session(session_id)
    all_attendance_records = ClassAttendance.query.filter_by(session_id=session_id).order_by(ClassAttendance.date, ClassAttendance.id).all()
    student_lookup = {stu.id: stu for stu in students}

    headers_with_meta = []
    if not all_attendance_records:
        # Show enrolled students even when no attendance has been taken yet.
        # This avoids an empty report experience after students are added.
        empty_student_rows = []
        for student in students:
            empty_student_rows.append({
                'info': student,
                'attendance_row': [],
                'total_classes': 0,
                'present_count': 0,
                'percentage': "0.00%",
                'marks': 0
            })
        return render_template(
            'class_management/view_attendance.html',
            session=session,
            students=students,
            headers=[],
            headers_with_meta=headers_with_meta,
            student_report_data=empty_student_rows,
            unique_dates=[],
            split_meta=_build_split_context(session, attendance_summary),
            attendance_summary=attendance_summary,
            part_total_classes=part_total_classes,
            is_split_course=False,
            teacher_attendance_data=[]
        )

    attendance_by_date = defaultdict(list)
    for record in all_attendance_records:
        attendance_by_date[record.date].append(record)

    unique_dates_for_modal = sorted(attendance_by_date.keys(), reverse=True)

    daily_class_counts = {}
    for date, records in attendance_by_date.items():
        student_records_on_date = defaultdict(list)
        for record in records:
            student_records_on_date[record.student_id].append(record)
        student_counts_on_date = {
            student_id: len(_records_by_slot(student_records))
            for student_id, student_records in student_records_on_date.items()
        }
        daily_class_counts[date] = _cap_classes_per_day(
            max(student_counts_on_date.values()) if student_counts_on_date else 0
        )

    headers = []
    headers_with_meta = []
    sorted_dates = sorted(daily_class_counts.keys())
    for date in sorted_dates:
        count = daily_class_counts.get(date, 0)
        if count == 1:
            headers.append(date.strftime('%b %d, %Y'))
            headers_with_meta.append({'label': date.strftime('%b %d, %Y'), 'date': date.strftime('%Y-%m-%d'), 'slot': 1})
        else:
            for i in range(1, count + 1):
                headers.append(f"{date.strftime('%b %d')} ({i})")
                headers_with_meta.append({'label': f"{date.strftime('%b %d')} ({i})", 'date': date.strftime('%Y-%m-%d'), 'slot': i})

    student_report_data = []
    agg_student_map = attendance_summary.get('per_student', {})
    agg_total_classes = attendance_summary.get('total_classes', sum(daily_class_counts.values()))

    for student in students:
        student_records = [r for r in all_attendance_records if r.student_id == student.id]
        agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})

        attendance_row = []
        student_attendance_by_date = defaultdict(list)
        for r in student_records:
            student_attendance_by_date[r.date].append(r)

        for date in sorted_dates:
            records_for_date = student_attendance_by_date[date]
            slot_map = _records_by_slot(records_for_date)
            num_classes_on_day = daily_class_counts.get(date, 0)
            for i in range(num_classes_on_day):
                cell = {
                    'status': '-',
                    'date': date.strftime('%Y-%m-%d'),
                    'slot': i + 1
                }
                record = slot_map.get(i + 1)
                if record:
                    cell['status'] = _attendance_label_from_record(record)
                    cell['record_id'] = record.id
                attendance_row.append(cell)

        student_data = {
            'info': student,
            'attendance_row': attendance_row,
            'total_classes': agg_stats.get('effective_total_classes', agg_total_classes),
            'present_count': agg_stats['present'],
            'percentage': f"{agg_stats['percentage']:.2f}%",
            'marks': agg_stats['marks']
        }
        student_report_data.append(student_data)

    unique_headers_with_metadata = []
    header_index = 0
    for date in sorted_dates:
        count = daily_class_counts.get(date, 0)
        if count <= 1:
            unique_headers_with_metadata.append({'label': date.strftime('%b %d, %Y'), 'date': date.strftime('%Y-%m-%d'), 'slot': 1})
        else:
            for i in range(1, count + 1):
                unique_headers_with_metadata.append({'label': f"{date.strftime('%b %d')} ({i})", 'date': date.strftime('%Y-%m-%d'), 'slot': i})
    return render_template(
        'class_management/view_attendance.html',
        session=session,
        headers=headers,
        headers_with_meta=unique_headers_with_metadata,
        student_report_data=student_report_data,
        unique_dates=unique_dates_for_modal,
        split_meta=_build_split_context(session, attendance_summary),
        attendance_summary=attendance_summary,
        part_total_classes=part_total_classes,
        is_split_course=False,
        teacher_attendance_data=[]
    )


@class_management_bp.route('/toggle_attendance_record/<int:record_id>', methods=['POST'])
@login_required
def toggle_attendance_record(record_id):
    record = ClassAttendance.query.get_or_404(record_id)
    session = get_or_404_for_window(Session, record.session_id)
    teacher = _ensure_current_teacher()
    if session.teacher_id != teacher.id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    data = request.get_json() or {}
    raw_target = data.get('target_status')
    current_label = _attendance_label_from_record(record)
    if raw_target in (None, ''):
        target_label = ATTENDANCE_CYCLE.get(current_label, 'A')
    else:
        target_label = _normalize_attendance_label(raw_target, allow_blank=True)
        if not target_label:
            return jsonify({'success': False, 'message': 'Invalid target_status'}), 400

    student_public_id = record.student.student_id
    student_db_id = record.student_id
    response_record_id = record.id
    _set_attendance_record_status(record, target_label)
    db.session.commit()

    attendance_summary = _build_attendance_summary(session)
    student_stats = attendance_summary.get('per_student', {}).get(
        student_public_id,
        {'present': 0, 'percentage': 0, 'marks': 0, 'effective_total_classes': 0}
    )

    # Emit WebSocket event for live update
    try:
        from utils.websocket_events import emit_attendance_update
        emit_attendance_update(record.session_id, {
            'record_id': response_record_id,
            'status': target_label,
            'student_id': student_public_id,
            'student_db_id': student_db_id,
            'present_count': student_stats.get('present', 0),
            'percentage': student_stats.get('percentage', 0),
            'marks': student_stats.get('marks', 0),
            'total_classes': student_stats.get('effective_total_classes', student_stats.get('base_total_classes', 0))
        })
    except Exception as e:
        current_app.logger.warning(f'Failed to emit attendance update event: {e}')

    return jsonify(_status_payload_for_response(
        student_stats=student_stats,
        status_label=target_label,
        record_id=response_record_id,
        student_db_id=student_db_id
    ))


@class_management_bp.route('/create_attendance_record/<int:session_id>', methods=['POST'])
@login_required
def create_attendance_record(session_id):
    """Create a new attendance record for a cell that currently shows '-' (no record)."""
    session = get_or_404_for_window(Session, session_id)
    teacher = _ensure_current_teacher()
    if session.teacher_id != teacher.id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    try:
        data = request.get_json() or {}
        class_student_id = data.get('class_student_id')
        date_str = data.get('date')
        slot_number = _normalize_slot_number(data.get('slot'))
        raw_target = data.get('target_status')
        if raw_target in (None, ''):
            target_label = 'P'
        else:
            target_label = _normalize_attendance_label(raw_target)
            if not target_label:
                return jsonify({'success': False, 'message': 'Invalid target_status'}), 400
        if not class_student_id or not date_str or slot_number is None:
            return jsonify({'success': False, 'message': 'class_student_id, date and valid slot required'}), 400
        from datetime import datetime
        try:
            date_val = datetime.strptime(date_str, '%Y-%m-%d').date()
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Invalid date format (use YYYY-MM-DD)'}), 400
        class_student = ClassStudent.query.filter_by(
            id=class_student_id,
            session_id=session_id
        ).first()
        if not class_student:
            return jsonify({'success': False, 'message': 'Student not found in this session'}), 404

        existing_records = ClassAttendance.query.filter_by(
            session_id=session_id,
            student_id=class_student_id,
            date=date_val
        ).order_by(ClassAttendance.id.asc()).all()
        slot_map = _records_by_slot(existing_records)
        existing_record_for_slot = slot_map.get(slot_number)

        if existing_record_for_slot:
            _set_attendance_record_status(existing_record_for_slot, target_label)
            if _normalize_slot_number(getattr(existing_record_for_slot, 'slot_number', None)) != slot_number:
                existing_record_for_slot.slot_number = slot_number
            db.session.commit()
            attendance_summary = _build_attendance_summary(session)
            student_public_id = class_student.student_id
            student_stats = attendance_summary.get('per_student', {}).get(
                student_public_id,
                {'present': 0, 'percentage': 0, 'marks': 0, 'effective_total_classes': 0}
            )
            return jsonify(_status_payload_for_response(
                student_stats=student_stats,
                status_label=target_label,
                record_id=existing_record_for_slot.id,
                student_db_id=class_student_id
            ))

        if len(slot_map) >= MAX_CLASSES_PER_DAY:
            return jsonify({
                'success': False,
                'message': f'Maximum {MAX_CLASSES_PER_DAY} attendance entries are allowed for this date.'
            }), 400

        new_record = ClassAttendance(
            session_id=session_id,
            student_id=class_student_id,
            teacher_id=session.teacher_id,
            date=date_val,
            is_present=True,
            slot_number=slot_number
        )
        _set_attendance_record_status(new_record, target_label)
        db.session.add(new_record)
        db.session.commit()
        attendance_summary = _build_attendance_summary(session)
        student_public_id = class_student.student_id
        student_stats = attendance_summary.get('per_student', {}).get(
            student_public_id,
            {'present': 0, 'percentage': 0, 'marks': 0, 'effective_total_classes': 0}
        )
        try:
            from utils.websocket_events import emit_attendance_update
            emit_attendance_update(session_id, {
                'record_id': new_record.id,
                'status': target_label,
                'student_id': student_public_id,
                'student_db_id': class_student_id,
                'present_count': student_stats.get('present', 0),
                'percentage': student_stats.get('percentage', 0),
                'marks': student_stats.get('marks', 0),
                'total_classes': student_stats.get('effective_total_classes', student_stats.get('base_total_classes', 0))
            })
        except Exception as e:
            current_app.logger.warning(f'Failed to emit attendance update event: {e}')
        return jsonify(_status_payload_for_response(
            student_stats=student_stats,
            status_label=target_label,
            record_id=new_record.id,
            student_db_id=class_student_id
        ))
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating attendance record: {e}", exc_info=True)
        return jsonify({'success': False, 'message': str(e)}), 500

@class_management_bp.route('/save_attendance_marks_manual/<int:session_id>', methods=['POST'])
@login_required
def save_attendance_marks_manual(session_id):
    """Save manual attendance marks override for students"""
    session = get_or_404_for_window(Session, session_id)
    teacher = _ensure_current_teacher()
    if session.teacher_id != teacher.id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    try:
        data = request.get_json()
        student_id = data.get('student_id')  # Public student_id (string)
        marks_value = data.get('marks')
        
        if not student_id:
            return jsonify({'success': False, 'message': 'Student ID required'}), 400
        
        # Find the student in this session
        student = ClassStudent.query.filter_by(
            session_id=session_id,
            student_id=student_id
        ).first()
        
        if not student:
            return jsonify({'success': False, 'message': 'Student not found'}), 404
        
        # For split courses, set manual marks on all related student records with the same student_id
        # This ensures manual marks are preserved across all sessions
        related_sessions = _get_related_sessions(session)
        session_ids = [s.id for s in related_sessions if s]
        all_students_with_same_id = ClassStudent.query.filter(
            ClassStudent.student_id == student_id,
            ClassStudent.session_id.in_(session_ids)
        ).all()
        
        # Set manual marks (None to clear manual override)
        if marks_value is not None and marks_value != '':
            try:
                marks_float = float(marks_value)
                # Set manual marks on all related student records
                for s in all_students_with_same_id:
                    s.attendance_marks_manual = marks_float
            except (ValueError, TypeError):
                return jsonify({'success': False, 'message': 'Invalid marks value'}), 400
        else:
            # Clear manual override from all related student records
            for s in all_students_with_same_id:
                s.attendance_marks_manual = None
        
        db.session.commit()
        
        # Rebuild attendance summary to get updated marks
        attendance_summary = _build_attendance_summary(session)
        student_stats = attendance_summary.get('per_student', {}).get(student_id, {'present': 0, 'percentage': 0, 'marks': 0, 'marks_manual': False})
        
        return jsonify({
            'success': True,
            'marks': student_stats.get('marks', 0),
            'marks_manual': student_stats.get('marks_manual', False),
            'present_count': student_stats.get('present', 0),
            'percentage': f"{student_stats.get('percentage', 0):.2f}%"
        })
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving manual attendance marks: {e}", exc_info=True)
        return jsonify({'success': False, 'message': str(e)}), 500

@class_management_bp.route('/students/<int:session_id>')
@login_required
def students_list(session_id):
    """View students list for a session"""
    session = get_or_404_for_window(Session, session_id)
    students = _class_students_for_session(session_id)
    
    # Get all batches for filter dropdown
    batches = []
    if Student:
        try:
            all_batches = db.session.query(Student.batch).distinct().filter(Student.batch.isnot(None)).order_by(Student.batch.desc()).all()
            batches = [batch[0] for batch in all_batches]
        except Exception:
            batches = []
    
    return render_template('class_management/students_list.html', 
                         session=session, students=students,
                         split_meta=_build_split_context(session),
                         batches=batches)

@class_management_bp.route('/api/students', methods=['GET'])
@login_required
def get_students_for_selection():
    """Get students from Students Management for selection (AJAX)"""
    try:
        if not Student:
            current_app.logger.warning('Student model not available in get_students_for_selection')
            return jsonify({'success': False, 'message': 'Students Management module not available'}), 503
        
        batch_filter = request.args.get('batch', '').strip()
        search = request.args.get('search', '').strip()
        
        current_app.logger.info(f'Fetching students - batch: {batch_filter}, search: {search}')
        
        query = Student.query
        
        if batch_filter:
            query = query.filter(Student.batch == batch_filter)
        
        if search:
            query = query.filter(
                or_(
                    Student.name.ilike(f'%{search}%'),
                    Student.student_id.ilike(f'%{search}%')
                )
            )
        
        # Increase limit to show more students (or remove limit entirely)
        students = query.order_by(Student.student_id.asc()).limit(500).all()
        
        current_app.logger.info(f'Found {len(students)} students')
        
        return jsonify({
            'success': True,
            'students': [{
                'id': s.id,
                'student_id': s.student_id,
                'name': s.name,
                'batch': s.batch,
                'email': s.email,
                'phone': s.phone
            } for s in students]
        })
    except Exception as e:
        current_app.logger.error(f'Error in get_students_for_selection: {str(e)}', exc_info=True)
        return jsonify({'success': False, 'message': f'Error loading students: {str(e)}'}), 500

@class_management_bp.route('/api/curricula', methods=['GET'])
@login_required
def get_curricula_by_batch():
    """Get curricula applicable to a specific batch (AJAX)"""
    if not Curriculum:
        current_app.logger.error('Curriculum model not available')
        return jsonify({'success': False, 'message': 'Curriculum Management module not available'}), 503
    
    batch = request.args.get('batch', '').strip()
    if not batch:
        return jsonify({'success': False, 'message': 'Batch is required'}), 400
    
    try:
        # Find curricula where the batch is in applicable_batches
        all_curricula = Curriculum.query.all()
        applicable_curricula = []
        
        # Normalize the input batch
        normalized_batch = str(batch).strip()
        current_app.logger.info(f'Searching curricula for batch: {normalized_batch}, Total curricula: {len(all_curricula)}')
        
        window_id = get_effective_window_id(admin_override=False) if get_effective_window_id else 1
        for curriculum in all_curricula:
            batches_list = curriculum.get_batches_list(window_id)
            current_app.logger.debug(f'Curriculum {curriculum.id} ({curriculum.name}) has batches: {batches_list}')
            
            # Check if the batch matches any batch in the list
            for b in batches_list:
                if str(b).strip() == normalized_batch:
                    applicable_curricula.append({
                        'id': curriculum.id,
                        'name': curriculum.name,
                        'date': curriculum.date
                    })
                    current_app.logger.info(f'Found matching curriculum: {curriculum.name} (ID: {curriculum.id})')
                    break  # Found a match, no need to check other batches for this curriculum
        
        current_app.logger.info(f'Found {len(applicable_curricula)} applicable curricula for batch {normalized_batch}')
        
        return jsonify({
            'success': True,
            'curricula': applicable_curricula,
            'batch_searched': normalized_batch,
            'total_curricula_checked': len(all_curricula),
            'debug': {
                'batch_received': batch,
                'normalized_batch': normalized_batch,
                'applicable_count': len(applicable_curricula)
            }
        })
    except Exception as e:
        current_app.logger.error(f'Error in get_curricula_by_batch: {str(e)}', exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error fetching curricula: {str(e)}',
            'error': str(e)
        }), 500

@class_management_bp.route('/api/curriculum/<int:curriculum_id>/years-terms', methods=['GET'])
@login_required
def get_years_terms_by_curriculum(curriculum_id):
    """Get distinct (display) years and terms from courses in a curriculum (AJAX)"""
    if not Course:
        return jsonify({'success': False, 'message': 'Course Management module not available'}), 503
    
    Curriculum.query.get_or_404(curriculum_id)
    window_id = get_effective_window_id(admin_override=False) if get_effective_window_id else 1
    course_query = Course.query.filter_by(curriculum_id=curriculum_id)
    if filter_offered_courses:
        course_query = filter_offered_courses(course_query, window_id=window_id)
    else:
        course_query = course_query.filter_by(offered=True)
    courses = course_query.all()  # Only offered courses in active window
    
    years = sorted({c.display_year for c in courses if getattr(c, 'display_year', None)}, key=lambda x: x or '')
    terms = sorted({c.display_term for c in courses if getattr(c, 'display_term', None)}, key=lambda x: x or '')
    
    return jsonify({
        'success': True,
        'years': years,
        'terms': terms
    })

@class_management_bp.route('/api/courses', methods=['GET'])
@login_required
def get_courses_by_filters():
    """Get courses filtered by curriculum, year, and term (AJAX)"""
    if not Course:
        return jsonify({'success': False, 'message': 'Course Management module not available'}), 503
    
    curriculum_id = request.args.get('curriculum_id', type=int)
    year = request.args.get('year', '').strip()
    term = request.args.get('term', '').strip()
    
    window_id = get_effective_window_id(admin_override=False) if get_effective_window_id else 1
    query = Course.query
    if filter_offered_courses:
        query = filter_offered_courses(query, window_id=window_id)
    else:
        query = query.filter_by(offered=True)
    if curriculum_id:
        query = query.filter_by(curriculum_id=curriculum_id)
    
    courses = query.order_by(Course.course_name.asc()).all()
    
    if year:
        courses = [c for c in courses if c.display_year == year]
    if term:
        courses = [c for c in courses if c.display_term == term]

    # Filter only Theory courses
    courses = [c for c in courses if (c.course_type or '').lower() == 'theory']
    
    return jsonify({
        'success': True,
        'courses': [{
            'id': c.id,
            'course_code': c.course_code,
            'course_name': c.course_name,
            'year': c.display_year,
            'term': c.display_term,
            'credit': c.credit,
            'course_type': c.course_type,
            'category': c.category
        } for c in courses]
    })

@class_management_bp.route('/api/academic-sessions', methods=['GET'])
@login_required
def get_academic_sessions():
    """Get academic sessions from CurriculumYearTerm based on curriculum, year, and term (AJAX)"""
    if not CurriculumYearTerm:
        return jsonify({'success': False, 'message': 'Course Management module not available'}), 503
    
    curriculum_id = request.args.get('curriculum_id', type=int)
    year = request.args.get('year', '').strip()
    term = request.args.get('term', '').strip()
    
    if not curriculum_id or not year or not term:
        return jsonify({
            'success': True,
            'academic_sessions': []
        })
    
    # Fetch academic session from CurriculumYearTerm
    config = query_for_window(CurriculumYearTerm).filter_by(
        curriculum_id=curriculum_id,
        year=year,
        term=term
    ).first()
    
    academic_sessions = []
    if config and config.academic_session:
        academic_sessions.append(config.academic_session)
    
    # Also fetch distinct academic sessions from all CurriculumYearTerm for this curriculum/year/term
    # in case there are multiple entries (shouldn't happen due to unique constraint, but for safety)
    all_sessions = query_for_window(CurriculumYearTerm).with_entities(
        CurriculumYearTerm.academic_session
    ).filter_by(
        curriculum_id=curriculum_id,
        year=year,
        term=term
    ).filter(CurriculumYearTerm.academic_session.isnot(None)).distinct().all()
    
    unique_sessions = sorted({s[0] for s in all_sessions if s[0]})
    
    return jsonify({
        'success': True,
        'academic_sessions': unique_sessions
    })

@class_management_bp.route('/add_student/<int:session_id>', methods=['POST'])
@login_required
def add_student(session_id):
    """Add students to a session from Students Management"""
    session = get_or_404_for_window(Session, session_id)
    teacher = _ensure_current_teacher()
    
    # Handle AJAX request (multiple students)
    if request.is_json:
        if not Student:
            return jsonify({'success': False, 'message': 'Students Management module not available'}), 503
        
        data = request.get_json()
        student_ids = data.get('student_ids', [])
        
        if not student_ids:
            return jsonify({'success': False, 'message': 'No students selected!'}), 400
        
        added_count = 0
        skipped_count = 0
        not_registered_count = 0
        
        # Get existing student IDs in this session
        existing_student_ids = {s.student_id for s in ClassStudent.query.filter_by(session_id=session_id).all()}
        
        for student_id in student_ids:
            # Get student from Students Management
            student = Student.query.get(student_id)
            if not student:
                continue
            
            # Check if already in session
            if student.student_id in existing_student_ids:
                skipped_count += 1
                continue
            
            # Check if student is registered for this course (finalized registration only)
            if not getattr(session, 'is_external_course', False) and StudentCourseRegistration and session.course_code and session.academic_session and session.year and session.term:
                registration = StudentCourseRegistration.query.filter_by(
                    student_id=student.id,
                    course_code=session.course_code,
                    academic_session=session.academic_session,
                    year=session.year,
                    term=session.term,
                    status='finalized'
                ).first()
                
                if not registration:
                    not_registered_count += 1
                    current_app.logger.info(f'Student {student.student_id} ({student.name}) not registered for course {session.course_code}, skipping...')
                    continue
            
            class_student = ClassStudent(
                student_id=student.student_id,
                name=student.name,
                session_id=session.id,
                teacher_id=session.teacher_id or teacher.id
            )
            db.session.add(class_student)
            db.session.flush()  # Flush to get class_student.id before carry on
            
            # Carry on assessment marks if enabled in registration
            _carry_on_assessment_marks(class_student, session)
            
            _replicate_student_to_peers(session, class_student)
            existing_student_ids.add(student.student_id)
            added_count += 1
        
        try:
            db.session.commit()
            message = f'Successfully added {added_count} student(s).'
            if skipped_count > 0:
                message += f' Skipped {skipped_count} existing student(s).'
            if not_registered_count > 0:
                message += f' Skipped {not_registered_count} student(s) not registered for this course.'
            return jsonify({'success': True, 'message': message})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': f'Error adding students: {str(e)}'}), 500
    
    # Handle form submission (backward compatibility)
    return_to = (request.form.get('return_to') or '').strip()
    student_id = (request.form.get('student_id') or '').strip()
    name = (request.form.get('name') or '').strip()

    if not student_id or not name:
        flash('Student ID and name are required!', 'error')
        return _student_upload_redirect(session_id, return_to=return_to)

    # Check if already exists
    existing = ClassStudent.query.filter_by(session_id=session_id, student_id=student_id).first()
    if existing:
        flash('Student already exists in this session!', 'error')
        return _student_upload_redirect(session_id, return_to=return_to)

    student = ClassStudent(
        student_id=student_id,
        name=name,
        session_id=session.id,
        teacher_id=session.teacher_id or teacher.id
    )
    db.session.add(student)
    db.session.flush()  # Flush to get student.id before carry on

    # Carry on assessment marks if enabled in registration
    _carry_on_assessment_marks(student, session)

    _replicate_student_to_peers(session, student)
    db.session.commit()
    flash('Student added successfully!', 'success')
    return _student_upload_redirect(session_id, return_to=return_to)

@class_management_bp.route('/edit_student/<int:student_id>', methods=['POST'])
@login_required
def edit_student(student_id):
    """Edit a student"""
    student = ClassStudent.query.get_or_404(student_id)
    session = student.session
    old_identifier = student.student_id
    student.student_id = request.form.get('student_id')
    student.name = request.form.get('name')
    _replicate_student_to_peers(session, student, old_identifier=old_identifier)
    db.session.commit()
    flash('Student updated successfully!', 'success')
    return redirect(url_for('class_management.students_list', session_id=student.session_id))

@class_management_bp.route('/delete_student/<int:student_id>', methods=['POST'])
@login_required
def delete_student(student_id):
    """Delete a student from a session"""
    try:
        student = ClassStudent.query.get_or_404(student_id)
        session_id = student.session_id
        session = student.session
        student_identifier = student.student_id
        
        # Delete student from peer sessions (split courses)
        _delete_student_from_peers(session, student_identifier)
        
        # Delete the student (cascade will handle related attendance records)
        db.session.delete(student)
        db.session.commit()
        flash('Student deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error deleting student {student_id}: {e}', exc_info=True)
        flash(f'Error deleting student: {str(e)}', 'danger')
        # Try to get session_id for redirect even if deletion failed
        try:
            student = ClassStudent.query.get(student_id)
            session_id = student.session_id if student else None
        except:
            session_id = None
    
    if session_id:
        return redirect(url_for('class_management.students_list', session_id=session_id))
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/delete_session/<int:session_id>', methods=['POST'])
@login_required
def delete_session(session_id):
    """Delete a session and all related records using SQLAlchemy ORM (database-agnostic)"""
    try:
        # Check authorization: admin/head can delete any session, regular teachers only their own
        user_roles = set(parse_roles(getattr(current_user, 'role', '')))
        if getattr(current_user, 'active_role', None):
            user_roles = set(parse_roles(current_user.active_role))
        can_delete_all = is_admin(current_user) or 'head' in user_roles or 'dean' in user_roles
        
        session = get_or_404_for_window(Session, session_id)
        
        # If not admin/head, check if this session belongs to the current teacher
        if not can_delete_all:
            current_teacher = _ensure_current_teacher()
            if not current_teacher or session.teacher_id != current_teacher.id:
                flash('You do not have permission to delete this session.', 'danger')
                return redirect(url_for('class_management.index'))
        
        # Import BatchCustomEvent if available
        try:
            from blueprints.academic_calendar.models import BatchCustomEvent
        except ImportError:
            BatchCustomEvent = None
        
        # Delete all related records in correct order (respecting foreign key constraints)
        try:
            # 1. Delete student feedback responses first (before feedback links)
            feedback_link_ids = [link.id for link in StudentFeedbackLink.query.filter_by(session_id=session_id).all()]
            if feedback_link_ids:
                StudentFeedbackResponse.query.filter(
                    StudentFeedbackResponse.feedback_link_id.in_(feedback_link_ids)
                ).delete(synchronize_session=False)
            
            # 2. Delete student feedback links
            StudentFeedbackLink.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 3. Delete batch custom events (if model exists)
            if BatchCustomEvent:
                BatchCustomEvent.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 4. Delete course outline
            CourseOutline.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 5. Delete evaluation submissions
            EvaluationSubmission.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 6. Delete evaluation invites
            EvaluationInvite.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 7. Delete course reviews
            CourseReview.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 8. Delete split course invites (where this session is the inviter)
            ClassSplitInvite.query.filter_by(inviter_session_id=session_id).delete(synchronize_session=False)
            
            # 9. Delete class attendance (cascade will handle if relationship is set up)
            ClassAttendance.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 10. Delete class students (cascade will handle if relationship is set up)
            ClassStudent.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 11. Finally delete the session itself
            db.session.delete(session)
            db.session.commit()
            
            current_app.logger.info(f'Session {session_id} and all related data deleted successfully')
            flash('Session deleted successfully!', 'success')
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error deleting session {session_id}: {e}', exc_info=True)
            flash(f'Error deleting session: {str(e)}', 'danger')
            
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error deleting session {session_id}: {e}', exc_info=True)
        flash(f'Error deleting session: {str(e)}', 'danger')
    
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/course_file/<int:session_id>')
@login_required
def course_file(session_id):
    """Course file management page"""
    session = get_or_404_for_window(Session, session_id)
    
    # Get or create course outline
    course_outline = None
    try:
        course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
        if not course_outline:
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if teacher:
                # Ensure all columns exist before creating
                try:
                    db.create_all()
                except Exception as e:
                    current_app.logger.warning(f"Could not create all tables/columns: {e}")
                
                course_outline = CourseOutline(session_id=session_id, teacher_id=teacher.id)
                db.session.add(course_outline)
                db.session.commit()
    except Exception as e:
        # If table doesn't exist, create it
        current_app.logger.warning(f"CourseOutline table might not exist: {e}")
        try:
            db.create_all()
            # Try again after creating tables
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if teacher:
                course_outline = CourseOutline(session_id=session_id, teacher_id=teacher.id)
                db.session.add(course_outline)
                db.session.commit()
        except Exception as e2:
            current_app.logger.error(f"Error creating CourseOutline: {e2}")
            flash('Course outline feature is not available. Please ensure database migration is complete.', 'warning')
    
    # Get course data from curriculum if available
    course_data = find_course_from_curriculum(session.course_code, session.course_name, session=session)
    
    # Get uploaded files for this session
    from blueprints.class_management.models import CourseFileUpload
    uploaded_files = CourseFileUpload.query.filter_by(session_id=session_id).order_by(CourseFileUpload.created_at.desc()).all()
    
    return render_template('class_management/course_file.html', 
                         session=session, 
                         course_outline=course_outline,
                         course_data=course_data,
                         uploaded_files=uploaded_files)

@class_management_bp.route('/course_file/<int:session_id>/save', methods=['POST'])
@login_required
def save_course_outline(session_id):
    """Save course outline data"""
    try:
        session = get_or_404_for_window(Session, session_id)
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        if not _is_course_outline_authorized(session, teacher):
            if request.is_json:
                return jsonify({'success': False, 'message': 'You are not authorized to edit this course outline.'}), 403
            flash('You are not authorized to edit this course outline.', 'danger')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
        if not course_outline:
            course_outline = CourseOutline(session_id=session_id, teacher_id=teacher.id)
            db.session.add(course_outline)
        
        # Ensure all columns exist (for database migrations)
        try:
            db.create_all()  # This will add missing columns if database supports it
        except Exception as e:
            current_app.logger.warning(f"Could not create all tables/columns: {e}")
        
        # Save all form data as JSON
        data = request.get_json() if request.is_json else request.form.to_dict()
        
        if not data:
            if request.is_json:
                return jsonify({'success': False, 'message': 'No data received.'}), 400
            flash('No data received.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        if 'course_objectives' in data:
            course_outline.course_objectives = json.dumps(data.get('course_objectives', [])) if isinstance(data.get('course_objectives'), list) else data.get('course_objectives')
        if 'course_summary' in data:
            course_outline.course_summary = data.get('course_summary')
        if 'prerequisites' in data:
            course_outline.prerequisites = data.get('prerequisites')
        if 'contact_hours' in data:
            course_outline.contact_hours = data.get('contact_hours')
        if 'cie_marks' in data:
            course_outline.cie_marks = data.get('cie_marks')
        if 'smee_marks' in data:
            course_outline.smee_marks = data.get('smee_marks')
        # Try to set new fields, but handle if columns don't exist yet
        try:
            if 'credit_value' in data:
                course_outline.credit_value = data.get('credit_value')
            if 'course_type' in data:
                course_outline.course_type = data.get('course_type')
            if 'level_term_section' in data:
                course_outline.level_term_section = data.get('level_term_section')
            if 'clo_data' in data:
                course_outline.clo_data = json.dumps(data.get('clo_data', []))
            if 'plo_mapping' in data:
                course_outline.plo_mapping = json.dumps(data.get('plo_mapping', {}))
        except AttributeError:
            current_app.logger.warning("Some new course_outline columns don't exist yet. Please run migration.")
        
        if 'course_content_summary' in data:
            # Store as JSON format (accept dict from AI apply or JSON string from form)
            content_summary = data.get('course_content_summary')
            if content_summary:
                content_data = None
                if isinstance(content_summary, dict):
                    content_data = content_summary
                elif isinstance(content_summary, str):
                    try:
                        content_data = json.loads(content_summary)
                    except (json.JSONDecodeError, TypeError):
                        course_outline.course_content_summary = content_summary
                if isinstance(content_data, dict):
                    classes_a = data.get('classes_a')
                    classes_b = data.get('classes_b')
                    if isinstance(classes_a, list) and classes_a and content_data.get('sectionA'):
                        for idx, item in enumerate(content_data['sectionA']):
                            if idx < len(classes_a) and classes_a[idx] not in (None, ''):
                                try:
                                    item['num_classes'] = max(1, int(classes_a[idx]) if isinstance(classes_a[idx], int) else int(classes_a[idx], 10))
                                except (TypeError, ValueError):
                                    pass
                    if isinstance(classes_b, list) and classes_b and content_data.get('sectionB'):
                        for idx, item in enumerate(content_data['sectionB']):
                            if idx < len(classes_b) and classes_b[idx] not in (None, ''):
                                try:
                                    item['num_classes'] = max(1, int(classes_b[idx]) if isinstance(classes_b[idx], int) else int(classes_b[idx], 10))
                                except (TypeError, ValueError):
                                    pass
                    course_outline.course_content_summary = json.dumps(content_data)
                elif content_data is not None and not isinstance(content_data, dict):
                    course_outline.course_content_summary = str(content_summary)
            else:
                course_outline.course_content_summary = None
        # Handle Classes: save to course_content_classes (separate column)
        # JSON save (fetch) sends classes_a, classes_b; traditional form sends course_content_*_classes[]
        if request.method == 'POST':
            if request.is_json:
                classes_a = data.get('classes_a', [])
                classes_b = data.get('classes_b', [])
                if isinstance(classes_a, list) and isinstance(classes_b, list):
                    try:
                        classes_data = {
                            'section_a': [max(1, int(v) if isinstance(v, (int, float)) else int(v or '1', 10)) for v in classes_a],
                            'section_b': [max(1, int(v) if isinstance(v, (int, float)) else int(v or '1', 10)) for v in classes_b]
                        }
                        course_outline.course_content_classes = json.dumps(classes_data)
                        current_app.logger.info("Saved Classes (JSON): A=%s, B=%s", len(classes_a), len(classes_b))
                    except (TypeError, ValueError) as e:
                        current_app.logger.warning("Could not parse classes_a/classes_b: %s", e)
            else:
                classes_a_raw = request.form.getlist('course_content_a_classes[]')
                classes_b_raw = request.form.getlist('course_content_b_classes[]')
                classes_data = {
                    'section_a': [max(1, int(v or '1')) for v in classes_a_raw],
                    'section_b': [max(1, int(v or '1')) for v in classes_b_raw]
                }
                course_outline.course_content_classes = json.dumps(classes_data)
                current_app.logger.info("Saved Classes (form): A=%s, B=%s", len(classes_a_raw), len(classes_b_raw))
        if 'clo_plo_mapping' in data:
            clo_plo = data.get('clo_plo_mapping')
            course_outline.clo_plo_mapping = json.dumps(clo_plo) if isinstance(clo_plo, (dict, list)) else clo_plo
        if 'evaluation_policy' in data:
            course_outline.evaluation_policy = json.dumps(data.get('evaluation_policy', {})) if isinstance(data.get('evaluation_policy'), dict) else data.get('evaluation_policy')
        if 'cie_breakdown' in data:
            cie_breakdown = data.get('cie_breakdown', [])
            if isinstance(cie_breakdown, (list, dict)):
                course_outline.cie_breakdown = json.dumps(cie_breakdown)
            else:
                course_outline.cie_breakdown = cie_breakdown
        if 'smee_breakdown' in data:
            smee_breakdown = data.get('smee_breakdown', [])
            if isinstance(smee_breakdown, (list, dict)):
                course_outline.smee_breakdown = json.dumps(smee_breakdown)
            else:
                course_outline.smee_breakdown = smee_breakdown
        
        if 'lesson_plan' in data:
            course_outline.lesson_plan = json.dumps(data.get('lesson_plan', [])) if isinstance(data.get('lesson_plan'), list) else data.get('lesson_plan')
        if 'assessment_strategy' in data:
            course_outline.assessment_strategy = json.dumps(data.get('assessment_strategy', {})) if isinstance(data.get('assessment_strategy'), dict) else data.get('assessment_strategy')
        if 'assessment_techniques' in data:
            course_outline.assessment_techniques = json.dumps(data.get('assessment_techniques', [])) if isinstance(data.get('assessment_techniques'), list) else data.get('assessment_techniques')
        if 'rubrics' in data:
            course_outline.rubrics = json.dumps(data.get('rubrics', [])) if isinstance(data.get('rubrics'), list) else data.get('rubrics')
        if 'grading_policy' in data:
            course_outline.grading_policy = json.dumps(data.get('grading_policy', [])) if isinstance(data.get('grading_policy'), list) else data.get('grading_policy')
        if 'textbooks' in data:
            course_outline.textbooks = json.dumps(data.get('textbooks', [])) if isinstance(data.get('textbooks'), list) else data.get('textbooks')
        if 'reference_books' in data:
            course_outline.reference_books = json.dumps(data.get('reference_books', [])) if isinstance(data.get('reference_books'), list) else data.get('reference_books')
        if 'other_resources' in data:
            course_outline.other_resources = json.dumps(data.get('other_resources', [])) if isinstance(data.get('other_resources'), list) else data.get('other_resources')
        try:
            if 'course_file_components' in data:
                course_outline.course_file_components = json.dumps(data.get('course_file_components', [])) if isinstance(data.get('course_file_components'), list) else data.get('course_file_components')
        except AttributeError:
            current_app.logger.warning("course_file_components column doesn't exist yet.")
        
        if 'make_up_procedures' in data:
            course_outline.make_up_procedures = data.get('make_up_procedures')
        if 'other_issues' in data:
            course_outline.other_issues = json.dumps(data.get('other_issues', {})) if isinstance(data.get('other_issues'), dict) else data.get('other_issues')
        
        try:
            # Always process student_access_enabled
            if 'student_access_enabled' in data:
                # Handle both boolean and string values
                student_access = data.get('student_access_enabled', False)
                if isinstance(student_access, str):
                    student_access = student_access.lower() in ('true', '1', 'on', 'yes')
                course_outline.student_access_enabled = bool(student_access)
                current_app.logger.info(f"Setting student_access_enabled to {course_outline.student_access_enabled} for session {session_id} (received value: {data.get('student_access_enabled')}, type: {type(data.get('student_access_enabled'))})")
            else:
                # If not provided, keep existing value (don't reset to False)
                current_app.logger.info(f"student_access_enabled not in data for session {session_id}, keeping existing value: {course_outline.student_access_enabled}")
        except AttributeError:
            current_app.logger.warning("student_access_enabled column doesn't exist yet.")
        except Exception as e:
            current_app.logger.error(f"Error setting student_access_enabled: {e}", exc_info=True)
        
        try:
            db.session.commit()
            current_app.logger.info(f"Course outline saved successfully for session {session_id}")
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error committing course outline for session {session_id}: {e}", exc_info=True)
            if request.is_json:
                return jsonify({'success': False, 'message': f'Database error: {str(e)}'}), 500
            flash(f'Error saving course outline: {str(e)}', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving course outline for session {session_id}: {e}", exc_info=True)
        if request.is_json:
            return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500
        flash(f'Error saving course outline: {str(e)}', 'error')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    if request.is_json:
        return jsonify({'success': True, 'message': 'Course outline saved successfully!'})
    flash('Course outline saved successfully!', 'success')
    return redirect(url_for('class_management.course_file', session_id=session_id))

@class_management_bp.route('/course_file/<int:session_id>/upload', methods=['POST'])
@login_required
def upload_course_file(session_id):
    """Upload course file"""
    try:
        import os
        from werkzeug.utils import secure_filename
        from blueprints.class_management.models import CourseFileUpload
        
        session = get_or_404_for_window(Session, session_id)
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        if not teacher:
            flash('Teacher not found.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Check authorization
        if teacher.id != session.teacher_id:
            flash('You are not authorized to upload files for this course.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Get form data
        file_name = request.form.get('file_name', '').strip()
        file_category = request.form.get('file_category', 'other')
        description = request.form.get('description', '').strip()
        
        if not file_name:
            flash('File name is required.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Get uploaded file
        if 'file' not in request.files:
            flash('No file selected.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        file = request.files['file']
        if file.filename == '':
            flash('No file selected.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Validate file extension
        allowed_extensions = {'.pdf', '.doc', '.docx', '.ppt', '.pptx', '.xls', '.xlsx',
                            '.epub', '.djvu', '.djv', '.mobi',
                            '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in allowed_extensions:
            flash('File type not allowed. Supported formats: PDF, DOC, DOCX, PPT, PPTX, XLS, XLSX, EPUB, DJVU, MOBI, images', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Create upload directory if it doesn't exist
        upload_dir = os.path.join(UPLOAD_FOLDER, 'course_files', str(session_id))
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate secure filename
        secure_name = secure_filename(file.filename)
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{secure_name}"
        file_path = os.path.join(upload_dir, filename)
        
        # Save file
        file.save(file_path)
        file_size = os.path.getsize(file_path)
        
        # Get file type (MIME type or extension)
        file_type = file_ext[1:] if file_ext else 'unknown'
        
        # Create CourseFileUpload record
        uploaded_file = CourseFileUpload(
            session_id=session_id,
            teacher_id=teacher.id,
            file_name=file_name,
            file_path=file_path,
            file_size=file_size,
            file_type=file_type,
            file_category=file_category or None,
            description=description,
            student_access_enabled=True
        )
        
        db.session.add(uploaded_file)
        db.session.flush()
        try:
            from utils.ai.rag_context import ensure_upload_extracted
            ensure_upload_extracted(uploaded_file)
        except Exception as extract_error:
            current_app.logger.warning(f'Could not extract text from upload {file_path}: {extract_error}')
        db.session.commit()
        current_app.logger.info(f"Course file uploaded successfully: {file_name} for session {session_id}")
        try:
            session_obj = get_for_window(Session, session_id)
            course_label = (session_obj.course_name or session_obj.course_code or 'Course') if session_obj else 'Course'
            title = f'New file shared: {file_name} ({course_label})'
            link_url = url_for('class_management.student_course_files')
            _notify_students_in_session(session_id, 'file_shared', title, link_url)
        except Exception as notif_e:
            current_app.logger.warning(f"Student notification (file shared): {notif_e}")
        flash(f'File "{file_name}" uploaded successfully!', 'success')
        return redirect(url_for('class_management.course_file', session_id=session_id))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error uploading course file for session {session_id}: {e}", exc_info=True)
        flash(f'Error uploading file: {str(e)}', 'error')
        return redirect(url_for('class_management.course_file', session_id=session_id))

@class_management_bp.route('/course_file/<int:file_id>/download')
@login_required
def download_course_file(file_id):
    """Download course file"""
    try:
        import os
        from flask import send_file
        from blueprints.class_management.models import CourseFileUpload
        
        uploaded_file = CourseFileUpload.query.get_or_404(file_id)
        session = get_or_404_for_window(Session, uploaded_file.session_id)
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        # Check authorization
        if not teacher or teacher.id != session.teacher_id:
            flash('You are not authorized to download this file.', 'error')
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        
        # Check if file exists
        if not os.path.exists(uploaded_file.file_path):
            flash('File not found.', 'error')
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        
        return send_file(
            uploaded_file.file_path,
            as_attachment=True,
            download_name=uploaded_file.file_name
        )
    except Exception as e:
        current_app.logger.error(f"Error downloading course file {file_id}: {e}", exc_info=True)
        flash(f'Error downloading file: {str(e)}', 'error')
        if 'uploaded_file' in locals():
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/course_file/<int:file_id>/delete', methods=['POST'])
@login_required
def delete_course_file(file_id):
    """Delete course file"""
    try:
        import os
        from blueprints.class_management.models import CourseFileUpload
        
        uploaded_file = CourseFileUpload.query.get_or_404(file_id)
        session = get_or_404_for_window(Session, uploaded_file.session_id)
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        # Check authorization
        if not teacher or teacher.id != session.teacher_id:
            flash('You are not authorized to delete this file.', 'error')
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        
        # Delete file from filesystem
        if os.path.exists(uploaded_file.file_path):
            try:
                os.remove(uploaded_file.file_path)
            except Exception as e:
                current_app.logger.warning(f"Could not delete file from filesystem: {e}")
        
        # Delete record from database
        db.session.delete(uploaded_file)
        db.session.commit()
        
        flash(f'File "{uploaded_file.file_name}" deleted successfully.', 'success')
        return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting course file {file_id}: {e}", exc_info=True)
        flash(f'Error deleting file: {str(e)}', 'error')
        if 'uploaded_file' in locals():
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        return redirect(url_for('class_management.index'))

def _is_course_outline_authorized(session, teacher):
    if not teacher:
        return False
    if teacher.id == session.teacher_id:
        return True
    if session.split_group_id:
        related_sessions = query_for_window(Session).filter_by(split_group_id=session.split_group_id).all()
        return any(related_session.teacher_id == teacher.id for related_session in related_sessions)
    return False


def _ai_outline_calendar_and_course(session):
    from datetime import date as date_cls
    from blueprints.academic_calendar.models import AcademicCalendarEvent

    year_start = date_cls(date_cls.today().year, 1, 1)
    year_end = date_cls(date_cls.today().year + 1, 12, 31)
    calendar_events = query_for_window(AcademicCalendarEvent).filter(
        AcademicCalendarEvent.event_date >= year_start,
        AcademicCalendarEvent.event_date <= year_end,
    ).order_by(AcademicCalendarEvent.event_date.asc()).all()
    course_data = find_course_from_curriculum(session.course_code, session.course_name, session=session)
    curriculum = course_data.curriculum if course_data and getattr(course_data, 'curriculum', None) else None
    return calendar_events, course_data, curriculum


@class_management_bp.route('/course_file/<int:session_id>/outline/generate-full-ai', methods=['POST'])
@login_required
def generate_full_outline_ai(session_id):
    """Generate Course Outline using AI (sync or async job)."""
    session = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not _is_course_outline_authorized(session, teacher):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        from utils.ai.client import AIClientError, get_active_provider_setting
        from utils.ai.job_service import normalize_parts
        from utils.ai.outline_service import generate_full_outline_for_session, start_async_outline_job
        from blueprints.class_management.models import CourseFileUpload

        get_active_provider_setting()

        data = request.get_json(silent=True) or {}
        parts = normalize_parts(data.get('parts'))
        use_async = data.get('async', True)

        calendar_events, course_data, curriculum = _ai_outline_calendar_and_course(session)
        teacher_name = teacher.name if teacher else current_user.full_name
        teacher_id = teacher.id if teacher else None

        from utils.ai.outline_guidelines import normalize_generation_options
        generation_options = normalize_generation_options(
            data.get('generation_options'), session=session, course_data=course_data,
        )

        if use_async:
            response = start_async_outline_job(
                session,
                user_id=current_user.id,
                teacher_id=teacher_id,
                teacher_name=teacher_name,
                course_data=course_data,
                curriculum=curriculum,
                calendar_events=calendar_events,
                Course=Course,
                CurriculumYearTerm=CurriculumYearTerm,
                query_for_window=query_for_window,
                parts=parts,
                CourseSessionAssignment=CourseSessionAssignment,
                CourseFileUpload=CourseFileUpload,
                generation_options=generation_options,
            )
            return jsonify(response)

        result = generate_full_outline_for_session(
            session,
            teacher_name=teacher_name,
            course_data=course_data,
            curriculum=curriculum,
            calendar_events=calendar_events,
            Course=Course,
            CurriculumYearTerm=CurriculumYearTerm,
            query_for_window=query_for_window,
            user_id=current_user.id,
            parts=parts,
            use_parts=True,
            CourseSessionAssignment=CourseSessionAssignment,
            CourseFileUpload=CourseFileUpload,
            generation_options=generation_options,
        )
        return jsonify({
            'success': True,
            'message': 'Course outline generated successfully. Review and save.',
            'payload': result['payload'],
            'context_summary': result['context_summary'],
            'parts': parts,
        })
    except AIClientError as exc:
        return jsonify({'success': False, 'message': str(exc)}), 400
    except Exception as exc:
        current_app.logger.error(f'AI full outline generation failed for session {session_id}: {exc}', exc_info=True)
        message = str(exc)
        status = 400 if 'AI' in message or 'Curriculum' in message or 'Semester' in message else 500
        return jsonify({'success': False, 'message': message}), status


@class_management_bp.route('/course_file/<int:session_id>/outline/generate-ai-job/<int:job_id>/tick', methods=['POST'])
@login_required
def tick_outline_ai_job(session_id, job_id):
    """Process one part of an async outline generation job (poll from browser)."""
    session = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not _is_course_outline_authorized(session, teacher):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        from utils.ai.models import AIOutlineGenerationJob
        from utils.ai.outline_service import tick_async_outline_job
        from blueprints.class_management.models import CourseFileUpload

        job = AIOutlineGenerationJob.query.get_or_404(job_id)
        if job.session_id != session_id or job.user_id != current_user.id:
            return jsonify({'success': False, 'message': 'Invalid generation job.'}), 403

        calendar_events, course_data, curriculum = _ai_outline_calendar_and_course(session)
        response = tick_async_outline_job(
            job,
            session,
            teacher_name=teacher.name if teacher else current_user.full_name,
            course_data=course_data,
            curriculum=curriculum,
            calendar_events=calendar_events,
            Course=Course,
            CurriculumYearTerm=CurriculumYearTerm,
            query_for_window=query_for_window,
            CourseSessionAssignment=CourseSessionAssignment,
            CourseFileUpload=CourseFileUpload,
        )
        status = 200 if response.get('success') else 400
        return jsonify(response), status
    except Exception as exc:
        current_app.logger.error(f'AI job tick failed for session {session_id} job {job_id}: {exc}', exc_info=True)
        return jsonify({'success': False, 'message': str(exc)}), 500


def _can_batch_generate_ai_outline():
    from role_utils import is_admin, parse_roles
    if is_admin(current_user):
        return True
    roles = parse_roles(getattr(current_user, 'role', '') or '')
    return 'head' in roles or 'dean' in roles


@class_management_bp.route('/outline/batch-generate', methods=['POST'])
@login_required
def batch_generate_outlines():
    """Start async outline jobs for all sessions in a semester (admin/head)."""
    if not _can_batch_generate_ai_outline():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        from utils.ai.batch_service import batch_job_to_response, create_batch_outline_jobs
        from utils.ai.job_service import normalize_parts
        from blueprints.class_management.models import CourseFileUpload

        data = request.get_json(silent=True) or {}
        academic_session = (data.get('academic_session') or '').strip()
        year = (data.get('year') or '').strip()
        term = (data.get('term') or '').strip()
        batch = (data.get('batch') or '').strip() or None
        parts = normalize_parts(data.get('parts'))

        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'academic_session, year, and term are required.'}), 400

        from datetime import date as date_cls
        from blueprints.academic_calendar.models import AcademicCalendarEvent
        year_start = date_cls(date_cls.today().year, 1, 1)
        year_end = date_cls(date_cls.today().year + 1, 12, 31)
        calendar_events = query_for_window(AcademicCalendarEvent).filter(
            AcademicCalendarEvent.event_date >= year_start,
            AcademicCalendarEvent.event_date <= year_end,
        ).order_by(AcademicCalendarEvent.event_date.asc()).all()

        batch_job, items = create_batch_outline_jobs(
            user_id=current_user.id,
            academic_session=academic_session,
            year=year,
            term=term,
            batch=batch,
            parts=parts,
            query_for_window=query_for_window,
            Session=Session,
            CourseSessionAssignment=CourseSessionAssignment,
            Course=Course,
            CurriculumYearTerm=CurriculumYearTerm,
            calendar_events=calendar_events,
            find_course_from_curriculum=find_course_from_curriculum,
            CourseFileUpload=CourseFileUpload,
        )
        response = batch_job_to_response(batch_job)
        response['message'] = f'Started outline generation for {len(items)} session(s).'
        return jsonify(response)
    except Exception as exc:
        current_app.logger.error(f'Batch outline generation failed: {exc}', exc_info=True)
        return jsonify({'success': False, 'message': str(exc)}), 400


@class_management_bp.route('/outline/batch-job/<int:batch_job_id>', methods=['GET'])
@login_required
def batch_outline_job_status(batch_job_id):
    if not _can_batch_generate_ai_outline():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    from utils.ai.batch_service import batch_job_to_response
    from utils.ai.models import AIOutlineBatchJob

    batch_job = AIOutlineBatchJob.query.get_or_404(batch_job_id)
    if batch_job.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Invalid batch job.'}), 403
    return jsonify(batch_job_to_response(batch_job))


@class_management_bp.route('/outline/batch-job/<int:batch_job_id>/tick-next', methods=['POST'])
@login_required
def batch_outline_job_tick_next(batch_job_id):
    """Tick the next pending child session job (one part per call)."""
    if not _can_batch_generate_ai_outline():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        from utils.ai.batch_service import batch_job_to_response, refresh_batch_job_status
        from utils.ai.models import AIOutlineBatchJob, AIOutlineGenerationJob
        from utils.ai.outline_service import tick_async_outline_job
        from blueprints.class_management.models import CourseFileUpload

        batch_job = AIOutlineBatchJob.query.get_or_404(batch_job_id)
        if batch_job.user_id != current_user.id:
            return jsonify({'success': False, 'message': 'Invalid batch job.'}), 403

        for item in batch_job.items_list():
            job = AIOutlineGenerationJob.query.get(item.get('job_id'))
            if not job or job.status in (
                AIOutlineGenerationJob.STATUS_COMPLETED,
                AIOutlineGenerationJob.STATUS_FAILED,
            ):
                continue
            session = Session.query.get(item.get('session_id'))
            if not session:
                continue
            calendar_events, course_data, curriculum = _ai_outline_calendar_and_course(session)
            tick_async_outline_job(
                job,
                session,
                teacher_name=current_user.full_name,
                course_data=course_data,
                curriculum=curriculum,
                calendar_events=calendar_events,
                Course=Course,
                CurriculumYearTerm=CurriculumYearTerm,
                query_for_window=query_for_window,
                CourseSessionAssignment=CourseSessionAssignment,
                CourseFileUpload=CourseFileUpload,
            )
            break

        refresh_batch_job_status(batch_job)
        return jsonify(batch_job_to_response(batch_job))
    except Exception as exc:
        current_app.logger.error(f'Batch tick failed for {batch_job_id}: {exc}', exc_info=True)
        return jsonify({'success': False, 'message': str(exc)}), 500


@class_management_bp.route('/course_file/<int:session_id>/outline/generate-ai', methods=['POST'])
@login_required
def generate_weekly_plan_ai(session_id):
    """Generate weekly plan using AI based on Course Content (section 14), Credit Value, and Academic Calendar"""
    from datetime import timedelta
    
    session = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    
    # Check if teacher is authorized (either main teacher or part of split group)
    is_authorized = False
    if teacher:
        if teacher.id == session.teacher_id:
            is_authorized = True
        elif session.split_group_id:
            related_sessions = query_for_window(Session).filter_by(split_group_id=session.split_group_id).all()
            for related_session in related_sessions:
                if related_session.teacher_id == teacher.id:
                    is_authorized = True
                    break
    
    if not is_authorized:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    data = request.get_json()
    course_content_summary = data.get('course_content_summary', '')  # JSON string from section 14
    credit_value = data.get('credit_value', '')  # Credit value from form
    part = data.get('part', None)  # 'A' or 'B' for split courses, None for full courses
    
    try:
        # Import Academic Calendar model
        try:
            from blueprints.academic_calendar.models import AcademicCalendarEvent
        except ImportError:
            AcademicCalendarEvent = None
            return jsonify({'success': False, 'message': 'Academic Calendar module not available'}), 503
        
        # Get credit value - try from data, then course outline, then course data
        credit = None
        if credit_value:
            try:
                credit = float(credit_value)
            except:
                pass
        
        if not credit:
            # Try to get from course outline
            course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
            if course_outline and course_outline.credit_value:
                try:
                    credit = float(course_outline.credit_value)
                except:
                    pass
        
        if not credit:
            # Try to get from course data
            course_data = find_course_from_curriculum(session.course_code, session.course_name, session=session)
            if course_data and course_data.credit:
                try:
                    credit = float(course_data.credit)
                except:
                    pass
        
        if not credit:
            return jsonify({'success': False, 'message': 'Credit value is required. Please fill in the Credit Value field.'}), 400
        
        # Parse Course Content (section 14) - Only use selected topics
        # For split courses, filter by part (A or B)
        course_contents = []
        if course_content_summary:
            try:
                content_data = json.loads(course_content_summary) if isinstance(course_content_summary, str) else course_content_summary
                if isinstance(content_data, dict):
                    # For split courses, use only the specified part (A or B)
                    # For full courses, combine both sections
                    if part == 'A':
                        section_a = content_data.get('sectionA', [])
                        all_contents = section_a
                    elif part == 'B':
                        section_b = content_data.get('sectionB', [])
                        all_contents = section_b
                    else:
                        # Full course: combine both sections
                        section_a = content_data.get('sectionA', [])
                        section_b = content_data.get('sectionB', [])
                        all_contents = section_a + section_b
                    
                    # Filter to only include items where selected is True (or undefined for backward compatibility)
                    course_contents = [
                        item for item in all_contents 
                        if item.get('selected', True)  # Default to True for backward compatibility
                    ]
                elif isinstance(content_data, list):
                    # Filter to only include selected topics
                    course_contents = [
                        item for item in content_data 
                        if item.get('selected', True)  # Default to True for backward compatibility
                    ]
            except Exception as e:
                current_app.logger.warning(f"Error parsing course content: {e}")
        
        # If no course content from section 14, try to get from curriculum
        if not course_contents:
            if course_data := find_course_from_curriculum(session.course_code, session.course_name, session=session):
                # For split courses, use only the specified part (A or B)
                # For full courses, combine both sections
                if part == 'A':
                    if course_data.content_section_a:
                        try:
                            content_a = json.loads(course_data.content_section_a) if isinstance(course_data.content_section_a, str) else course_data.content_section_a
                            if isinstance(content_a, list):
                                course_contents.extend(content_a)
                        except:
                            pass
                elif part == 'B':
                    if course_data.content_section_b:
                        try:
                            content_b = json.loads(course_data.content_section_b) if isinstance(course_data.content_section_b, str) else course_data.content_section_b
                            if isinstance(content_b, list):
                                course_contents.extend(content_b)
                        except:
                            pass
                else:
                    # Full course: combine both sections
                    if course_data.content_section_a:
                        try:
                            content_a = json.loads(course_data.content_section_a) if isinstance(course_data.content_section_a, str) else course_data.content_section_a
                            if isinstance(content_a, list):
                                course_contents.extend(content_a)
                        except:
                            pass
                    if course_data.content_section_b:
                        try:
                            content_b = json.loads(course_data.content_section_b) if isinstance(course_data.content_section_b, str) else course_data.content_section_b
                            if isinstance(content_b, list):
                                course_contents.extend(content_b)
                        except:
                            pass
        
        # Get CLOs
        clos = []
        course_data = find_course_from_curriculum(session.course_code, session.course_name, session=session)
        if course_data:
            clos = course_data.get_clos_list()
            
        # Normalize session year and term for matching
        session_year = str(session.year).strip().lower() if session.year else ''
        session_term = str(session.term).strip().lower() if session.term else ''
        
        # Map year/term to common formats for matching
        year_mapping = {
            '1': 'first', 'first': 'first', '1st': 'first',
            '2': 'second', 'second': 'second', '2nd': 'second',
            '3': 'third', 'third': 'third', '3rd': 'third',
            '4': 'fourth', 'fourth': 'fourth', '4th': 'fourth'
        }
        term_mapping = {
            '1': 'first', 'first': 'first', '1st': 'first',
            '2': 'second', 'second': 'second', '2nd': 'second'
        }
        
        normalized_year = year_mapping.get(session_year, session_year)
        normalized_term = term_mapping.get(session_term, session_term)
        
        # Get Academic Calendar events
        holidays = set()
        semester_start_date = None
        semester_end_date = None
        
        # Get current year and next year for broader search
        current_year = datetime.now().year
        year_start = date(current_year, 1, 1)
        year_end = date(current_year + 1, 12, 31)
        
        calendar_events = query_for_window(AcademicCalendarEvent).filter(
            AcademicCalendarEvent.event_date >= year_start,
            AcademicCalendarEvent.event_date <= year_end
        ).order_by(AcademicCalendarEvent.event_date.asc()).all()
        
        # Collect all holidays
        for event in calendar_events:
            if event.event_type == 'holiday':
                # Add all dates in the range if end_date exists
                if event.end_date and event.end_date > event.event_date:
                    current_date = event.event_date
                    while current_date <= event.end_date:
                        holidays.add(current_date)
                        current_date += timedelta(days=1)
        else:
                    holidays.add(event.event_date)
        
        # Add recurring Friday and Saturday holidays
        current_date = year_start
        while current_date <= year_end:
            if current_date.weekday() == 4:  # Friday
                holidays.add(current_date)
            elif current_date.weekday() == 5:  # Saturday
                holidays.add(current_date)
            current_date += timedelta(days=1)
        
        # Find semester start and end dates
        # First, check which session, year, and term this subject belongs to
        session_academic_session = session.academic_session or ''
        session_year = normalized_year
        session_term = normalized_term
        
        current_app.logger.info(f"Looking for semester dates for Session: {session_academic_session}, Year: {session_year}, Term: {session_term}")
        
        semester_start_events = []
        semester_end_events = []
        
        for event in calendar_events:
            if event.event_type == 'semester_start':
                semester_start_events.append(event)
            elif event.event_type == 'semester_end':
                semester_end_events.append(event)
        
        # Try to find matching semester based on session, year, term, and academic_session
        # Priority: 1. Match by academic_session + year + term in title/description
        #           2. Match by year + term in title/description
        #           3. Use most appropriate date based on session academic_session or current date
        if semester_start_events:
            matched_start = None
            
            # Priority 1: Match by academic_session + year + term
            if session_academic_session:
                for event in semester_start_events:
                    title_lower = (event.title or '').lower()
                    description_lower = (event.description or '').lower()
                    event_text = title_lower + ' ' + description_lower
                    
                    # Check if academic_session, year, and term all appear in the event
                    if (session_academic_session.lower() in event_text and 
                        normalized_year in event_text and 
                        normalized_term in event_text):
                        matched_start = event
                        current_app.logger.info(f"Matched semester start by academic_session+year+term: {event.title} on {event.event_date}")
                        break
            
            # Priority 2: Match by year + term in title/description
            if not matched_start:
                for event in semester_start_events:
                    title_lower = (event.title or '').lower()
                    description_lower = (event.description or '').lower()
                    event_text = title_lower + ' ' + description_lower
                    
                    if normalized_year in event_text and normalized_term in event_text:
                        matched_start = event
                        current_app.logger.info(f"Matched semester start by year+term: {event.title} on {event.event_date}")
                        break
            
            # Priority 3: Use the most recent past or upcoming start date
            if not matched_start:
                today = date.today()
                upcoming_starts = [e for e in semester_start_events if e.event_date >= today]
                if upcoming_starts:
                    matched_start = min(upcoming_starts, key=lambda x: x.event_date)
                    current_app.logger.info(f"Using upcoming semester start: {matched_start.title} on {matched_start.event_date}")
                else:
                    # Use most recent past
                    matched_start = max(semester_start_events, key=lambda x: x.event_date)
                    current_app.logger.info(f"Using most recent semester start: {matched_start.title} on {matched_start.event_date}")
            
            if matched_start:
                semester_start_date = matched_start.event_date
        
        if semester_end_events:
            matched_end = None
            
            # Priority 1: Match by academic_session + year + term
            if session_academic_session:
                for event in semester_end_events:
                    title_lower = (event.title or '').lower()
                    description_lower = (event.description or '').lower()
                    event_text = title_lower + ' ' + description_lower
                    
                    # Check if academic_session, year, and term all appear in the event
                    if (session_academic_session.lower() in event_text and 
                        normalized_year in event_text and 
                        normalized_term in event_text):
                        matched_end = event
                        current_app.logger.info(f"Matched semester end by academic_session+year+term: {event.title} on {event.event_date}")
                        break
            
            # Priority 2: Match by year + term in title/description
            if not matched_end:
                for event in semester_end_events:
                    title_lower = (event.title or '').lower()
                    description_lower = (event.description or '').lower()
                    event_text = title_lower + ' ' + description_lower
                    
                    if normalized_year in event_text and normalized_term in event_text:
                        matched_end = event
                        current_app.logger.info(f"Matched semester end by year+term: {event.title} on {event.event_date}")
                        break
            
            # Priority 3: Use the end date that comes after the start date
            if not matched_end:
                if semester_start_date:
                    future_ends = [e for e in semester_end_events if e.event_date > semester_start_date]
                    if future_ends:
                        matched_end = min(future_ends, key=lambda x: x.event_date)
                        current_app.logger.info(f"Using future semester end after start: {matched_end.title} on {matched_end.event_date}")
                else:
                    # Use most recent past or upcoming
                    today = date.today()
                    upcoming_ends = [e for e in semester_end_events if e.event_date >= today]
                    if upcoming_ends:
                        matched_end = min(upcoming_ends, key=lambda x: x.event_date)
                        current_app.logger.info(f"Using upcoming semester end: {matched_end.title} on {matched_end.event_date}")
                    else:
                        matched_end = max(semester_end_events, key=lambda x: x.event_date)
                        current_app.logger.info(f"Using most recent semester end: {matched_end.title} on {matched_end.event_date}")
            
            if matched_end:
                semester_end_date = matched_end.event_date
        
        # Validate dates
        if not semester_start_date:
            return jsonify({
                'success': False,
                'message': 'Semester Start Date not found in Academic Calendar. Please add a "Semester Start Date" event.'
            }), 400
        
        if not semester_end_date:
            return jsonify({
                'success': False,
                'message': 'Semester End Date not found in Academic Calendar. Please add a "Semester End Date" event.'
            }), 400
        
        if semester_end_date <= semester_start_date:
            return jsonify({
                'success': False,
                'message': 'Semester End Date must be after Semester Start Date.'
            }), 400
        
        # ========================================
        # NEW LOGIC: Generate Plan with Week Ranges
        # ========================================
        
        # Step 1: Calculate classes_per_week based on credit
        # Credit = Classes per week (e.g., 3 credit = 3 classes/week)
        if part in ['A', 'B']:
            # Split course: each part gets half the classes
            classes_per_week = int(credit) // 2
            if classes_per_week == 0:
                classes_per_week = 1  # Minimum 1 class per week
        else:
            # Full course: all classes
            classes_per_week = int(credit)
        
        # Also load Classes data from course_content_classes (separate column)
        # This is where the teacher's manual Classes input is stored
        classes_data = {'section_a': [], 'section_b': []}
        course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
        if course_outline and getattr(course_outline, 'course_content_classes', None):
            try:
                classes_data = json.loads(course_outline.course_content_classes)
                if not isinstance(classes_data, dict):
                    classes_data = {'section_a': [], 'section_b': []}
                classes_data.setdefault('section_a', [])
                classes_data.setdefault('section_b', [])
            except (TypeError, ValueError):
                pass
        
        # Step 2: Build topic_slots with classes count from Section 13
        # Each topic gets its num_classes from the saved classes_data
        topic_slots = []
        total_topic_classes = 0
        
        # Process based on part (A, B, or full course)
        if part == 'A':
            section_items = course_contents  # Already filtered to sectionA
            section_classes = classes_data.get('section_a', [])
        elif part == 'B':
            section_items = course_contents  # Already filtered to sectionB
            section_classes = classes_data.get('section_b', [])
        else:
            # Full course - need to process both sections
            # Re-parse to get section info if available
            section_a_items = []
            section_b_items = []
            if course_content_summary:
                try:
                    content_data = json.loads(course_content_summary) if isinstance(course_content_summary, str) else course_content_summary
                    if isinstance(content_data, dict):
                        section_a_items = [item for item in content_data.get('sectionA', []) if item.get('selected', True)]
                        section_b_items = [item for item in content_data.get('sectionB', []) if item.get('selected', True)]
                except:
                    pass
            section_items = section_a_items + section_b_items
            section_classes = classes_data.get('section_a', []) + classes_data.get('section_b', [])
        
        # Build topic slots - each topic with its classes count
        for idx, content_item in enumerate(course_contents if part else section_items):
            if isinstance(content_item, dict) and content_item.get('selected', True):
                content_text = (content_item.get('content') or '').strip()
                if not content_text:
                    continue
                
                # Get num_classes: first from classes_data (saved separately), then from item itself
                if part:
                    if idx < len(section_classes):
                        num_classes = max(1, int(section_classes[idx] or 1))
                    else:
                        num_classes = max(1, int(content_item.get('num_classes') or content_item.get('classes') or 1))
                else:
                    if idx < len(section_classes):
                        num_classes = max(1, int(section_classes[idx] or 1))
                    else:
                        num_classes = max(1, int(content_item.get('num_classes') or content_item.get('classes') or 1))
                
                clo_val = content_item.get('clo') or ''
                clo_str = str(clo_val).strip() if clo_val is not None else ''
                
                topic_slots.append({
                    'topic': content_text,
                    'clo': clo_str,
                    'classes': num_classes
                })
                total_topic_classes += num_classes
        
        # Step 3: Add 3 Assessments (Class Tests only, no Mid-term)
        TOTAL_ASSESSMENTS = 4
        total_classes_needed = total_topic_classes + TOTAL_ASSESSMENTS
        
        current_app.logger.info(f"Generate Plan: {total_topic_classes} topic classes + {TOTAL_ASSESSMENTS} assessments = {total_classes_needed} total classes")
        
        # Step 4: Collect working days (excluding Friday, Saturday, and holidays)
        working_days = []
        check_date = semester_start_date
        while check_date <= semester_end_date:
            # Skip Friday (4) and Saturday (5) - Bangladesh weekend
            if check_date.weekday() not in (4, 5) and check_date not in holidays:
                working_days.append(check_date)
            check_date += timedelta(days=1)
        
        # Step 5: Group working days into weeks with date ranges
        # Bangladesh: Week starts on Sunday (weekday 6), ends on Thursday (weekday 3)
        # Friday (4) and Saturday (5) are holidays
        weeks = []
        current_week_days = []
        current_week_start = None
        
        for day in working_days:
            if not current_week_days:
                # First working day
                current_week_start = day
                current_week_days.append(day)
            elif day.weekday() == 6:
                # Sunday = new week starts in Bangladesh
                # Save current week first
                weeks.append({
                    'week_num': len(weeks) + 1,
                    'start_date': current_week_start,
                    'end_date': current_week_days[-1],
                    'working_days': current_week_days.copy(),
                    'date_range': f"{current_week_start.strftime('%d %b %Y')} to {current_week_days[-1].strftime('%d %b %Y')}"
                })
                # Start new week
                current_week_start = day
                current_week_days = [day]
            else:
                # Same week - add to current week
                current_week_days.append(day)
        
        # Add last week
        if current_week_days:
            weeks.append({
                'week_num': len(weeks) + 1,
                'start_date': current_week_start,
                'end_date': current_week_days[-1],
                'working_days': current_week_days.copy(),
                'date_range': f"{current_week_start.strftime('%d %b %Y')} to {current_week_days[-1].strftime('%d %b %Y')}"
            })
        
        total_weeks = len(weeks)
        
        # Limit to maximum 18 weeks
        MAX_WEEKS = 18
        if total_weeks > MAX_WEEKS:
            current_app.logger.info(f"Semester has {total_weeks} weeks, limiting to {MAX_WEEKS} weeks")
            total_weeks = MAX_WEEKS
            weeks = weeks[:MAX_WEEKS]
        
        # Step 6: Calculate assessment positions (3 Class Tests evenly distributed)
        assessment_positions = []
        if total_weeks >= 3:
            # Distribute 3 assessments evenly across semester
            # Assessment at approximately 25%, 50%, 75% of semester
            interval = total_weeks // 4
            assessment_positions = [
                interval,           # ~25% - CT1
                interval * 2,       # ~50% - CT2
                interval * 3        # ~75% - CT3
            ]
            # Ensure positions are valid
            assessment_positions = [min(p, total_weeks) for p in assessment_positions]
            assessment_positions = list(set(assessment_positions))  # Remove duplicates
            assessment_positions.sort()
            # If we have fewer than 3 unique positions, adjust
            while len(assessment_positions) < 3 and assessment_positions[-1] < total_weeks:
                assessment_positions.append(assessment_positions[-1] + 1)
        else:
            # Less than 3 weeks - spread assessments across available weeks
            assessment_positions = list(range(1, min(total_weeks + 1, 4)))
        
        current_app.logger.info(f"Assessment positions (weeks): {assessment_positions}")
        
        # Step 7: Generate lesson plan with topics distributed by week
        lesson_plan = []
        topic_index = 0
        classes_used_for_current_topic = 0
        assessment_labels = ['Assessment 1', 'Assessment 2', 'Assessment 3', 'Assessment 4']
        assessment_idx = 0
        
        for week in weeks:
            week_num = week['week_num']
            date_range = week['date_range']
            available_classes = min(classes_per_week, len(week['working_days']))
            
            # Check if this week has an assessment
            is_assessment_week = week_num in assessment_positions and assessment_idx < len(assessment_labels)
            
            if is_assessment_week:
                # Add assessment entry
                lesson_plan.append({
                    'week': f"Week {week_num}",
                    'date': date_range,
                    'topic': assessment_labels[assessment_idx],
                    'outcome': '',
                    'teaching_assessment': '',
                    'clo_alignment': '',
                    'is_assessment': True
                })
                assessment_idx += 1
                available_classes -= 1  # Assessment takes 1 class slot
            
            # Fill remaining classes with topics
            if available_classes > 0 and topic_index < len(topic_slots):
                classes_filled = 0
                topics_this_week = []
                
                while classes_filled < available_classes and topic_index < len(topic_slots):
                    current_topic = topic_slots[topic_index]
                    remaining_for_topic = current_topic['classes'] - classes_used_for_current_topic
                    classes_to_use = min(remaining_for_topic, available_classes - classes_filled)
                    
                    topics_this_week.append({
                        'topic': current_topic['topic'],
                        'clo': current_topic['clo'],
                        'classes': classes_to_use
                    })
                    
                    classes_filled += classes_to_use
                    classes_used_for_current_topic += classes_to_use
                    
                    if classes_used_for_current_topic >= current_topic['classes']:
                        topic_index += 1
                        classes_used_for_current_topic = 0
                
                # Combine topics for this week into single entry
                if topics_this_week:
                    # Format: "Topic A (2 classes), Topic B (1 class)"
                    combined_topic = ', '.join([
                        f"{t['topic']} ({t['classes']} class{'es' if t['classes'] > 1 else ''})" 
                        for t in topics_this_week
                    ])
                    combined_clo = ', '.join(sorted(set(t['clo'] for t in topics_this_week if t['clo'])))
                    
                    lesson_plan.append({
                        'week': f"Week {week_num}",
                        'date': date_range,
                        'topic': combined_topic,
                        'outcome': '',
                        'teaching_assessment': '',
                        'clo_alignment': combined_clo or '1',
                        'is_assessment': False
                    })
            elif available_classes > 0 and topic_index >= len(topic_slots):
                # No more topics - add buffer/revision
                lesson_plan.append({
                    'week': f"Week {week_num}",
                    'date': date_range,
                    'topic': 'Revision / Buffer',
                    'outcome': '',
                    'teaching_assessment': '',
                    'clo_alignment': '1',
                    'is_assessment': False
                })
        
        # Ensure remaining assessments are added if not all were placed
        while assessment_idx < len(assessment_labels):
            lesson_plan.append({
                'week': f"Week {total_weeks}",
                'date': weeks[-1]['date_range'] if weeks else '',
                'topic': assessment_labels[assessment_idx],
                'outcome': '',
                'teaching_assessment': '',
                'clo_alignment': '',
                'is_assessment': True
            })
            assessment_idx += 1
        
        part_text = f" (Part {part})" if part else ""
        topics_scheduled = topic_index
        topics_remaining = len(topic_slots) - topic_index
        
        msg = f'Generated plan{part_text}: {total_topic_classes} topic classes + {TOTAL_ASSESSMENTS} assessments distributed across {total_weeks} weeks ({classes_per_week} classes/week based on {int(credit)} credit). Semester: {semester_start_date.strftime("%d %b %Y")} to {semester_end_date.strftime("%d %b %Y")}'
        
        if topics_remaining > 0:
            msg += f'. Warning: {topics_remaining} topics could not be scheduled within the semester.'
        
        return jsonify({
            'success': True,
            'lesson_plan': lesson_plan,
            'message': msg
        })
    except Exception as e:
        current_app.logger.error(f"Error generating AI plan: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error generating plan: {str(e)}'
        }), 500


def _build_lesson_plan_from_topic_slots_and_dates(slot_dates_with_week, topic_slots):
    """Build lesson_plan by mapping topic_slots (content, clo) to calendar slot_dates.
    slot_dates_with_week: list of (date, week_num). topic_slots: list of (content, clo).
    If more slots than topics, fill with 'Revision / Buffer'. If more topics than slots, use only first len(slot_dates) topics."""
    lesson_plan = []
    for i, (slot_date, week_num) in enumerate(slot_dates_with_week):
        if i < len(topic_slots):
            content_text, clo_str = topic_slots[i]
            topic = content_text
            clo_alignment = clo_str if clo_str else '1'
        else:
            topic = 'Revision / Buffer'
            clo_alignment = '1'
        date_str = slot_date.strftime('%d-%b-%Y')
        lesson_plan.append({
            'week': f'Week {week_num}',
            'date': date_str,
            'topic': topic,
            'outcome': '',
            'teaching_assessment': '',
            'clo_alignment': clo_alignment
        })
    return lesson_plan


def _generate_rule_based_plan(session, credit, course_contents, week_groups, holidays, 
                               semester_start_date, semester_end_date, classes_per_week, 
                               total_weeks, all_topics, clos, part=None):
    """Fallback rule-based plan generation when AI is not available - Always within 14 weeks maximum
    For split courses, both Part A and Part B start from Week 1 and run simultaneously"""
    from datetime import timedelta
    
    # Limit to maximum 14 weeks
    MAX_WEEKS = 14
    limited_weeks = min(total_weeks, MAX_WEEKS)
    limited_week_groups = week_groups[:limited_weeks]
    
    lesson_plan = []
    content_index = 0
    total_classes = limited_weeks * classes_per_week
    
    # For split courses, both parts start from Week 1
    # The classes_per_week is already divided (half for each part)
    
    # Distribute course contents across all classes
    if all_topics and total_classes > 0:
        topics_per_class = max(1, len(all_topics) // total_classes)
    else:
        topics_per_class = 1
    
    # Track which weeks have assessments (for 3-4 credit courses)
    # For split courses, divide assessments equally between Part A and Part B
    assessment_weeks = set()
    if credit in [3, 4]:
        # Distribute 4 assessments across 14 weeks
        if limited_weeks >= 14:
            all_assessment_week_indices = [3, 6, 9, 12]  # Weeks 4, 7, 10, 13
        elif limited_weeks >= 10:
            all_assessment_week_indices = [2, 4, 6, 8]
        else:
            all_assessment_week_indices = [1, 2, 3, 4]
        
        # For split courses, divide assessments equally
        if part in ['A', 'B']:
            # Split the assessments: Part A gets first half, Part B gets second half
            mid_point = len(all_assessment_week_indices) // 2
            if part == 'A':
                assessment_week_indices = all_assessment_week_indices[:mid_point]  # First half
            else:  # part == 'B'
                assessment_week_indices = all_assessment_week_indices[mid_point:]  # Second half
        else:
            # Full course: use all assessments
            assessment_week_indices = all_assessment_week_indices
        
        for week_idx in assessment_week_indices:
            if week_idx < limited_weeks:
                assessment_weeks.add(week_idx)
    
    # Generate plan for each week (maximum 14 weeks)
    week_num = 1
    for week_idx, week_days in enumerate(limited_week_groups):
        # Safety check: never exceed 14 weeks
        if week_num > MAX_WEEKS:
            break
            
        if not week_days:
            continue
        
        week_start = week_days[0]
        week_end = week_days[-1]
        week_classes_count = min(classes_per_week, len(week_days))
        
        # Check if this week has an assessment (for 3-4 credit courses)
        has_assessment = week_idx in assessment_weeks
        
        # If this week has an assessment, reduce regular classes by 1 to make room
        regular_classes = week_classes_count - 1 if has_assessment else week_classes_count
        
        # Generate regular classes for this week
        for class_num in range(regular_classes):
            if regular_classes == 1:
                class_date = week_days[0]
            else:
                day_index = int((class_num / regular_classes) * len(week_days))
                class_date = week_days[min(day_index, len(week_days) - 1)]
            
            # Get topics for this class
            week_contents = []
            week_clos = set()
            
            if all_topics:
                for i in range(topics_per_class):
                    if content_index < len(all_topics):
                        week_contents.append(all_topics[content_index])
                        content_index += 1
            
            # Get CLO from course contents if available
            if course_contents and content_index < len(course_contents):
                content_item = course_contents[min(content_index, len(course_contents) - 1)]
                if isinstance(content_item, dict):
                    clo_value = content_item.get('clo', '')
                    if clo_value:
                        if isinstance(clo_value, str):
                            for clo in clo_value.split(','):
                                clo = clo.strip()
                                if clo:
                                    week_clos.add(clo)
                        elif isinstance(clo_value, (int, float)):
                            week_clos.add(str(int(clo_value)))
            
            date_str = f"{week_start.strftime('%d-%b-%Y')} to {week_end.strftime('%d-%b-%Y')}"
            
            if week_contents:
                topic = ', '.join(week_contents[:2])
                if len(week_contents) > 2:
                    topic += f' and {len(week_contents) - 2} more'
            else:
                topic = f'Week {week_num} - Class {class_num + 1} Content'
            
            # Specific Outcome and Teaching & Assessment should be empty
            outcome = ''
            teaching_assessment = ''
            
            if week_clos:
                clo_alignment = ', '.join(sorted(week_clos, key=lambda x: int(x) if x.isdigit() else 999))
            elif clos:
                clo_alignment = ', '.join([str(i+1) for i in range(min(2, len(clos)))])
            else:
                clo_alignment = '1'
            
            lesson_plan.append({
                'week': f'Week {week_num}',
                'date': date_str,
                'topic': topic,
                'outcome': outcome,
                'teaching_assessment': teaching_assessment,
                'clo_alignment': clo_alignment
            })
        
        # Add assessment for this week if scheduled
        if has_assessment:
            # For split courses, renumber assessments starting from 1 for each part
            # Part A: Assessment 1, 2; Part B: Assessment 1, 2
            # For full course: Assessment 1, 2, 3, 4
            sorted_assessment_weeks = sorted(list(assessment_weeks))
            assessment_num = sorted_assessment_weeks.index(week_idx) + 1
            if part in ['A', 'B']:
                # For split courses, add part label
                topic_text = f'Assessment {assessment_num} (Part {part})'
            else:
                topic_text = f'Assessment {assessment_num}'
            
            date_str = f"{week_start.strftime('%d-%b-%Y')} to {week_end.strftime('%d-%b-%Y')}"
            
            # Specific Outcome and Teaching & Assessment should be empty
            lesson_plan.append({
                'week': f'Week {week_num}',
                'date': date_str,
                'topic': topic_text,
                'outcome': '',
                'teaching_assessment': '',
                'clo_alignment': '1, 2, 3, 4'
            })
        
        week_num += 1
        # Ensure we never exceed 14 weeks
        if week_num > MAX_WEEKS:
            break
    
    # Final safety check: ensure we never return more than 14 weeks worth of classes
    # Calculate max entries: 14 weeks * classes_per_week + assessments (max 4)
    max_entries = (MAX_WEEKS * classes_per_week) + (4 if credit in [3, 4] else 0)
    if len(lesson_plan) > max_entries:
        current_app.logger.warning(f"Generated {len(lesson_plan)} entries, limiting to {max_entries} (14 weeks max)")
        lesson_plan = lesson_plan[:max_entries]
    
    # Final safety: ensure we never exceed max_entries (14 weeks worth of classes)
    if len(lesson_plan) > max_entries:
        current_app.logger.warning(f"Generated {len(lesson_plan)} entries, limiting to {max_entries} (14 weeks max)")
        lesson_plan = lesson_plan[:max_entries]
    
    return lesson_plan

@class_management_bp.route('/course_file/<int:session_id>/outline/edit')
@login_required
def edit_course_outline(session_id):
    """Edit course outline page"""
    session = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    
    # Check if teacher is authorized (either main teacher or part of split group)
    is_authorized = False
    if teacher:
        if teacher.id == session.teacher_id:
            is_authorized = True
        elif session.split_group_id:
            # Check if teacher is part of the split group
            related_sessions = query_for_window(Session).filter_by(split_group_id=session.split_group_id).all()
            for related_session in related_sessions:
                if related_session.teacher_id == teacher.id:
                    is_authorized = True
                    break
    
    if not is_authorized:
        flash('You are not authorized to edit this course outline.', 'danger')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    # Get all teachers for this course (if split group exists)
    course_teachers = [session.teacher]
    if session.split_group_id:
        related_sessions = query_for_window(Session).filter_by(split_group_id=session.split_group_id).all()
        for related_session in related_sessions:
            if related_session.teacher and related_session.teacher not in course_teachers:
                course_teachers.append(related_session.teacher)
    
    course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
    if not course_outline:
        course_outline = CourseOutline(session_id=session_id, teacher_id=teacher.id)
        db.session.add(course_outline)
        db.session.commit()
    
    # Get course data from curriculum using improved matching
    # IMPORTANT: Same course can exist in multiple curricula, so we need to find the right one
    course_data = None
    try:
        if Course:
            import re
            current_app.logger.info(f"Searching for course - session.course_code: '{session.course_code}', session.course_name: '{session.course_name}', session.year: '{session.year}', session.term: '{session.term}', session.academic_session: '{session.academic_session}'")
            
            def extract_core_code(code_str):
                """Extract the core code pattern (e.g., 'Law4103' without space) from various formats"""
                if not code_str:
                    return None
                # Try to find pattern like "Law 4103", "Law4103", "CSE 1101", etc.
                match = re.search(r'([A-Za-z]+)\s*(\d{4})', code_str)
                if match:
                    return f"{match.group(1)}{match.group(2)}"  # No space: "Law4103"
                return None
            
            def has_course_info(course):
                """Check if a course has Course Information (rationale, CLO, or content)"""
                return bool(course.rationale or course.clo or course.content_section_a or course.content_section_b)
            
            def normalize_year_term(value):
                """Normalize year/term values for comparison"""
                if not value:
                    return ''
                value = str(value).strip().lower()
                # Map numeric to text
                year_map = {'1': 'first', '2': 'second', '3': 'third', '4': 'fourth', '5': 'fifth'}
                term_map = {'1': 'first', '2': 'second'}
                if value in year_map:
                    return year_map[value]
                if value in term_map:
                    return term_map[value]
                # Remove "year" or "term" suffix
                for suffix in [' year', ' term']:
                    if value.endswith(suffix):
                        value = value[:-len(suffix)]
                return value
            
            # Extract core code pattern from session course code
            session_core_code = extract_core_code(session.course_code)
            if session_core_code:
                current_app.logger.info(f"Extracted core code: '{session_core_code}' from '{session.course_code}'")
            
            # Normalize session year/term for matching
            session_year_norm = normalize_year_term(session.year)
            session_term_norm = normalize_year_term(session.term)
            
            def find_matching_courses(code_filter_func):
                """Find all courses matching the code filter, then prioritize by curriculum match and course info"""
                all_courses = Course.query.all()
                matching_courses = [c for c in all_courses if c.course_code and code_filter_func(c.course_code)]
                
                if not matching_courses:
                    return None
                
                current_app.logger.info(f"Found {len(matching_courses)} courses matching code filter")
                
                # Score courses: higher score = better match
                scored_courses = []
                for course in matching_courses:
                    score = 0
                    reasons = []
                    
                    # Priority 1: Has Course Information (most important)
                    if has_course_info(course):
                        score += 100
                        reasons.append("has_course_info")
                    
                    # Priority 2: Year/Term match from course's derived or stored year/term
                    course_year_norm = normalize_year_term(course.year or course.derived_year)
                    course_term_norm = normalize_year_term(course.term or course.derived_term)
                    
                    if session_year_norm and course_year_norm and session_year_norm == course_year_norm:
                        score += 20
                        reasons.append(f"year_match({course_year_norm})")
                    if session_term_norm and course_term_norm and session_term_norm == course_term_norm:
                        score += 10
                        reasons.append(f"term_match({course_term_norm})")
                    
                    # Priority 3: Academic session match via CurriculumYearTerm
                    if CurriculumYearTerm and session.academic_session and course.curriculum_id:
                        try:
                            config = query_for_window(CurriculumYearTerm).filter_by(
                                curriculum_id=course.curriculum_id,
                                academic_session=session.academic_session
                            ).first()
                            if config:
                                score += 50
                                reasons.append(f"academic_session_match({session.academic_session})")
                        except Exception as e:
                            current_app.logger.warning(f"Error checking CurriculumYearTerm: {e}")
                    
                    scored_courses.append((course, score, reasons))
                    current_app.logger.info(f"  Course: {course.course_code} (curriculum_id={course.curriculum_id}), score={score}, reasons={reasons}")
                
                # Sort by score (highest first)
                scored_courses.sort(key=lambda x: x[1], reverse=True)
                
                if scored_courses:
                    best_course, best_score, best_reasons = scored_courses[0]
                    current_app.logger.info(f"Selected best match: {best_course.course_code} with score {best_score} ({best_reasons})")
                    return best_course
                
                return None
            
            # Try exact match by course code
            if session.course_code:
                course_data = find_matching_courses(
                    lambda code: code == session.course_code
                )
                if course_data:
                    current_app.logger.info(f"Found by exact course_code match: {course_data.course_code}")
            
            # If not found, try case-insensitive match by course code
            if not course_data and session.course_code:
                course_data = find_matching_courses(
                    lambda code: code.lower() == session.course_code.lower()
                )
                if course_data:
                    current_app.logger.info(f"Found by case-insensitive course_code match: {course_data.course_code}")
            
            # If not found, try whitespace-normalized match
            if not course_data and session.course_code:
                session_code_normalized = ' '.join(session.course_code.strip().split()).lower()
                course_data = find_matching_courses(
                    lambda code: ' '.join(code.strip().split()).lower() == session_code_normalized
                )
                if course_data:
                    current_app.logger.info(f"Found by whitespace-normalized match: {course_data.course_code}")
            
            # If not found, try with extracted course code pattern (with space)
            if not course_data and session_core_code:
                extracted_with_space = re.sub(r'([A-Za-z]+)(\d{4})', r'\1 \2', session_core_code).lower()
                course_data = find_matching_courses(
                    lambda code: code.lower() == extracted_with_space
                )
                if course_data:
                    current_app.logger.info(f"Found by extracted code with space: {course_data.course_code}")
            
            # If not found, try with extracted course code pattern (without space)
            if not course_data and session_core_code:
                course_data = find_matching_courses(
                    lambda code: code.lower() == session_core_code.lower()
                )
                if course_data:
                    current_app.logger.info(f"Found by extracted code without space: {course_data.course_code}")
            
            # If not found, try normalized core code matching
            if not course_data and session_core_code:
                session_core_lower = session_core_code.lower()
                course_data = find_matching_courses(
                    lambda code: extract_core_code(code) and extract_core_code(code).lower() == session_core_lower
                )
                if course_data:
                    current_app.logger.info(f"Found by normalized core code match: {course_data.course_code}")
            
            # If not found, try partial match - check if curriculum course code is contained in session code
            if not course_data and session.course_code:
                session_code_lower = session.course_code.lower()
                session_code_no_space = session_code_lower.replace(' ', '')
                course_data = find_matching_courses(
                    lambda code: code.lower() in session_code_lower or code.lower().replace(' ', '') in session_code_no_space
                )
                if course_data:
                    current_app.logger.info(f"Found by partial code match: {course_data.course_code}")
            
            # If not found, try exact match by course name
            if not course_data and session.course_name:
                course_data = find_matching_courses(
                    lambda code: True  # Match all, but filter by name below
                )
                # Re-filter by name since find_matching_courses filters by code
                if not course_data:
                    all_courses = Course.query.filter_by(course_name=session.course_name).all()
                    if all_courses:
                        # Pick the one with course info
                        for c in all_courses:
                            if has_course_info(c):
                                course_data = c
                                break
                        if not course_data:
                            course_data = all_courses[0]
                        if course_data:
                            current_app.logger.info(f"Found by exact course_name match: {course_data.course_name}")
            
            # If not found, try case-insensitive partial match by course name
            if not course_data and session.course_name:
                all_courses = Course.query.filter(func.lower(Course.course_name).like(f'%{session.course_name.lower()}%')).all()
                if all_courses:
                    # Pick the one with course info
                    for c in all_courses:
                        if has_course_info(c):
                            course_data = c
                            break
                    if not course_data:
                        course_data = all_courses[0]
                    if course_data:
                        current_app.logger.info(f"Found by partial course_name match: {course_data.course_name}")
            
            # If still not found, try reverse match (session name contains course name)
            if not course_data and session.course_name:
                all_courses = Course.query.all()
                session_name_lower = session.course_name.lower()
                matching = []
                for course in all_courses:
                    if course.course_name:
                        course_name_lower = course.course_name.lower()
                        if session_name_lower in course_name_lower or course_name_lower in session_name_lower:
                            matching.append(course)
                if matching:
                    # Pick the one with course info
                    for c in matching:
                        if has_course_info(c):
                            course_data = c
                            break
                    if not course_data:
                        course_data = matching[0]
                    if course_data:
                        current_app.logger.info(f"Found by reverse course_name match: {course_data.course_name}")
            
            # Log the result for debugging
            if course_data:
                current_app.logger.info(f"✓ Found course_data: {course_data.course_code} - {course_data.course_name}, core_optional: {course_data.core_optional}, course_type: {course_data.course_type}, category: {course_data.category}")
                # Log additional fields for debugging
                current_app.logger.info(f"  - rationale: {course_data.rationale[:100] if course_data.rationale else 'EMPTY'}")
                current_app.logger.info(f"  - clo: {course_data.clo[:100] if course_data.clo else 'EMPTY'}")
                current_app.logger.info(f"  - content_section_a: {course_data.content_section_a[:100] if course_data.content_section_a else 'EMPTY'}")
                current_app.logger.info(f"  - content_section_b: {course_data.content_section_b[:100] if course_data.content_section_b else 'EMPTY'}")
            else:
                current_app.logger.warning(f"✗ Course data NOT found for session - course_code: '{session.course_code}', course_name: '{session.course_name}'")
                # List all courses for debugging with their core codes
                all_courses = Course.query.all()
                current_app.logger.info(f"Available courses in database ({len(all_courses)} total):")
                for c in all_courses[:15]:  # Show first 15
                    c_core = extract_core_code(c.course_code)
                    current_app.logger.info(f"  - Code: '{c.course_code}', Core: '{c_core}', Name: '{c.course_name}'")
    except Exception as e:
        current_app.logger.error(f"Error fetching course data: {e}", exc_info=True)
        current_app.logger.warning(f"Session course_code: {session.course_code}, course_name: {session.course_name}")
        course_data = None
    
    # Get CLO data from course if not already saved in outline
    clo_data_from_course = []
    if course_data and hasattr(course_data, 'get_clos_list'):
        try:
            course_clos = course_data.get_clos_list()
            if course_clos:
                # Convert course CLO format to outline CLO format
                for idx, clo in enumerate(course_clos, 1):
                    # Parse PLO from curriculum
                    plos_list = []
                    if clo.get('plo'):
                        plo_value = clo.get('plo', '')
                        if isinstance(plo_value, str) and plo_value.strip():
                            plos_list = [p.strip() for p in plo_value.split(',') if p.strip()]
                        elif isinstance(plo_value, list):
                            plos_list = [str(p).strip() for p in plo_value if p]
                    
                    clo_data_from_course.append({
                        'number': idx,
                        'description': clo.get('text', ''),
                        'plos': plos_list
                    })
        except Exception as e:
            current_app.logger.warning(f"Error parsing course CLOs: {e}")
    
    # Parse course contents from curriculum for import
    course_contents_a = []
    course_contents_b = []
    if course_data:
        try:
            if course_data.content_section_a:
                content_a = course_data.content_section_a
                try:
                    content_a_data = json.loads(content_a) if isinstance(content_a, str) else content_a
                    if isinstance(content_a_data, list):
                        course_contents_a = content_a_data
                except:
                    pass
            if course_data.content_section_b:
                content_b = course_data.content_section_b
                try:
                    content_b_data = json.loads(content_b) if isinstance(content_b, str) else content_b
                    if isinstance(content_b_data, list):
                        course_contents_b = content_b_data
                except:
                    pass
        except Exception as e:
            current_app.logger.warning(f"Error parsing course contents: {e}")
    
    # Load Classes separately (stored in course_content_classes)
    existing_classes = {'section_a': [], 'section_b': []}
    if getattr(course_outline, 'course_content_classes', None):
        try:
            existing_classes = json.loads(course_outline.course_content_classes)
            if not isinstance(existing_classes, dict):
                existing_classes = {'section_a': [], 'section_b': []}
            existing_classes.setdefault('section_a', [])
            existing_classes.setdefault('section_b', [])
        except (TypeError, ValueError):
            pass

    # Parse existing course content summary (fallback to curriculum for display)
    existing_content_summary = None
    if course_outline.course_content_summary:
        try:
            existing_content_summary = json.loads(course_outline.course_content_summary)
        except (TypeError, ValueError, json.JSONDecodeError):
            existing_content_summary = course_outline.course_content_summary

    from utils.ai.curriculum_anchor import resolve_course_content_summary
    existing_content_summary = resolve_course_content_summary(
        existing_content_summary, course_data=course_data, classes_data=existing_classes,
    )

    # Parse JSON fields
    outline_data = {
        'course_objectives': json.loads(course_outline.course_objectives) if course_outline.course_objectives else [],
        'course_content_summary': existing_content_summary or '',
        'clo_plo_mapping': course_outline.clo_plo_mapping or '',
        'clo_data': json.loads(course_outline.clo_data) if hasattr(course_outline, 'clo_data') and course_outline.clo_data else (clo_data_from_course if clo_data_from_course else []),
        'plo_mapping': json.loads(course_outline.plo_mapping) if hasattr(course_outline, 'plo_mapping') and course_outline.plo_mapping else {},
        'lesson_plan': json.loads(course_outline.lesson_plan) if course_outline.lesson_plan else [],
        'assessment_strategy': json.loads(course_outline.assessment_strategy) if course_outline.assessment_strategy else {},
        'assessment_techniques': json.loads(course_outline.assessment_techniques) if course_outline.assessment_techniques else [],
        'rubrics': json.loads(course_outline.rubrics) if course_outline.rubrics else [],
        'grading_policy': json.loads(course_outline.grading_policy) if course_outline.grading_policy else [],
        'cie_breakdown': json.loads(course_outline.cie_breakdown) if hasattr(course_outline, 'cie_breakdown') and course_outline.cie_breakdown else [],
        'smee_breakdown': json.loads(course_outline.smee_breakdown) if hasattr(course_outline, 'smee_breakdown') and course_outline.smee_breakdown else [],
        'textbooks': json.loads(course_outline.textbooks) if course_outline.textbooks else [],
        'reference_books': json.loads(course_outline.reference_books) if course_outline.reference_books else [],
        'other_resources': json.loads(course_outline.other_resources) if course_outline.other_resources else [],
        'course_file_components': json.loads(course_outline.course_file_components) if course_outline.course_file_components else [],
        'other_issues': json.loads(course_outline.other_issues) if course_outline.other_issues else {},
        'make_up_procedures': course_outline.make_up_procedures or '',
    }
    from utils.ai.outline_parser import normalize_assessment_strategy
    outline_data['assessment_strategy'] = normalize_assessment_strategy(outline_data.get('assessment_strategy'))

    # Parse make-up procedures into a list for the sessional Part C form
    make_up_raw = outline_data.get('make_up_procedures') or ''
    make_up_list = []
    if isinstance(make_up_raw, list):
        make_up_list = [str(x).strip() for x in make_up_raw if str(x).strip()]
    elif isinstance(make_up_raw, str) and make_up_raw.strip():
        stripped = make_up_raw.strip()
        if stripped.startswith('['):
            try:
                parsed = json.loads(stripped)
                if isinstance(parsed, list):
                    make_up_list = [str(x).strip() for x in parsed if str(x).strip()]
            except (TypeError, ValueError, json.JSONDecodeError):
                make_up_list = [ln.strip() for ln in make_up_raw.splitlines() if ln.strip()]
        else:
            make_up_list = [ln.strip() for ln in make_up_raw.splitlines() if ln.strip()]
    outline_data['make_up_procedures_list'] = make_up_list

    ai_outline_enabled = False
    try:
        from utils.ai.models import AIProviderSetting
        ai_outline_enabled = bool(AIProviderSetting.get_active_default())
    except Exception:
        ai_outline_enabled = False

    ai_default_classes_per_week = ''
    ai_default_total_classes = ''
    try:
        credit_val = float(course_data.credit) if course_data and course_data.credit else None
    except (TypeError, ValueError):
        credit_val = None
    if credit_val:
        cpw = int(credit_val)
        if getattr(session, 'course_scope', None) in ('part_a', 'part_b'):
            cpw = max(1, cpw // 2)
        ai_default_classes_per_week = cpw
        ai_default_total_classes = int(credit_val * 14)
    if course_outline and course_outline.contact_hours:
        try:
            ai_default_total_classes = int(float(str(course_outline.contact_hours).strip()))
        except (TypeError, ValueError):
            pass
    
    return render_template('class_management/edit_course_outline.html',
                         session=session,
                         course_outline=course_outline,
                         course_data=course_data,
                         outline_data=outline_data,
                         existing_classes=existing_classes,
                         course_contents_a=course_contents_a,
                         course_contents_b=course_contents_b,
                         course_teachers=course_teachers,
                         ai_outline_enabled=ai_outline_enabled,
                         session_delivery_type=(session.course_type or 'theory'),
                         ai_default_total_classes=ai_default_total_classes,
                         ai_default_classes_per_week=ai_default_classes_per_week)

def _generate_course_outline_docx(session_id):
    """Generate course outline as DOCX document"""
    session = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    
    if not teacher or teacher.id != session.teacher_id:
        flash('You are not authorized to download this course outline.', 'danger')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
    if not course_outline:
        flash('Course outline not found. Please create it first.', 'warning')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    # Get course data
    course_data = find_course_from_curriculum(session.course_code, session.course_name, session=session)
    
    # Parse JSON fields
    course_objectives = json.loads(course_outline.course_objectives) if course_outline.course_objectives else []
    lesson_plan = json.loads(course_outline.lesson_plan) if course_outline.lesson_plan else []
    textbooks = json.loads(course_outline.textbooks) if course_outline.textbooks else []
    reference_books = json.loads(course_outline.reference_books) if course_outline.reference_books else []
    other_resources = json.loads(course_outline.other_resources) if course_outline.other_resources else []
    
    # Lazy import docx to prevent startup hang
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Create DOCX document
    doc = Document()
    
    # Cover Page
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_para.add_run('Course Outline')
    run.font.size = Pt(18)
    run.font.bold = True
    
    if session.course_code:
        code_para = doc.add_paragraph()
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = code_para.add_run(session.course_code + ':')
        run.font.size = Pt(16)
        run.font.bold = True
    
    if session.course_name:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(session.course_name.upper())
        run.font.size = Pt(20)
        run.font.bold = True
    
    doc.add_page_break()
    
    # Part A: Introduction
    doc.add_heading('PART A: INTRODUCTION', level=1)
    
    # Course Identification Table
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    info_table.cell(0, 0).text = 'Course No:'
    info_table.cell(0, 1).text = session.course_code or '—'
    # Add more rows...
    
    # Course Objectives
    if course_objectives:
        doc.add_heading('Course Objectives', level=2)
        for obj in course_objectives:
            doc.add_paragraph(obj, style='List Bullet')
    
    # Course Summary
    if course_outline.course_summary:
        doc.add_heading('Course Summary', level=2)
        doc.add_paragraph(course_outline.course_summary)
    
    # Lesson Plan
    if lesson_plan:
        doc.add_page_break()
        doc.add_heading('Class Schedule/Lesson Plan/Weekly plan', level=1)
        lesson_table = doc.add_table(rows=1, cols=7)
        lesson_table.style = 'Table Grid'
        headers = ['Week', 'Date', 'Topic', 'Specific Outcome', 'Suggested Activities', 'Teaching and Assessment', 'Alignment with CLO']
        for i, header in enumerate(headers):
            lesson_table.cell(0, i).text = header
            lesson_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        
        for lesson in lesson_plan:
            row = lesson_table.add_row().cells
            row[0].text = lesson.get('week', '')
            row[1].text = lesson.get('date', '')
            row[2].text = lesson.get('topic', '')
            row[3].text = lesson.get('outcome', '')
            row[4].text = lesson.get('activities', '')
            row[5].text = lesson.get('teaching_assessment', '')
            row[6].text = lesson.get('clo_alignment', '')
    
    # Learning Resources
    if textbooks or reference_books or other_resources:
        doc.add_page_break()
        doc.add_heading('PART D: LEARNING RESOURCES', level=1)
        
        if textbooks:
            doc.add_heading('Textbooks', level=2)
            for book in textbooks:
                doc.add_paragraph(book, style='List Bullet')
        
        if reference_books:
            doc.add_heading('Reference Books', level=2)
            for book in reference_books:
                doc.add_paragraph(book, style='List Bullet')
        
        if other_resources:
            doc.add_heading('Other Resources', level=2)
            for resource in other_resources:
                doc.add_paragraph(resource, style='List Bullet')
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_data = buffer.getvalue()
    buffer.close()
    
    filename = f"course_outline_{session.course_code or 'course'}.docx"
    return Response(
        docx_data,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(docx_data)),
        },
    )

def _generate_course_outline_pdf(session_id, skip_auth_check=False):
    """Generate comprehensive course outline as PDF document with cover page and page numbers using WeasyPrint
    
    Args:
        session_id: The session ID for the course outline
        skip_auth_check: If True, skip authorization checks (for student downloads)
    """
    try:
        from weasyprint import HTML, CSS
    except ImportError:
        flash('WeasyPrint is not installed. Please install it to generate PDFs.', 'error')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    session = get_or_404_for_window(Session, session_id)
    
    # Skip authorization check if requested (for student downloads)
    if not skip_auth_check:
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        # Check if user is authorized (teacher or part of split group)
        is_authorized = False
        if teacher and teacher.id == session.teacher_id:
            is_authorized = True
        elif session.split_group_id:
            related_sessions = query_for_window(Session).filter_by(split_group_id=session.split_group_id).all()
            for related_session in related_sessions:
                if related_session.teacher and related_session.teacher.id == teacher.id:
                    is_authorized = True
                    break
        
        if not is_authorized:
            flash('You are not authorized to download this course outline.', 'danger')
            return redirect(url_for('class_management.course_file', session_id=session_id))
    
    course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
    if not course_outline:
        flash('Course outline not found. Please create it first.', 'warning')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    # Get course data from curriculum if available
    course_data = find_course_from_curriculum(session.course_code, session.course_name, session=session)
    
    # Parse all JSON fields
    def safe_json_parse(data, default=None):
        if not data:
            return default if default is not None else []
        try:
            return json.loads(data) if isinstance(data, str) else data
        except:
            return default if default is not None else []
    
    course_objectives = safe_json_parse(course_outline.course_objectives, [])
    lesson_plan = safe_json_parse(course_outline.lesson_plan, [])
    clo_data_raw = safe_json_parse(course_outline.clo_data, [])
    # Ensure plos is always a list in each CLO entry
    clo_data = []
    for clo in clo_data_raw:
        clo_entry = dict(clo)
        plos = clo_entry.get('plos', [])
        if not isinstance(plos, list):
            if isinstance(plos, str):
                clo_entry['plos'] = [plos] if plos else []
            else:
                clo_entry['plos'] = []
        clo_data.append(clo_entry)
    course_content_summary = safe_json_parse(course_outline.course_content_summary, {})
    # Merge Classes from course_content_classes into content summary for PDF
    classes_data = {'section_a': [], 'section_b': []}
    if getattr(course_outline, 'course_content_classes', None):
        try:
            classes_data = json.loads(course_outline.course_content_classes)
            if not isinstance(classes_data, dict):
                classes_data = {'section_a': [], 'section_b': []}
            classes_data.setdefault('section_a', [])
            classes_data.setdefault('section_b', [])
        except (TypeError, ValueError):
            pass
    from utils.ai.curriculum_anchor import resolve_course_content_summary
    course_content_summary = resolve_course_content_summary(
        course_content_summary, course_data=course_data, classes_data=classes_data,
    )
    assessment_strategy = safe_json_parse(course_outline.assessment_strategy, {})
    from utils.ai.outline_parser import normalize_assessment_strategy
    assessment_strategy = normalize_assessment_strategy(assessment_strategy)
    assessment_techniques = safe_json_parse(course_outline.assessment_techniques, [])
    cie_breakdown = safe_json_parse(course_outline.cie_breakdown, []) if hasattr(course_outline, 'cie_breakdown') else []
    smee_breakdown = safe_json_parse(course_outline.smee_breakdown, []) if hasattr(course_outline, 'smee_breakdown') else []
    rubrics = safe_json_parse(course_outline.rubrics, [])
    # Group rubrics by type for easier template rendering
    rubrics_by_type = {}
    for rubric in rubrics:
        rubric_type = rubric.get('type', '') or ''
        if rubric_type not in rubrics_by_type:
            rubrics_by_type[rubric_type] = []
        rubrics_by_type[rubric_type].append(rubric)
    
    grading_policy = safe_json_parse(course_outline.grading_policy, [])
    evaluation_policy = safe_json_parse(course_outline.evaluation_policy, {})
    textbooks = safe_json_parse(course_outline.textbooks, [])
    reference_books = safe_json_parse(course_outline.reference_books, [])
    other_resources = safe_json_parse(course_outline.other_resources, [])
    course_file_components = safe_json_parse(course_outline.course_file_components, [])
    other_issues = safe_json_parse(course_outline.other_issues, {})
    make_up_raw = getattr(course_outline, 'make_up_procedures', None) or ''
    make_up_procedures_list = []
    if isinstance(make_up_raw, list):
        make_up_procedures_list = [str(x).strip() for x in make_up_raw if str(x).strip()]
    elif isinstance(make_up_raw, str) and make_up_raw.strip():
        stripped = make_up_raw.strip()
        if stripped.startswith('['):
            try:
                parsed = json.loads(stripped)
                if isinstance(parsed, list):
                    make_up_procedures_list = [str(x).strip() for x in parsed if str(x).strip()]
            except (TypeError, ValueError, json.JSONDecodeError):
                make_up_procedures_list = [ln.strip() for ln in make_up_raw.splitlines() if ln.strip()]
        else:
            make_up_procedures_list = [ln.strip() for ln in make_up_raw.splitlines() if ln.strip()]
    session_delivery_type = (session.course_type or 'theory')
    course_teachers = [session.teacher]
    course_teachers_pdf = [session.teacher]
    if session.split_group_id:
        related_sessions = query_for_window(Session).filter_by(split_group_id=session.split_group_id).all()
        for related_session in related_sessions:
            if related_session.teacher and related_session.teacher not in course_teachers:
                course_teachers.append(related_session.teacher)
            if related_session.teacher and related_session.teacher not in course_teachers_pdf:
                course_teachers_pdf.append(related_session.teacher)
    
    formal_fonts = _resolve_formal_pdf_fonts()
    if not formal_fonts:
        flash(
            'PDF fonts missing. Upload LiberationSerif-Regular.ttf and LiberationSerif-Bold.ttf to static/fonts/.',
            'error',
        )
        return redirect(url_for('class_management.course_file', session_id=session_id))

    # Render HTML template
    html_content = render_template(
        'class_management/course_outline_pdf.html',
        session=session,
        course_outline=course_outline,
        course_data=course_data,
        course_objectives=course_objectives,
        lesson_plan=lesson_plan,
        clo_data=clo_data,
        course_content_summary=course_content_summary,
        assessment_strategy=assessment_strategy,
        assessment_techniques=assessment_techniques,
        cie_breakdown=cie_breakdown,
        smee_breakdown=smee_breakdown,
        rubrics=rubrics,
        rubrics_by_type=rubrics_by_type,
        grading_policy=grading_policy,
        evaluation_policy=evaluation_policy,
        textbooks=textbooks,
        reference_books=reference_books,
        other_resources=other_resources,
        course_file_components=course_file_components,
        other_issues=other_issues,
        make_up_procedures=make_up_raw if isinstance(make_up_raw, str) else '',
        make_up_procedures_list=make_up_procedures_list,
        session_delivery_type=session_delivery_type,
        course_teachers=course_teachers,
        course_teachers_pdf=course_teachers_pdf,
        pdf_font_regular=formal_fonts['regular'],
        pdf_font_bold=formal_fonts['bold'],
        pdf_font_italic=formal_fonts.get('italic'),
        pdf_font_bold_italic=formal_fonts.get('bold_italic'),
    )
    
    # Generate PDF using WeasyPrint with embedded Liberation Serif (Times-compatible)
    buffer = io.BytesIO()
    HTML(string=html_content, base_url=formal_fonts['fonts_dir'].as_uri() + '/').write_pdf(buffer)
    buffer.seek(0)
    
    filename = f"course_outline_{session.course_code or 'course'}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return Response(
        buffer.getvalue(),
        mimetype='application/pdf',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(buffer.getvalue())),
        },
    )

@class_management_bp.route('/course_file/<int:session_id>/outline/download/docx')
@login_required
def download_course_outline_docx(session_id):
    """Download course outline as DOCX"""
    return _generate_course_outline_docx(session_id)

@class_management_bp.route('/course_file/<int:session_id>/outline/download/pdf')
@login_required
def download_course_outline_pdf(session_id):
    """Download course outline as PDF"""
    return _generate_course_outline_pdf(session_id)

@class_management_bp.route('/archive_session/<int:session_id>', methods=['POST'])
@login_required
def archive_session(session_id):
    """Archive a session"""
    session = get_or_404_for_window(Session, session_id)
    session.archived = True
    db.session.commit()
    flash('Session archived successfully!', 'success')
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/unarchive_session/<int:session_id>', methods=['POST'])
@login_required
def unarchive_session(session_id):
    """Unarchive a session"""
    session = get_or_404_for_window(Session, session_id)
    session.archived = False
    db.session.commit()
    flash('Session unarchived successfully!', 'success')
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/edit_session/<int:session_id>', methods=['GET', 'POST'])
@login_required
def edit_session(session_id):
    """Edit a session"""
    session = get_or_404_for_window(Session, session_id)
    
    if request.method == 'POST':
        try:
            if getattr(session, 'is_external_course', False):
                academic_session = (request.form.get('academic_session') or '').strip()
                year = (request.form.get('year') or '').strip()
                term = (request.form.get('term') or '').strip()
                session.academic_session = academic_session or session.academic_session
                session.year = year or session.year
                session.term = term or session.term
            else:
                session.year = request.form.get('year')
                session.term = request.form.get('term')
                session.academic_session = request.form.get('academic_session')
            session.course_code = request.form.get('course_code')
            session.course_name = request.form.get('course_name')
            session.course_type = request.form.get('course_type', 'theory')
            session.category = request.form.get('category', 'ug')

            if not session.year or not session.term:
                flash('Year and term are required!', 'error')
                return redirect(url_for('class_management.edit_session', session_id=session_id))

            if getattr(session, 'is_external_course', False):
                conflict_message = _external_course_conflict(
                    session.course_code,
                    session.teacher_id,
                    exclude_session_id=session.id,
                )
                if conflict_message:
                    flash(conflict_message, 'error')
                    return redirect(url_for('class_management.edit_session', session_id=session_id))
            
            db.session.commit()
            flash('Session updated successfully!', 'success')
            return redirect(url_for('class_management.index'))
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating session {session_id}: {e}")
            flash(f'Error updating session: {str(e)}', 'error')
            return redirect(url_for('class_management.edit_session', session_id=session_id))
    
    return render_template('class_management/edit_session.html', session=session)

@class_management_bp.route('/archive')
@login_required
def archive():
    """View archived sessions"""
    # Get or create teacher for current user
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher:
        teacher = Teacher(name=current_user.full_name, institute=current_tenant().institute_label)
        db.session.add(teacher)
        db.session.commit()
    
    archived_sessions = query_for_window(Session).filter_by(teacher_id=teacher.id, archived=True).order_by(Session.created_at.desc()).all()
    
    # Build assignment map for template to access batch and academic_session from CourseSessionAssignment
    assignment_map = {}
    if CourseSessionAssignment and Course:
        try:
            for session in archived_sessions:
                # Try to find assignment by session_id first
                assignment = CourseSessionAssignment.query.filter_by(session_id=session.id).first()
                
                # If not found by session_id, try to find by course_code, teacher_id, year, term
                if not assignment and session.course_code and session.teacher_id and session.year and session.term:
                    try:
                        # Try to match by course_code, teacher_id, year, term
                        assignment = CourseSessionAssignment.query.filter_by(
                            teacher_id=session.teacher_id,
                            year=session.year,
                            term=session.term
                        ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                            Course.course_code == session.course_code
                        ).first()
                        
                        # If not found, try without section matching (for full course sessions)
                        if not assignment:
                            assignment = CourseSessionAssignment.query.filter_by(
                                teacher_id=session.teacher_id,
                                year=session.year,
                                term=session.term
                            ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                                Course.course_code == session.course_code
                            ).filter(
                                or_(
                                    CourseSessionAssignment.section.is_(None),
                                    CourseSessionAssignment.section == ''
                                )
                            ).first()
                    except Exception as query_error:
                        current_app.logger.warning(f'Error querying assignment for archived session {session.id}: {query_error}')
                
                if assignment:
                    # If assignment doesn't have batch/academic_session, try to get from curriculum year-term config
                    batch = assignment.batch
                    academic_session = assignment.academic_session
                    
                    if Curriculum and CurriculumYearTerm and (not batch or not academic_session):
                        try:
                            if assignment.curriculum_id:
                                curriculum = Curriculum.query.get(assignment.curriculum_id)
                                if curriculum:
                                    year_term_config = curriculum.get_year_term_config(assignment.year, assignment.term)
                                    if year_term_config:
                                        if not batch and year_term_config.batch and year_term_config.batch != 'None':
                                            batch = year_term_config.batch
                                        if not academic_session and year_term_config.academic_session:
                                            academic_session = year_term_config.academic_session
                        except Exception as config_error:
                            current_app.logger.warning(f'Error getting year-term config for assignment {assignment.id}: {config_error}')
                    
                    assignment_map[session.id] = {
                        'batch': batch or '',
                        'academic_session': academic_session or ''
                    }
        except Exception as e:
            current_app.logger.error(f'Error building assignment map for archived sessions: {str(e)}', exc_info=True)
    
    return render_template('class_management/archive.html', 
                         sessions=archived_sessions, 
                         assignment_map=assignment_map if assignment_map else {})

@class_management_bp.route('/delete_attendance/<int:session_id>/<string:date_str>', methods=['POST'])
@login_required
def delete_attendance_by_date(session_id, date_str):
    """Delete all attendance records for a specific date."""
    try:
        attendance_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        
        session = get_or_404_for_window(Session, session_id)
        
        if not user_owns_class_session(current_user, session):
            flash('You are not authorized to delete attendance for this session.', 'danger')
            return redirect(url_for('class_management.index'))
        
        # Count records before deletion
        records_to_delete = ClassAttendance.query.filter_by(
            session_id=session_id,
            date=attendance_date
        )
        
        if records_to_delete.count() == 0:
            flash(f'No attendance records found for {attendance_date.strftime("%b %d, %Y")}.', 'warning')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))
            
        # Delete the records
        deleted_count = records_to_delete.delete()
        
        db.session.commit()
        
        flash(f'Successfully deleted {deleted_count} attendance records for {attendance_date.strftime("%b %d, %Y")}.', 'success')
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting attendance for session {session_id} on date {date_str}: {str(e)}")
        flash(f'Error deleting attendance: {str(e)}', 'danger')
        
    return redirect(url_for('class_management.view_attendance', session_id=session_id))

@class_management_bp.route('/download_attendance_excel/<int:session_id>')
@login_required
def download_attendance_excel(session_id):
    """Generate and download an Excel report of the attendance."""
    try:
        # Import error handler for detailed logging
        from error_handler import log_error
        
        current_app.logger.info(f"Starting Excel generation for session {session_id}")
        
        # Check if required modules are available
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter
            current_app.logger.info("Required modules available for Excel")
        except ImportError as e:
            current_app.logger.error(f"Missing required module for Excel: {e}")
            flash(f'Missing required module for Excel: {e}', 'error')
            return redirect(url_for('class_management.index'))
            
        session = get_or_404_for_window(Session, session_id)
        students = _class_students_for_session(session_id)
        attendance_summary = _build_attendance_summary(session)
        combined_assessment_map = _collect_combined_assessment_marks(session)
        attendance_summary = _build_attendance_summary(session)
        combined_assessment_map = _collect_combined_assessment_marks(session)
        attendance_summary = _build_attendance_summary(session)
        all_attendance_records = ClassAttendance.query.filter_by(session_id=session_id).order_by(ClassAttendance.date, ClassAttendance.id).all()

        if not all_attendance_records:
            flash('No attendance data to download.', 'warning')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))

        # This logic is similar to view_attendance, consider refactoring in a real app
        attendance_by_date = defaultdict(list)
        for record in all_attendance_records:
            attendance_by_date[record.date].append(record)
        
        daily_class_counts = {}
        for date, records in attendance_by_date.items():
            student_records_on_date = defaultdict(list)
            for record in records:
                student_records_on_date[record.student_id].append(record)
            student_counts_on_date = {
                student_id: len(_records_by_slot(student_records))
                for student_id, student_records in student_records_on_date.items()
            }
            if student_counts_on_date:
                daily_class_counts[date] = _cap_classes_per_day(max(student_counts_on_date.values()))
                
        headers = []
        sorted_dates = sorted(daily_class_counts.keys())
        for dt in sorted_dates:
            count = daily_class_counts.get(dt, 0)
            if count == 1:
                headers.append(dt.strftime('%b %d, %Y'))
            else:
                for i in range(1, count + 1):
                    headers.append(f"{dt.strftime('%b %d, %Y')} ({i})")

        # Prepare structured data for Excel
        local_classes_held = sum(daily_class_counts.values())
        data_rows = []

        agg_student_map = attendance_summary.get('per_student', {})
        agg_total_classes = attendance_summary.get('total_classes', local_classes_held)

        for index, student in enumerate(students, start=1):
            student_attendance_records = [r for r in all_attendance_records if r.student_id == student.id]
            agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})
            present_count = agg_stats['present']
            percentage = agg_stats['percentage']
            marks = agg_stats['marks']
            
            # Initialize attendance row with placeholders
            attendance_statuses = ['-'] * len(headers)
            
            col_idx = 0
            for dt in sorted_dates:
                records_on_date = [r for r in student_attendance_records if r.date == dt]
                slot_map = _records_by_slot(records_on_date)
                num_classes_on_date = daily_class_counts.get(dt, 0)
                
                for class_num in range(num_classes_on_date):
                    record = slot_map.get(class_num + 1)
                    if record:
                        attendance_statuses[col_idx] = _attendance_label_from_record(record)
                    col_idx += 1

            data_rows.append({
                'serial': index,
                'student_id': student.student_id,
                'name': student.name,
                'attendance': attendance_statuses,
                'total_classes': agg_stats.get('effective_total_classes', agg_total_classes),
                'present': present_count,
                'percentage': percentage,
                'marks': marks
            })

        # Create workbook with styled layout
        wb = Workbook()
        ws = wb.active
        ws.title = 'Attendance Report'

        total_columns = 3 + len(headers) + 4  # SL + ID + Name + attendance + summary columns
        last_column_letter = get_column_letter(total_columns)

        # Title row
        ws.merge_cells(start_row=1, start_column=2, end_row=1, end_column=total_columns)
        title_cell = ws.cell(row=1, column=2)
        title_cell.value = 'Attendance Sheet'
        title_cell.font = Font(size=16, bold=True)
        title_cell.alignment = Alignment(horizontal='center', vertical='center')

        # Subject information
        subject_text = 'Subject: '
        if session.course_code and session.course_name:
            subject_text += f"{session.course_code} {session.course_name}"
        elif session.course_name:
            subject_text += session.course_name
        elif session.course_code:
            subject_text += session.course_code
        else:
            subject_text += 'N/A'

        ws.merge_cells(start_row=2, start_column=2, end_row=2, end_column=total_columns)
        subject_cell = ws.cell(row=2, column=2)
        subject_cell.value = subject_text
        subject_cell.font = Font(bold=True)
        subject_cell.alignment = Alignment(horizontal='left', vertical='center')

        # Additional metadata row
        teacher_name = None
        if session.teacher and getattr(session.teacher, 'name', None):
            teacher_name = session.teacher.name
        else:
            teacher_name = getattr(current_user, 'full_name', None) or getattr(current_user, 'username', 'N/A')

        metadata_items = [
            f"Year: {session.year or 'N/A'}",
            f"Term: {session.term or 'N/A'}",
            f"Session: {session.academic_session or 'N/A'}",
            f"Course Teacher: {teacher_name}"
        ]

        metadata_start_col = 2
        for item in metadata_items:
            if metadata_start_col > total_columns:
                break
            cell = ws.cell(row=3, column=metadata_start_col)
            cell.value = item
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='left', vertical='center')
            metadata_start_col += 4

        # Header row for table
        header_row = 5
        headers_for_sheet = ['Sl', 'Student ID', 'Name'] + headers + ['Total Classes', 'Present', 'Percentage', 'Marks']

        header_fill = PatternFill(start_color='D8E4BC', end_color='D8E4BC', fill_type='solid')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        for col_index, header_label in enumerate(headers_for_sheet, start=1):
            cell = ws.cell(row=header_row, column=col_index)
            cell.value = header_label
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=False, text_rotation=90, shrink_to_fit=False)
            cell.border = thin_border

        # Data rows
        data_start_row = header_row + 1
        attendance_start_col = 4
        percentage_col = attendance_start_col + len(headers) + 2  # Serial(1) + ID(2) + Name(3) + attendance + Total + Present => +2, Percentage column index
        marks_col = percentage_col + 1

        for row_offset, row_data in enumerate(data_rows):
            row_number = data_start_row + row_offset
            ws.cell(row=row_number, column=1, value=row_data['serial']).alignment = Alignment(horizontal='center')
            ws.cell(row=row_number, column=2, value=row_data['student_id']).alignment = Alignment(horizontal='center')
            ws.cell(row=row_number, column=3, value=row_data['name']).alignment = Alignment(horizontal='left')

            for idx, status in enumerate(row_data['attendance']):
                cell = ws.cell(row=row_number, column=attendance_start_col + idx, value=status)
                cell.alignment = Alignment(horizontal='center', vertical='center')

            total_col = attendance_start_col + len(headers)
            ws.cell(row=row_number, column=total_col, value=row_data['total_classes']).alignment = Alignment(horizontal='center')
            ws.cell(row=row_number, column=total_col + 1, value=row_data['present']).alignment = Alignment(horizontal='center')

            percentage_cell = ws.cell(row=row_number, column=percentage_col, value=(row_data['percentage'] / 100 if agg_total_classes else 0))
            percentage_cell.number_format = '0.00%'
            percentage_cell.alignment = Alignment(horizontal='center')

            ws.cell(row=row_number, column=marks_col, value=row_data['marks']).alignment = Alignment(horizontal='center')

        # Apply borders to data cells
        max_row = max(header_row, data_start_row + len(data_rows) - 1)
        for row in ws.iter_rows(min_row=header_row, max_row=max_row, min_col=1, max_col=total_columns):
            for cell in row:
                cell.border = thin_border

        # Auto fit all columns based on content - exact fit without extra space
        for col_idx in range(1, total_columns + 1):
            column_letter = get_column_letter(col_idx)
            max_length = 0
            
            # Check all cells in this column to find maximum content length
            for row in ws.iter_rows(min_col=col_idx, max_col=col_idx, values_only=False):
                for cell in row:
                    if cell.value is not None:
                        # Calculate exact length of cell content
                        cell_value = str(cell.value)
                        # Count actual character length
                        actual_length = len(cell_value)
                        max_length = max(max_length, actual_length)
            
            # Set column width to exact content width with minimal padding
            # openpyxl column width is in character units (approximate)
            if max_length > 0:
                # Very minimal padding: 0.5-1 character for readability
                # For narrow columns (like attendance P/A), use less padding
                if max_length <= 2:  # Single character columns (P, A, -)
                    adjusted_width = max_length + 0.5
                else:
                    adjusted_width = max_length + 1.0
            else:
                # If column is completely empty, set minimal width
                adjusted_width = 2
            
            ws.column_dimensions[column_letter].width = adjusted_width

        ws.freeze_panes = ws.cell(row=data_start_row, column=attendance_start_col)

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        current_app.logger.info(f"Excel generated successfully for session {session_id}")
        
        # Use Response instead of send_file for better cPanel compatibility
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={
                'Content-Disposition': f'attachment; filename="attendance_report_{session.course_name}_{date.today()}.xlsx"',
                'Content-Length': str(len(output.getvalue()))
            }
        )
        
    except Exception as e:
        # Log detailed error information
        log_error(e, {
            'session_id': session_id,
            'function': 'download_attendance_excel',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        
        current_app.logger.error(f"Error generating Excel for session {session_id}: {e}")
        flash(f'Error generating Excel: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/download_pdf_report/<int:session_id>')
@login_required
def download_pdf_report(session_id):
    """Generate and download a PDF summary report with headers and page numbers."""
    try:
        # Import error handler for detailed logging
        from error_handler import log_error
        from flask import Response
        
        current_app.logger.info(f"Starting PDF generation for session {session_id}")
        
        # Check if required modules are available
        try:
            import reportlab
            import pandas
            current_app.logger.info("Required modules available")
        except ImportError as e:
            current_app.logger.error(f"Missing required module: {e}")
            flash(f'Missing required module: {e}', 'error')
            return redirect(url_for('class_management.index'))
        session = get_or_404_for_window(Session, session_id)
        students = _class_students_for_session(session_id)
        combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session)
        # Keep attendance aggregation identical to on-screen views
        # so marks stay consistent between display and combined PDF.
        attendance_summary = _build_attendance_summary(session, include_archived=True)
        combined_assessment_map = _collect_combined_assessment_marks(session)
        combined_sessional_values, _combined_sessional_absent = (
            _build_combined_sessional_values(session)
            if session.course_type == 'sessional'
            else ({}, {})
        )

        buffer = io.BytesIO()
        
        # Define margins
        top_margin = 2.5 * inch
        bottom_margin = 0.7 * inch
        
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=0.5*inch, leftMargin=0.5*inch,
                                topMargin=top_margin, bottomMargin=bottom_margin)
        
        elements = []
        
        # --- Data Calculation ---
        student_data_for_pdf = []
        per_student_attendance = attendance_summary.get('per_student', {})
        
        def format_one_decimal(value):
            """Display marks with at most one decimal place (hide trailing .0)."""
            try:
                numeric_value = round(float(value or 0), 1)
                if numeric_value.is_integer():
                    return str(int(numeric_value))
                return f"{numeric_value:.1f}"
            except (TypeError, ValueError):
                return "0"
        
        for student in students:
            attendance_stats = per_student_attendance.get(student.student_id, {'marks': 0})
            attendance_marks = attendance_stats.get('marks', 0)

            if session.course_type == 'sessional':
                merged_values = combined_sessional_values.get(student.student_id, {})
                sessional_report = (
                    student.sessional_report
                    if student.sessional_report is not None
                    else merged_values.get('sessional_report')
                ) or 0
                sessional_viva = (
                    student.sessional_viva
                    if student.sessional_viva is not None
                    else merged_values.get('sessional_viva')
                ) or 0
                total_marks = attendance_marks + sessional_report + sessional_viva
                total_marks = _maybe_round_assessment_total(session, total_marks)
                student_data_for_pdf.append({
                    'id': student.student_id,
                    'attendance': format_one_decimal(attendance_marks),
                    'sessional_report': format_one_decimal(sessional_report),
                    'sessional_viva': format_one_decimal(sessional_viva),
                    'total': format_one_decimal(total_marks),
                })
            else:
                # For theory courses, use combined assessment
                assessment_marks_display = 0
                if session.course_type == 'theory':
                    if _is_external_theory_session(session):
                        display_total, _ = _get_external_assessment_display_total(
                            session,
                            student.student_id,
                            combined_values,
                            combined_best3=combined_best3,
                            combined_pg_total=combined_pg_total,
                        )
                        assessment_marks_display = display_total if display_total is not None else 0
                    elif session.category == 'pg':
                        combined = combined_values.get(student.student_id, {})
                        valid_marks = [v for v in combined.values() if v is not None]
                        valid_marks.sort(reverse=True)
                        if valid_marks:
                            best_three = take_best_marks(valid_marks)
                            scaled_pg_mark = scale_pg_total(sum(best_three))
                            # PG combined PDF requires integer rounding with .5 always rounding up.
                            assessment_marks_display = int(
                                Decimal(str(scaled_pg_mark)).quantize(
                                    Decimal('1'),
                                    rounding=ROUND_HALF_UP
                                )
                            )
                        else:
                            assessment_marks_display = 0
                    else:
                        # UG: combined_best3 already respects the round-total toggle
                        ug_total = combined_best3.get(student.student_id) or 0
                        assessment_marks_display = ug_total if ug_total else 0

                student_data_for_pdf.append({
                    'id': student.student_id,
                    'attendance': format_one_decimal(attendance_marks),
                    'assessment': format_one_decimal(assessment_marks_display)
                })

        # --- Table Creation ---
        if session.course_type == 'sessional':
            sess = result_split('sessional')
            table_data = [[
                'ID',
                f'Attendance ({sess["attendance"]})',
                f'Sessional Report ({sess["report"]})',
                f'Sessional Viva ({sess["viva"]})',
                'Total Marks (100)',
            ]]
            for s_data in student_data_for_pdf:
                table_data.append([
                    s_data['id'],
                    s_data['attendance'],
                    s_data['sessional_report'],
                    s_data['sessional_viva'],
                    s_data['total'],
                ])
            table = Table(
                table_data,
                colWidths=[1.1 * inch, 1.3 * inch, 1.6 * inch, 1.6 * inch, 1.4 * inch],
                repeatRows=0,
            )
        else:
            # For theory courses: single assessment column
            assessment_header_text = _combined_pdf_assessment_header(session)
            
            table_data = [['ID', f'Attendance ({result_split("theory_ug")["attendance"]})', assessment_header_text]]
            for s_data in student_data_for_pdf:
                table_data.append([s_data['id'], s_data['attendance'], s_data['assessment']])
            table = Table(table_data, colWidths=[2*inch, 2*inch, 2.5*inch], repeatRows=0)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F4F4F')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F2F2F2')),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
        ]))
        
        elements.append(table)
        
        # --- PDF Build with header/footer ---
        def _draw_header_footer(canvas_obj, doc_obj, include_header=False):
            canvas_obj.saveState()
            width, height = doc_obj.pagesize
            if include_header:
                canvas_obj.setFont('Helvetica-Bold', 18)
                canvas_obj.drawCentredString(width / 2.0, height - 1.0 * inch, current_tenant().university_name)

                canvas_obj.setFont('Helvetica', 12)
                canvas_obj.drawCentredString(width / 2.0, height - 1.35 * inch, current_tenant().name)
                canvas_obj.drawCentredString(width / 2.0, height - 1.60 * inch, _combined_pdf_title(session))

                canvas_obj.setFont('Helvetica-Bold', 10)
                course_info = f"Course: {session.course_name} ({session.course_code or 'N/A'})  |  Type: {session.course_type.capitalize()}"
                year_term_info = f"Year: {session.year}, Term: {session.term}  |  Session: {session.academic_session}"
                canvas_obj.drawCentredString(width / 2.0, height - 1.95 * inch, course_info)
                canvas_obj.drawCentredString(width / 2.0, height - 2.15 * inch, year_term_info)

            canvas_obj.setFont('Helvetica', 9)
            page_text = f"Page {doc_obj.page}"
            canvas_obj.drawRightString(width - 0.5 * inch, 0.5 * inch, page_text)
            canvas_obj.restoreState()

        def first_page(canvas_obj, doc_obj):
            _draw_header_footer(canvas_obj, doc_obj, include_header=True)

        def later_pages(canvas_obj, doc_obj):
            _draw_header_footer(canvas_obj, doc_obj, include_header=False)

        doc.build(elements, onFirstPage=first_page, onLaterPages=later_pages)
        
        buffer.seek(0)
        filename = f"Report_{session.course_name}_{session.year}.pdf"
        
        current_app.logger.info(f"PDF generated successfully for session {session_id}")
        
        # Use Response instead of send_file for better cPanel compatibility
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
    except Exception as e:
        # Log detailed error information
        log_error(e, {
            'session_id': session_id,
            'function': 'download_pdf_report',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        
        current_app.logger.error(f"Error generating PDF for session {session_id}: {e}")
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))


@class_management_bp.route('/download_attendance_sheet/<int:session_id>')
@login_required
def download_attendance_sheet(session_id):
    try:
        from error_handler import log_error
        from reportlab.lib.pagesizes import landscape, letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        session = get_or_404_for_window(Session, session_id)
        attendance_summary = _build_attendance_summary(session, include_archived=True)

        related_sessions = _resolve_attendance_related_sessions(session, include_archived=True)
        session_ids = [s.id for s in related_sessions if s]
        attendance_records = ClassAttendance.query.filter(
            ClassAttendance.session_id.in_(session_ids)
        ).order_by(ClassAttendance.date.asc(), ClassAttendance.id.asc()).all()

        all_students = _class_students_for_sessions(session_ids)
        class_student_to_public = {stu.id: stu.student_id for stu in all_students}
        students_by_public = {}
        for stu in all_students:
            selected = students_by_public.get(stu.student_id)
            if selected is None:
                students_by_public[stu.student_id] = stu
            elif selected.session_id != session_id and stu.session_id == session_id:
                # Prefer the current session row for display metadata/order.
                students_by_public[stu.student_id] = stu
        students = sorted(students_by_public.values(), key=lambda s: s.student_id)

        attendance_by_date = defaultdict(list)
        for record in attendance_records:
            public_id = class_student_to_public.get(record.student_id)
            if not public_id:
                continue
            attendance_by_date[record.date].append((public_id, record))

        daily_class_counts = {}
        for date, records in attendance_by_date.items():
            student_records_on_date = defaultdict(list)
            for public_id, record in records:
                student_records_on_date[public_id].append(record)
            student_counts = {
                public_id: len(_records_by_slot(student_records))
                for public_id, student_records in student_records_on_date.items()
            }
            daily_class_counts[date] = _cap_classes_per_day(max(student_counts.values()) if student_counts else 0)

        sorted_dates = sorted(daily_class_counts.keys())
        headers = []
        header_keys = []
        for date in sorted_dates:
            count = daily_class_counts.get(date, 0)
            day_label = date.strftime('%d %b')
            if count <= 1:
                headers.append(day_label)
                header_keys.append((date, 1))
            else:
                for i in range(1, count + 1):
                    headers.append(f"{day_label}-{i}")
                    header_keys.append((date, i))

        if not headers:
            flash('No attendance data to download.', 'warning')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))

        student_records_by_public = defaultdict(list)
        for record in attendance_records:
            public_id = class_student_to_public.get(record.student_id)
            if public_id:
                student_records_by_public[public_id].append(record)
        for public_id in student_records_by_public:
            student_records_by_public[public_id].sort(key=lambda x: (x.date, x.id))

        # Prepare data for template
        data_rows = []
        agg_student_map = attendance_summary.get('per_student', {})
        total_classes = attendance_summary.get('total_classes', sum(daily_class_counts.values()))

        for idx, student in enumerate(students, start=1):
            student_records = student_records_by_public.get(student.student_id, [])
            student_attendance_by_date = defaultdict(list)
            for r in student_records:
                student_attendance_by_date[r.date].append(r)
            attendance_list = []
            for date, slot in header_keys:
                slot_map = _records_by_slot(student_attendance_by_date[date])
                record = slot_map.get(slot)
                if record:
                    attendance_list.append(_attendance_label_from_record(record))
                else:
                    attendance_list.append('-')
            agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})
            data_rows.append({
                'idx': idx,
                'student_id': str(student.student_id),
                'name': student.name,
                'attendance': attendance_list,
                'total_classes': str(agg_stats.get('effective_total_classes', total_classes)),
                'present_days': str(agg_stats['present']),
                'percentage': f"{agg_stats['percentage']:.2f}",
                'marks': str(agg_stats['marks'])
            })

        # Format course scope label
        scope_label = 'Full'
        if session.course_scope == 'part_a':
            scope_label = 'Part A'
        elif session.course_scope == 'part_b':
            scope_label = 'Part B'
        elif session.course_scope:
            scope_label = session.course_scope.replace('_', ' ').title()

        buffer = io.BytesIO()
        # Landscape orientation with minimal margins for maximum space
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=landscape(letter), 
            leftMargin=0.15*inch, 
            rightMargin=0.15*inch, 
            topMargin=0.2*inch, 
            bottomMargin=0.15*inch
        )
        styles = getSampleStyleSheet()
        elements = []

        # Create custom centered styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=14,
            textColor=colors.HexColor('#000000'),
            spaceAfter=8,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        meta_style = ParagraphStyle(
            'CustomMeta',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            spaceAfter=4,
            alignment=TA_LEFT,
            fontName='Helvetica',
            leftIndent=5
        )

        # Build header text - Excel style
        header_text = "Attendance Sheet"
        
        # Build metadata - Excel style (separate lines)
        course_code = session.course_code or ''
        course_name = session.course_name or ''
        if course_code and course_name:
            subject_text = f"Subject: {course_code} {course_name}"
        elif course_name:
            subject_text = f"Subject: {course_name}"
        elif course_code:
            subject_text = f"Subject: {course_code}"
        else:
            subject_text = "Subject: N/A"
        
        teacher_name = session.teacher.name if session.teacher else 'N/A'
        year_text = f"Year: {session.year}" if session.year else "Year: N/A"
        term_text = f"Term: {session.term}" if session.term else "Term: N/A"
        
        session_text = ""
        if session.academic_session:
            session_text = f"Session: {session.academic_session}"
        session_text += f" Course Teacher: {teacher_name}" if session_text else f"Course Teacher: {teacher_name}"

        elements.append(Spacer(1, 5))
        elements.append(Paragraph(header_text, title_style))
        elements.append(Paragraph(subject_text, meta_style))
        elements.append(Paragraph(year_text, meta_style))
        elements.append(Paragraph(term_text, meta_style))
        elements.append(Paragraph(session_text, meta_style))
        elements.append(Spacer(1, 5))

        # Create styles for text wrapping
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['Normal'],
            fontSize=6,
            leading=7,
            alignment=0,  # LEFT
            wordWrap='CJK'
        )
        
        # Header row
        header_row = ['SI', 'Student ID', 'Name']
        for h in headers:
            header_row.append(h)
        header_row.extend(['Total Classes', 'Present', 'Percentage', 'Marks'])
        
        # Convert name column to Paragraph objects for text wrapping
        table_data = [header_row]
        for row in data_rows:
            wrapped_row = [str(row['idx']), str(row['student_id']), Paragraph(row['name'], name_style)]
            wrapped_row.extend(row['attendance'])
            wrapped_row.extend([row['total_classes'], row['present_days'], f"{row['percentage']}%", row['marks']])
            table_data.append(wrapped_row)

        # Calculate dynamic column widths - optimized for ONE page with rotated headers
        available_width = doc.width
        si_width = 0.15*inch
        student_id_width = 0.35*inch
        name_width = 0.5*inch
        summary_col_width = 0.3*inch
        summary_total_width = summary_col_width * 4

        # Calculate remaining width for attendance columns
        fixed_width = si_width + student_id_width + name_width + summary_total_width
        remaining_width = max(0.15*inch, available_width - fixed_width)
        
        # Distribute remaining width among attendance columns - narrow for rotated headers
        if len(headers) > 0:
            attendance_col_width = max(0.1*inch, remaining_width / len(headers))
        else:
            attendance_col_width = 0.1*inch

        column_widths = (
            [si_width, student_id_width, name_width] +
            [attendance_col_width] * len(headers) +
            [summary_col_width] * 4
        )

        # Ensure total width fits within available width
        total_width = sum(column_widths)
        if total_width > available_width * 0.95:
            scale_factor = (available_width * 0.95) / total_width
            column_widths = [w * scale_factor for w in column_widths]
        
        table = Table(table_data, repeatRows=1, colWidths=column_widths, hAlign='CENTER')
        
        # Build table style - optimized for ONE page with rotated headers
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D8E4BC')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTSIZE', (0, 0), (-1, 0), 6),  # Header font
            ('FONTSIZE', (0, 1), (-1, -1), 5),  # Data font
            ('ALIGN', (2, 1), (2, -1), 'LEFT'),
            ('LEFTPADDING', (0, 0), (-1, -1), 1),
            ('RIGHTPADDING', (0, 0), (-1, -1), 1),
            ('TOPPADDING', (0, 0), (-1, -1), 1),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
        ])
        
        # Rotate attendance header columns (-90 degrees) to save horizontal space
        if len(headers) > 0:
            attendance_start_col = 3
            attendance_end_col = 3 + len(headers) - 1
            table_style.add('TEXTROTATION', (attendance_start_col, 0), (attendance_end_col, 0), -90)
            table_style.add('FONTSIZE', (attendance_start_col, 0), (attendance_end_col, 0), 5.5)
            table_style.add('TOPPADDING', (attendance_start_col, 0), (attendance_end_col, 0), 50)
            table_style.add('BOTTOMPADDING', (attendance_start_col, 0), (attendance_end_col, 0), 50)
            table_style.add('LEFTPADDING', (attendance_start_col, 0), (attendance_end_col, 0), 1)
            table_style.add('RIGHTPADDING', (attendance_start_col, 0), (attendance_end_col, 0), 1)
        
        table.setStyle(table_style)
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)

        filename = f"attendance_sheet_{session.course_code or session.id}.pdf"
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename=\"{filename}\"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
    except Exception as e:
        log_error(e, {
            'session_id': session_id,
            'function': 'download_attendance_sheet',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        current_app.logger.error(f"Error generating attendance sheet PDF for session {session_id}: {e}")
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.view_attendance', session_id=session_id))

@class_management_bp.route('/download_attendance_sheet_weasyprint/<int:session_id>')
@login_required
def download_attendance_sheet_weasyprint(session_id):
    """Generate attendance sheet PDF using WeasyPrint in legal landscape format."""
    # Lazy import WeasyPrint - only when actually needed
    HTML = _get_weasyprint_html()
    if HTML is None:
        from error_handler import log_error
        error_msg = 'Error generating PDF: WeasyPrint is not available. '
        error_msg += 'Please ensure WeasyPrint dependencies are installed. '
        error_msg += 'On macOS, run: brew install cairo pango gdk-pixbuf gobject-introspection'
        flash(error_msg, 'error')
        current_app.logger.error("WeasyPrint not available for PDF generation")
        current_app.logger.error(f"Current availability status: {_WEASYPRINT_AVAILABLE}")
        return redirect(url_for('class_management.view_attendance', session_id=session_id))
    
    try:
        from error_handler import log_error
        
        session = get_or_404_for_window(Session, session_id)
        attendance_summary = _build_attendance_summary(session, include_archived=True)
        
        # Get related sessions for split courses
        related_sessions = _resolve_attendance_related_sessions(session, include_archived=True)
        session_ids = [s.id for s in related_sessions if s]
        attendance_records = ClassAttendance.query.filter(
            ClassAttendance.session_id.in_(session_ids)
        ).order_by(ClassAttendance.date.asc(), ClassAttendance.id.asc()).all()

        all_students = _class_students_for_sessions(session_ids)
        class_student_to_public = {stu.id: stu.student_id for stu in all_students}
        students_by_public = {}
        for stu in all_students:
            selected = students_by_public.get(stu.student_id)
            if selected is None:
                students_by_public[stu.student_id] = stu
            elif selected.session_id != session_id and stu.session_id == session_id:
                students_by_public[stu.student_id] = stu
        students = sorted(students_by_public.values(), key=lambda s: s.student_id)
        
        attendance_by_date = defaultdict(list)
        for record in attendance_records:
            public_id = class_student_to_public.get(record.student_id)
            if not public_id:
                continue
            attendance_by_date[record.date].append((public_id, record))
        
        daily_class_counts = {}
        for date, records in attendance_by_date.items():
            student_records_on_date = defaultdict(list)
            for public_id, record in records:
                student_records_on_date[public_id].append(record)
            student_counts = {
                public_id: len(_records_by_slot(student_records))
                for public_id, student_records in student_records_on_date.items()
            }
            daily_class_counts[date] = _cap_classes_per_day(max(student_counts.values()) if student_counts else 0)
        
        sorted_dates = sorted(daily_class_counts.keys())
        headers = []
        header_keys = []
        for date in sorted_dates:
            count = daily_class_counts.get(date, 0)
            # Include year; compact for vertical SVG headers: "19 Jul 2024" / "19 Jul 2024·1"
            day_label = f"{date.day} {date.strftime('%b')} {date.year}"
            if count <= 1:
                headers.append(day_label)
                header_keys.append((date, 1))
            else:
                for i in range(1, count + 1):
                    headers.append(f"{day_label}·{i}")
                    header_keys.append((date, i))
        
        if not headers:
            flash('No attendance data to download.', 'warning')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))
        
        student_records_by_public = defaultdict(list)
        for record in attendance_records:
            public_id = class_student_to_public.get(record.student_id)
            if public_id:
                student_records_by_public[public_id].append(record)
        for public_id in student_records_by_public:
            student_records_by_public[public_id].sort(key=lambda x: (x.date, x.id))
        
        # Prepare data for template
        data_rows = []
        agg_student_map = attendance_summary.get('per_student', {})
        total_classes = attendance_summary.get('total_classes', sum(daily_class_counts.values()))
        
        for idx, student in enumerate(students, start=1):
            student_records = student_records_by_public.get(student.student_id, [])
            
            student_attendance_by_date = defaultdict(list)
            for r in student_records:
                student_attendance_by_date[r.date].append(r)
            
            # Sort records by id for each date to ensure correct slot order
            for date in student_attendance_by_date:
                student_attendance_by_date[date].sort(key=lambda x: x.id)
            
            attendance_list = []
            for date, slot in header_keys:
                slot_map = _records_by_slot(student_attendance_by_date.get(date, []))
                record = slot_map.get(slot)
                if record:
                    attendance_list.append(_attendance_label_from_record(record))
                else:
                    attendance_list.append('-')
            agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})
            data_rows.append({
                'idx': idx,
                'student_id': str(student.student_id),
                'name': student.name,
                'attendance': attendance_list,
                'total_classes': str(agg_stats.get('effective_total_classes', total_classes)),
                'present_days': str(agg_stats['present']),
                'percentage': f"{agg_stats['percentage']:.1f}",
                'marks': str(agg_stats['marks'])
            })
        
        # Get assignment data for Session, Year, Term, Course Teacher
        assignment = None
        if CourseSessionAssignment:
            assignment = CourseSessionAssignment.query.filter_by(session_id=session_id).first()
            # If not found by session_id, try to find by course_code, teacher_id, year, term
            if not assignment and session.course_code and session.teacher_id and session.year and session.term:
                try:
                    if Course:
                        assignment = CourseSessionAssignment.query.filter_by(
                            teacher_id=session.teacher_id,
                            year=session.year,
                            term=session.term
                        ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                            Course.course_code == session.course_code
                        ).first()
                except Exception as query_error:
                    current_app.logger.warning(f'Error querying CourseSessionAssignment: {query_error}')
        
        # Get course information
        course_code = session.course_code or ''
        course_name = session.course_name or ''
        
        # Get session, year, term, course teacher
        academic_session = ''
        if assignment and assignment.academic_session:
            academic_session = assignment.academic_session
        elif session.academic_session:
            academic_session = session.academic_session
        
        year = session.year or ''
        term = session.term or ''
        
        course_teacher = 'N/A'
        if assignment and assignment.teacher:
            course_teacher = assignment.teacher.name
        elif session.teacher:
            course_teacher = session.teacher.name
        
        # Layout scales with date-column count so the sheet always fits legal landscape
        n_dates = max(len(headers), 1)
        if n_dates <= 10:
            date_font_size = 9
            pdf_layout = {
                'body_font_size': 9,
                'header_font_size': 8,
                'date_font_size': date_font_size,
                'mark_font_size': 9,
                'name_font_size': 8.5,
                'summary_font_size': 8.5,
                'summary_label_size': 7.5,
                'si_width': '3%',
                'id_width': '7%',
                'name_col_width': '16%',
                'date_col_width': f'{52 / n_dates:.2f}%',
                'summary_width': '5.5%',
            }
        elif n_dates <= 20:
            date_font_size = 8
            pdf_layout = {
                'body_font_size': 8,
                'header_font_size': 7.5,
                'date_font_size': date_font_size,
                'mark_font_size': 8,
                'name_font_size': 7.5,
                'summary_font_size': 7.5,
                'summary_label_size': 7,
                'si_width': '2.5%',
                'id_width': '6.5%',
                'name_col_width': '14%',
                'date_col_width': f'{56 / n_dates:.2f}%',
                'summary_width': '5.25%',
            }
        else:
            date_font_size = 7
            pdf_layout = {
                'body_font_size': 7,
                'header_font_size': 7,
                'date_font_size': date_font_size,
                'mark_font_size': 7.5,
                'name_font_size': 7,
                'summary_font_size': 7,
                'summary_label_size': 6.5,
                'si_width': '2.2%',
                'id_width': '6%',
                'name_col_width': '12%',
                'date_col_width': f'{60 / n_dates:.2f}%',
                'summary_width': '4.95%',
            }

        # Absolute file:// font URIs — required on cPanel where system fonts are Type1-only
        from utils.pdf_fonts import resolve_dejavu_pdf_fonts
        pdf_font_regular, pdf_font_bold, fonts_dir = resolve_dejavu_pdf_fonts()
        if not pdf_font_regular or not pdf_font_bold or not fonts_dir:
            flash(
                'PDF fonts missing. Upload DejaVuSans.ttf and DejaVuSans-Bold.ttf to static/fonts/.',
                'error'
            )
            return redirect(url_for('class_management.view_attendance', session_id=session_id))

        # Pre-render Rotate Text Up headers as PNGs (CSS/SVG rotate fails in WeasyPrint)
        bold_font_path = fonts_dir / 'DejaVuSans-Bold.ttf'
        header_images = []
        max_header_h = 0
        for label in headers:
            src, w, h = _attendance_rotate_up_header_png(
                label,
                bold_font_path,
                font_size=date_font_size,
                scale=4,
            )
            header_images.append({'src': src, 'width': w, 'height': h, 'label': label})
            if h > max_header_h:
                max_header_h = h
        pdf_layout['date_header_height'] = max(max_header_h + 4, 64)

        # Render template
        html_content = render_template(
            'class_management/attendance_sheet_weasyprint.html',
            course_code=course_code,
            course_name=course_name,
            academic_session=academic_session,
            year=year,
            term=term,
            course_teacher=course_teacher,
            headers=headers,
            header_images=header_images,
            data_rows=data_rows,
            pdf_font_regular=pdf_font_regular,
            pdf_font_bold=pdf_font_bold,
            **pdf_layout,
        )
        
        # Generate PDF with WeasyPrint (lazy import already done above)
        try:
            pdf_buffer = io.BytesIO()
            HTML(string=html_content, base_url=fonts_dir.as_uri() + '/').write_pdf(pdf_buffer)
            pdf_buffer.seek(0)
        except Exception as e:
            current_app.logger.error(f"Error generating PDF with WeasyPrint: {e}", exc_info=True)
            flash(f'Error generating PDF: {str(e)}', 'error')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))
        
        filename = f"attendance_sheet_{course_code or session.id}.pdf"
        
        current_app.logger.info(f"WeasyPrint PDF generated successfully for session {session_id}")
        return Response(
            pdf_buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(pdf_buffer.getvalue()))
            }
        )
        
    except Exception as e:
        log_error(e, {
            'session_id': session_id,
            'function': 'download_attendance_sheet_weasyprint',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        current_app.logger.error(f"Error generating WeasyPrint attendance sheet PDF for session {session_id}: {e}", exc_info=True)
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.view_attendance', session_id=session_id))

@class_management_bp.route('/assessment/<int:session_id>', methods=['GET', 'POST'])
@login_required
def assessment(session_id):
    """Assessment management for a session"""
    try:
        session = get_or_404_for_window(Session, session_id)

        if not user_owns_class_session(current_user, session):
            flash('You are not authorized to manage assessment for this session.', 'danger')
            return redirect(url_for('class_management.index'))

        students = _class_students_for_session(session_id)
        is_split_theory = session.course_type == 'theory' and session.course_scope in SPLIT_PARTS
        editable_indices = _get_editable_assessment_indices(session)
        editable_sessional_fields = _get_editable_sessional_fields(session)
        current_teacher = _ensure_current_teacher()
        
        # Load reveal status
        import json
        reveal_status = {}
        if session.assessment_revealed:
            try:
                reveal_status = json.loads(session.assessment_revealed)
            except:
                reveal_status = {}
        current_teacher_reveals = reveal_status.get(str(current_teacher.id), {})
        
        # Default attendance marks to revealed if not set
        if 'attendance' not in current_teacher_reveals:
            current_teacher_reveals['attendance'] = True
            # Save the default if session doesn't have reveal status yet
            if str(current_teacher.id) not in reveal_status:
                reveal_status[str(current_teacher.id)] = current_teacher_reveals
                session.assessment_revealed = json.dumps(reveal_status)
                db.session.commit()
        
        if request.method == 'POST':
            try:
                import json
                if session.course_type == 'theory':
                    for student in students:
                        # Load existing absent status
                        absent_status = {}
                        if student.assessment_absent:
                            try:
                                absent_status = json.loads(student.assessment_absent)
                            except:
                                absent_status = {}
                        
                        for i in range(1, 5):
                            if i in editable_indices:
                                absent_key = f'absent_{i}_{student.id}'
                                is_absent = absent_key in request.form
                                absent_status[f'assessment{i}'] = is_absent
                                
                                if is_absent:
                                    setattr(student, f'assessment{i}', None)
                                else:
                                    value = request.form.get(f'assessment{i}_{student.id}')
                                    if _is_external_theory_session(session):
                                        setattr(student, f'assessment{i}', _parse_external_assessment_value(value, session))
                                    else:
                                        setattr(student, f'assessment{i}', float(value) if value else None)
                        
                        # Save absent status
                        student.assessment_absent = json.dumps(absent_status) if absent_status else None
                    _recalculate_assessment_totals(session)

                elif session.course_type == 'sessional' and session.category == 'ug':
                    for student in students:
                        # Load existing absent status
                        absent_status = {}
                        if student.assessment_absent:
                            try:
                                absent_status = json.loads(student.assessment_absent)
                            except:
                                absent_status = {}
                        
                        report_absent_key = f'sessional_absent_report_{student.id}'
                        viva_absent_key = f'sessional_absent_viva_{student.id}'

                        if 'sessional_report' in editable_sessional_fields:
                            report_absent = report_absent_key in request.form
                            absent_status['sessional_report'] = report_absent

                            if report_absent:
                                student.sessional_report = None
                            else:
                                report = request.form.get(f'sessional_report_{student.id}')
                                student.sessional_report = float(report) if report else None

                        if 'sessional_viva' in editable_sessional_fields:
                            viva_absent = viva_absent_key in request.form
                            absent_status['sessional_viva'] = viva_absent

                            if viva_absent:
                                student.sessional_viva = None
                            else:
                                viva = request.form.get(f'sessional_viva_{student.id}')
                                student.sessional_viva = float(viva) if viva else None
                        
                        # Save absent status
                        student.assessment_absent = json.dumps(absent_status) if absent_status else None
                else:
                    flash('Unsupported course type for assessment entry.', 'error')
                    return redirect(url_for('class_management.assessment', session_id=session_id))


                db.session.commit()
                flash('Assessment marks saved successfully!', 'success')
                return redirect(url_for('class_management.assessment', session_id=session_id))
                
            except Exception as e:
                db.session.rollback()
                flash(f'Error saving assessment: {str(e)}', 'error')
                return redirect(url_for('class_management.assessment', session_id=session_id))
        
        # Build combined assessment values from all related sessions (for split courses)
        combined_assessment_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session)
        combined_sessional_values, combined_sessional_absent = _build_combined_sessional_values(session)

        labels = _assessment_column_labels()
        external_assessment_mode = DEFAULT_EXTERNAL_ASSESSMENT_MODE
        external_total_column_label = labels[DEFAULT_EXTERNAL_ASSESSMENT_MODE]
        combined_external_total = {}
        if _is_external_theory_session(session):
            external_assessment_mode = _normalize_external_assessment_mode(getattr(session, 'external_assessment_mode', None))
            external_total_column_label = labels.get(
                external_assessment_mode,
                labels[DEFAULT_EXTERNAL_ASSESSMENT_MODE],
            )
            for student in students:
                combined_vals = combined_assessment_values.get(student.student_id, {})
                result = _compute_external_assessment_total(combined_vals, external_assessment_mode)
                combined_external_total[student.student_id] = result['display_total']
        
        # Debug logging for split courses
        if session.split_group_id:
            current_app.logger.info(f"Assessment page - Session {session_id}: split_group_id={session.split_group_id}, course_scope={session.course_scope}, editable_indices={editable_indices}")
            if students:
                first_student = students[0]
                first_vals = combined_assessment_values.get(first_student.student_id, {})
                current_app.logger.info(f"First student {first_student.student_id} combined values: A1={first_vals.get(1)}, A2={first_vals.get(2)}, A3={first_vals.get(3)}, A4={first_vals.get(4)}")
        
        # Build absent status map for template - combine from all related sessions for split courses
        import json
        absent_status_map = {}
        
        # Get all related sessions for split courses
        related_sessions, student_map = _gather_split_student_map(session)
        all_student_records = []
        
        # Collect all student records from related sessions
        for student_id, entries in student_map.items():
            for entry in entries:
                all_student_records.append(entry)
        
        # Build absent status map combining from all sessions
        sessional_display_map = {}
        for student in students:
            absent_status = {}
            # Start with current student's absent status
            if student.assessment_absent:
                try:
                    absent_status = json.loads(student.assessment_absent)
                except:
                    absent_status = {}
            
            # For split courses, also check absent status from related sessions for the same student_id
            if session.split_group_id:
                for entry in all_student_records:
                    if entry.student_id == student.student_id and entry.id != student.id:
                        if entry.assessment_absent:
                            try:
                                other_absent = json.loads(entry.assessment_absent)
                                # Merge absent status (if any assessment is absent in any session, mark as absent)
                                for key, value in other_absent.items():
                                    if key not in absent_status or not absent_status.get(key):
                                        absent_status[key] = value
                            except:
                                pass
            
            absent_status_map[student.id] = absent_status

            if session.course_type == 'sessional' and session.category == 'ug':
                merged_values = combined_sessional_values.get(student.student_id, {})
                merged_absent = combined_sessional_absent.get(student.student_id, {})

                report_value = student.sessional_report if student.sessional_report is not None else merged_values.get('sessional_report')
                viva_value = student.sessional_viva if student.sessional_viva is not None else merged_values.get('sessional_viva')

                report_absent = bool(absent_status.get('sessional_report', False)) or bool(merged_absent.get('sessional_report', False))
                viva_absent = bool(absent_status.get('sessional_viva', False)) or bool(merged_absent.get('sessional_viva', False))

                sessional_display_map[student.id] = {
                    'report_value': report_value,
                    'viva_value': viva_value,
                    'report_absent': report_absent,
                    'viva_absent': viva_absent,
                    'report_editable': 'sessional_report' in editable_sessional_fields,
                    'viva_editable': 'sessional_viva' in editable_sessional_fields
                }
        
        return render_template(
            'class_management/assessment.html',
            session=session,
            students=students,
            split_meta=_build_split_context(session),
            editable_indices=editable_indices,
            combined_assessment_values=combined_assessment_values,
            combined_best3=combined_best3,
            combined_pg_avg=combined_pg_avg,
            combined_pg_total=combined_pg_total,
            combined_external_total=combined_external_total,
            external_assessment_mode=external_assessment_mode,
            external_assessment_modes=_external_assessment_modes(),
            external_total_column_label=external_total_column_label,
            sessional_display_map=sessional_display_map,
            absent_status_map=absent_status_map,
            current_teacher_id=current_teacher.id,
            reveal_status=current_teacher_reveals,
            assessment_slot_count=len(editable_indices),
            round_assessment_total=_session_rounds_assessment_total(session),
        )
        
    except Exception as e:
        flash(f'Error loading assessment page: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/assessment/<int:session_id>/toggle-reveal', methods=['POST'])
@login_required
def toggle_assessment_reveal(session_id):
    """Toggle reveal status for assessment scores"""
    try:
        import json
        session = get_or_404_for_window(Session, session_id)
        current_teacher = _ensure_current_teacher()
        
        # Ensure current teacher owns this session or is part of split course
        if session.teacher_id != current_teacher.id:
            # Check if it's a split course and teacher is partner
            split_meta = _build_split_context(session)
            if not split_meta or not split_meta.peers:
                return jsonify({'success': False, 'message': 'Unauthorized'}), 403
            partner_ids = [peer.teacher_id for peer in split_meta.peers]
            if current_teacher.id not in partner_ids:
                return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json()
        assessment_type = data.get('assessment_type')
        revealed = data.get('revealed', False)
        teacher_id = data.get('teacher_id', current_teacher.id)
        
        # Load existing reveal status
        reveal_status = {}
        if session.assessment_revealed:
            try:
                reveal_status = json.loads(session.assessment_revealed)
            except:
                reveal_status = {}
        
        # Initialize teacher's reveal status if not exists
        teacher_key = str(teacher_id)
        if teacher_key not in reveal_status:
            reveal_status[teacher_key] = {}
        
        # Update reveal status for this assessment type
        reveal_status[teacher_key][assessment_type] = revealed
        
        # Save to database
        session.assessment_revealed = json.dumps(reveal_status)
        db.session.commit()
        # Notify students when marks are revealed
        if revealed:
            course_label = session.course_name or session.course_code or 'Course'
            title = f'Assessment marks revealed: {course_label}'
            link_url = url_for('class_management.student_view_scores')
            try:
                _notify_students_in_session(session_id, 'marks_revealed', title, link_url)
            except Exception as notif_e:
                current_app.logger.warning(f"Student in-app notification (marks revealed): {notif_e}")
            try:
                _send_marks_revealed_email_to_session_students(
                    session_id=session_id,
                    course_label=course_label,
                    assessment_type=assessment_type
                )
            except Exception as mail_e:
                current_app.logger.error(
                    f"Student email (marks revealed) failed: {mail_e}",
                    exc_info=True,
                )
        return jsonify({'success': True, 'message': 'Reveal status updated'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@class_management_bp.route('/assessment/<int:session_id>/toggle-round-total', methods=['POST'])
@login_required
def toggle_round_assessment_total(session_id):
    """Enable/disable whole-number rounding for assessment totals (UI + PDF/Excel)."""
    try:
        session = get_or_404_for_window(Session, session_id)
        current_teacher = _ensure_current_teacher()
        if session.teacher_id != current_teacher.id:
            split_meta = _build_split_context(session)
            if not split_meta or not split_meta.peers:
                return jsonify({'success': False, 'message': 'Unauthorized'}), 403
            partner_ids = [peer.teacher_id for peer in split_meta.peers]
            if current_teacher.id not in partner_ids:
                return jsonify({'success': False, 'message': 'Unauthorized'}), 403

        data = request.get_json(silent=True) or {}
        enabled = bool(data.get('enabled'))
        _sync_round_assessment_total(session, enabled)
        if session.course_type == 'theory':
            _recalculate_assessment_totals(session)
        db.session.commit()

        students = _class_students_for_session(session_id)
        _, combined_best3, _, combined_pg_total = _build_combined_assessment_values(session)
        totals = {}
        for student in students:
            if session.course_type == 'sessional' and session.category == 'ug':
                report = student.sessional_report or 0
                viva = student.sessional_viva or 0
                totals[str(student.id)] = _maybe_round_assessment_total(session, report + viva)
            elif _is_external_theory_session(session):
                mode = _normalize_external_assessment_mode(getattr(session, 'external_assessment_mode', None))
                if mode == 'best_three_40':
                    totals[str(student.id)] = combined_pg_total.get(student.student_id)
                else:
                    totals[str(student.id)] = combined_best3.get(student.student_id)
            elif session.category == 'pg':
                totals[str(student.id)] = combined_pg_total.get(student.student_id)
            else:
                totals[str(student.id)] = combined_best3.get(student.student_id)

        return jsonify({
            'success': True,
            'enabled': enabled,
            'totals': totals,
            'message': 'Round total setting updated',
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


@class_management_bp.route('/assessment/<int:session_id>/set-external-mode', methods=['POST'])
@login_required
def set_external_assessment_mode(session_id):
    """Save external course assessment total calculation mode."""
    try:
        session = get_or_404_for_window(Session, session_id)
        if not _is_external_theory_session(session):
            return jsonify({'success': False, 'message': 'This option is only available for external theory courses.'}), 400

        data = request.get_json(silent=True) or {}
        mode = _normalize_external_assessment_mode(data.get('mode'))
        session.external_assessment_mode = mode
        _recalculate_assessment_totals(session)
        db.session.commit()

        students = _class_students_for_session(session_id)
        _, combined_best3, _, combined_pg_total = _build_combined_assessment_values(session)
        totals = {}
        for student in students:
            if mode == 'best_three_40':
                display = combined_pg_total.get(student.student_id)
            else:
                display = combined_best3.get(student.student_id)
            totals[str(student.id)] = display

        return jsonify({
            'success': True,
            'mode': mode,
            'column_label': _assessment_column_labels().get(mode),
            'totals': totals,
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@class_management_bp.route('/assessment/<int:session_id>/auto-save', methods=['POST'])
@login_required
def auto_save_assessment(session_id):
    """Auto-save assessment marks via AJAX"""
    try:
        import json
        session = get_or_404_for_window(Session, session_id)

        if not user_owns_class_session(current_user, session):
            return jsonify({'success': False, 'message': 'You are not authorized to manage assessment for this session.'}), 403

        students = _class_students_for_session(session_id)
        editable_indices = _get_editable_assessment_indices(session)
        editable_sessional_fields = _get_editable_sessional_fields(session)
        
        if request.is_json:
            data = request.get_json()
        else:
            data = request.form.to_dict()
        
        try:
            if session.course_type == 'theory':
                for student in students:
                    # Load existing absent status
                    absent_status = {}
                    if student.assessment_absent:
                        try:
                            absent_status = json.loads(student.assessment_absent)
                        except:
                            absent_status = {}
                    
                    for i in range(1, 5):
                        if i in editable_indices:
                            key = f'assessment{i}_{student.id}'
                            absent_key = f'absent_{i}_{student.id}'
                            
                            # Check if absent checkbox is checked
                            is_absent = data.get(absent_key) == 'on' or data.get(absent_key) == True
                            absent_status[f'assessment{i}'] = is_absent
                            
                            # If absent, set mark to None, otherwise save the value
                            if is_absent:
                                setattr(student, f'assessment{i}', None)
                            else:
                                value = data.get(key, '')
                                if _is_external_theory_session(session):
                                    setattr(student, f'assessment{i}', _parse_external_assessment_value(value, session))
                                else:
                                    setattr(student, f'assessment{i}', float(value) if value else None)
                    
                    # Save absent status
                    student.assessment_absent = json.dumps(absent_status) if absent_status else None
                    
                _recalculate_assessment_totals(session)

            elif session.course_type == 'sessional' and session.category == 'ug':
                for student in students:
                    # Load existing absent status
                    absent_status = {}
                    if student.assessment_absent:
                        try:
                            absent_status = json.loads(student.assessment_absent)
                        except:
                            absent_status = {}
                    
                    report_absent_key = f'sessional_absent_report_{student.id}'
                    viva_absent_key = f'sessional_absent_viva_{student.id}'

                    if 'sessional_report' in editable_sessional_fields:
                        report_absent = data.get(report_absent_key) == 'on' or data.get(report_absent_key) == True
                        absent_status['sessional_report'] = report_absent

                        if report_absent:
                            student.sessional_report = None
                        else:
                            report = data.get(f'sessional_report_{student.id}', '')
                            student.sessional_report = float(report) if report else None

                    if 'sessional_viva' in editable_sessional_fields:
                        viva_absent = data.get(viva_absent_key) == 'on' or data.get(viva_absent_key) == True
                        absent_status['sessional_viva'] = viva_absent

                        if viva_absent:
                            student.sessional_viva = None
                        else:
                            viva = data.get(f'sessional_viva_{student.id}', '')
                            student.sessional_viva = float(viva) if viva else None
                    
                    # Save absent status
                    student.assessment_absent = json.dumps(absent_status) if absent_status else None
            else:
                return jsonify({'success': False, 'message': 'Unsupported course type'}), 400

            db.session.commit()
            return jsonify({'success': True, 'message': 'Assessment marks saved automatically'})
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error auto-saving assessment for session {session_id}: {e}", exc_info=True)
            return jsonify({'success': False, 'message': f'Error saving: {str(e)}'}), 500
            
    except Exception as e:
        current_app.logger.error(f"Error in auto_save_assessment for session {session_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Server error'}), 500

@class_management_bp.route('/student/view-scores')
@login_required
def student_view_scores():
    """Student view for revealed assessment and attendance scores"""
    from utils.dashboard_settings import require_student_dashboard_card
    blocked = require_student_dashboard_card('my_scores')
    if blocked:
        return blocked
    try:
        import json
        from flask_login import current_user
        
        # Get student ID from current user (assuming username is student_id)
        student_id = current_user.username if hasattr(current_user, 'username') else None
        if not student_id:
            flash('Student ID not found.', 'error')
            return redirect(url_for('index'))
        
        # Find all ClassStudent records for this student
        student_records = ClassStudent.query.filter_by(student_id=student_id).all()
        
        # First pass: collect all sessions and identify split groups
        # Filter: Only show Theory courses (course_type == 'theory')
        session_records_map = {}  # session_id -> Session
        split_group_processed = set()  # Track processed split course identity to avoid duplicate processing
        processed_sessions = set()  # Track which sessions we've processed
        course_map = {}  # unique_course_key -> {'session_ids': set(), 'student_record_ids': set(), 'student_records': []}
        
        # Collect all sessions first
        for student_record in student_records:
            session_obj = get_for_window(Session, student_record.session_id)
            if not session_obj or session_obj.archived:
                continue
            
            # Filter: Only process Theory courses
            # Also check course_name to catch cases where course_type might be incorrectly set
            if session_obj.course_type != 'theory':
                continue
            
            # Additional check: Filter out courses with "Sessional" in the name (case-insensitive)
            # This catches cases where course_type might be incorrectly set to 'theory'
            if session_obj.course_name and 'sessional' in session_obj.course_name.lower():
                continue
            
            session_records_map[session_obj.id] = session_obj
        
        # Second pass: Group sessions by course (handle split courses properly)
        # Iterate over a snapshot because we may add related sessions into the map below.
        # Direct iteration on a mutating dict raises:
        # "dictionary changed size during iteration".
        for session_id, session_obj in list(session_records_map.items()):
            if session_id in processed_sessions:
                continue
            
            # Create unique key for grouping by course
            if session_obj.course_scope in SPLIT_PARTS:
                # Split course: group by course identity (more stable than split_group_id
                # after unassign/reassign cycles).
                split_identity = (
                    str(session_obj.course_code or '').strip().lower(),
                    str(session_obj.year or '').strip(),
                    str(session_obj.term or '').strip(),
                    str(session_obj.academic_session or '').strip(),
                )
                if split_identity in split_group_processed:
                    continue  # Already processed this split group
                split_group_processed.add(split_identity)
                
                course_key = f"split_{split_identity[0]}_{split_identity[1]}_{split_identity[2]}_{split_identity[3]}"
                related_sessions = _resolve_attendance_related_sessions(session_obj, include_archived=True)
                
                if course_key not in course_map:
                    course_map[course_key] = {
                        'session_ids': set(),
                        'student_record_ids': set(),
                        'student_records': []
                    }
                
                # Add all related sessions and their student records
                for related_session in related_sessions:
                    if related_session.course_type != 'theory':
                        continue
                    # Additional check: Filter out courses with "Sessional" in the name
                    if related_session.course_name and 'sessional' in related_session.course_name.lower():
                        continue
                    if related_session.id in processed_sessions:
                        continue
                    
                    session_records_map[related_session.id] = related_session
                    course_map[course_key]['session_ids'].add(related_session.id)
                    processed_sessions.add(related_session.id)
                    
                    # Add student records from this related session
                    related_student_records = ClassStudent.query.filter_by(
                        session_id=related_session.id,
                        student_id=student_id
                    ).all()
                    for related_rec in related_student_records:
                        if related_rec.id not in course_map[course_key]['student_record_ids']:
                            course_map[course_key]['student_record_ids'].add(related_rec.id)
                            course_map[course_key]['student_records'].append(related_rec)
            else:
                # Regular course: group by course_name, course_code, year, term, academic_session
                course_key = (
                    str(session_obj.course_name or '').strip().lower(),
                    str(session_obj.course_code or '').strip().lower(),
                    str(session_obj.year or '').strip(),
                    str(session_obj.term or '').strip(),
                    str(session_obj.academic_session or '').strip()
                )
                
                if course_key not in course_map:
                    course_map[course_key] = {
                        'session_ids': set(),
                        'student_record_ids': set(),
                        'student_records': []
                    }
                
                course_map[course_key]['session_ids'].add(session_obj.id)
                processed_sessions.add(session_obj.id)
                
                # Add student records for this session
                session_student_records = ClassStudent.query.filter_by(
                    session_id=session_obj.id,
                    student_id=student_id
                ).all()
                for rec in session_student_records:
                    if rec.id not in course_map[course_key]['student_record_ids']:
                        course_map[course_key]['student_record_ids'].add(rec.id)
                        course_map[course_key]['student_records'].append(rec)
        
        teacher_callsign_cache = {}

        # Second pass: build reveal status by combining all sessions with same course key
        for course_key, course_data in course_map.items():
            reveal_status = {}
            reveal_callsigns = defaultdict(set)
            # Default attendance to revealed
            reveal_status['attendance'] = True
            
            # Combine reveal status from all sessions with this course key
            for session_id in course_data['session_ids']:
                session_obj = session_records_map.get(session_id)
                if session_obj and session_obj.assessment_revealed:
                    try:
                        all_reveals = json.loads(session_obj.assessment_revealed)
                        for teacher_id, teacher_reveals in all_reveals.items():
                            callsign = teacher_callsign_cache.get(teacher_id)
                            if teacher_id not in teacher_callsign_cache:
                                teacher_obj = Teacher.query.get(int(teacher_id)) if str(teacher_id).isdigit() else None
                                callsign = ''
                                if teacher_obj:
                                    callsign = (teacher_obj.call_sign or teacher_obj.short_name or '').strip()
                                teacher_callsign_cache[teacher_id] = callsign
                            for assessment_type, is_revealed in teacher_reveals.items():
                                if is_revealed:
                                    reveal_status[assessment_type] = True
                                    if callsign:
                                        reveal_callsigns[assessment_type].add(callsign)
                    except:
                        pass
            
            course_data['reveal_status'] = reveal_status
            course_data['reveal_callsigns'] = {
                assessment_type: ', '.join(sorted(callsigns))
                for assessment_type, callsigns in reveal_callsigns.items()
            }
            # Convert session_ids set to list and get primary session
            session_ids_list = sorted(list(course_data['session_ids']))
            course_data['primary_session'] = session_records_map[session_ids_list[0]]  # Use first session as primary
        
        # Build combined data for each unique course
        courses_data = []
        for course_key, course_data in course_map.items():
            session_obj = course_data['primary_session']  # Use primary session for display
            
            # Double-check: Only process Theory courses (skip Sessional courses)
            if session_obj.course_type != 'theory':
                continue
            
            # Additional check: Filter out courses with "Sessional" in the name (case-insensitive)
            # This catches cases where course_type might be incorrectly set to 'theory'
            if session_obj.course_name and 'sessional' in session_obj.course_name.lower():
                continue
            
            student_records = course_data['student_records']
            reveal_status = course_data['reveal_status']
            reveal_callsigns = course_data.get('reveal_callsigns', {})
            all_session_ids = course_data['session_ids']
            is_split_course = any(
                (session_records_map.get(session_id) and session_records_map.get(session_id).course_scope in SPLIT_PARTS)
                for session_id in all_session_ids
            )
            
            # Use the first student record as primary (prioritize non-null values)
            primary_record = student_records[0]
            
            # Build assessment scores based on reveal status
            # For split courses, use _build_combined_assessment_values to properly combine from all related sessions
            assessment_scores = {}
            highest_assessment_scores = {}
            best3_total = None
            pg_total = None
            
            if _is_external_theory_session(session_obj):
                combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session_obj)
                student_combined_values = combined_values.get(student_id, {})
                for i in range(1, 5):
                    max_value = None
                    for combined_student_values in combined_values.values():
                        current_value = combined_student_values.get(i)
                        if current_value is not None and (max_value is None or current_value > max_value):
                            max_value = current_value
                    highest_assessment_scores[f'assessment{i}'] = max_value

                for i in range(1, 5):
                    assessment_key = f'assessment{i}'
                    if reveal_status.get(assessment_key, False):
                        value = student_combined_values.get(i, None)
                        assessment_scores[assessment_key] = value

                revealed_combined = {i: assessment_scores.get(f'assessment{i}') for i in range(1, 5)}
                mode = _normalize_external_assessment_mode(getattr(session_obj, 'external_assessment_mode', None))
                result = _compute_external_assessment_total(revealed_combined, mode)
                if mode == 'best_three_40':
                    pg_total = result['display_total']
                else:
                    best3_total = result['display_total']
            elif session_obj.course_type == 'theory' and session_obj.category == 'ug':
                # Use _build_combined_assessment_values which properly handles split courses
                combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session_obj)
                student_combined_values = combined_values.get(student_id, {})
                for i in range(1, 5):
                    max_value = None
                    for combined_student_values in combined_values.values():
                        current_value = combined_student_values.get(i)
                        if current_value is not None and (max_value is None or current_value > max_value):
                            max_value = current_value
                    highest_assessment_scores[f'assessment{i}'] = max_value
                
                # Build assessment_scores only for revealed assessments (students must not see unrevealed marks)
                for i in range(1, 5):
                    assessment_key = f'assessment{i}'
                    if reveal_status.get(assessment_key, False):
                        value = student_combined_values.get(i, None)
                        assessment_scores[assessment_key] = value
                
                # Total (Best 3) must be computed from REVEALED assessments only – do not include unrevealed marks
                revealed_values = [assessment_scores[k] for k in ['assessment1', 'assessment2', 'assessment3', 'assessment4'] if assessment_scores.get(k) is not None]
                if revealed_values:
                    best3 = sorted([float(v) for v in revealed_values], reverse=True)[:3]
                    best3_total = sum(best3) if best3 else None
                else:
                    best3_total = None
            elif session_obj.course_type == 'theory' and session_obj.category == 'pg':
                combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session_obj)
                student_combined_values = combined_values.get(student_id, {})
                for i in range(1, 5):
                    max_value = None
                    for combined_student_values in combined_values.values():
                        current_value = combined_student_values.get(i)
                        if current_value is not None and (max_value is None or current_value > max_value):
                            max_value = current_value
                    highest_assessment_scores[f'assessment{i}'] = max_value
                
                # Build assessment_scores only for revealed assessments
                for i in range(1, 5):
                    assessment_key = f'assessment{i}'
                    if reveal_status.get(assessment_key, False):
                        value = student_combined_values.get(i, None)
                        assessment_scores[assessment_key] = value
                
                # PG total (on 40 scale) must be computed from REVEALED assessments only
                revealed_values = [assessment_scores[k] for k in ['assessment1', 'assessment2', 'assessment3', 'assessment4'] if assessment_scores.get(k) is not None]
                if revealed_values:
                    best3 = sorted([float(v) for v in revealed_values], reverse=True)[:3]
                    best3_sum = sum(best3)
                    pg_total = int(round((best3_sum / 30.0) * 40)) if best3_sum else None
                else:
                    pg_total = None
            # Note: Sessional courses are filtered out - only Theory courses should reach this point
            
            
            # Get attendance marks if revealed
            # For split courses, aggregate attendance from all related sessions
            attendance_data = None
            if reveal_status.get('attendance', True):  # Default to True if not set
                # Use _build_attendance_summary which already handles split courses correctly
                # It aggregates attendance from all related sessions automatically
                attendance_summary = _build_attendance_summary(session_obj, include_archived=True)
                
                # Find the student's attendance stats from the summary
                # Use student_id (string) for lookup
                student_stats = attendance_summary.get('per_student', {}).get(student_id, {})
                
                if student_stats:
                    attendance_data = {
                        'present_count': student_stats.get('present', 0),
                        'total_classes': attendance_summary.get('total_classes', 0),
                        'percentage': student_stats.get('percentage', 0),
                        'marks': student_stats.get('marks', 0),
                        'part_breakdown': (
                            _split_attendance_part_breakdown(attendance_summary, student_id)
                            if is_split_course else None
                        ),
                    }
            
            # Build teacher options (split courses may have multiple teachers)
            teacher_options = []
            try:
                related_sessions = _resolve_attendance_related_sessions(session_obj, include_archived=False)
                for related_session in related_sessions:
                    if related_session and not related_session.archived and related_session.teacher:
                        teacher_options.append({
                            'session_id': related_session.id,
                            'teacher_id': related_session.teacher.id,
                            'teacher_name': related_session.teacher.name,
                            'teacher_short': related_session.teacher.short_name,
                            'scope_label': COURSE_SCOPE_LABELS.get(related_session.course_scope, 'Part')
                        })
            except Exception as teacher_error:
                current_app.logger.warning(f"Error building teacher options: {teacher_error}")

            # Load Q&A threads for this course (student-specific)
            qa_threads = []
            qa_new_reply_count = 0
            try:
                from sqlalchemy.orm import selectinload
                session_ids = list(course_data['session_ids'])
                if session_ids:
                    qa_threads = CourseQuestionThread.query.options(
                        selectinload(CourseQuestionThread.messages).selectinload(CourseQuestionMessage.attachments)
                    ).filter(
                        CourseQuestionThread.session_id.in_(session_ids),
                        CourseQuestionThread.student_id == student_id
                    ).order_by(CourseQuestionThread.created_at.desc()).all()
                    student_seen_updates = False
                    seen_at = datetime.utcnow()
                    # Annotate threads with latest message role for notifications
                    for thread in qa_threads:
                        last_message = None
                        if thread.messages:
                            last_message = max(
                                thread.messages,
                                key=lambda m: (m.created_at or datetime.min, m.id or 0)
                            )
                        thread.last_message_role = last_message.sender_role if last_message else None
                        # Student opening the thread marks all teacher messages as seen.
                        for msg in thread.messages:
                            if msg.sender_role == 'teacher' and msg.seen_by_student_at is None:
                                msg.seen_by_student_at = seen_at
                                student_seen_updates = True
                        thread.has_unseen_teacher_message = any(
                            m.sender_role == 'teacher' and m.seen_by_student_at is None
                            for m in thread.messages
                        )
                        if thread.has_unseen_teacher_message:
                            qa_new_reply_count += 1
                    if student_seen_updates:
                        db.session.commit()
            except Exception as qa_error:
                db.session.rollback()
                current_app.logger.warning(f"Error loading Q&A threads for student {student_id}: {qa_error}")

            courses_data.append({
                'session': session_obj,
                'student_record': primary_record,
                'assessment_scores': assessment_scores,
                'highest_assessment_scores': highest_assessment_scores,
                'best3_total': best3_total,
                'pg_total': pg_total,
                'attendance_data': attendance_data,
                'reveal_status': reveal_status,
                'reveal_callsigns': reveal_callsigns,
                'is_split_course': is_split_course,
                'qa_threads': qa_threads,
                'qa_new_reply_count': qa_new_reply_count,
                'teacher_options': teacher_options
            })
        
        # Sort courses by course_name, then by year-term for consistent display
        courses_data.sort(key=lambda x: (
            x['session'].course_name or '',
            x['session'].year or '',
            x['session'].term or ''
        ))
        
        return render_template(
            'class_management/student_view_scores.html',
            student_id=student_id,
            courses_data=courses_data
        )
        
    except Exception as e:
        current_app.logger.error(f"Error loading student view scores: {e}", exc_info=True)
        flash(f'Error loading scores: {str(e)}', 'error')
        return redirect(url_for('index'))


@class_management_bp.route('/student/course-questions/<int:session_id>/create', methods=['POST'])
@login_required
def student_create_course_question(session_id):
    """Create a new Q&A thread from student."""
    student_id = current_user.username if hasattr(current_user, 'username') else None
    if not student_id:
        flash('Student ID not found.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    session_obj = get_or_404_for_window(Session, session_id)
    related_sessions = _get_related_sessions(session_obj)
    related_session_ids = [s.id for s in related_sessions if s]

    selected_session_id = request.form.get('selected_session_id', type=int) or session_id
    if selected_session_id not in related_session_ids:
        flash('Invalid teacher selection for this course.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    # Ensure student is enrolled in any related session
    student_record = ClassStudent.query.filter(
        ClassStudent.student_id == student_id,
        ClassStudent.session_id.in_(related_session_ids)
    ).first()
    if not student_record:
        flash('You are not enrolled in this course.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    selected_session = get_or_404_for_window(Session, selected_session_id)

    subject = request.form.get('subject', '').strip()
    message_body = request.form.get('message', '').strip()
    files = request.files.getlist('attachments')

    if not subject:
        flash('Please provide a subject for your question.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    if not message_body and not files:
        flash('Please write a message or attach a file.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    try:
        thread = CourseQuestionThread(
            session_id=selected_session_id,
            student_id=student_id,
            student_name=student_record.name or getattr(current_user, 'full_name', '') or student_id,
            teacher_id=selected_session.teacher_id,
            subject=subject,
            status='open'
        )
        db.session.add(thread)
        db.session.flush()

        message = CourseQuestionMessage(
            thread_id=thread.id,
            sender_role='student',
            sender_user_id=None,
            body=message_body
        )
        db.session.add(message)
        db.session.flush()

        saved_attachments = _save_qa_attachments(files, thread.id)
        for attachment in saved_attachments:
            db.session.add(CourseQuestionAttachment(
                message_id=message.id,
                file_name=attachment['file_name'],
                file_path=attachment['file_path'],
                file_size=attachment.get('file_size'),
                file_type=attachment.get('file_type')
            ))

        db.session.commit()
        flash('Your question has been sent to the teacher.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating question thread: {e}", exc_info=True)
        flash('Failed to send your question. Please try again.', 'error')

    return redirect(url_for('class_management.student_view_scores'))


@class_management_bp.route('/student/course-questions/<int:thread_id>/reply', methods=['POST'])
@login_required
def student_reply_course_question(thread_id):
    """Reply to an existing Q&A thread from student."""
    student_id = current_user.username if hasattr(current_user, 'username') else None
    if not student_id:
        flash('Student ID not found.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    thread = CourseQuestionThread.query.get_or_404(thread_id)
    if thread.student_id != student_id:
        flash('You are not authorized to reply to this thread.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    message_body = request.form.get('message', '').strip()
    files = request.files.getlist('attachments')

    if not message_body and not files:
        flash('Please write a message or attach a file.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    try:
        message = CourseQuestionMessage(
            thread_id=thread.id,
            sender_role='student',
            sender_user_id=None,
            body=message_body
        )
        db.session.add(message)
        db.session.flush()

        saved_attachments = _save_qa_attachments(files, thread.id)
        for attachment in saved_attachments:
            db.session.add(CourseQuestionAttachment(
                message_id=message.id,
                file_name=attachment['file_name'],
                file_path=attachment['file_path'],
                file_size=attachment.get('file_size'),
                file_type=attachment.get('file_type')
            ))

        thread.updated_at = datetime.utcnow()
        thread.teacher_read_at = None  # so teacher sees new reply in bell (unread)
        db.session.commit()
        flash('Reply sent successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error replying to question thread: {e}", exc_info=True)
        flash('Failed to send reply. Please try again.', 'error')

    return redirect(url_for('class_management.student_view_scores'))


@class_management_bp.route('/course-questions/<int:session_id>')
@login_required
def course_questions(session_id):
    """Teacher view of course Q&A threads for a session."""
    session_obj = get_or_404_for_window(Session, session_id)
    teacher = _ensure_current_teacher()

    if not teacher or (teacher.id != session_obj.teacher_id and not is_admin(current_user)):
        flash('You are not authorized to view these questions.', 'error')
        return redirect(url_for('class_management.index'))

    try:
        from sqlalchemy.orm import selectinload
        threads = CourseQuestionThread.query.options(
            selectinload(CourseQuestionThread.messages).selectinload(CourseQuestionMessage.attachments)
        ).filter_by(session_id=session_id).order_by(CourseQuestionThread.created_at.desc()).all()
        teacher_seen_updates = False
        seen_at = datetime.utcnow()
        for thread in threads:
            last_message = None
            if thread.messages:
                last_message = max(
                    thread.messages,
                    key=lambda m: (m.created_at or datetime.min, m.id or 0)
                )
            thread.last_message_role = last_message.sender_role if last_message else None
            # Teacher opening the thread marks all student messages as seen.
            for msg in thread.messages:
                if msg.sender_role == 'student' and msg.seen_by_teacher_at is None:
                    msg.seen_by_teacher_at = seen_at
                    thread.teacher_read_at = seen_at
                    teacher_seen_updates = True
            thread.is_unread_for_teacher = any(
                m.sender_role == 'student' and m.seen_by_teacher_at is None
                for m in thread.messages
            )
        if teacher_seen_updates:
            db.session.commit()
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error loading course questions: {e}", exc_info=True)
        threads = []

    return render_template(
        'class_management/course_questions.html',
        session=session_obj,
        threads=threads
    )


@class_management_bp.route('/course-questions/<int:thread_id>/reply', methods=['POST'])
@login_required
def teacher_reply_course_question(thread_id):
    """Teacher reply to a Q&A thread."""
    thread = CourseQuestionThread.query.get_or_404(thread_id)
    session_obj = get_or_404_for_window(Session, thread.session_id)
    teacher = _ensure_current_teacher()

    if not teacher or (teacher.id != session_obj.teacher_id and not is_admin(current_user)):
        flash('You are not authorized to reply to this thread.', 'error')
        return redirect(url_for('class_management.index'))

    message_body = request.form.get('message', '').strip()
    files = request.files.getlist('attachments')

    if not message_body and not files:
        flash('Please write a message or attach a file.', 'error')
        return redirect(url_for('class_management.course_questions', session_id=session_obj.id))

    try:
        message = CourseQuestionMessage(
            thread_id=thread.id,
            sender_role='teacher',
            sender_user_id=teacher.id,
            body=message_body
        )
        db.session.add(message)
        db.session.flush()

        saved_attachments = _save_qa_attachments(files, thread.id)
        for attachment in saved_attachments:
            db.session.add(CourseQuestionAttachment(
                message_id=message.id,
                file_name=attachment['file_name'],
                file_path=attachment['file_path'],
                file_size=attachment.get('file_size'),
                file_type=attachment.get('file_type')
            ))

        thread.updated_at = datetime.utcnow()
        thread.teacher_read_at = datetime.utcnow()
        db.session.commit()
        flash('Reply sent successfully.', 'success')
        # Notify student about teacher reply
        try:
            link_url = url_for('class_management.student_view_scores')
            _notify_student_by_username(
                thread.student_id,
                'question_reply',
                f'Teacher replied to your question: {thread.subject[:50]}{"..." if len(thread.subject) > 50 else ""}',
                link_url
            )
        except Exception as notif_e:
            current_app.logger.warning(f"Student notification (question reply): {notif_e}")
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error replying to question thread: {e}", exc_info=True)
        flash('Failed to send reply. Please try again.', 'error')

    return redirect(url_for('class_management.course_questions', session_id=session_obj.id))


@class_management_bp.route('/course-questions/<int:thread_id>/mark-read', methods=['POST'])
@login_required
def mark_course_question_thread_read(thread_id):
    """Mark a single course question thread as read for the teacher."""
    thread = CourseQuestionThread.query.get_or_404(thread_id)
    session_obj = get_or_404_for_window(Session, thread.session_id)
    teacher = _ensure_current_teacher()

    if not teacher or (teacher.id != session_obj.teacher_id and not is_admin(current_user)):
        flash('You are not authorized to update this thread.', 'error')
        return redirect(url_for('class_management.index'))

    try:
        seen_at = datetime.utcnow()
        thread.teacher_read_at = seen_at
        for msg in thread.messages:
            if msg.sender_role == 'student' and msg.seen_by_teacher_at is None:
                msg.seen_by_teacher_at = seen_at
        db.session.commit()
        flash('Thread marked as read.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error marking thread as read: {e}", exc_info=True)
        flash('Failed to mark thread as read.', 'error')

    return redirect(url_for('class_management.course_questions', session_id=session_obj.id))


@class_management_bp.route('/course-questions/attachments/<int:attachment_id>')
@login_required
def download_course_question_attachment(attachment_id):
    """Download attachment for course Q&A."""
    try:
        attachment = CourseQuestionAttachment.query.get_or_404(attachment_id)
        message = CourseQuestionMessage.query.get_or_404(attachment.message_id)
        thread = CourseQuestionThread.query.get_or_404(message.thread_id)
        session_obj = get_or_404_for_window(Session, thread.session_id)

        student_id = current_user.username if hasattr(current_user, 'username') else None
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()

        is_student_allowed = bool(student_id and thread.student_id == student_id)
        is_teacher_allowed = bool(teacher and teacher.id == session_obj.teacher_id)
        if not (is_student_allowed or is_teacher_allowed or is_admin(current_user)):
            flash('You are not authorized to access this file.', 'error')
            return redirect(url_for('index'))

        if not os.path.exists(attachment.file_path):
            flash('File not found.', 'error')
            return redirect(url_for('index'))

        return send_file(
            attachment.file_path,
            as_attachment=True,
            download_name=attachment.file_name
        )
    except Exception as e:
        current_app.logger.error(f"Error downloading question attachment: {e}", exc_info=True)
        flash('Error downloading file.', 'error')
        return redirect(url_for('index'))


@class_management_bp.route('/course-questions/<int:thread_id>/delete', methods=['POST'])
@login_required
def delete_course_question_thread(thread_id):
    """Delete a Q&A thread (question)."""
    thread = CourseQuestionThread.query.get_or_404(thread_id)
    session_obj = get_or_404_for_window(Session, thread.session_id)
    student_id = current_user.username if hasattr(current_user, 'username') else None
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()

    is_student_owner = bool(student_id and thread.student_id == student_id)
    is_teacher_owner = bool(teacher and teacher.id == session_obj.teacher_id)
    if not (is_student_owner or is_teacher_owner or is_admin(current_user)):
        flash('You are not authorized to delete this question.', 'error')
        return redirect(url_for('index'))

    try:
        # Delete files for all attachments
        for msg in thread.messages:
            _delete_qa_attachments(msg.attachments)
        db.session.delete(thread)
        db.session.commit()
        flash('Question deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting question thread: {e}", exc_info=True)
        flash('Failed to delete question.', 'error')

    if is_teacher_owner:
        return redirect(url_for('class_management.course_questions', session_id=session_obj.id))
    return redirect(url_for('class_management.student_view_scores'))


@class_management_bp.route('/course-questions/message/<int:message_id>/delete', methods=['POST'])
@login_required
def delete_course_question_message(message_id):
    """Delete a single message in a Q&A thread."""
    message = CourseQuestionMessage.query.get_or_404(message_id)
    thread = CourseQuestionThread.query.get_or_404(message.thread_id)
    session_obj = get_or_404_for_window(Session, thread.session_id)
    student_id = current_user.username if hasattr(current_user, 'username') else None
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()

    is_student_owner = bool(student_id and thread.student_id == student_id and message.sender_role == 'student')
    is_teacher_owner = bool(teacher and message.sender_role == 'teacher' and message.sender_user_id == teacher.id)
    if not (is_student_owner or is_teacher_owner or is_admin(current_user)):
        flash('You are not authorized to delete this message.', 'error')
        return redirect(url_for('index'))

    try:
        _delete_qa_attachments(message.attachments)
        db.session.delete(message)
        db.session.commit()
        flash('Message deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting question message: {e}", exc_info=True)
        flash('Failed to delete message.', 'error')

    if is_teacher_owner:
        return redirect(url_for('class_management.course_questions', session_id=session_obj.id))
    return redirect(url_for('class_management.student_view_scores'))

@class_management_bp.route('/download_assessment_excel/<int:session_id>')
@login_required
def download_assessment_excel(session_id):
    """Download assessment data as Excel file"""
    try:
        import json
        session = get_or_404_for_window(Session, session_id)
        students = _class_students_for_session(session_id)
        combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session)
        
        # Helper function to get assessment value with absent check
        def get_assessment_value(student, assessment_num):
            """Get assessment value, showing 'Absent' if marked absent"""
            absent_status = {}
            if student.assessment_absent:
                try:
                    absent_status = json.loads(student.assessment_absent)
                except:
                    absent_status = {}
            
            is_absent = absent_status.get(f'assessment{assessment_num}', False)
            if is_absent:
                return 'Absent'
            
            combined = combined_values.get(student.student_id, {})
            value = combined.get(assessment_num)
            return value if value is not None else ''
        
        # Helper function to get sessional value with absent check
        def get_sessional_value(student, sessional_type):
            """Get sessional value, showing 'Absent' if marked absent"""
            absent_status = {}
            if student.assessment_absent:
                try:
                    absent_status = json.loads(student.assessment_absent)
                except:
                    absent_status = {}
            
            is_absent = absent_status.get(f'sessional_{sessional_type}', False)
            if is_absent:
                return 'Absent'
            
            if sessional_type == 'report':
                value = student.sessional_report
            else:  # viva
                value = student.sessional_viva
            
            return int(round(value)) if value is not None else ''
        
        # Build data for DataFrame
        data = []
        if _is_external_theory_session(session):
            mode = _normalize_external_assessment_mode(getattr(session, 'external_assessment_mode', None))
            total_column = _assessment_column_labels().get(mode, 'Total')
            columns = [
                'Student ID', 'Name',
                _external_assessment_column_header(1, mode),
                _external_assessment_column_header(2, mode),
                _external_assessment_column_header(3, mode),
                _external_assessment_column_header(4, mode),
                total_column,
            ]
            for s in students:
                total_value, _ = _get_external_assessment_display_total(
                    session,
                    s.student_id,
                    combined_values,
                    combined_best3=combined_best3,
                    combined_pg_total=combined_pg_total,
                )
                data.append([
                    s.student_id,
                    s.name,
                    get_assessment_value(s, 1),
                    get_assessment_value(s, 2),
                    get_assessment_value(s, 3),
                    get_assessment_value(s, 4),
                    total_value,
                ])
        elif session.course_type == 'theory' and session.category == 'ug':
            columns = ['Student ID', 'Name', 'Assessment 1', 'Assessment 2', 'Assessment 3', 'Assessment 4', 'Total of Best 3']
            for s in students:
                best3_total = combined_best3.get(s.student_id) if combined_best3 else None
                data.append([
                    s.student_id,
                    s.name,
                    get_assessment_value(s, 1),
                    get_assessment_value(s, 2),
                    get_assessment_value(s, 3),
                    get_assessment_value(s, 4),
                    best3_total
                ])
        elif session.course_type == 'theory' and session.category == 'pg':
            columns = ['Student ID', 'Name', 'Assessment 1', 'Assessment 2', 'Assessment 3', 'Assessment 4', f'Total ({assessment_cfg()["pg_out_of"]})']
            for s in students:
                data.append([
                    s.student_id,
                    s.name,
                    get_assessment_value(s, 1),
                    get_assessment_value(s, 2),
                    get_assessment_value(s, 3),
                    get_assessment_value(s, 4),
                    combined_pg_total.get(s.student_id)
                ])
        elif session.course_type == 'sessional' and session.category == 'ug':
            columns = ['Student ID', 'Name', 'Sessional Report (60)', 'Sessional Viva (30)', 'Total (Sessional: 90)']
            for s in students:
                report_val = get_sessional_value(s, 'report')
                viva_val = get_sessional_value(s, 'viva')
                
                # Calculate total (skip if either is 'Absent')
                if report_val == 'Absent' or viva_val == 'Absent':
                    total = 'Absent'
                else:
                    total = (s.sessional_report or 0) + (s.sessional_viva or 0)
                    if total:
                        rounded = _maybe_round_assessment_total(session, total)
                        total = rounded if rounded is not None else ''
                    else:
                        total = ''
                
                data.append([
                    s.student_id,
                    s.name,
                    report_val,
                    viva_val,
                    total
                ])
        else:
            flash('Unsupported course type for assessment export', 'error')
            return redirect(url_for('class_management.assessment', session_id=session_id))
        
        # Create DataFrame and Excel file
        df = pd.DataFrame(data, columns=columns)
        
        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Assessment', index=False)
            
            # Auto-adjust column widths
            worksheet = writer.sheets['Assessment']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        
        filename = f"assessment_{session.course_name or session.term}_{session.year}_{session.term}.xlsx"
        
        # Use Response instead of send_file for better cPanel compatibility
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(output.getvalue()))
            }
        )
        
    except Exception as e:
        flash(f'Error downloading assessment Excel: {str(e)}', 'error')
        return redirect(url_for('class_management.assessment', session_id=session_id))

@class_management_bp.route('/download_assessment_pdf/<int:session_id>')
@login_required
def download_assessment_pdf(session_id):
    """Download assessment marks as PDF"""
    try:
        import json
        session = get_or_404_for_window(Session, session_id)
        students = _class_students_for_session(session_id)
        combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session)
        
        def format_mark_for_pdf(value):
            """Format numeric marks without forcing rounded integers."""
            if value is None or value == '-':
                return '-'
            if isinstance(value, str):
                return value
            try:
                numeric_value = float(value)
            except (TypeError, ValueError):
                return str(value)
            if numeric_value.is_integer():
                return str(int(numeric_value))
            return f"{numeric_value:.2f}".rstrip('0').rstrip('.')
        
        # Helper function to get assessment value with absent check
        def get_assessment_value(student, assessment_num):
            """Get assessment value, showing 'Absent' if marked absent"""
            absent_status = {}
            if student.assessment_absent:
                try:
                    absent_status = json.loads(student.assessment_absent)
                except:
                    absent_status = {}
            
            is_absent = absent_status.get(f'assessment{assessment_num}', False)
            if is_absent:
                return 'Absent'
            
            combined = combined_values.get(student.student_id, {})
            value = combined.get(assessment_num)
            return value if value is not None else '-'
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=0.3*inch,
            rightMargin=0.3*inch,
            topMargin=0.3*inch,
            bottomMargin=0.3*inch
        )
        styles = getSampleStyleSheet()
        elements = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.black,
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        title = Paragraph("Continuous Assessment Marks", title_style)
        elements.append(title)
        elements.append(Spacer(1, 0.15*inch))
        
        # Course Information
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=5,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        # Course Information with spacing between items
        info_items = []
        if session.course_code:
            info_items.append(f"<b>Subject Code:</b> {session.course_code}")
        if session.course_name:
            info_items.append(f"<b>Subject Name:</b> {session.course_name}")
        if session.academic_session:
            info_items.append(f"<b>Session:</b> {session.academic_session}")
        if session.year:
            info_items.append(f"<b>Year:</b> {session.year}")
        if session.term:
            info_items.append(f"<b>Term:</b> {session.term}")
        
        # Add items with spacing between them
        for i, info_line in enumerate(info_items):
            elements.append(Paragraph(info_line, info_style))
            if i < len(info_items) - 1:  # Add space between items, but not after the last one
                elements.append(Spacer(1, 0.1*inch))
        
        elements.append(Spacer(1, 0.2*inch))
        
        # Table data
        if _is_external_theory_session(session):
            mode = _normalize_external_assessment_mode(getattr(session, 'external_assessment_mode', None))
            total_header = _assessment_column_labels().get(mode, 'Total')
            calc_label = _external_assessment_modes().get(mode, mode)
            elements.append(Paragraph(f"<b>Calculation:</b> {calc_label}", info_style))
            elements.append(Spacer(1, 0.1*inch))
            table_data = [[
                'SI', 'Student ID',
                _external_assessment_column_header(1, mode),
                _external_assessment_column_header(2, mode),
                _external_assessment_column_header(3, mode),
                _external_assessment_column_header(4, mode),
                total_header,
            ]]
            for idx, student in enumerate(students, start=1):
                total_value, _ = _get_external_assessment_display_total(
                    session,
                    student.student_id,
                    combined_values,
                    combined_best3=combined_best3,
                    combined_pg_total=combined_pg_total,
                )
                row = [
                    str(idx),
                    str(student.student_id),
                    format_mark_for_pdf(get_assessment_value(student, 1)),
                    format_mark_for_pdf(get_assessment_value(student, 2)),
                    format_mark_for_pdf(get_assessment_value(student, 3)),
                    format_mark_for_pdf(get_assessment_value(student, 4)),
                    format_mark_for_pdf(total_value),
                ]
                table_data.append(row)
        elif session.course_type == 'theory' and session.category == 'ug' and not _is_external_theory_session(session):
            # UG Theory: Assessment 1-4 and Total of Best 3
            cfg = assessment_cfg()
            slot = cfg['slot_max']
            table_data = [['SI', 'Student ID', f'Assessment 1 ({slot})', f'Assessment 2 ({slot})', f'Assessment 3 ({slot})', f'Assessment 4 ({slot})', f'Total of Best {cfg["take_best"]} ({cfg["ug_out_of"]})']]
            
            for idx, student in enumerate(students, start=1):
                best3_total = combined_best3.get(student.student_id) if combined_best3 else '-'
                formatted_total = format_mark_for_pdf(best3_total)
                row = [
                    str(idx),
                    str(student.student_id),
                    format_mark_for_pdf(get_assessment_value(student, 1)),
                    format_mark_for_pdf(get_assessment_value(student, 2)),
                    format_mark_for_pdf(get_assessment_value(student, 3)),
                    format_mark_for_pdf(get_assessment_value(student, 4)),
                    formatted_total
                ]
                table_data.append(row)
        
        elif session.course_type == 'theory' and session.category == 'pg' and not _is_external_theory_session(session):
            # PG Theory: Assessment 1-4 and Total (40)
            cfg = assessment_cfg()
            slot = cfg['slot_max']
            table_data = [['SI', 'Student ID', f'Assessment 1 ({slot})', f'Assessment 2 ({slot})', f'Assessment 3 ({slot})', f'Assessment 4 ({slot})', f'Total ({cfg["pg_out_of"]})']]
            
            for idx, student in enumerate(students, start=1):
                combined = combined_values.get(student.student_id, {})
                valid_marks = [v for v in combined.values() if v is not None]
                valid_marks.sort(reverse=True)
                if valid_marks:
                    best_three = take_best_marks(valid_marks)
                    pg_total_unrounded = scale_pg_total(sum(best_three))
                    # PG Theory: total on 40 scale is always a whole-number display (half-up).
                    rounded_total = _round_half_up_int(pg_total_unrounded)
                    formatted_total = str(rounded_total) if rounded_total is not None else '-'
                else:
                    formatted_total = '-'
                row = [
                    str(idx),
                    str(student.student_id),
                    format_mark_for_pdf(get_assessment_value(student, 1)),
                    format_mark_for_pdf(get_assessment_value(student, 2)),
                    format_mark_for_pdf(get_assessment_value(student, 3)),
                    format_mark_for_pdf(get_assessment_value(student, 4)),
                    formatted_total
                ]
                table_data.append(row)
        
        else:
            flash('Assessment PDF is only available for theory courses.', 'error')
            return redirect(url_for('class_management.assessment', session_id=session_id))
        
        # Create table
        table = Table(table_data, repeatRows=1)
        
        # Table style - Black and white
        table_style = TableStyle([
            # Header row - White background, black text
            ('BACKGROUND', (0, 0), (-1, 0), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            
            # Data rows
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1.5, colors.black),  # Thicker borders
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            
            # All rows white background (black and white)
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            
            # Column widths - optimized for one page (without Name column)
            ('COLWIDTH', (0, 0), (0, -1), 0.4*inch),  # SI
            ('COLWIDTH', (1, 0), (1, -1), 1.1*inch),  # Student ID
            ('COLWIDTH', (2, 0), (2, -1), 1.0*inch),  # Assessment 1 (10)
            ('COLWIDTH', (3, 0), (3, -1), 1.0*inch),  # Assessment 2 (10)
            ('COLWIDTH', (4, 0), (4, -1), 1.0*inch),  # Assessment 3 (10)
            ('COLWIDTH', (5, 0), (5, -1), 1.0*inch),  # Assessment 4 (10)
            ('COLWIDTH', (6, 0), (6, -1), 1.1*inch),  # Total
        ])
        
        table.setStyle(table_style)
        elements.append(table)
        
        # Page number callback function
        def add_page_number(canvas_obj, doc_obj):
            """Add page number to each page"""
            canvas_obj.saveState()
            page_num = canvas_obj.getPageNumber()
            text = f"Page {page_num}"
            width, height = letter
            canvas_obj.setFont('Helvetica', 9)
            canvas_obj.drawRightString(width - 0.3*inch, 0.2*inch, text)
            canvas_obj.restoreState()
        
        # Build PDF with page numbers
        doc.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)
        buffer.seek(0)
        
        filename = f"assessment_{session.course_code or 'marks'}_{session.year}_{session.term}.pdf"
        
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
        
    except Exception as e:
        current_app.logger.error(f"Error generating assessment PDF for session {session_id}: {e}", exc_info=True)
        flash(f'Error generating assessment PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.assessment', session_id=session_id))

# Template filter for dynamic attribute access
@class_management_bp.app_template_filter('getattr')
def jinja_getattr(obj, name):
    return getattr(obj, name)


@class_management_bp.app_template_filter('qa_message_body')
def qa_message_body_filter(value):
    """Q&A message body: treat literal <br> as line breaks, preserve \\n, linkify URLs (XSS-safe)."""
    import re
    from markupsafe import Markup
    from jinja2.utils import urlize

    if value is None or value == '':
        return Markup('')
    text = str(value)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    linked = urlize(text, trim_url_limit=True)
    return Markup(str(linked).replace('\n', '<br />'))

@class_management_bp.route('/evaluation/<int:session_id>')
@login_required
def evaluation(session_id):
    """Placeholder Evaluation page for a session. Forms will be added later."""
    session = get_or_404_for_window(Session, session_id)
    students = _class_students_for_session(session_id)
    return render_template('class_management/evaluation.html', session=session, students=students)

@class_management_bp.route('/evaluation/<int:session_id>/course-assessment', methods=['GET', 'POST'])
@login_required
def course_assessment(session_id):
    """Invite other teachers to evaluate this course; show existing invitations."""
    session = get_or_404_for_window(Session, session_id)
    inviter_teacher = Teacher.query.filter_by(id=session.teacher_id).first()

    if request.method == 'POST':
        try:
            evaluator_teacher_id = int(request.form.get('evaluator_teacher_id'))
            if evaluator_teacher_id == session.teacher_id:
                flash('You cannot invite yourself to evaluate.', 'warning')
                return redirect(url_for('class_management.course_assessment', session_id=session_id))

            # prevent duplicate invites
            existing = EvaluationInvite.query.filter_by(
                session_id=session_id,
                evaluator_teacher_id=evaluator_teacher_id
            ).first()
            if existing:
                if existing.status == 'cancelled':
                    # remove any previous submission
                    EvaluationSubmission.query.filter_by(invite_id=existing.id).delete()
                    existing.status = 'invited'
                    existing.created_at = datetime.utcnow()
                    if getattr(existing, 'window_id', None) is None:
                        existing.window_id = getattr(session, 'window_id', None) or get_effective_window_id()
                    if stamp_window_id:
                        stamp_window_id(existing)
                    db.session.commit()
                    flash('Invitation re-activated.', 'success')
                else:
                    flash('This teacher is already invited for this course.', 'info')
            else:
                invite = EvaluationInvite(
                    session_id=session_id,
                    inviter_teacher_id=session.teacher_id,
                    evaluator_teacher_id=evaluator_teacher_id,
                    status='invited',
                    window_id=getattr(session, 'window_id', None) or get_effective_window_id(),
                )
                if stamp_window_id:
                    stamp_window_id(invite)
                db.session.add(invite)
                db.session.commit()
                flash('Invitation created successfully.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating invitation: {str(e)}', 'danger')
        return redirect(url_for('class_management.course_assessment', session_id=session_id))

    # List of other teachers to invite (exclude self, head, and teaching assistants)
    from role_utils import get_teachers_excluding_head
    all_teachers = get_teachers_excluding_head()
    other_teachers = [t for t in all_teachers if t.id != session.teacher_id]
    # Existing invites for this session
    invites = EvaluationInvite.query.filter_by(session_id=session_id).all()
    session_ids = [inv.session_id for inv in invites]
    sessions_by_id = {s.id: s for s in query_for_window(Session).filter(Session.id.in_(session_ids)).all()} if session_ids else {}
    teacher_ids = {inv.evaluator_teacher_id for inv in invites}
    teachers_by_id = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(teacher_ids)).all()} if teacher_ids else {}

    return render_template(
        'class_management/evaluation_course_assessment.html',
        session=session,
        inviter_teacher=inviter_teacher,
        other_teachers=other_teachers,
        invites=invites,
        sessions_by_id=sessions_by_id,
        teachers_by_id=teachers_by_id
    )

def _pending_teacher_question_threads_query(teacher_id):
    """Threads where the teacher has not acknowledged (teacher_read_at is null) and the latest message is from the student."""
    last_msg_sq = (
        db.session.query(
            CourseQuestionMessage.thread_id.label('t_id'),
            func.max(CourseQuestionMessage.id).label('max_msg_id'),
        )
        .group_by(CourseQuestionMessage.thread_id)
        .subquery()
    )
    latest_msg = aliased(CourseQuestionMessage)
    return (
        CourseQuestionThread.query.join(
            last_msg_sq, CourseQuestionThread.id == last_msg_sq.c.t_id
        )
        .join(latest_msg, latest_msg.id == last_msg_sq.c.max_msg_id)
        .filter(
            CourseQuestionThread.teacher_id == teacher_id,
            CourseQuestionThread.teacher_read_at.is_(None),
            latest_msg.sender_role == 'student',
        )
        .order_by(CourseQuestionThread.updated_at.desc())
    )


# Context processor to inject pending invites count, question notifications, and student notifications
@class_management_bp.app_context_processor
def inject_invites_count():
    from role_utils import parse_roles
    out = {
        'pending_invites_count': 0,
        'question_notification_count': 0,
        'question_notifications': [],
        'student_notification_count': 0,
        'student_notifications': [],
    }
    try:
        if current_user.is_authenticated and has_teacher_privileges(current_user):
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if teacher:
                count = query_for_window(EvaluationInvite).filter_by(
                    evaluator_teacher_id=teacher.id, status='invited'
                ).count()
                exam_count = query_for_window(ExamScrutinizerInvite).filter_by(scrutinizer_teacher_id=teacher.id, status='invited').count()
                split_count = ClassSplitInvite.query.filter_by(invited_teacher_id=teacher.id, status='pending').count()
                out['pending_invites_count'] = count + exam_count + split_count
                pending_q = _pending_teacher_question_threads_query(teacher.id)
                out['question_notification_count'] = pending_q.count()
                unread = pending_q.limit(20).all()
                session_ids = list({t.session_id for t in unread})
                sessions_by_id = {s.id: s for s in query_for_window(Session).filter(Session.id.in_(session_ids)).all()} if session_ids else {}
                out['question_notifications'] = [
                    {
                        'thread_id': t.id,
                        'subject': t.subject,
                        'student_name': t.student_name,
                        'created_at': t.created_at,
                        'session_id': t.session_id,
                        'course_label': (sessions_by_id.get(t.session_id).course_name or sessions_by_id.get(t.session_id).course_code or 'Course') if sessions_by_id.get(t.session_id) else 'Course',
                    }
                    for t in unread
                ]
        if current_user.is_authenticated and parse_roles(current_user.role) and 'student' in parse_roles(current_user.role):
            unread_student = StudentNotification.query.filter_by(
                user_id=current_user.id
            ).filter(
                StudentNotification.read_at.is_(None)
            ).order_by(StudentNotification.created_at.desc()).limit(20).all()
            out['student_notification_count'] = len(unread_student)
            out['student_notifications'] = [
                {'id': n.id, 'type': n.type, 'title': n.title, 'link_url': n.link_url, 'created_at': n.created_at}
                for n in unread_student
            ]
    except Exception:
        pass
    return out


@class_management_bp.route('/notification/<int:notification_id>/read')
@login_required
def student_notification_read(notification_id):
    """Mark a student notification as read and redirect to its target URL."""
    if not (current_user.is_authenticated and parse_roles(current_user.role) and 'student' in parse_roles(current_user.role)):
        return redirect(url_for('index'))
    notification = StudentNotification.query.filter_by(
        id=notification_id,
        user_id=current_user.id
    ).first()
    if notification:
        try:
            notification.read_at = datetime.utcnow()
            db.session.commit()
        except Exception:
            db.session.rollback()
    target = (notification and notification.link_url and notification.link_url.strip()) or None
    if target and (target.startswith('/') or target.startswith('http')):
        return redirect(target)
    return redirect(url_for('class_management.student_view_scores'))


@class_management_bp.route('/student/notifications')
@login_required
def student_notifications_page():
    """Single page for students to see all notifications (read + unread) and open/mark as read."""
    if not (current_user.is_authenticated and parse_roles(current_user.role) and 'student' in parse_roles(current_user.role)):
        return redirect(url_for('index'))
    notifications = (
        StudentNotification.query.filter_by(user_id=current_user.id)
        .order_by(StudentNotification.created_at.desc())
        .limit(50)
        .all()
    )
    unread_count = sum(1 for n in notifications if n.read_at is None)
    return render_template(
        'class_management/student_notifications.html',
        notifications=notifications,
        unread_count=unread_count,
    )


@class_management_bp.route('/student/notifications/mark-all-read', methods=['POST'])
@login_required
def student_notifications_mark_all_read():
    """Mark all student notifications as read and redirect back."""
    if not (current_user.is_authenticated and parse_roles(current_user.role) and 'student' in parse_roles(current_user.role)):
        return redirect(url_for('index'))
    try:
        StudentNotification.query.filter_by(
            user_id=current_user.id,
            read_at=None
        ).update({StudentNotification.read_at: datetime.utcnow()}, synchronize_session=False)
        db.session.commit()
    except Exception:
        db.session.rollback()
    return redirect(url_for('class_management.student_notifications_page'))


@class_management_bp.route('/invitations')
@login_required
def my_invitations():
    """List invitations for the logged-in teacher."""
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher:
        flash('No teacher profile found.', 'warning')
        return redirect(url_for('index'))
    invites = (
        query_for_window(EvaluationInvite)
        .filter_by(evaluator_teacher_id=teacher.id)
        .order_by(EvaluationInvite.created_at.desc())
        .all()
    )
    session_ids = [inv.session_id for inv in invites]
    sessions_by_id = {s.id: s for s in query_for_window(Session).filter(Session.id.in_(session_ids)).all()} if session_ids else {}
    inviter_ids = [inv.inviter_teacher_id for inv in invites]
    inviter_by_id = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(inviter_ids)).all()} if inviter_ids else {}

    exam_invites = query_for_window(ExamScrutinizerInvite).filter_by(scrutinizer_teacher_id=teacher.id).order_by(ExamScrutinizerInvite.created_at.desc()).all()
    exam_entry_ids = [inv.exam_entry_id for inv in exam_invites]
    exam_entries_by_id = {e.id: e for e in query_for_window(ExamPaperEvaluation).filter(ExamPaperEvaluation.id.in_(exam_entry_ids)).all()} if exam_entry_ids else {}
    exam_inviter_ids = [inv.inviter_teacher_id for inv in exam_invites]
    exam_inviter_by_id = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(exam_inviter_ids)).all()} if exam_inviter_ids else {}

    split_invites = ClassSplitInvite.query.filter_by(invited_teacher_id=teacher.id).order_by(ClassSplitInvite.created_at.desc()).all()
    split_sessions_ids = {inv.inviter_session_id for inv in split_invites}
    split_sessions_by_id = {s.id: s for s in query_for_window(Session).filter(Session.id.in_(split_sessions_ids)).all()} if split_sessions_ids else {}
    split_inviter_ids = {inv.inviter_teacher_id for inv in split_invites}
    split_inviter_by_id = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(split_inviter_ids)).all()} if split_inviter_ids else {}
    has_cancelled_invites = (
        any(inv.status == 'cancelled' for inv in invites)
        or any(inv.status == 'cancelled' for inv in exam_invites)
        or any(inv.status == 'cancelled' for inv in split_invites)
    )

    return render_template(
        'class_management/invitations.html',
        invites=invites,
        sessions_by_id=sessions_by_id,
        inviter_by_id=inviter_by_id,
        exam_invites=exam_invites,
        exam_entries_by_id=exam_entries_by_id,
        exam_inviter_by_id=exam_inviter_by_id,
        split_invites=split_invites,
        split_sessions_by_id=split_sessions_by_id,
        split_inviter_by_id=split_inviter_by_id,
        course_scope_labels=COURSE_SCOPE_LABELS,
        has_cancelled_invites=has_cancelled_invites,
    )


@class_management_bp.route('/invitations/clear-cancelled', methods=['POST'])
@login_required
def clear_cancelled_invitations():
    """Remove cancelled invitations from the current teacher's invitation list."""
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher:
        flash('No teacher profile found.', 'warning')
        return redirect(url_for('index'))

    try:
        peer_deleted = query_for_window(EvaluationInvite).filter_by(
            evaluator_teacher_id=teacher.id,
            status='cancelled'
        ).delete(synchronize_session=False)
        exam_deleted = query_for_window(ExamScrutinizerInvite).filter_by(
            scrutinizer_teacher_id=teacher.id,
            status='cancelled'
        ).delete(synchronize_session=False)
        split_deleted = ClassSplitInvite.query.filter_by(
            invited_teacher_id=teacher.id,
            status='cancelled'
        ).delete(synchronize_session=False)
        total_deleted = peer_deleted + exam_deleted + split_deleted
        db.session.commit()

        if total_deleted:
            flash(f'Cleared {total_deleted} cancelled invitation(s).', 'success')
        else:
            flash('No cancelled invitations found to clear.', 'info')
    except Exception as e:
        db.session.rollback()
        flash(f'Error clearing cancelled invitations: {str(e)}', 'danger')

    return redirect(url_for('class_management.my_invitations'))

@class_management_bp.route('/evaluation/<int:session_id>/course-assessment/open/<int:invite_id>', methods=['GET', 'POST'])
@login_required
def course_assessment_form(session_id, invite_id):
    """Invitee fills the class observation report form."""
    invite = EvaluationInvite.query.get_or_404(invite_id)
    if invite.session_id != session_id or invite.status == 'cancelled':
        flash('Invalid invitation.', 'danger')
        return redirect(url_for('index'))

    # Current user must be evaluator
    evaluator = Teacher.query.filter_by(name=current_user.full_name).first()
    if not evaluator or evaluator.id != invite.evaluator_teacher_id:
        flash('You are not authorized for this form.', 'danger')
        return redirect(url_for('index'))

    session = get_or_404_for_window(Session, session_id)
    current_teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not current_teacher or current_teacher.id not in {invite.inviter_teacher_id, invite.evaluator_teacher_id}:
        flash('You are not authorized to view this submission.', 'danger')
        return redirect(url_for('index'))

    submission = EvaluationSubmission.query.filter_by(invite_id=invite.id).first()
    general_data = {}
    score_data = {}
    section_totals = {}
    if submission:
        import json
        try:
            general_data = json.loads(submission.general_info or '{}')
            score_data = json.loads(submission.scores or '{}')
            section_totals = {
                'b1': sum(score_data.get(k, 0) for k in ['b1_a','b1_b','b1_c']),
                'b2': sum(score_data.get(k, 0) for k in ['b2_a','b2_b','b2_c','b2_d','b2_e','b2_f']),
                'b3': sum(score_data.get(k, 0) for k in ['b3_a','b3_b','b3_c','b3_d','b3_e','b3_f','b3_g','b3_h']),
                'b4': sum(score_data.get(k, 0) for k in ['b4_a','b4_b','b4_c'])
            }
            # Backward compatibility for older data
            if 'venue' not in general_data and general_data.get('venue_time'):
                general_data.setdefault('venue', general_data.get('venue_time'))
            general_data.setdefault('session_date', '')
            general_data.setdefault('session_time', '')
        except Exception:
            general_data = {}
            score_data = {}
            section_totals = {}

    general_data.setdefault('program_name', '')
    general_data['teacher_name'] = session.teacher.name if session.teacher else ''
    general_data.setdefault('observer_name', current_user.full_name)
    general_data.setdefault('course_name', session.course_name or '')
    general_data.setdefault('course_code', session.course_code or '')
    general_data.setdefault('course_year', session.year or '')
    general_data.setdefault('course_term', session.term or '')
    general_data.setdefault('academic_session', session.academic_session or '')
    general_data.setdefault('venue', '')
    general_data.setdefault('session_date', '')
    general_data.setdefault('session_time', '')

    if request.method == 'POST':
        try:
            import json
            general = {
                'program_name': request.form.get('program_name'),
                'teacher_name': session.teacher.name if session.teacher else '',
                'observer_name': request.form.get('observer_name'),
                'course_name': request.form.get('course_name') or session.course_name,
                'course_code': request.form.get('course_code') or session.course_code,
                'course_year': request.form.get('course_year') or session.year,
                'course_term': request.form.get('course_term') or session.term,
                'academic_session': request.form.get('academic_session') or session.academic_session,
                'session_date': request.form.get('session_date'),
                'session_time': request.form.get('session_time'),
                'venue': request.form.get('venue')
            }
            # Collect scores
            score_keys = [
                'b1_a','b1_b','b1_c',
                'b2_a','b2_b','b2_c','b2_d','b2_e','b2_f',
                'b3_a','b3_b','b3_c','b3_d','b3_e','b3_f','b3_g','b3_h',
                'b4_a','b4_b','b4_c'
            ]
            scores = {}
            total = 0
            for k in score_keys:
                v = request.form.get(k)
                if v:
                    scores[k] = int(v)
                    total += int(v)
            comments_observer = request.form.get('comments_observer')
            comments_presenter = submission.comments_presenter if submission else None

            if not submission:
                submission = EvaluationSubmission(
                    invite_id=invite.id,
                    session_id=session_id,
                    evaluator_teacher_id=evaluator.id
                )
                db.session.add(submission)

            submission.general_info = json.dumps(general)
            submission.scores = json.dumps(scores)
            submission.comments_observer = comments_observer
            submission.comments_presenter = comments_presenter
            submission.total_score = total
            invite.status = 'submitted'
            db.session.commit()
            flash('Assessment form submitted.', 'success')
            return redirect(url_for('class_management.my_invitations'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error submitting form: {str(e)}', 'danger')

    _bd = bd_now()
    current_date_str = _bd.date().isoformat()
    current_time_str = _bd.strftime('%H:%M')
    return render_template(
        'class_management/evaluation_course_assessment_form.html',
        session=session,
        invite=invite,
        submission=submission,
        general_data=general_data,
        score_data=score_data,
        section_totals=section_totals,
        current_date=current_date_str,
        current_time=current_time_str
    )


@class_management_bp.route('/evaluation/<int:session_id>/student-feedback', methods=['GET', 'POST'])
@login_required
def student_feedback_manage(session_id):
    """Manage anonymous student feedback link and view submissions."""
    session = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session.teacher_id:
        flash('শুধুমাত্র কোর্সটির শিক্ষকই ফিডব্যাক সেটআপ করতে পারবেন।', 'danger')
        return redirect(url_for('class_management.evaluation', session_id=session_id))

    feedback_link = (
        StudentFeedbackLink.query.filter_by(session_id=session_id)
        .order_by(StudentFeedbackLink.created_at.desc())
        .first()
    )

    if request.method == 'POST':
        action = request.form.get('action')
        expires_at = None
        expires_raw = request.form.get('expires_at') or ''
        if expires_raw:
            try:
                expires_at = datetime.strptime(expires_raw, '%Y-%m-%d')
            except ValueError:
                flash('ভ্যালিড মেয়াদ তারিখ দিন (YYYY-MM-DD).', 'warning')
                return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

        allow_multiple = bool(request.form.get('allow_multiple'))
        title = request.form.get('title') or f"{session.course_name or 'Course'} Feedback"
        description = request.form.get('description') or ''

        try:
            if action == 'create':
                if feedback_link:
                    flash('ফিডব্যাক লিংক ইতিমধ্যে বিদ্যমান।', 'info')
                else:
                    feedback_link = StudentFeedbackLink(
                        session_id=session_id,
                        access_code=_generate_feedback_code(),
                        title=title,
                        description=description,
                        expires_at=expires_at,
                        allow_multiple=allow_multiple,
                    )
                    db.session.add(feedback_link)
                    db.session.commit()
                    flash('ফিডব্যাক লিংক তৈরি হয়েছে।', 'success')
            elif action == 'update' and feedback_link:
                feedback_link.title = title
                feedback_link.description = description
                feedback_link.expires_at = expires_at
                feedback_link.allow_multiple = allow_multiple
                db.session.commit()
                flash('সেটিংস আপডেট হয়েছে।', 'success')
            elif action == 'regenerate' and feedback_link:
                feedback_link.access_code = _generate_feedback_code()
                db.session.commit()
                flash('নতুন অ্যাক্সেস কোড তৈরি হয়েছে।', 'info')
            elif action == 'delete' and feedback_link:
                db.session.delete(feedback_link)
                db.session.commit()
                flash('ফিডব্যাক লিংক ও সমস্ত উত্তর মুছে ফেলা হয়েছে।', 'info')
                return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))
            else:
                flash('অজানা অ্যাকশন।', 'danger')
        except Exception as exc:
            db.session.rollback()
            flash(f'ফিডব্যাক সেটআপ পরিবর্তন ব্যর্থ: {exc}', 'danger')

        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    feedback_responses = []
    if feedback_link:
        responses = (
            StudentFeedbackResponse.query.filter_by(feedback_link_id=feedback_link.id)
            .order_by(StudentFeedbackResponse.submitted_at.desc())
            .all()
        )
        for item in responses:
            try:
                answers = json.loads(item.payload or '{}')
            except json.JSONDecodeError:
                answers = {}
            feedback_responses.append(
                {
                    'id': item.id,
                    'is_read': bool(item.is_read),
                    'submitted_at': item.submitted_at,
                    'data': answers,
                }
            )

    feedback_url = (
        url_for('student_feedback_form', code=feedback_link.access_code, _external=True)
        if feedback_link
        else None
    )

    return render_template(
        'class_management/student_feedback_manage.html',
        session=session,
        feedback_link=feedback_link,
        feedback_url=feedback_url,
        feedback_responses=feedback_responses,
        section_a_labels=FEEDBACK_SECTION_A,
        section_b_likert=FEEDBACK_SECTION_B_LIKERT,
        method_options=FEEDBACK_METHOD_OPTIONS,
        effort_options=FEEDBACK_EFFORT_OPTIONS,
    )


@class_management_bp.route('/evaluation/<int:session_id>/student-feedback/responses/<int:response_id>/mark-read', methods=['POST'])
@login_required
def mark_student_feedback_response_read(session_id, response_id):
    """Mark one student feedback response as read by the course teacher."""
    session_obj = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session_obj.teacher_id:
        return jsonify({'ok': False, 'message': 'Unauthorized'}), 403

    feedback_link = StudentFeedbackLink.query.filter_by(session_id=session_id).first()
    if not feedback_link:
        return jsonify({'ok': False, 'message': 'Feedback link not found'}), 404

    response = StudentFeedbackResponse.query.filter_by(
        id=response_id,
        feedback_link_id=feedback_link.id
    ).first()
    if not response:
        return jsonify({'ok': False, 'message': 'Response not found'}), 404

    if not response.is_read:
        try:
            response.is_read = True
            db.session.commit()
        except Exception:
            db.session.rollback()
            return jsonify({'ok': False, 'message': 'Could not update read status'}), 500

    return jsonify({'ok': True})


@class_management_bp.route('/evaluation/<int:session_id>/course-review', methods=['GET', 'POST'])
@login_required
def course_review_form(session_id):
    """Course teacher documents reflections and improvement plans."""
    session = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()

    if not teacher:
        flash('You must be registered as a teacher to access the course review form.', 'danger')
        return redirect(url_for('class_management.index'))

    if session.teacher_id != teacher.id:
        flash('Only the course teacher can access the course review form.', 'danger')
        return redirect(url_for('class_management.evaluation', session_id=session_id))

    review = CourseReview.query.filter_by(session_id=session_id, teacher_id=teacher.id).first()
    review_data = {}
    if review and review.data:
        try:
            review_data = json.loads(review.data)
        except Exception:
            review_data = {}
    if not isinstance(review_data, dict):
        review_data = {}

    if 'discipline' not in review_data and review_data.get('department'):
        review_data['discipline'] = review_data.pop('department')
    if 'school' not in review_data and review_data.get('faculty'):
        review_data['school'] = review_data.pop('faculty')
    if 'course_term' not in review_data and review_data.get('semester_term'):
        review_data['course_term'] = review_data.pop('semester_term')
    if 'course_year' not in review_data and review_data.get('level'):
        review_data['course_year'] = review_data.pop('level')

    review_data.setdefault('discipline', '')
    review_data.setdefault('school', '')
    review_data.setdefault('course_title', session.course_name or '')
    review_data.setdefault('course_code', session.course_code or '')
    review_data.setdefault('session_name', session.academic_session or session.year or '')
    review_data.setdefault('course_year', session.year or '')
    review_data.setdefault('course_term', session.term or '')
    review_data.setdefault('credit_value', '')
    review_data.setdefault('instructor_name', session.teacher.name if session.teacher else '')
    review_data.setdefault('enrollment_count', '')
    for row in COURSE_REVIEW_GRADE_ROWS:
        review_data.setdefault(f"{row['key']}_number", '')
        review_data.setdefault(f"{row['key']}_percentage", '')
    review_data.setdefault('grade_total_number', '')
    review_data.setdefault('grade_total_percentage', '')
    if request.method == 'POST':
        try:
            def _clean(field):
                return (request.form.get(field) or '').strip()
            def _to_number(value):
                try:
                    if value is None or value == '':
                        return None
                    return float(value)
                except (TypeError, ValueError):
                    return None

            def _format_number(value, decimals=None):
                if value is None:
                    return ''
                if decimals is not None:
                    return f"{value:.{decimals}f}"
                if float(value).is_integer():
                    return str(int(round(value)))
                return f"{value:.2f}".rstrip('0').rstrip('.')

            form_data = {
                'discipline': _clean('discipline'),
                'school': _clean('school'),
                'course_code': _clean('course_code'),
                'course_title': _clean('course_title'),
                'session_name': _clean('session_name'),
                'course_year': _clean('course_year'),
                'course_term': _clean('course_term'),
                'credit_value': _clean('credit_value'),
                'instructor_name': _clean('instructor_name'),
                'contact_hours': _clean('contact_hours'),
                'lecture_hours': _clean('lecture_hours'),
                'seminar_hours': _clean('seminar_hours'),
                'other_instruction': _clean('other_instruction'),
                'assessment_methods': _clean('assessment_methods'),
                'enrollment_count': _clean('enrollment_count'),
            }

            enrollment_value = _to_number(form_data['enrollment_count'])
            total_number = 0.0
            has_grade_values = False

            for row in COURSE_REVIEW_GRADE_ROWS:
                num_key = f"{row['key']}_number"
                pct_key = f"{row['key']}_percentage"
                number_value = _to_number(_clean(num_key))
                if number_value is not None:
                    has_grade_values = True
                    total_number += number_value
                    form_data[num_key] = _format_number(number_value)
                    if enrollment_value and enrollment_value > 0:
                        percentage_value = (number_value / enrollment_value) * 100
                        form_data[pct_key] = _format_number(percentage_value, decimals=2)
                    else:
                        form_data[pct_key] = ''
                else:
                    form_data[num_key] = ''
                    form_data[pct_key] = ''

            if has_grade_values:
                form_data['grade_total_number'] = _format_number(total_number)
                if enrollment_value and enrollment_value > 0:
                    total_percentage = (total_number / enrollment_value) * 100
                    form_data['grade_total_percentage'] = _format_number(total_percentage, decimals=2)
                else:
                    form_data['grade_total_percentage'] = ''
            else:
                form_data['grade_total_number'] = ''
                form_data['grade_total_percentage'] = ''

            for item in COURSE_REVIEW_COMMENT_FIELDS:
                form_data[item['key']] = _clean(item['key'])

            if review is None:
                review = CourseReview(session_id=session_id, teacher_id=teacher.id)
                db.session.add(review)

            review.data = json.dumps(form_data)
            db.session.commit()
            flash('Course review saved successfully.', 'success')
            return redirect(url_for('class_management.course_review_form', session_id=session_id))
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error saving course review for session {session_id}: {e}")
            flash(f'Error saving course review: {str(e)}', 'danger')

    has_saved_review = bool(review and review.data)

    return render_template(
        'class_management/evaluation_course_review_form.html',
        session=session,
        review_data=review_data,
        grade_rows=COURSE_REVIEW_GRADE_ROWS,
        comment_fields=COURSE_REVIEW_COMMENT_FIELDS,
        has_saved_review=has_saved_review
    )


@class_management_bp.route('/evaluation/<int:session_id>/course-review/pdf')
@login_required
def course_review_pdf(session_id):
    """Download the Faculty Course Review Report as a PDF."""
    session = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()

    if not teacher or teacher.id != session.teacher_id:
        flash('Only the course teacher can download the course review PDF.', 'danger')
        return redirect(url_for('class_management.evaluation', session_id=session_id))

    review = CourseReview.query.filter_by(session_id=session_id, teacher_id=teacher.id).first()
    if not review or not review.data:
        flash('No saved course review found to generate PDF.', 'warning')
        return redirect(url_for('class_management.course_review_form', session_id=session_id))

    try:
        stored_data = json.loads(review.data)
    except Exception:
        flash('Stored course review data is invalid.', 'danger')
        return redirect(url_for('class_management.course_review_form', session_id=session_id))

    from xml.sax.saxutils import escape

    def get_value(key):
        return escape(str(stored_data.get(key, '') or ''))

    def get_multiline_value(key):
        text = get_value(key)
        if not text:
            return '&nbsp;'
        return text.replace('\n', '<br/>')

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name='CourseReviewTitle',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        fontSize=14,
        leading=16,
        spaceAfter=12
    )
    subtitle_style = ParagraphStyle(
        name='CourseReviewSubtitle',
        parent=styles['BodyText'],
        alignment=TA_CENTER,
        fontSize=9,
        leading=12,
        spaceAfter=18
    )
    label_style = ParagraphStyle(
        name='CourseReviewLabel',
        parent=styles['BodyText'],
        fontSize=9,
        leading=11,
        spaceAfter=4,
        spaceBefore=2
    )
    comment_style = ParagraphStyle(
        name='CourseReviewComment',
        parent=styles['BodyText'],
        fontSize=9,
        leading=12
    )

    elements = []
    elements.append(Paragraph('FACULTY COURSE REVIEW REPORT', title_style))
    elements.append(Paragraph('(To be filled by each teacher at the time of Course Completion)', subtitle_style))

    table_width = doc.width

    other_paragraph = Paragraph(get_multiline_value('other_instruction'), comment_style)
    assessment_paragraph = Paragraph(get_multiline_value('assessment_methods'), comment_style)
    enrollment_count_text = get_value('enrollment_count') or '&nbsp;'
    enrollment_paragraph = Paragraph(
        f'<b>Number of Enrolled Students:</b> {enrollment_count_text}',
        label_style
    )

    info_col_widths = [table_width * 0.2, table_width * 0.3, table_width * 0.2, table_width * 0.3]
    info_data = [
        [Paragraph('Discipline', label_style), Paragraph(get_value('discipline'), comment_style), Paragraph('School', label_style), Paragraph(get_value('school'), comment_style)],
        [Paragraph('Course Code', label_style), Paragraph(get_value('course_code'), comment_style), Paragraph('Title', label_style), Paragraph(get_value('course_title'), comment_style)],
        [Paragraph('Session', label_style), Paragraph(get_value('session_name'), comment_style), Paragraph('Year', label_style), Paragraph(get_value('course_year'), comment_style)],
        [Paragraph('Term', label_style), Paragraph(get_value('course_term'), comment_style), Paragraph('Credit Value', label_style), Paragraph(get_value('credit_value'), comment_style)],
        [Paragraph('Name of Course Instructor', label_style), Paragraph(get_value('instructor_name'), comment_style), Paragraph('No. of Students Contact Hour', label_style), Paragraph(get_value('contact_hours'), comment_style)],
        [Paragraph('Lectures', label_style), Paragraph(get_value('lecture_hours'), comment_style), Paragraph('Seminar', label_style), Paragraph(get_value('seminar_hours'), comment_style)],
        [Paragraph('Other (Please State)', label_style), other_paragraph, '', ''],
        [Paragraph('Assessment Methods: give precise details (no &amp; length of assignments, exams, weightings etc)', label_style), assessment_paragraph, '', '']
    ]

    info_table = Table(info_data, colWidths=info_col_widths, hAlign='CENTER')
    info_table.setStyle(TableStyle([
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.whitesmoke),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('SPAN', (1, 6), (3, 6)),
        ('SPAN', (1, 7), (3, 7)),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('WORDWRAP', (0, 0), (-1, -1), True),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 16))
    elements.append(Paragraph('Distribution of Grade/Marks and Other Outcomes', label_style))
    elements.append(enrollment_paragraph)
    elements.append(Spacer(1, 8))

    grade_table_data = [['Scale', 'Letter Grade', 'Number of Students', '%']]
    for row in COURSE_REVIEW_GRADE_ROWS:
        grade_table_data.append([
            row['scale'],
            row['letter'],
            get_value(f"{row['key']}_number"),
            get_value(f"{row['key']}_percentage"),
        ])
    grade_table_data.append([
        'Total', '',
        get_value('grade_total_number'),
        get_value('grade_total_percentage'),
    ])

    grade_col_widths = [
        table_width * 0.36,
        table_width * 0.14,
        table_width * 0.25,
        table_width * 0.25,
    ]
    grade_table = Table(grade_table_data, colWidths=grade_col_widths, hAlign='CENTER')
    grade_table.setStyle(TableStyle([
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('BACKGROUND', (0, -1), (-1, -1), colors.whitesmoke),
        ('ALIGN', (1, 1), (-1, -2), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('WORDWRAP', (0, 0), (-1, -1), True),
    ]))
    elements.append(grade_table)
    elements.append(Spacer(1, 18))

    elements.append(Paragraph('Overview/Evaluation (Course Coordinator\'s Comments)', label_style))

    comment_table_data = []
    for item in COURSE_REVIEW_COMMENT_FIELDS:
        comment_table_data.append([
            Paragraph(item['label'], label_style),
            Paragraph(get_value(item['key']) or '&nbsp;', comment_style)
        ])

    comment_table = Table(comment_table_data, colWidths=[table_width / 2, table_width / 2], hAlign='CENTER')
    comment_table.setStyle(TableStyle([
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BACKGROUND', (0, 0), (-1, -1), colors.whitesmoke),
        ('BACKGROUND', (1, 0), (1, -1), colors.white),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('WORDWRAP', (0, 0), (-1, -1), True),
    ]))
    elements.append(comment_table)
    elements.append(Spacer(1, 20))

    signature_width = table_width / 2
    signature_table_data = [
        [
            Paragraph('Head of the Discipline Signature', label_style),
            Paragraph('Course Instructor (s) Signature', label_style)
        ],
        [
            Paragraph('&nbsp;' * 4, comment_style),
            Paragraph('&nbsp;' * 4, comment_style)
        ],
        [
            Paragraph('Date: ___________________', label_style),
            Paragraph('Date: ___________________', label_style)
        ]
    ]

    signature_table = Table(
        signature_table_data,
        colWidths=[signature_width, signature_width],
        hAlign='CENTER'
    )
    signature_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('LEFTPADDING', (0, 0), (-1, -1), 24),
        ('RIGHTPADDING', (0, 0), (-1, -1), 24),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LINEABOVE', (0, 1), (0, 1), 0.8, colors.black),
        ('LINEABOVE', (1, 1), (1, 1), 0.8, colors.black),
        ('BOTTOMPADDING', (0, 1), (1, 1), 24),
    ]))

    elements.append(signature_table)

    def _add_page_number(canvas_obj, doc_obj):
        canvas_obj.saveState()
        page_num = canvas_obj.getPageNumber()
        text = f"Page {page_num}"
        canvas_obj.setFont('Helvetica', 9)
        canvas_obj.drawRightString(
            doc_obj.pagesize[0] - doc_obj.rightMargin,
            doc_obj.bottomMargin - 20,
            text
        )
        canvas_obj.restoreState()

    doc.build(elements, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    buffer.seek(0)

    filename = f"course_review_{session.course_code or session.id}.pdf"
    return Response(
        buffer.getvalue(),
        mimetype='application/pdf',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(buffer.getvalue()))
        }
    )


@class_management_bp.route('/evaluation/<int:session_id>/course-assessment/view/<int:invite_id>', methods=['GET', 'POST'])
@login_required
def course_assessment_view(session_id, invite_id):
    """Inviter views a submission and can download PDF."""
    invite = EvaluationInvite.query.get_or_404(invite_id)
    if invite.session_id != session_id:
        flash('Invalid invitation.', 'danger')
        return redirect(url_for('index'))
    session = get_or_404_for_window(Session, session_id)
    current_teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not current_teacher or current_teacher.id not in {invite.inviter_teacher_id, invite.evaluator_teacher_id}:
        flash('You are not authorized to view this submission.', 'danger')
        return redirect(url_for('index'))

    submission = EvaluationSubmission.query.filter_by(invite_id=invite.id).first()
    if not submission:
        flash('No submission yet.', 'info')
        return redirect(url_for('class_management.course_assessment', session_id=session_id))
    import json
    score_data = json.loads(submission.scores or '{}')
    general_data = json.loads(submission.general_info or '{}')
    if 'venue' not in general_data and general_data.get('venue_time'):
        general_data.setdefault('venue', general_data.get('venue_time'))
    general_data.setdefault('session_date', '')
    general_data.setdefault('session_time', '')
    general_data.setdefault('program_name', '')
    general_data['teacher_name'] = session.teacher.name if session.teacher else ''
    general_data.setdefault('observer_name', '')
    general_data.setdefault('course_name', session.course_name or '')
    general_data.setdefault('course_code', session.course_code or '')
    general_data.setdefault('course_year', session.year or '')
    general_data.setdefault('course_term', session.term or '')
    general_data.setdefault('academic_session', session.academic_session or '')
    section_totals = {
        'b1': sum(score_data.get(k, 0) for k in ['b1_a','b1_b','b1_c']),
        'b2': sum(score_data.get(k, 0) for k in ['b2_a','b2_b','b2_c','b2_d','b2_e','b2_f']),
        'b3': sum(score_data.get(k, 0) for k in ['b3_a','b3_b','b3_c','b3_d','b3_e','b3_f','b3_g','b3_h']),
        'b4': sum(score_data.get(k, 0) for k in ['b4_a','b4_b','b4_c'])
    }
    can_edit_presenter = current_teacher and current_teacher.id == session.teacher_id

    if request.method == 'POST':
        if not can_edit_presenter:
            flash('You are not authorized to update presenter comments.', 'danger')
            return redirect(url_for('class_management.course_assessment_view', session_id=session_id, invite_id=invite_id))
        submission.comments_presenter = request.form.get('comments_presenter')
        try:
            db.session.commit()
            flash("Presenter's comments updated.", 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating comments: {str(e)}', 'danger')
        return redirect(url_for('class_management.course_assessment_view', session_id=session_id, invite_id=invite_id))

    return render_template('class_management/evaluation_course_assessment_view.html', session=session, invite=invite, submission=submission, general_data=general_data, score_data=score_data, section_totals=section_totals, can_edit_presenter=can_edit_presenter)

@class_management_bp.route('/evaluation/<int:session_id>/course-assessment/pdf/<int:invite_id>')
@login_required
def course_assessment_pdf(session_id, invite_id):
    """Generate PDF of the submitted assessment."""
    try:
        import json
        from xml.sax.saxutils import escape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        invite = EvaluationInvite.query.get_or_404(invite_id)
        if invite.session_id != session_id:
            flash('Invalid invitation.', 'danger')
            return redirect(url_for('index'))
        session = get_or_404_for_window(Session, session_id)
        current_teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not current_teacher or current_teacher.id not in {invite.inviter_teacher_id, invite.evaluator_teacher_id}:
            flash('You are not authorized to download this report.', 'danger')
            return redirect(url_for('index'))

        submission = EvaluationSubmission.query.filter_by(invite_id=invite.id).first()
        if not submission:
            flash('No submission to export.', 'warning')
            return redirect(url_for('class_management.course_assessment', session_id=session_id))

        general = json.loads(submission.general_info or '{}')
        if 'venue' not in general and general.get('venue_time'):
            general.setdefault('venue', general.get('venue_time'))
        general.setdefault('session_date', '')
        general.setdefault('session_time', '')
        general.setdefault('program_name', '')
        general.setdefault('teacher_name', session.teacher.name if session.teacher else '')
        general.setdefault('observer_name', '')
        general.setdefault('course_name', session.course_name or '')
        general.setdefault('course_code', session.course_code or '')
        general.setdefault('course_year', session.year or '')
        general.setdefault('course_term', session.term or '')
        general.setdefault('academic_session', session.academic_session or '')
        scores = json.loads(submission.scores or '{}')
        section_totals = {
            'b1': sum(scores.get(k, 0) for k in ['b1_a','b1_b','b1_c']),
            'b2': sum(scores.get(k, 0) for k in ['b2_a','b2_b','b2_c','b2_d','b2_e','b2_f']),
            'b3': sum(scores.get(k, 0) for k in ['b3_a','b3_b','b3_c','b3_d','b3_e','b3_f','b3_g','b3_h']),
            'b4': sum(scores.get(k, 0) for k in ['b4_a','b4_b','b4_c'])
        }

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('ReportTitle', parent=styles['Title'], fontSize=16, spaceAfter=12)
        label_style = ParagraphStyle('Label', parent=styles['BodyText'], fontName='Helvetica-Bold', fontSize=10)
        value_style = ParagraphStyle('Value', parent=styles['BodyText'], fontSize=10)
        center_value_style = ParagraphStyle('CenterValue', parent=value_style, alignment=1)
        section_header_style = ParagraphStyle('SectionHeader', parent=styles['BodyText'], fontName='Helvetica-Bold', fontSize=10)
        subtotal_style = ParagraphStyle('Subtotal', parent=styles['BodyText'], fontName='Helvetica-Oblique', fontSize=10)
        comment_header_style = ParagraphStyle('CommentHeader', parent=styles['Heading4'], fontSize=11, spaceBefore=12, spaceAfter=6)

        elements = []
        elements.append(Paragraph('Classroom Teaching Observation Report', title_style))

        info_data = [
            [Paragraph('Program', label_style), Paragraph(escape(general.get('program_name') or '-') , value_style),
             Paragraph('Teacher', label_style), Paragraph(escape(general.get('teacher_name') or '-') , value_style)],
            [Paragraph('Observer', label_style), Paragraph(escape(general.get('observer_name') or '-') , value_style),
             Paragraph('Course Name', label_style), Paragraph(escape(general.get('course_name') or '-') , value_style)],
            [Paragraph('Course Code', label_style), Paragraph(escape(general.get('course_code') or '-') , value_style),
             Paragraph('Academic Session', label_style), Paragraph(escape(general.get('academic_session') or '-') , value_style)],
            [Paragraph('Year', label_style), Paragraph(escape(general.get('course_year') or '-') , value_style),
             Paragraph('Term', label_style), Paragraph(escape(general.get('course_term') or '-') , value_style)],
            [Paragraph('Date', label_style), Paragraph(escape(general.get('session_date') or '-') , value_style),
             Paragraph('Time', label_style), Paragraph(escape(general.get('session_time') or '-') , value_style)],
            [Paragraph('Venue', label_style), Paragraph(escape(general.get('venue') or '-') , value_style),
             Paragraph('', label_style), Paragraph('', value_style)]
        ]
        info_table = Table(info_data, colWidths=[80, 180, 80, 180])
        info_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f7f7f7')),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('SPAN', (2,5), (3,5)),
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 10))
        elements.append(Paragraph('Score Scale: 5 = Excellent, 4 = Very Good, 3 = Good, 2 = Fair, 1 = Poor', styles['BodyText']))
        elements.append(Spacer(1, 8))

        score_rows = [[Paragraph('<b>Description</b>', label_style), Paragraph('<b>Score</b>', center_value_style)]]
        section_headers = []
        span_rows = []
        subtotal_rows = []

        def add_section(title, items, subtotal_key):
            header_idx = len(score_rows)
            score_rows.append([Paragraph(f'<b>{escape(title)}</b>', section_header_style), ''])
            section_headers.append(header_idx)
            span_rows.append(header_idx)
            for text, key in items:
                score_rows.append([
                    Paragraph(escape(text), value_style),
                    Paragraph(str(scores.get(key)) if key in scores else '-', center_value_style)
                ])
            subtotal_idx = len(score_rows)
            score_rows.append([
                Paragraph(f'<i>Subtotal for {escape(title)}</i>', subtotal_style),
                Paragraph(f"<b>{section_totals.get(subtotal_key, 0)}</b>", center_value_style)
            ])
            subtotal_rows.append(subtotal_idx)

        add_section('Section 1: Set Induction', [
            ('a) Clarity of objectives', 'b1_a'),
            ('b) Relevance to topic', 'b1_b'),
            ('c) Appropriateness of introduction', 'b1_c')
        ], 'b1')

        add_section('Section 2: Content', [
            ('a) Knowledge', 'b2_a'),
            ('b) Extend of coverage', 'b2_b'),
            ('c) Level of interest generated', 'b2_c'),
            ('d) Logical flow of presentation', 'b2_d'),
            ('e) Correctness of language used', 'b2_e'),
            ('f) Clear and relevant use of analogies/examples', 'b2_f')
        ], 'b2')

        add_section('Section 3: Presentation', [
            ('a) Appropriate pacing', 'b3_a'),
            ('b) Confidence', 'b3_b'),
            ('c) Enthusiasm', 'b3_c'),
            ('d) Provoking students to think', 'b3_d'),
            ('e) Clarity of presentation', 'b3_e'),
            ('f) Interaction with students', 'b3_f'),
            ('g) Effective use of teaching/learning aids', 'b3_g'),
            ('h) Effective class management', 'b3_h')
        ], 'b3')

        add_section('Section 4: Closure', [
            ('a) Appropriateness of closure', 'b4_a'),
            ('b) Effective questions for feedback', 'b4_b'),
            ('c) Appropriate links to the next lesson', 'b4_c')
        ], 'b4')

        score_rows.append([
            Paragraph('<b>Total Score</b>', label_style),
            Paragraph(f"<b>{submission.total_score or 0}</b>", center_value_style)
        ])

        score_table = Table(score_rows, colWidths=[360, 80])
        score_style = TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0f0f0')),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ALIGN', (1,0), (1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
        ])
        for idx in section_headers:
            score_style.add('SPAN', (0, idx), (1, idx))
            score_style.add('BACKGROUND', (0, idx), (-1, idx), colors.HexColor('#e5e5e5'))
            score_style.add('FONTNAME', (0, idx), (-1, idx), 'Helvetica-Bold')
        for idx in subtotal_rows:
            score_style.add('BACKGROUND', (0, idx), (-1, idx), colors.HexColor('#f9f9f9'))
        score_style.add('BACKGROUND', (0, len(score_rows)-1), (-1, len(score_rows)-1), colors.HexColor('#e8f4ff'))
        score_table.setStyle(score_style)
        elements.append(score_table)

        elements.append(Spacer(1, 12))

        interpret_data = [
            [Paragraph('<b>Score</b>', label_style), Paragraph('<b>Interpretation</b>', label_style)],
            ['90 to 100', 'Excellent'],
            ['80 to less than 90', 'Very Good'],
            ['70 to less than 80', 'Good'],
            ['60 to less than 70', 'Fair'],
            ['50 to less than 60', 'Poor'],
            ['40 to less than 50', 'Very Poor']
        ]
        interpret_table = Table(interpret_data, colWidths=[140, 200])
        interpret_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0f0f0')),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER')
        ]))
        elements.append(interpret_table)

        observer_comment = (submission.comments_observer or '').strip()
        presenter_comment = (submission.comments_presenter or '').strip()
        elements.append(Paragraph("Observer's Overall Comments and Suggestions for Improvement", comment_header_style))
        elements.append(Paragraph(escape(observer_comment).replace('\n', '<br/>') or '-', styles['BodyText']))
        elements.append(Paragraph("Presenter's Comments", comment_header_style))
        elements.append(Paragraph(escape(presenter_comment).replace('\n', '<br/>') or '-', styles['BodyText']))

        elements.append(Spacer(1, 60))
        line_width = (A4[0] - 72) / 2
        signature_table = Table(
            [['', ''], ["Presenter's Signature", "Observer's Signature"]],
            colWidths=[line_width, line_width]
        )
        signature_table.setStyle(TableStyle([
            ('LINEABOVE', (0,0), (0,0), 0.7, colors.black),
            ('LINEABOVE', (1,0), (1,0), 0.7, colors.black),
            ('TOPPADDING', (0,0), (-1,0), 30),
            ('ALIGN', (0,1), (-1,1), 'CENTER'),
            ('TOPPADDING', (0,1), (-1,1), 8)
        ]))
        elements.append(signature_table)

        class NumberedCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self._saved_page_states = []

            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                num_pages = len(self._saved_page_states)
                for state in self._saved_page_states:
                    self.__dict__.update(state)
                    self.draw_page_number(num_pages)
                    super().showPage()
                super().save()

            def draw_page_number(self, page_count):
                self.setFont('Helvetica', 9)
                text = f"Page {self._pageNumber} of {page_count}"
                self.drawRightString(A4[0] - 36, 30, text)

        doc.build(elements, canvasmaker=NumberedCanvas)
        buffer.seek(0)
        filename = f"course_assessment_{session.course_code or session.id}.pdf"
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
    except Exception as e:
        flash(f'Error generating PDF: {str(e)}', 'danger')
        return redirect(url_for('class_management.course_assessment', session_id=session_id))

@class_management_bp.route('/invitation/<int:invite_id>/cancel', methods=['POST'])
@login_required
def cancel_invitation(invite_id):
    """Allow either inviter or invitee to cancel an invitation."""
    invite = EvaluationInvite.query.get_or_404(invite_id)

    # Resolve current teacher for the logged-in user
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher:
        flash('No teacher profile found.', 'warning')
        return redirect(url_for('index'))

    if teacher.id not in [invite.inviter_teacher_id, invite.evaluator_teacher_id]:
        flash('You are not authorized to cancel this invitation.', 'danger')
        return redirect(url_for('index'))

    try:
        invite.status = 'cancelled'
        EvaluationSubmission.query.filter_by(invite_id=invite.id).delete()
        db.session.commit()
        flash('Invitation has been cancelled.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error cancelling invitation: {str(e)}', 'danger')

    # Redirect back to sensible page
    if teacher.id == invite.inviter_teacher_id:
        return redirect(url_for('class_management.course_assessment', session_id=invite.session_id))
    else:
        return redirect(url_for('class_management.my_invitations'))

FEEDBACK_SECTION_A = [
    ('course_structure', 'Course content is structured in a comprehensible manner.'),
    ('course_goals', 'The goals of the course are clear.'),
    ('course_content_guidance', 'The course contents are explained in an understandable fashion.'),
    ('course_interest', 'The course fosters my interest in the discussed topics.'),
]

FEEDBACK_SECTION_B_LIKERT = [
    ('course_plan_discussed', 'Course plan (assessment criteria/ content) was discussed in advance.'),
    ('guidelines_received', 'Received oral instruction and written guidelines for continuous assessment.'),
    ('assessment_helpful', 'Course projects/assignments/tests were helpful to demonstrate an understanding of the course material.'),
    ('feedback_timely', 'Received feedback and grades of continuous assessments from course teacher in due time.'),
]

FEEDBACK_METHOD_OPTIONS = [
    'Lectures (including online lectures)',
    'Class discussions (including online discussion boards)',
    'In-class learning activities (other than discussion)',
    'In-class clickers or other quick response methods',
    'Homework (readings and assignments)',
    'Labs',
    'Projects or portfolios',
    'Teamwork or group activities',
    'Student presentations',
    'Guest lecturers',
    'Fieldwork/field trips',
    'Mentoring outside of the classroom',
    'Support from Teaching/ Research Assistants',
    'Others'
]

FEEDBACK_EFFORT_OPTIONS = [
    'Memorizing facts and repeating ideas from the readings and lectures.',
    'Making judgments about the value of information, arguments, or methods.',
    'Applying basic elements of an idea, experience, or theory.',
    'Applying theories or concepts to practical problems or in new situations.',
    'Synthesizing and organizing ideas, information, or experiences.',
    'Solving problems.',
    'Thinking creatively or critically.',
    'Teamwork or group activities.',
    'Doing lab work.',
    'Presenting in person or via a recording.',
    'Reading and writing for deep understanding.',
    'Others'
]

@class_management_bp.route('/evaluation/<int:session_id>/student-feedback/responses/pdf')
@login_required
def student_feedback_responses_pdf(session_id):
    session_obj = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session_obj.teacher_id:
        flash('You are not authorized to access this download.', 'danger')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    feedback_link = StudentFeedbackLink.query.filter_by(session_id=session_id).first()
    if not feedback_link:
        flash('No feedback responses found.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    responses = (
        StudentFeedbackResponse.query.filter_by(feedback_link_id=feedback_link.id)
        .order_by(StudentFeedbackResponse.submitted_at.asc())
        .all()
    )
    if not responses:
        flash('No feedback responses to download.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    buffer = io.BytesIO()
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    response_font = 'Helvetica'
    response_bold_font = 'Helvetica-Bold'
    kalpurush_available = False
    try:
        font_root = os.path.join(current_app.root_path, 'static', 'fonts')
        regular_candidates = ['Kalpurush.ttf', 'Kalpurush-Regular.ttf']
        bold_candidates = ['Kalpurush-Bold.ttf', 'Kalpurush Bold.ttf']
        regular_path = next((os.path.join(font_root, f) for f in regular_candidates if os.path.exists(os.path.join(font_root, f))), None)
        bold_path = next((os.path.join(font_root, f) for f in bold_candidates if os.path.exists(os.path.join(font_root, f))), None)
        if regular_path:
            pdfmetrics.registerFont(TTFont('Kalpurush', regular_path))
            response_font = 'Kalpurush'
            kalpurush_available = True
        if bold_path:
            pdfmetrics.registerFont(TTFont('Kalpurush-Bold', bold_path))
            response_bold_font = 'Kalpurush-Bold'
        elif regular_path:
            response_bold_font = 'Kalpurush'
    except Exception as exc:  # pragma: no cover
        current_app.logger.warning('Kalpurush font registration failed: %s', exc)

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=28,
        bottomMargin=28,
        leftMargin=32,
        rightMargin=32,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], alignment=1, fontSize=16, leading=18, spaceAfter=6, textTransform='uppercase', fontName='Helvetica-Bold')
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Heading2'], alignment=1, fontSize=13, leading=16, spaceAfter=12, textTransform='uppercase', fontName='Helvetica-Bold')
    section_header_style = ParagraphStyle('SectionHeader', parent=styles['Heading3'], fontSize=11, leading=13, spaceBefore=8, spaceAfter=4, textTransform='uppercase', fontName='Helvetica-Bold')
    label_style = ParagraphStyle('Label', parent=styles['Normal'], fontSize=9, leading=11, fontName='Helvetica-Bold', wordWrap='CJK')
    instruction_style = ParagraphStyle('Instruction', parent=styles['Normal'], fontSize=9, leading=11, alignment=1, textTransform='uppercase', spaceBefore=6, spaceAfter=6, wordWrap='CJK', fontName='Helvetica-Bold')
    value_style = ParagraphStyle('Value', parent=styles['Normal'], fontSize=9.5, leading=11, fontName=response_font, wordWrap='CJK')
    value_bold_style = ParagraphStyle('ValueBold', parent=value_style, fontName=response_bold_font)
    # Style for Praise and Suggestions section - always use Kalpurush if available
    praise_suggestions_font = 'Kalpurush' if kalpurush_available else response_font
    praise_suggestions_style = ParagraphStyle('PraiseSuggestions', parent=styles['Normal'], fontSize=9.5, leading=11, fontName=praise_suggestions_font, wordWrap='CJK')

    likert_header = ['Statement', 'Strongly disagree', 'Disagree', 'Neutral', 'Agree', 'Strongly Agree']

    def likert_table(section_values, labels):
        data = [[Paragraph(text, label_style) for text in likert_header]]
        for key, question in labels:
            selected = (section_values or {}).get(key)
            row = [Paragraph(question, value_style)]
            for option in likert_header[1:]:
                mark = '✓' if selected and selected.lower() == option.lower() else ''
                row.append(Paragraph(mark, value_style))
            data.append(row)
        table = Table(data, colWidths=[240, 52, 52, 52, 52, 52])
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEADING', (0, 0), (-1, -1), 11),
        ]))
        return table

    def checklist_table(options, selected):
        rows = []
        selected = selected or []
        for idx in range(0, len(options), 2):
            row = []
            for offset in (0, 1):
                if idx + offset < len(options):
                    option = options[idx + offset]
                    mark = '✓' if option in selected else ''
                    row.append(Paragraph(f"{mark} {option}", value_style))
                else:
                    row.append('')
            rows.append(row)
        table = Table(rows, colWidths=[255, 255])
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        return table

    elements = []
    for idx, item in enumerate(responses, start=1):
        try:
            data = json.loads(item.payload or '{}')
        except json.JSONDecodeError:
            data = {}

        academic = data.get('academic_info', {}) or {}
        section_a = data.get('section_a', {}) or {}
        section_b = data.get('section_b', {}) or {}
        section_c = data.get('section_c', {}) or {}
        section_d = data.get('section_d', {}) or {}

        if idx > 1:
            elements.append(PageBreak())

        elements.append(Paragraph('Student Feedback Form', title_style))
        elements.append(Paragraph('Khulna University', subtitle_style))
        elements.append(Paragraph(f"Response {idx} - {format_bd(item.submitted_at, '%Y-%m-%d %H:%M')}", value_bold_style))
        elements.append(Spacer(1, 6))

        info_data = [
            [Paragraph('Academic Session', label_style), Paragraph(academic.get('academic_session') or '—', value_style)],
            [Paragraph('Title of the Course', label_style), Paragraph(academic.get('course_title') or session_obj.course_name or '—', value_style)],
            [Paragraph('Course Code', label_style), Paragraph(academic.get('course_code') or session_obj.course_code or '—', value_style)],
        ]
        info_table = Table(info_data, colWidths=[140, 360])
        info_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (-1, -1), colors.whitesmoke),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
        ]))
        elements.append(info_table)
        elements.append(Paragraph('Please tick/cross in the blank space which best describes how much you agree with the following statements', instruction_style))

        elements.append(Paragraph('A. Satisfaction with the Course', section_header_style))
        elements.append(likert_table(section_a, FEEDBACK_SECTION_A))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('B. Teaching-Learning Methods', section_header_style))
        elements.append(likert_table(section_b, [FEEDBACK_SECTION_B_LIKERT[0]]))
        elements.append(Spacer(1, 4))

        methods = section_b.get('teaching_methods') or []
        if methods:
            elements.append(Paragraph('Teaching methods that contributed significantly (selected):', value_bold_style))
            elements.append(checklist_table(FEEDBACK_METHOD_OPTIONS, methods))
            elements.append(Spacer(1, 4))

        elements.append(likert_table(section_b, FEEDBACK_SECTION_B_LIKERT[1:]))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('C. Engagement and Workload', section_header_style))
        engagement_table = Table(
            [
                [Paragraph('How much time do you devote to this course before and after each lecture?', label_style), Paragraph(section_c.get('study_time') or '—', value_style)],
                [Paragraph('About what percent of the class meetings (including discussions) did you attend?', label_style), Paragraph(section_c.get('attendance_percent') or '—', value_style)],
            ],
            colWidths=[360, 140]
        )
        engagement_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(engagement_table)
        elements.append(Spacer(1, 4))

        efforts = section_c.get('effort_focus') or []
        elements.append(Paragraph('Significant aspects of your effort (selected):', value_bold_style))
        elements.append(checklist_table(FEEDBACK_EFFORT_OPTIONS, efforts))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('D. Praise and Suggestions', section_header_style))
        def open_block(title, content):
            # Use Kalpurush font for Praise and Suggestions content
            content_style = praise_suggestions_style
            table = Table(
                [
                    [Paragraph(title, label_style)],
                    [Paragraph(content or '—', content_style)]
                ],
                colWidths=[500]
            )
            table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 12)
            ]))
            return table

        elements.append(open_block('What did you like especially about this course?', section_d.get('likes')))
        elements.append(Spacer(1, 4))
        elements.append(open_block('What are the challenges you have faced in attending the course?', section_d.get('challenges')))
        elements.append(Spacer(1, 4))
        elements.append(open_block('Suggestions on how to improve the course:', section_d.get('suggestions')))

    doc.build(elements)
    pdf_data = buffer.getvalue()
    buffer.close()

    filename = f"student_feedback_responses_{session_obj.course_code or 'course'}.pdf"
    return Response(
        pdf_data,
        mimetype='application/pdf',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(pdf_data)),
        },
    )

@class_management_bp.route('/evaluation/<int:session_id>/student-feedback/responses/pdf-weasyprint')
@login_required
def student_feedback_responses_pdf_weasyprint(session_id):
    """Generate student feedback PDF using WeasyPrint with Kalpurush for Praise and Suggestions."""
    # Lazy import WeasyPrint - only when actually needed
    HTML = _get_weasyprint_html()
    if HTML is None:
        error_msg = 'Error generating PDF: WeasyPrint is not available. '
        error_msg += 'Please ensure WeasyPrint dependencies are installed. '
        error_msg += 'On macOS, run: brew install cairo pango gdk-pixbuf gobject-introspection'
        flash(error_msg, 'error')
        current_app.logger.error("WeasyPrint not available for PDF generation")
        current_app.logger.error(f"Current availability status: {_WEASYPRINT_AVAILABLE}")
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))
    
    session_obj = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session_obj.teacher_id:
        flash('You are not authorized to access this download.', 'danger')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    feedback_link = StudentFeedbackLink.query.filter_by(session_id=session_id).first()
    if not feedback_link:
        flash('No feedback responses found.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    responses = (
        StudentFeedbackResponse.query.filter_by(feedback_link_id=feedback_link.id)
        .order_by(StudentFeedbackResponse.submitted_at.asc())
        .all()
    )
    if not responses:
        flash('No feedback responses to download.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    try:
        from error_handler import log_error
        import os
        
        # Prepare data for template
        feedback_data = []
        for idx, item in enumerate(responses, start=1):
            try:
                data = json.loads(item.payload or '{}')
            except json.JSONDecodeError:
                data = {}
            
            academic = data.get('academic_info', {}) or {}
            section_a = data.get('section_a', {}) or {}
            section_b = data.get('section_b', {}) or {}
            section_c = data.get('section_c', {}) or {}
            section_d = data.get('section_d', {}) or {}
            
            feedback_data.append({
                'index': idx,
                'submitted_at': format_bd(item.submitted_at, '%Y-%m-%d'),
                'academic': academic,
                'section_a': section_a,
                'section_b': section_b,
                'section_c': section_c,
                'section_d': section_d,
                'course_name': academic.get('course_title') or session_obj.course_name or '—',
                'course_code': academic.get('course_code') or session_obj.course_code or '—',
                'academic_session': academic.get('academic_session') or '—',
            })
        
        # Get fonts for WeasyPrint (Liberation for English body; Kalpurush for Bengali comments)
        from utils.pdf_fonts import resolve_formal_pdf_fonts
        import os

        formal_fonts = resolve_formal_pdf_fonts()
        if not formal_fonts:
            flash(
                'PDF fonts missing. Upload LiberationSerif-*.ttf to static/fonts/.',
                'error',
            )
            return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

        font_path = os.path.join(current_app.root_path, 'static', 'Fonts', 'kalpurush.ttf')
        if not os.path.exists(font_path):
            font_path = os.path.join(current_app.root_path, 'static', 'fonts', 'kalpurush.ttf')
        
        # Render template
        html_content = render_template(
            'class_management/student_feedback_weasyprint.html',
            feedback_data=feedback_data,
            feedback_section_a=FEEDBACK_SECTION_A,
            feedback_section_b_likert=FEEDBACK_SECTION_B_LIKERT,
            feedback_method_options=FEEDBACK_METHOD_OPTIONS,
            feedback_effort_options=FEEDBACK_EFFORT_OPTIONS,
            likert_options=['Strongly disagree', 'Disagree', 'Neutral', 'Agree', 'Strongly Agree'],
            kalpurush_font_path=font_path if os.path.exists(font_path) else None,
            pdf_font_regular=formal_fonts['regular'],
            pdf_font_bold=formal_fonts['bold'],
            pdf_font_italic=formal_fonts.get('italic'),
            pdf_font_bold_italic=formal_fonts.get('bold_italic'),
        )
        
        # Generate PDF with WeasyPrint (lazy import already done above)
        try:
            pdf_buffer = io.BytesIO()
            HTML(string=html_content, base_url=formal_fonts['fonts_dir'].as_uri() + '/').write_pdf(pdf_buffer)
            pdf_buffer.seek(0)
        except Exception as e:
            current_app.logger.error(f"Error generating PDF with WeasyPrint: {e}", exc_info=True)
            flash(f'Error generating PDF: {str(e)}', 'error')
            return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))
        
        filename = f"student_feedback_responses_{session_obj.course_code or 'course'}.pdf"
        
        current_app.logger.info(f"WeasyPrint PDF generated successfully for student feedback session {session_id}")
        return Response(
            pdf_buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(pdf_buffer.getvalue()))
            }
        )
        
    except Exception as e:
        log_error(e, {
            'session_id': session_id,
            'function': 'student_feedback_responses_pdf_weasyprint',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        current_app.logger.error(f"Error generating WeasyPrint student feedback PDF for session {session_id}: {e}", exc_info=True)
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

@class_management_bp.route('/evaluation/<int:session_id>/student-feedback/responses/docx')
@login_required
def student_feedback_responses_docx(session_id):
    session_obj = get_or_404_for_window(Session, session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session_obj.teacher_id:
        flash('You are not authorized to access this download.', 'danger')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    feedback_link = StudentFeedbackLink.query.filter_by(session_id=session_id).first()
    if not feedback_link:
        flash('No feedback responses found.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    responses = (
        StudentFeedbackResponse.query.filter_by(feedback_link_id=feedback_link.id)
        .order_by(StudentFeedbackResponse.submitted_at.asc())
        .all()
    )
    if not responses:
        flash('No feedback responses to download.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    document = Document()
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Kalpurush'
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Kalpurush')

    heading1 = document.styles['Heading 1']
    heading1.font.name = 'Helvetica'
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = document.styles['Heading 2']
    heading2.font.name = 'Helvetica'
    heading2.font.size = Pt(13)
    heading2.font.bold = True

    for idx, item in enumerate(responses, start=1):
        try:
            data = json.loads(item.payload or '{}')
        except json.JSONDecodeError:
            data = {}

        academic = data.get('academic_info', {}) or {}
        section_a = data.get('section_a', {}) or {}
        section_b = data.get('section_b', {}) or {}
        section_c = data.get('section_c', {}) or {}
        section_d = data.get('section_d', {}) or {}

        if idx > 1:
            document.add_page_break()

        title_para = document.add_paragraph('STUDENT FEEDBACK FORM')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(16)

        uni_para = document.add_paragraph('KHULNA UNIVERSITY')
        uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        uni_para.runs[0].bold = True
        uni_para.runs[0].font.size = Pt(12)

        meta_para = document.add_paragraph(f"Response {idx} - {format_bd(item.submitted_at, '%Y-%m-%d %H:%M')}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.runs[0].font.size = Pt(10)

        add_header_table([
            ('Academic Session', academic.get('academic_session') or '—'),
            ('Title of the Course', academic.get('course_title') or session_obj.course_name or '—'),
            ('Course Code', academic.get('course_code') or session_obj.course_code or '—'),
        ])

        instruction_para = document.add_paragraph(
            'PLEASE TICK/CROSS IN THE BLANK SPACE WHICH BEST DESCRIBES HOW MUCH YOU AGREE WITH THE FOLLOWING STATEMENTS'
        )
        instruction_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        instruction_para.runs[0].font.size = Pt(9)

        document.add_paragraph('A. SATISFACTION WITH THE COURSE', style='Heading 3')
        section_a_rows = []
        for key, question in FEEDBACK_SECTION_A:
            selected = section_a.get(key)
            row = [question]
            for option in likert_header[1:]:
                row.append('✓' if selected and selected.lower() == option.lower() else '')
            section_a_rows.append(row)
        add_likert_table(section_a_rows, likert_header)

        document.add_paragraph('B. TEACHING-LEARNING METHODS', style='Heading 3')
        first_question = FEEDBACK_SECTION_B_LIKERT[0]
        first_selected = section_b.get(first_question[0])
        first_rows = [[first_question[1]] + [('✓' if first_selected and first_selected.lower() == opt.lower() else '') for opt in likert_header[1:]]]
        add_likert_table(first_rows, likert_header)

        methods_heading = document.add_paragraph('Teaching methods that contributed significantly (selected):')
        methods_heading.runs[0].bold = True
        methods = section_b.get('teaching_methods') or []
        if methods:
            for method in methods:
                bullet = document.add_paragraph(method, style='List Bullet')
                bullet.paragraph_format.space_after = Pt(1)
        else:
            document.add_paragraph('—', style='List Bullet')

        remaining_rows = []
        for key, question in FEEDBACK_SECTION_B_LIKERT[1:]:
            selected = section_b.get(key)
            row = [question]
            for option in likert_header[1:]:
                row.append('✓' if selected and selected.lower() == option.lower() else '')
            remaining_rows.append(row)
        add_likert_table(remaining_rows, likert_header)

        document.add_paragraph('C. ENGAGEMENT AND WORKLOAD', style='Heading 3')
        engagement_rows = [
            ('How much time do you devote to this course before and after each lecture?', section_c.get('study_time') or '—'),
            ('About what percent of the class meetings (including discussions) did you attend?', section_c.get('attendance_percent') or '—'),
        ]
        engagement_table = document.add_table(rows=len(engagement_rows), cols=2)
        engagement_table.style = 'Table Grid'
        for row_idx, (label, value) in enumerate(engagement_rows):
            set_cell_shading(engagement_table.cell(row_idx, 0), 'D9D9D9')
            set_cell_text(engagement_table.cell(row_idx, 0), label, bold=True, align='left')
            set_cell_text(engagement_table.cell(row_idx, 1), value, align='left')

        effort_heading = document.add_paragraph('Significant aspects of your effort (selected):')
        effort_heading.runs[0].bold = True
        efforts = section_c.get('effort_focus') or []
        if efforts:
            for effort in efforts:
                bullet = document.add_paragraph(effort, style='List Bullet')
                bullet.paragraph_format.space_after = Pt(1)
        else:
            document.add_paragraph('—', style='List Bullet')

        document.add_paragraph('D. PRAISE AND SUGGESTIONS', style='Heading 3')
        add_open_block('What did you like especially about this course?', section_d.get('likes'))
        add_open_block('What are the challenges you have faced in attending the course?', section_d.get('challenges'))
        add_open_block('Suggestions on how to improve the course:', section_d.get('suggestions'))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    docx_data = buffer.getvalue()
    buffer.close()

    filename = f"student_feedback_responses_{session_obj.course_code or 'course'}.docx"
    return Response(
        docx_data,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(docx_data)),
        },
    )

@class_management_bp.route('/student/course-files')
@login_required
def student_course_files():
    """Student view for course files - course outlines and teacher-uploaded files"""
    from utils.dashboard_settings import require_student_dashboard_card
    blocked = require_student_dashboard_card('course_files')
    if blocked:
        return blocked
    try:
        # Get student ID from current user (assuming username is student_id)
        student_id = current_user.username if hasattr(current_user, 'username') else None
        if not student_id:
            flash('Student ID not found.', 'error')
            return redirect(url_for('index'))
        
        # Find all ClassStudent records for this student
        student_records = ClassStudent.query.filter_by(student_id=student_id).all()
        
        # Get all sessions where student is enrolled
        enrolled_sessions = []
        for student_record in student_records:
            session_obj = get_for_window(Session, student_record.session_id)
            if session_obj and not session_obj.archived:
                enrolled_sessions.append(session_obj)
        
        # Import CourseFileUpload model
        from blueprints.class_management.models import CourseFileUpload
        
        # Get course outlines that are enabled for student access
        course_files = []
        for session in enrolled_sessions:
            course_outline = CourseOutline.query.filter_by(session_id=session.id).first()
            if course_outline and course_outline.student_access_enabled:
                course_files.append({
                    'session_id': session.id,
                    'course_code': session.course_code,
                    'course_name': session.course_name,
                    'teacher_name': session.teacher.name if session.teacher else 'Unknown',
                    'academic_session': session.academic_session,
                    'year': session.year,
                    'term': session.term,
                    'type': 'course_outline',
                    'file_name': f"{session.course_code or 'Course'}_Outline.pdf"
                })
            
            # Get teacher-uploaded files for this session
            uploaded_files = CourseFileUpload.query.filter_by(
                session_id=session.id,
                student_access_enabled=True
            ).all()
            
            for uploaded_file in uploaded_files:
                course_files.append({
                    'session_id': session.id,
                    'course_code': session.course_code,
                    'course_name': session.course_name,
                    'teacher_name': session.teacher.name if session.teacher else 'Unknown',
                    'academic_session': session.academic_session,
                    'year': session.year,
                    'term': session.term,
                    'type': 'uploaded_file',
                    'file_id': uploaded_file.id,
                    'file_name': uploaded_file.file_name,
                    'description': uploaded_file.description
                })
        
        # Sort by academic session, year, term
        course_files.sort(key=lambda x: (
            x.get('academic_session', ''),
            x.get('year', ''),
            x.get('term', '')
        ), reverse=True)
        
        return render_template('class_management/student_course_files.html',
                             course_files=course_files)
    except Exception as e:
        current_app.logger.error(f"Error in student_course_files: {e}", exc_info=True)
        flash('An error occurred while loading course files.', 'error')
        return redirect(url_for('index'))

@class_management_bp.route('/student/course-files/<int:session_id>/download-pdf')
@login_required
def student_download_course_outline_pdf(session_id):
    """Student download course outline PDF"""
    try:
        # Get student ID from current user
        student_id = current_user.username if hasattr(current_user, 'username') else None
        if not student_id:
            flash('Student ID not found.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Check if student is enrolled in this session
        student_record = ClassStudent.query.filter_by(
            session_id=session_id,
            student_id=student_id
        ).first()
        
        if not student_record:
            flash('You are not enrolled in this course.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Get course outline
        course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
        if not course_outline:
            flash('Course outline not found.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Check if student access is enabled
        if not course_outline.student_access_enabled:
            flash('This course outline is not available for download.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Generate and return PDF (skip auth check since we already verified student access)
        return _generate_course_outline_pdf(session_id, skip_auth_check=True)
    except Exception as e:
        current_app.logger.error(f"Error in student_download_course_outline_pdf: {e}", exc_info=True)
        flash('An error occurred while downloading the course outline.', 'error')
        return redirect(url_for('class_management.student_course_files'))

@class_management_bp.route('/student/course-files/<int:file_id>/download')
@login_required
def student_download_uploaded_file(file_id):
    """Student download teacher-uploaded file"""
    try:
        from blueprints.class_management.models import CourseFileUpload
        import os
        from flask import send_file
        
        # Get student ID from current user
        student_id = current_user.username if hasattr(current_user, 'username') else None
        if not student_id:
            flash('Student ID not found.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Get uploaded file
        uploaded_file = CourseFileUpload.query.get_or_404(file_id)
        
        # Check if student access is enabled
        if not uploaded_file.student_access_enabled:
            flash('This file is not available for download.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Check if student is enrolled in this session
        student_record = ClassStudent.query.filter_by(
            session_id=uploaded_file.session_id,
            student_id=student_id
        ).first()
        
        if not student_record:
            flash('You are not enrolled in this course.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Check if file exists
        file_path = uploaded_file.file_path
        if not os.path.exists(file_path):
            flash('File not found.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Send file for download
        return send_file(
            file_path,
            as_attachment=True,
            download_name=uploaded_file.file_name
        )
    except Exception as e:
        current_app.logger.error(f"Error in student_download_uploaded_file: {e}", exc_info=True)
        flash('An error occurred while downloading the file.', 'error')
        return redirect(url_for('class_management.student_course_files'))


@class_management_bp.route('/question-bank')
@login_required
def question_bank():
    """Past question bank page grouped by folder name."""
    from utils.dashboard_settings import require_student_dashboard_card
    blocked = require_student_dashboard_card('question_bank')
    if blocked:
        return blocked
    # Optional filters via querystring (GET)
    folder_filter = (request.args.get('folder_name') or '').strip()
    year_filter = (request.args.get('question_year') or '').strip()

    query = QuestionBankFile.query
    if folder_filter:
        query = query.filter(QuestionBankFile.subject_name.ilike(f"%{folder_filter}%"))
    if year_filter:
        query = query.filter(QuestionBankFile.question_year == year_filter)

    files = query.order_by(
        QuestionBankFile.subject_name.asc(),
        QuestionBankFile.question_year.desc(),
        QuestionBankFile.created_at.desc()
    ).all()
    folders = QuestionBankFolder.query.order_by(QuestionBankFolder.name.asc()).all()
    can_upload = current_user.is_authenticated
    can_manage = has_teacher_privileges(current_user) or is_admin(current_user)
    grouped_files = {}
    folder_options = []
    for folder in folders:
        folder_name = _qb_normalize_folder_path(folder.name)
        if not folder_name:
            continue
        folder_parts = folder_name.split('/')
        folder_options.append({
            'value': folder_name,
            'label': folder_name,
            'depth': max(len(folder_parts) - 1, 0),
            'basename': folder_parts[-1],
        })
        grouped_files[folder_name] = {
            'folder_name': folder_name,
            'folder_basename': folder_parts[-1] if folder_parts else folder_name,
            'folder_depth': max(len(folder_parts) - 1, 0),
            'files': []
        }
    for f in files:
        folder_label = _qb_normalize_folder_path(f.subject_name) or 'Untitled Folder'
        folder_key = folder_label
        if folder_key not in grouped_files:
            folder_parts = folder_label.split('/')
            folder_options.append({
                'value': folder_label,
                'label': folder_label,
                'depth': max(len(folder_parts) - 1, 0),
                'basename': folder_parts[-1],
            })
            grouped_files[folder_key] = {
                'folder_name': folder_label,
                'folder_basename': folder_parts[-1] if folder_parts else folder_label,
                'folder_depth': max(len(folder_parts) - 1, 0),
                'files': []
            }
        grouped_files[folder_key]['files'].append(f)
    grouped_items = list(grouped_files.values())
    grouped_items.sort(key=lambda item: (item.get('folder_name') or '').lower())

    # Build nested folder tree from slash-delimited folder paths.
    tree_nodes = {}
    for item in grouped_items:
        folder_path = _qb_normalize_folder_path(item.get('folder_name'))
        if not folder_path:
            continue
        dom_id = ''.join(ch.lower() if ch.isalnum() else '-' for ch in folder_path).strip('-')
        if not dom_id:
            dom_id = f"qb-folder-{len(tree_nodes) + 1}"
        tree_nodes[folder_path] = {
            'folder_name': folder_path,
            'folder_basename': item.get('folder_basename') or folder_path.split('/')[-1],
            'folder_depth': item.get('folder_depth', 0),
            'files': item.get('files', []),
            'children': [],
            'dom_id': dom_id,
        }

    folder_tree = []
    for folder_path, node in sorted(tree_nodes.items(), key=lambda x: (len(x[0].split('/')), x[0].lower())):
        parent_path = '/'.join(folder_path.split('/')[:-1])
        parent_node = tree_nodes.get(parent_path)
        if parent_node:
            parent_node['children'].append(node)
        else:
            folder_tree.append(node)

    def _sort_children(nodes):
        def _qb_year_sort_key(node):
            label = (node.get('folder_basename') or node.get('folder_name') or '').lower()
            # Normalize common inputs
            label = label.replace('year', '').replace('term', '').strip()

            # Order requested:
            # First Year -> Second Year -> Third Year -> Fourth Year -> LLM
            if 'llm' in label:
                order = 4
            elif 'first' in label or label.startswith('1'):
                order = 0
            elif 'second' in label or label.startswith('2'):
                order = 1
            elif 'third' in label or label.startswith('3'):
                order = 2
            elif 'fourth' in label or label.startswith('4'):
                order = 3
            else:
                order = 5
            return (order, label)
        for n in nodes:
            _sort_children(n.get('children', []))

        nodes.sort(key=_qb_year_sort_key)

    _sort_children(folder_tree)

    # Compute total file count per folder including everything in its subfolders,
    # so a parent folder (e.g. "Third Year") reflects files stored deeper in the
    # tree instead of showing "0 file(s)".
    def _compute_total_files(node):
        total = len(node.get('files') or [])
        for child in node.get('children', []):
            total += _compute_total_files(child)
        node['total_files'] = total
        return total

    for _root_node in folder_tree:
        _compute_total_files(_root_node)

    folder_options = sorted(
        {f"{opt['label']}::{opt['value']}": opt for opt in folder_options}.values(),
        key=lambda opt: (opt.get('label') or '').lower()
    )
    return render_template(
        'class_management/question_bank.html',
        grouped_files=grouped_items,
        folder_tree=folder_tree,
        folder_options=folder_options,
        can_upload=can_upload,
        can_manage=can_manage,
        filters={
            'folder_name': folder_filter,
            'question_year': year_filter
        }
    )


def _qb_normalize_folder_path(raw_value):
    value = str(raw_value or '').replace('\\', '/').strip()
    if not value:
        return ''
    parts = [p.strip() for p in value.split('/') if p and p.strip()]
    return '/'.join(parts)


def _qb_is_descendant(path, maybe_ancestor):
    path = _qb_normalize_folder_path(path).lower()
    maybe_ancestor = _qb_normalize_folder_path(maybe_ancestor).lower()
    if not path or not maybe_ancestor:
        return False
    return path == maybe_ancestor or path.startswith(maybe_ancestor + '/')


def _qb_subtree_filter(column_attr, folder_path):
    normalized = _qb_normalize_folder_path(folder_path)
    lowered = normalized.lower()
    return or_(
        func.lower(column_attr) == lowered,
        func.lower(column_attr).like(f"{lowered}/%")
    )


@class_management_bp.route('/question-bank/folder/create', methods=['POST'])
@login_required
def create_question_bank_folder():
    """Teacher/admin can create an empty folder without uploading files."""
    if not (has_teacher_privileges(current_user) or is_admin(current_user)):
        flash('You are not authorized to create folders.', 'error')
        return redirect(url_for('class_management.question_bank'))

    folder_name = (request.form.get('folder_name') or '').strip()
    parent_folder = _qb_normalize_folder_path(request.form.get('parent_folder'))
    if not folder_name:
        flash('Folder name is required.', 'error')
        return redirect(url_for('class_management.question_bank'))

    full_folder_name = _qb_normalize_folder_path(f"{parent_folder}/{folder_name}" if parent_folder else folder_name)
    if not full_folder_name:
        flash('Folder name is invalid.', 'error')
        return redirect(url_for('class_management.question_bank'))

    existing = QuestionBankFolder.query.filter(
        func.lower(QuestionBankFolder.name) == full_folder_name.lower()
    ).first()
    if existing:
        flash('Folder already exists.', 'info')
        return redirect(url_for('class_management.question_bank'))

    try:
        db.session.add(
            QuestionBankFolder(
                name=full_folder_name,
                created_by_user_id=getattr(current_user, 'id', None)
            )
        )
        db.session.commit()
        flash('Folder created successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating question bank folder: {e}", exc_info=True)
        flash('Failed to create folder.', 'error')

    return redirect(url_for('class_management.question_bank'))


@class_management_bp.route('/question-bank/folder/rename', methods=['POST'])
@login_required
def rename_question_bank_folder():
    """Rename a folder path and cascade rename to child folders/files."""
    if not (has_teacher_privileges(current_user) or is_admin(current_user)):
        flash('You are not authorized to rename folders.', 'error')
        return redirect(url_for('class_management.question_bank'))

    folder_name = _qb_normalize_folder_path(request.form.get('folder_name'))
    new_name_input = (request.form.get('new_folder_name') or '').strip()
    if not folder_name or not new_name_input:
        flash('Both current folder and new folder name are required.', 'error')
        return redirect(url_for('class_management.question_bank'))

    parent_path = '/'.join(folder_name.split('/')[:-1])
    new_root = _qb_normalize_folder_path(f"{parent_path}/{new_name_input}" if parent_path else new_name_input)
    if not new_root:
        flash('New folder name is invalid.', 'error')
        return redirect(url_for('class_management.question_bank'))
    if new_root.lower() == folder_name.lower():
        flash('Folder name is unchanged.', 'info')
        return redirect(url_for('class_management.question_bank'))

    subtree_folders = QuestionBankFolder.query.filter(_qb_subtree_filter(QuestionBankFolder.name, folder_name)).all()
    if not subtree_folders:
        flash('Folder not found.', 'error')
        return redirect(url_for('class_management.question_bank'))

    all_folder_names = {
        _qb_normalize_folder_path(f.name).lower()
        for f in QuestionBankFolder.query.all()
        if _qb_normalize_folder_path(f.name)
    }
    old_subtree_names = {
        _qb_normalize_folder_path(f.name).lower()
        for f in subtree_folders
        if _qb_normalize_folder_path(f.name)
    }

    for folder in subtree_folders:
        old_path = _qb_normalize_folder_path(folder.name)
        suffix = old_path[len(folder_name):]
        candidate = _qb_normalize_folder_path(f"{new_root}{suffix}")
        if candidate.lower() in all_folder_names and candidate.lower() not in old_subtree_names:
            flash(f'Cannot rename: target path "{candidate}" already exists.', 'error')
            return redirect(url_for('class_management.question_bank'))

    subtree_files = QuestionBankFile.query.filter(_qb_subtree_filter(QuestionBankFile.subject_name, folder_name)).all()

    try:
        for folder in subtree_folders:
            old_path = _qb_normalize_folder_path(folder.name)
            suffix = old_path[len(folder_name):]
            folder.name = _qb_normalize_folder_path(f"{new_root}{suffix}")

        for qb_file in subtree_files:
            old_path = _qb_normalize_folder_path(qb_file.subject_name)
            suffix = old_path[len(folder_name):]
            qb_file.subject_name = _qb_normalize_folder_path(f"{new_root}{suffix}")

        db.session.commit()
        flash('Folder renamed successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error renaming question bank folder: {e}", exc_info=True)
        flash('Failed to rename folder.', 'error')

    return redirect(url_for('class_management.question_bank'))


@class_management_bp.route('/question-bank/folder/delete', methods=['POST'])
@login_required
def delete_question_bank_folder():
    """Delete a folder recursively (subfolders + files)."""
    if not (has_teacher_privileges(current_user) or is_admin(current_user)):
        flash('You are not authorized to delete folders.', 'error')
        return redirect(url_for('class_management.question_bank'))

    folder_name = _qb_normalize_folder_path(request.form.get('folder_name'))
    if not folder_name:
        flash('Folder name is required.', 'error')
        return redirect(url_for('class_management.question_bank'))

    folders_to_delete = QuestionBankFolder.query.filter(_qb_subtree_filter(QuestionBankFolder.name, folder_name)).all()
    files_to_delete = QuestionBankFile.query.filter(_qb_subtree_filter(QuestionBankFile.subject_name, folder_name)).all()
    if not folders_to_delete and not files_to_delete:
        flash('Folder not found.', 'warning')
        return redirect(url_for('class_management.question_bank'))

    for qb_file in files_to_delete:
        try:
            if qb_file.file_path and os.path.exists(qb_file.file_path):
                os.remove(qb_file.file_path)
        except Exception:
            pass

    try:
        for qb_file in files_to_delete:
            db.session.delete(qb_file)
        for folder in folders_to_delete:
            db.session.delete(folder)
        db.session.commit()
        flash('Folder and its contents deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting question bank folder: {e}", exc_info=True)
        flash('Failed to delete folder.', 'error')

    return redirect(url_for('class_management.question_bank'))


@class_management_bp.route('/question-bank/file/move', methods=['POST'])
@login_required
def move_question_bank_files():
    """Move selected files to another folder."""
    if not (has_teacher_privileges(current_user) or is_admin(current_user)):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    payload = request.get_json(silent=True) or request.form
    target_folder = _qb_normalize_folder_path((payload.get('target_folder') or '').strip())
    raw_file_ids = payload.get('file_ids') or ''
    if isinstance(raw_file_ids, str):
        file_ids = [int(x) for x in raw_file_ids.split(',') if x.strip().isdigit()]
    elif isinstance(raw_file_ids, list):
        file_ids = [int(x) for x in raw_file_ids if str(x).isdigit()]
    else:
        file_ids = []

    if not target_folder:
        return jsonify({'success': False, 'message': 'Target folder is required.'}), 400
    if not file_ids:
        return jsonify({'success': False, 'message': 'No file selected.'}), 400

    target_exists = QuestionBankFolder.query.filter(
        func.lower(QuestionBankFolder.name) == target_folder.lower()
    ).first()
    if not target_exists:
        return jsonify({'success': False, 'message': 'Target folder does not exist.'}), 400

    files = QuestionBankFile.query.filter(QuestionBankFile.id.in_(sorted(set(file_ids)))).all()
    if not files:
        return jsonify({'success': False, 'message': 'Selected files were not found.'}), 404

    try:
        for qb_file in files:
            qb_file.subject_name = target_folder
        db.session.commit()
        return jsonify({'success': True, 'message': f'Moved {len(files)} file(s).'})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error moving question bank files: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to move files.'}), 500


@class_management_bp.route('/question-bank/folder/move', methods=['POST'])
@login_required
def move_question_bank_folders():
    """Move selected folder(s) under another folder."""
    if not (has_teacher_privileges(current_user) or is_admin(current_user)):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    payload = request.get_json(silent=True) or request.form
    target_folder = _qb_normalize_folder_path((payload.get('target_folder') or '').strip())
    raw_folder_paths = payload.get('folder_paths') or ''
    if isinstance(raw_folder_paths, str):
        folder_paths = [_qb_normalize_folder_path(x) for x in raw_folder_paths.split(',') if _qb_normalize_folder_path(x)]
    elif isinstance(raw_folder_paths, list):
        folder_paths = [_qb_normalize_folder_path(x) for x in raw_folder_paths if _qb_normalize_folder_path(x)]
    else:
        folder_paths = []

    if not target_folder:
        return jsonify({'success': False, 'message': 'Target folder is required.'}), 400
    if not folder_paths:
        return jsonify({'success': False, 'message': 'No folder selected.'}), 400

    target_exists = QuestionBankFolder.query.filter(
        func.lower(QuestionBankFolder.name) == target_folder.lower()
    ).first()
    if not target_exists:
        return jsonify({'success': False, 'message': 'Target folder does not exist.'}), 400

    unique_sources = []
    for source in sorted(set(folder_paths), key=lambda p: len(p.split('/'))):
        if source.lower() == target_folder.lower() or _qb_is_descendant(target_folder, source):
            return jsonify({'success': False, 'message': f'Cannot move "{source}" into itself or its child.'}), 400
        if any(_qb_is_descendant(source, kept) for kept in unique_sources):
            continue
        unique_sources.append(source)

    all_folder_names = {
        _qb_normalize_folder_path(f.name).lower()
        for f in QuestionBankFolder.query.all()
        if _qb_normalize_folder_path(f.name)
    }

    move_plan = []
    for source in unique_sources:
        source_exists = QuestionBankFolder.query.filter(
            func.lower(QuestionBankFolder.name) == source.lower()
        ).first()
        if not source_exists:
            return jsonify({'success': False, 'message': f'Source folder not found: {source}'}), 404
        basename = source.split('/')[-1]
        destination_root = _qb_normalize_folder_path(f"{target_folder}/{basename}")
        if destination_root.lower() in all_folder_names and destination_root.lower() != source.lower():
            return jsonify({'success': False, 'message': f'Target already contains "{basename}".'}), 400
        move_plan.append((source, destination_root))

    try:
        for source, destination_root in move_plan:
            subtree_folders = QuestionBankFolder.query.filter(_qb_subtree_filter(QuestionBankFolder.name, source)).all()
            subtree_files = QuestionBankFile.query.filter(_qb_subtree_filter(QuestionBankFile.subject_name, source)).all()
            for folder in subtree_folders:
                old_path = _qb_normalize_folder_path(folder.name)
                suffix = old_path[len(source):]
                folder.name = _qb_normalize_folder_path(f"{destination_root}{suffix}")
            for qb_file in subtree_files:
                old_path = _qb_normalize_folder_path(qb_file.subject_name)
                suffix = old_path[len(source):]
                qb_file.subject_name = _qb_normalize_folder_path(f"{destination_root}{suffix}")
        db.session.commit()
        return jsonify({'success': True, 'message': f'Moved {len(move_plan)} folder(s).'})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error moving question bank folders: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to move folder(s).'}), 500


@class_management_bp.route('/question-bank/upload', methods=['POST'])
@login_required
def upload_question_bank_file():
    """Upload one or multiple question PDFs (any logged-in user)."""

    folder_name = (request.form.get('folder_name') or '').strip()
    question_year = (request.form.get('question_year') or '').strip()
    title = (request.form.get('title') or '').strip()  # optional (used as common title/prefix)

    if not folder_name or not question_year:
        flash('Folder name and year are required.', 'error')
        return redirect(url_for('class_management.question_bank'))

    # Upload allowed only inside existing folders (for all users).
    existing_folder = QuestionBankFolder.query.filter(
        func.lower(QuestionBankFolder.name) == folder_name.lower()
    ).first()
    if not existing_folder:
        flash('Please upload inside an existing folder.', 'error')
        return redirect(url_for('class_management.question_bank'))

    try:
        files = request.files.getlist('files')
        if not files:
            # backward compatibility: old single input name
            legacy_file = request.files.get('file')
            if legacy_file:
                files = [legacy_file]

        valid_files = [f for f in files if f and f.filename]
        if not valid_files:
            flash('Please choose at least one PDF file.', 'error')
            return redirect(url_for('class_management.question_bank'))

        uploaded_count = 0
        for file in valid_files:
            ext = os.path.splitext(file.filename)[1].lower()
            if ext != '.pdf':
                flash(f'Skipped non-PDF file: {file.filename}', 'warning')
                continue

            safe_name = secure_filename(file.filename)
            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
            stored_name = f"{timestamp}_{uuid4().hex}_{safe_name}"
            file_path = os.path.join(QUESTION_BANK_UPLOAD_FOLDER, stored_name)
            file.save(file_path)

            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else None
            file_stem = os.path.splitext(safe_name)[0]
            effective_title = title if len(valid_files) == 1 and title else (f"{title} - {file_stem}" if title else file_stem)

            entry = QuestionBankFile(
                subject_name=folder_name,
                course_code=None,
                question_year=question_year,
                title=effective_title,
                file_path=file_path,
                file_size=file_size,
                uploaded_by_user_id=getattr(current_user, 'id', None)
            )
            db.session.add(entry)
            uploaded_count += 1

        if uploaded_count == 0:
            flash('No valid PDF files were uploaded.', 'error')
            return redirect(url_for('class_management.question_bank'))

        db.session.commit()
        flash(f'{uploaded_count} question PDF file(s) uploaded successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error uploading question bank PDF: {e}", exc_info=True)
        flash('Failed to upload question PDF.', 'error')

    return redirect(url_for('class_management.question_bank'))


@class_management_bp.route('/question-bank/download/<int:file_id>')
@login_required
def download_question_bank_file(file_id):
    """Download a question bank PDF (available to logged-in users)."""
    qb_file = QuestionBankFile.query.get_or_404(file_id)
    if not qb_file.file_path or not os.path.exists(qb_file.file_path):
        flash('File not found.', 'error')
        return redirect(url_for('class_management.question_bank'))
    return send_file(
        qb_file.file_path,
        as_attachment=True,
        download_name=f"{qb_file.title}.pdf" if not qb_file.title.lower().endswith('.pdf') else qb_file.title
    )


@class_management_bp.route('/question-bank/folder/download')
@login_required
def download_question_bank_folder():
    """Bulk download all PDFs from a folder as a ZIP file."""
    try:
        folder_name = (request.args.get('folder_name') or '').strip()
        if not folder_name:
            flash('Folder name is required for bulk download.', 'error')
            return redirect(url_for('class_management.question_bank'))

        files = QuestionBankFile.query.filter(
            func.lower(QuestionBankFile.subject_name) == folder_name.lower()
        ).order_by(
            QuestionBankFile.question_year.desc(),
            QuestionBankFile.created_at.desc()
        ).all()

        if not files:
            flash('No files found in this folder.', 'warning')
            return redirect(url_for('class_management.question_bank'))

        import zipfile
        zip_buffer = io.BytesIO()
        added = 0

        # Some cPanel builds lack zlib for ZIP_DEFLATED; fallback to ZIP_STORED.
        compression_method = zipfile.ZIP_DEFLATED
        try:
            with zipfile.ZipFile(zip_buffer, mode='w', compression=compression_method) as zipf:
                for f in files:
                    if not f.file_path or not os.path.exists(f.file_path):
                        continue
                    base_title = (f.title or f"question_{f.id}").strip()
                    safe_base = secure_filename(base_title) or f"question_{f.id}"
                    year_part = (f.question_year or '').strip()
                    arcname = f"{year_part}_{safe_base}.pdf" if year_part else f"{safe_base}.pdf"
                    if arcname in zipf.namelist():
                        arcname = f"{os.path.splitext(arcname)[0]}_{f.id}.pdf"
                    zipf.write(f.file_path, arcname=arcname)
                    added += 1
        except Exception:
            zip_buffer = io.BytesIO()
            added = 0
            with zipfile.ZipFile(zip_buffer, mode='w', compression=zipfile.ZIP_STORED) as zipf:
                for f in files:
                    if not f.file_path or not os.path.exists(f.file_path):
                        continue
                    base_title = (f.title or f"question_{f.id}").strip()
                    safe_base = secure_filename(base_title) or f"question_{f.id}"
                    year_part = (f.question_year or '').strip()
                    arcname = f"{year_part}_{safe_base}.pdf" if year_part else f"{safe_base}.pdf"
                    if arcname in zipf.namelist():
                        arcname = f"{os.path.splitext(arcname)[0]}_{f.id}.pdf"
                    zipf.write(f.file_path, arcname=arcname)
                    added += 1

        if added == 0:
            flash('No downloadable files found in this folder.', 'warning')
            return redirect(url_for('class_management.question_bank'))

        folder_safe = secure_filename(folder_name) or "question_bank_folder"
        zip_buffer.seek(0)
        zip_filename = f"{folder_safe}.zip"

        # Flask compatibility: try modern, then legacy, then raw Response.
        try:
            return send_file(
                zip_buffer,
                as_attachment=True,
                download_name=zip_filename,
                mimetype='application/zip'
            )
        except TypeError:
            try:
                return send_file(
                    zip_buffer,
                    as_attachment=True,
                    attachment_filename=zip_filename,
                    mimetype='application/zip'
                )
            except TypeError:
                data = zip_buffer.getvalue()
                return Response(
                    data,
                    mimetype='application/zip',
                    headers={
                        'Content-Disposition': f'attachment; filename="{zip_filename}"',
                        'Content-Length': str(len(data)),
                    }
                )
    except Exception as e:
        current_app.logger.error(f"Error bulk-downloading question bank folder: {e}", exc_info=True)
        flash('Bulk download failed. Please try again.', 'error')
        return redirect(url_for('class_management.question_bank'))


@class_management_bp.route('/question-bank/download-zip', methods=['POST'])
@login_required
def download_question_bank_zip():
    """Download selected/all question files as a single ZIP (cPanel-safe response)."""
    try:
        raw_ids = (request.form.get('file_ids') or '').strip()
        folder_name = (request.form.get('folder_name') or '').strip()
        if not raw_ids:
            flash('No files selected for ZIP download.', 'warning')
            return redirect(url_for('class_management.question_bank'))

        # Parse CSV ids safely.
        parsed_ids = []
        for tok in raw_ids.split(','):
            tok = tok.strip()
            if tok.isdigit():
                parsed_ids.append(int(tok))
        file_ids = sorted(set(parsed_ids))
        if not file_ids:
            flash('No valid files selected for ZIP download.', 'warning')
            return redirect(url_for('class_management.question_bank'))

        files = QuestionBankFile.query.filter(QuestionBankFile.id.in_(file_ids)).all()
        if not files:
            flash('Selected files were not found.', 'warning')
            return redirect(url_for('class_management.question_bank'))

        import zipfile
        zip_buffer = io.BytesIO()
        added = 0
        with zipfile.ZipFile(zip_buffer, mode='w', compression=zipfile.ZIP_STORED) as zipf:
            for f in files:
                if not f.file_path or not os.path.exists(f.file_path):
                    continue
                base_title = (f.title or f"question_{f.id}").strip()
                safe_base = secure_filename(base_title) or f"question_{f.id}"
                year_part = (f.question_year or '').strip()
                arcname = f"{year_part}_{safe_base}.pdf" if year_part else f"{safe_base}.pdf"
                if arcname in zipf.namelist():
                    arcname = f"{os.path.splitext(arcname)[0]}_{f.id}.pdf"
                zipf.write(f.file_path, arcname=arcname)
                added += 1

        if added == 0:
            flash('No downloadable files found in selection.', 'warning')
            return redirect(url_for('class_management.question_bank'))

        zip_data = zip_buffer.getvalue()
        name_base = secure_filename(folder_name) or "question_bank_selection"
        zip_filename = f"{name_base}.zip"
        return Response(
            zip_data,
            mimetype='application/zip',
            headers={
                'Content-Disposition': f'attachment; filename="{zip_filename}"',
                'Content-Length': str(len(zip_data)),
            }
        )
    except Exception as e:
        current_app.logger.error(f"Error generating question bank ZIP: {e}", exc_info=True)
        flash('ZIP download failed. Please try again.', 'error')
        return redirect(url_for('class_management.question_bank'))


@class_management_bp.route('/question-bank/file/<int:file_id>/delete', methods=['POST'])
@login_required
def delete_question_bank_file(file_id):
    """Delete a question bank file entry and remove its PDF from disk."""
    if not (has_teacher_privileges(current_user) or is_admin(current_user)):
        flash('You are not authorized to delete question bank files.', 'error')
        return redirect(url_for('class_management.question_bank'))

    qb_file = QuestionBankFile.query.get_or_404(file_id)
    try:
        if qb_file.file_path and os.path.exists(qb_file.file_path):
            os.remove(qb_file.file_path)
    except Exception:
        # Even if disk delete fails, keep DB consistent by removing the record.
        pass

    try:
        db.session.delete(qb_file)
        db.session.commit()
        flash('Question paper deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting question bank file: {e}", exc_info=True)
        flash('Failed to delete question paper.', 'error')

    return redirect(url_for('class_management.question_bank'))


@class_management_bp.route('/question-bank/file/<int:file_id>/edit-title', methods=['POST'])
@login_required
def edit_question_bank_file_title(file_id):
    """Edit the display title (name) of a question bank file."""
    if not (has_teacher_privileges(current_user) or is_admin(current_user)):
        flash('You are not authorized to edit question bank files.', 'error')
        return redirect(url_for('class_management.question_bank'))

    qb_file = QuestionBankFile.query.get_or_404(file_id)
    new_title = (request.form.get('title') or '').strip()

    if not new_title:
        flash('Title cannot be empty.', 'error')
        return redirect(url_for('class_management.question_bank'))

    try:
        qb_file.title = new_title
        db.session.commit()
        flash('Title updated successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error editing question bank title: {e}", exc_info=True)
        flash('Failed to update title.', 'error')

    return redirect(url_for('class_management.question_bank'))